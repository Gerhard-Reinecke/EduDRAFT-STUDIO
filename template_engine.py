# ======================================================================================
# template_engine.py
# ======================================================================================
# Module: Institutional Style Transfer & Document Blueprint Engine
#
# System: EduDraft Studio
# Version: 5.5 (BS5 Identity Fidelity Milestone)
#
# --------------------------------------------------------------------------------------
# PURPOSE
# --------------------------------------------------------------------------------------
# This module is the core document intelligence and style-transfer engine of EduDraft Studio.
#
# It constructs new, institution-aligned teaching documents by:
# - extracting visual and structural identity from donor documents, and
# - applying that identity to draft-authored educational content.
#
# The donor document is treated strictly as a source of institutional style.
# The draft is treated as the sole authority of educational meaning and structure.
#
# The system does not transform donor documents — it builds new documents from first principles
# using extracted institutional signals and draft-defined intent.
#
# --------------------------------------------------------------------------------------
# CORE RESPONSIBILITIES
# --------------------------------------------------------------------------------------
# 1. Donor Style Extraction
#    - Extracts structural, visual, and branding signals from donor files
#    - Supports DOCX, PDF, and PPTX ingestion
#    - Captures page setup, typography, headers/footers, tables, and images (binary-safe)
#
# 2. Donor Interpretation Engine (DIE)
#    - Classifies extracted signals into identity, framing, and structural categories
#    - Separates institutional identity from subject-specific academic content
#    - Produces inheritance-safe interpretation layers
#
# 3. Profile Normalisation
#    - Converts raw donor evidence into a stable, reusable schema
#    - Preserves binary assets (e.g. logos) across serialization boundaries
#    - Provides deterministic structure for downstream blueprint construction
#
# 4. Draft Analysis
#    - Interprets teacher-authored content structure and intent
#    - Detects question formats, subparts, marks, visuals, and answer-space requirements
#    - Builds the Draft Content Model (source of document truth)
#
# 5. Appropriateness & Inheritance Control
#    - Determines which donor features may be applied to the draft
#    - Prevents subject-specific layout leakage (e.g. maths column structures)
#    - Applies structure-based decision logic rather than subject assumptions
#
# 6. Blueprint Construction (BS5 Contract Surface)
#    - Combines donor style intelligence with draft structure
#    - Produces a deterministic document blueprint
#    - Includes identity blocks, layout plans, question structures, and visual slots
#
# 7. Rendering Integration (Front-Matter Ownership Model)
#    - Ensures institutional identity (including logos) is rendered within the correct layout zones
#    - Front matter is constructed explicitly from blueprint data, not donor copying
#    - Maintains separation between:
#         • first-page institutional identity
#         • running headers/footers
#         • draft-driven content
#
# --------------------------------------------------------------------------------------
# PIPELINE STAGES
# --------------------------------------------------------------------------------------
# Stage 1: Donor Extraction
#     Raw signal capture (including binary image extraction)
#
# Stage 2: Donor Interpretation Engine (DIE)
#     Classification into identity, framing, and structure
#
# Stage 3: Profile Normalisation
#     Stable schema construction with persistence-safe encoding
#
# Stage 4: Draft Analysis
#     Structural interpretation of teacher-authored content
#
# Stage 5: Appropriateness Engine
#     Filtering and decision-making on feature inheritance
#
# Stage 6: Blueprint Construction
#     Creation of a new document plan (authoritative contract surface)
#
# Stage 7: Rendering & Validation
#     Controlled rendering into DOCX/PPTX outputs with identity fidelity
#
# --------------------------------------------------------------------------------------
# DESIGN PRINCIPLES
# --------------------------------------------------------------------------------------
# - Draft is truth:
#     All educational meaning, structure, and purpose come from the draft
#
# - Donor is style:
#     The donor contributes only institutional identity and presentation logic
#
# - Construct, do not transform:
#     Documents are built from blueprint logic, not edited donor templates
#
# - Identity fidelity:
#     Institutional branding (including logos) must survive the full pipeline intact
#
# - Structure-driven decisions:
#     Layout is determined by draft structure, not donor subject patterns
#
# - Deterministic and explainable:
#     All decisions are traceable and logged with reasoning
#
# - Safe inheritance:
#     Institutional signals may be reused; academic content must never leak
#
# --------------------------------------------------------------------------------------
# KEY SUBSYSTEMS
# --------------------------------------------------------------------------------------
# - Donor extraction and parsing helpers
# - Branding and binary image persistence system
# - Donor Interpretation Engine (DIE)
# - Profile normalisation and bundle safety layer
# - Draft analysis and question modelling
# - Appropriateness and inheritance decision engine
# - Blueprint construction (BS5 contract surface)
# - Front-matter rendering integration
#
# --------------------------------------------------------------------------------------
# DEPENDENCIES
# --------------------------------------------------------------------------------------
# - Gradio → UI workflow integration
# - Supabase → template storage and persistence
# - llm.py → optional reasoning support
# - ingest_docx / ingest_pdf / ingest_pptx → donor ingestion
# - python-docx → DOCX rendering and style application
# - PIL / pytesseract → optional OCR-assisted identity extraction
# - Standard library → parsing, typing, logging, schema handling
#
# --------------------------------------------------------------------------------------
# SYSTEM STATUS
# --------------------------------------------------------------------------------------
# - BS5 blueprint architecture implemented
# - Identity fidelity pipeline (including logos) operational end-to-end
# - Front-matter rendering aligned with blueprint ownership model
# - Appropriateness engine integrated with structure-based decisions
# - Ongoing refinement focused on:
#     • layout generalisation across diverse donors
#     • rendering fidelity (print-ready quality)
#     • answer-space accuracy and visual alignment
#
# --------------------------------------------------------------------------------------
# NOTES
# --------------------------------------------------------------------------------------
# - This module is the core intelligence layer of EduDraft Studio
# - It governs how institutional identity is learned, preserved, and applied
# - Any changes must be validated against the Design Contract and BS5 architecture
#
# ======================================================================================


import json
import os
import re
import uuid
import tempfile
import logging
from collections import Counter
from typing import Any, Dict, List, Tuple, Optional, Union
from enum import Enum
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

import gradio as gr

from auth import _require_session
from config import supabase
from llm import call_llm
from ingest_pdf import extract_text_from_pdf
from ingest_docx import extract_text_from_docx
from ingest_pptx import extract_text_from_pptx

# Setup logging for decision transparency
logger = logging.getLogger(__name__)

# Try to import python-docx for style extraction
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
    logger.info("python-docx loaded successfully - full style extraction available")
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")

# Optional OCR support for branding-image institution recovery
try:
    from PIL import Image
    OCR_IMAGE_AVAILABLE = True
except ImportError:
    OCR_IMAGE_AVAILABLE = False

try:
    import pytesseract
    OCR_TEXT_AVAILABLE = True
except ImportError:
    OCR_TEXT_AVAILABLE = False


def _extract_text_candidates_from_branding_image(extracted_images: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Best-effort OCR on kept branding images.

    Goal:
    - recover institution name when the donor stores it inside the logo image
    - never crash if OCR libraries are unavailable
    - return conservative, filtered candidates only
    """
    result = {
        "ocr_lines": [],
        "school_candidates": [],
        "likely_institution_name": "",
    }

    if not extracted_images or not OCR_IMAGE_AVAILABLE or not OCR_TEXT_AVAILABLE:
        return result

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _dedupe_keep_order(items):
        seen = set()
        out = []
        for item in items:
            s = _clean(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    SCHOOL_STOPWORDS = {
        "OFFICIAL", "NAME", "TEACHER", "DATE", "CLASS", "STUDENT",
        "YEAR", "EXAM", "TEST", "QUESTION", "ANSWER", "BOOKLET",
        "CALCULATOR", "ASSUMED", "FREE"
    }

    likely_lines = []

    for img in extracted_images:
        if not isinstance(img, dict):
            continue
        if img.get("image_role") != "branding":
            continue

        blob = img.get("binary_data")
        if not blob:
            continue

        try:
            import io
            from PIL import ImageOps, ImageFilter

            if isinstance(blob, str):
                try:
                    import base64
                    raw = base64.b64decode(blob)
                except Exception:
                    raw = blob.encode("utf-8")
            else:
                raw = bytes(blob)

            base_img = Image.open(io.BytesIO(raw)).convert("L")

            # --------------------------------------------------
            # KEY FIX:
            # The left side of the logo contains the large red "R"
            # and dolphin mark, which pollutes OCR badly.
            # So we OCR text-focused crops from the right-hand side.
            # --------------------------------------------------
            width, height = base_img.size

            crops = []

            # Whole image (keep as fallback)
            crops.append(base_img)

            # Right-heavy crop: removes most of the logo mark
            crops.append(base_img.crop((int(width * 0.28), 0, width, height)))

            # Stronger right crop: mostly text block
            crops.append(base_img.crop((int(width * 0.38), 0, width, height)))

            # Middle-right band where school name / motto usually sits
            crops.append(base_img.crop((int(width * 0.28), int(height * 0.08), width, int(height * 0.92))))

            variants = []

            for crop in crops:
                # Original crop
                variants.append(crop)

                # Autocontrast
                auto = ImageOps.autocontrast(crop)
                variants.append(auto)

                # Upscaled
                big = auto.resize((auto.width * 4, auto.height * 4))
                variants.append(big)

                # Upscaled + sharpen
                sharp = big.filter(ImageFilter.SHARPEN)
                variants.append(sharp)

                # Binary thresholds
                bw1 = sharp.point(lambda p: 255 if p > 170 else 0)
                bw2 = sharp.point(lambda p: 255 if p > 145 else 0)
                bw3 = sharp.point(lambda p: 255 if p > 120 else 0)

                variants.append(bw1)
                variants.append(bw2)
                variants.append(bw3)

            ocr_text_chunks = []
            configs = [
                "--psm 6",
                "--psm 11",
                "--psm 4",
                "--psm 7",
            ]

            for variant in variants:
                for cfg in configs:
                    try:
                        text = pytesseract.image_to_string(variant, config=cfg) or ""
                        if text.strip():
                            ocr_text_chunks.append(text)
                    except Exception:
                        continue

            text = "\n".join(ocr_text_chunks)

        except Exception:
            continue

        raw_lines = [_clean(x) for x in text.splitlines() if _clean(x)]
        raw_lines = _dedupe_keep_order(raw_lines)

        filtered_lines = []
        for line in raw_lines:
            low = line.lower()

            if len(line) < 3:
                continue
            if re.fullmatch(r"[_\-\=\.\%\s]+", line):
                continue
            if re.search(r"\b(question|answer|booklet|semester|calculator|marks|teacher|name)\b", low):
                continue

            filtered_lines.append(line)

        likely_lines.extend(filtered_lines)

    likely_lines = _dedupe_keep_order(likely_lines)
    result["ocr_lines"] = likely_lines[:20]

    school_candidates = []
    for line in likely_lines:
        low = line.lower()

        # For OCR from branding images, only trust explicit institution phrases.
        # Do NOT trust acronym-only OCR results like "AYP".
        if re.search(r"\b(high school|primary school|secondary school|college|grammar|academy|institute|university|school)\b", low, re.I):
            school_candidates.append(line)

    school_candidates = _dedupe_keep_order(school_candidates)
    result["school_candidates"] = school_candidates[:10]

    if school_candidates:
        result["likely_institution_name"] = school_candidates[0]

    return result


# ======================================================================================
# CONSTANTS
# ======================================================================================

BUNDLE_START = "<!-- TEMPLATE_BUNDLE_START -->"
BUNDLE_END = "<!-- TEMPLATE_BUNDLE_END -->"
CONTENT_PLACEHOLDER = "{{TEACHER_CONTENT}}"

SUPPORTED_TEMPLATE_EXTS = {".docx", ".pdf", ".ppt", ".pptx"}

PROFILE_SCHEMA_VERSION = "style_profile_v3"
BLUEPRINT_SCHEMA_VERSION = "document_blueprint_v3"

# Educational levels - for appropriateness tuning
class EducationLevel(Enum):
    PRIMARY = "primary"          # Years K-6
    SECONDARY = "secondary"      # Years 7-12
    UNDERGRADUATE = "undergraduate"
    POSTGRADUATE = "postgraduate"
    PROFESSIONAL = "professional"
    UNKNOWN = "unknown"

# Feature categories for intelligent layout inheritance
class FeatureCategory(Enum):
    ALWAYS_INHERIT = "ALWAYS_INHERIT"        # Logo, margins, fonts, header/footer
    CONTEXTUAL_INHERIT = "CONTEXTUAL_INHERIT" # Columns, tables, answer lines
    NEVER_INHERIT = "NEVER_INHERIT"          # Math two-column, equation formatting

# Document types for type-specific rules
class DocumentType(Enum):
    WORKSHEET = "worksheet"
    EXAM = "exam"
    TEST = "test"
    MEMO = "memo"
    RUBRIC = "rubric"
    ASSIGNMENT = "assignment"
    LESSON = "lesson"
    INVESTIGATION = "investigation"
    CUSTOM = "custom"

# Answer space styles
class AnswerSpaceStyle(Enum):
    SHORT_RESPONSE = "short_response"
    PARAGRAPH_RESPONSE = "paragraph_response"
    SHOW_WORKING = "show_working"
    LABEL_DIAGRAM = "label_diagram"
    TABLE_RESPONSE = "table_response"
    SUBPARTS = "subparts"

@dataclass
class Decision:
    """Structured decision with confidence and reasoning."""
    apply: bool
    confidence: float  # 0.0 to 1.0
    reason: str
    alternative: Optional[str] = None  # What to do if uncertain
    requires_teacher_confirmation: bool = False
    source: str = "engine"  # "engine", "teacher_hint", "rule", "default"


@dataclass
class AppropriatenessResult:
    """Complete appropriateness analysis result."""
    decisions: Dict[str, Decision]
    structure_analysis: Dict[str, Any]
    user_hints_applied: Dict[str, bool]
    education_level: EducationLevel
    document_type: DocumentType
    uncertainty_flags: List[str] = field(default_factory=list)
    quality_score: Optional[float] = None
    llm_used: bool = False


@dataclass
class ExtractedImage:
    """Represents an image extracted from donor document."""
    image_id: str
    binary_data: bytes
    width_inches: float
    height_inches: float
    position: str  # "header", "footer", "body", "first_page"
    alignment: str  # "left", "center", "right"
    alt_text: Optional[str] = None


@dataclass
class TableStructure:
    """Represents a table structure for preservation."""
    rows: int
    cols: int
    has_header_row: bool
    borders: Dict[str, bool]
    column_widths: List[float]
    cell_merges: List[Dict[str, int]]
    cell_styles: Dict[str, Any]
    content_placeholders: List[Tuple[int, int, str]]  # (row, col, placeholder)


# ======================================================================================
# BASIC HELPERS
# ======================================================================================

# Test _resolve_upload_path
# Should handle:
# - _resolve_upload_path(None) -> ("", "")
# - _resolve_upload_path({"name": "/tmp/file.docx"}) -> ("/tmp/file.docx", "file.docx")
# - _resolve_upload_path("/tmp/file.docx") -> ("/tmp/file.docx", "file.docx")

# Test _safe_ext
# - _safe_ext("document.docx") -> ".docx"
# - _safe_ext("file.PDF") -> ".pdf"
# - _safe_ext("noextension") -> ""

# Test _norm_spaces
# - _norm_spaces("Hello\r\nWorld") -> "Hello\nWorld"
# - _norm_spaces("Too   many   spaces") -> "Too many spaces"
# - _norm_spaces("Line1\n\n\n\nLine2") -> "Line1\n\nLine2"

# Test _short
# - _short("Short text", 100) -> "Short text"
# - _short("This is a very long string that needs truncation", 20) -> "This is a very long..."

# Test _try_int
# - _try_int("123") -> 123
# - _try_int(3.7) -> 3 (truncates, doesn't round)
# - _try_int("abc", 0) -> 0

# Test _try_float
# - _try_float("123.45") -> 123.45
# - _try_float(5) -> 5.0
# - _try_float("abc", 0.0) -> 0.0

# Test _dominant
# - from collections import Counter
# - c = Counter(['a', 'a', 'b', 'b', 'b', 'c'])
# - _dominant(c, 2) -> [{"value": "b", "count": 3}, {"value": "a", "count": 2}]

def _resolve_upload_path(upload_obj) -> Tuple[str, str]:
    """
    Resolve Gradio upload to (path, filename).
    
    Handles:
        - Gradio file object with 'name' attribute
        - Dictionary with 'name' or 'path' key
        - String path
        - None or empty input
    
    Returns:
        Tuple of (absolute_path_or_path, basename)
    """
    if not upload_obj:
        return "", ""
    
    # Handle Gradio file object (has 'name' attribute)
    if hasattr(upload_obj, 'name'):
        path = upload_obj.name
        return path, os.path.basename(path)
    
    # Handle dictionary (from some Gradio versions)
    if isinstance(upload_obj, dict):
        path = upload_obj.get("name") or upload_obj.get("path") or ""
        return path, os.path.basename(path) if path else ""
    
    # Handle string path
    if isinstance(upload_obj, str):
        return upload_obj, os.path.basename(upload_obj)
    
    # Fallback: convert to string
    path = str(upload_obj)
    return path, os.path.basename(path) if path else ""


def _safe_ext(filename: str) -> str:
    """
    Return lowercase file extension including the leading dot.
    
    Examples:
        _safe_ext("document.docx") -> ".docx"
        _safe_ext("file.PDF") -> ".pdf"
        _safe_ext("noextension") -> ""
        _safe_ext("") -> ""
    """
    if not filename:
        return ""
    
    ext = os.path.splitext(filename)[1].lower().strip()
    return ext


def _norm_spaces(text: str) -> str:
    """
    Normalize line endings and collapse excessive whitespace.
    
    Does:
        - Converts \r\n and \r to \n
        - Removes trailing spaces before newlines
        - Collapses 3+ consecutive newlines to 2
        - Strips leading/trailing whitespace
    
    Example:
        "Hello   World\r\n\r\n\r\nHow are you?" -> "Hello World\n\nHow are you?"
    """
    if not text:
        return ""
    
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Remove trailing spaces before newlines
    text = re.sub(r"[ \t]+\n", "\n", text)
    
    # Collapse multiple spaces (but not newlines)
    text = re.sub(r"[ \t]{2,}", " ", text)
    
    # Collapse 3+ consecutive newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    # Strip leading/trailing whitespace
    return text.strip()


def _short(text: str, n: int = 300) -> str:
    """
    Truncate text to n characters for preview/debug display.
    
    Adds "..." if truncated.
    Preserves whole words when possible within limit.
    
    Example:
        _short("This is a very long string that needs truncation", 20)
        -> "This is a very long..."
    """
    if not text:
        return ""
    
    text = str(text).strip()
    
    if len(text) <= n:
        return text
    
    # Try to cut at word boundary
    truncated = text[:n]
    last_space = truncated.rfind(' ')
    
    if last_space > n * 0.8:  # Only cut at word if we're not losing too much
        truncated = truncated[:last_space]
    
    return truncated + "..."


def _try_int(x, default=0) -> int:
    """
    Safely convert a value to int, returning default on failure.
    
    Handles:
        - int values
        - float values (truncates)
        - numeric strings
        - None
        - any other type
    
    Example:
        _try_int("123") -> 123
        _try_int("abc", 0) -> 0
        _try_int(None, 5) -> 5
    """
    if x is None:
        return default
    
    try:
        # Handle float by truncating (not rounding)
        if isinstance(x, float):
            return int(x)
        return int(x)
    except (ValueError, TypeError):
        return default


def _try_float(x, default=0.0) -> float:
    """
    Safely convert a value to float, returning default on failure.
    
    Handles:
        - float values
        - int values
        - numeric strings
        - None
        - any other type
    
    Example:
        _try_float("123.45") -> 123.45
        _try_float("abc", 0.0) -> 0.0
        _try_float(None, 1.5) -> 1.5
    """
    if x is None:
        return default
    
    try:
        return float(x)
    except (ValueError, TypeError):
        return default


def _dominant(counter_obj: Counter, top_n: int = 5) -> List[Dict[str, Any]]:
    """
    Get top-N items from a Counter as a list of dictionaries.
    
    Each dictionary has:
        - "value": the item
        - "count": its frequency
    
    Example:
        counter = Counter(['a', 'a', 'b', 'b', 'b', 'c'])
        _dominant(counter, 2) -> [
            {"value": "b", "count": 3},
            {"value": "a", "count": 2}
        ]
    """
    if not counter_obj or not isinstance(counter_obj, Counter):
        return []
    
    result = []
    for value, count in counter_obj.most_common(top_n):
        result.append({
            "value": value,
            "count": count
        })
    
    return result


# ======================================================================================
# BUNDLE HELPERS
# ======================================================================================

def pack_template_bundle(clean_md: str, profile: Dict[str, Any]) -> str:
    """
    Pack preview markdown + profile into bundle with markers.
    Format: <!-- BUNDLE_START -->{json}<!-- BUNDLE_END -->\n\n{markdown}
    
    Example output:
        <!-- TEMPLATE_BUNDLE_START -->
        {"clean_md": "# Preview...", "profile": {...}}
        <!-- TEMPLATE_BUNDLE_END -->
        
        # Preview markdown here...
    """
    clean_md = clean_md or ""
    profile = profile or {}
    
    safe_profile = _make_profile_json_safe(profile)

    # Create payload with both components
    payload = {
        "clean_md": clean_md,
        "profile": safe_profile,
        "schema_version": PROFILE_SCHEMA_VERSION,
        "packed_at": datetime.now().isoformat()
    }
    
    # Serialize to JSON with nice formatting
    json_str = json.dumps(payload, ensure_ascii=False, indent=2)
    
    # Build bundle with markers
    bundle = (
        f"{BUNDLE_START}\n"
        f"{json_str}\n"
        f"{BUNDLE_END}\n\n"
        f"{clean_md}"
    )
    
    return bundle


def _make_profile_json_safe(obj):
    """
    Recursively convert a profile object into a JSON-safe version.

    BS5 logo rule:
    - bytes are preserved as base64 payloads so branding/logo images survive
      template save/load cycles.
    """
    if isinstance(obj, bytes):
        import base64
        return {
            "__binary_base64__": True,
            "encoding": "base64",
            "data": base64.b64encode(obj).decode("ascii"),
            "byte_length": len(obj),
        }

    if isinstance(obj, dict):
        return {k: _make_profile_json_safe(v) for k, v in obj.items()}

    if isinstance(obj, list):
        return [_make_profile_json_safe(v) for v in obj]

    return obj


def unpack_template_bundle(template_md: str) -> Tuple[str, Dict[str, Any]]:
    """
    Extract (clean_md, profile) from bundle.
    Must handle legacy formats gracefully.
    
    Handles:
        1. New format with BUNDLE_START/BUNDLE_END markers
        2. Legacy format where entire string is JSON
        3. Legacy format where string is plain markdown
        4. Corrupted or malformed bundles (graceful fallback)
    
    Returns:
        Tuple of (clean_markdown_preview, profile_dict)
    """
    raw = template_md or ""
    
    # ================================================================
    # Method 1: Extract from bundle markers (NEW format)
    # ================================================================
    marker_pattern = re.escape(BUNDLE_START) + r"\s*(\{.*?\})\s*" + re.escape(BUNDLE_END)
    m = re.search(marker_pattern, raw, flags=re.DOTALL)
    
    if m:
        try:
            payload = json.loads(m.group(1))
            clean_md = payload.get("clean_md", "")
            profile = payload.get("profile", {})
            
            # Validate we got something usable
            if isinstance(profile, dict) and profile:
                logger.debug("Unpacked bundle from markers")
                return clean_md, profile
            elif isinstance(profile, dict):
                # Empty profile but valid structure
                return clean_md, profile
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse bundle JSON: {e}")
        except Exception as e:
            logger.warning(f"Unexpected error parsing bundle: {e}")
    
    # ================================================================
    # Method 2: Try to parse entire string as JSON (legacy format)
    # ================================================================
    try:
        direct = json.loads(raw)
        if isinstance(direct, dict):
            # Check if this looks like a valid profile
            if "profile" in direct:
                return direct.get("clean_md", ""), direct.get("profile", {})
            elif "source_type" in direct or "layout_features" in direct:
                # This appears to be a profile directly
                return "", direct
    except json.JSONDecodeError:
        pass
    
    # ================================================================
    # Method 3: Look for JSON object anywhere in the text (embedded)
    # ================================================================
    try:
        # Find any JSON object in the text
        json_match = re.search(r'(\{.*\})', raw, re.DOTALL)
        if json_match:
            potential = json.loads(json_match.group(1))
            if isinstance(potential, dict):
                if "profile" in potential:
                    return potential.get("clean_md", ""), potential.get("profile", {})
                elif "clean_md" in potential:
                    return potential.get("clean_md", ""), potential.get("profile", {})
                elif "source_type" in potential:
                    return "", potential
    except json.JSONDecodeError:
        pass
    
    # ================================================================
    # Method 4: Check for legacy marker without proper JSON
    # ================================================================
    if BUNDLE_START in raw and BUNDLE_END in raw:
        # Try to extract content between markers as plain text
        try:
            between = re.search(
                re.escape(BUNDLE_START) + r"\s*(.*?)\s*" + re.escape(BUNDLE_END),
                raw,
                flags=re.DOTALL
            )
            if between:
                content = between.group(1).strip()
                # Try to parse as JSON
                try:
                    payload = json.loads(content)
                    return payload.get("clean_md", ""), payload.get("profile", {})
                except:
                    # Content might be plain text profile description
                    return content, {}
        except:
            pass
    
    # ================================================================
    # Method 5: Legacy plain markdown (no bundle structure)
    # ================================================================
    # Remove any bundle markers if they exist but malformed
    clean_md = re.sub(
        re.escape(BUNDLE_START) + r".*?" + re.escape(BUNDLE_END),
        "",
        raw,
        flags=re.DOTALL
    ).strip()
    
    # If we have content, treat as markdown preview
    if clean_md:
        logger.debug("Unpacked as legacy markdown (no profile)")
        return clean_md, {}
    
    # ================================================================
    # Method 6: Ultimate fallback - return raw as markdown
    # ================================================================
    logger.warning("Could not parse template bundle, returning as plain markdown")
    return raw, {}


# Optional: Helper to validate bundle integrity
def _validate_bundle(bundle: str) -> bool:
    """
    Validate that a bundle is properly formatted and can be unpacked.
    Returns True if valid, False otherwise.
    """
    try:
        clean_md, profile = unpack_template_bundle(bundle)
        # A valid bundle has either clean_md or profile (or both)
        return bool(clean_md) or bool(profile)
    except Exception:
        return False






# ======================================================================================
# BS5 FIRST-PAGE LAYOUT CONTRACT HELPERS
# ======================================================================================

def _safe_len_value(value):
    try:
        return float(value.inches) if value is not None and hasattr(value, "inches") else None
    except Exception:
        return None

def _extract_run_style_signature(run) -> Dict[str, Any]:
    """Return a small, JSON-safe run style signature. Missing properties are omitted."""
    sig: Dict[str, Any] = {}
    try:
        font = getattr(run, "font", None)
        if font is not None:
            if getattr(font, "name", None):
                sig["font_name"] = font.name
            if getattr(font, "size", None):
                sig["font_size_pt"] = font.size.pt
            for attr in ("bold", "italic", "underline"):
                val = getattr(font, attr, None)
                if val is not None:
                    sig[attr] = bool(val)
            try:
                if font.color and font.color.rgb:
                    sig["color_rgb"] = str(font.color.rgb)
            except Exception:
                pass
    except Exception:
        pass
    try:
        if getattr(run, "style", None) and run.style.name:
            sig["style_name"] = run.style.name
    except Exception:
        pass
    return sig

def _extract_paragraph_style_signature(paragraph) -> Dict[str, Any]:
    """Return a conservative paragraph style signature plus dominant run style."""
    sig: Dict[str, Any] = {}
    try:
        if getattr(paragraph, "style", None) and paragraph.style.name:
            sig["style_name"] = paragraph.style.name
    except Exception:
        pass
    try:
        if getattr(paragraph, "alignment", None) is not None:
            sig["alignment"] = paragraph.alignment.name
    except Exception:
        pass
    try:
        pf = paragraph.paragraph_format
        for key, attr in (("space_before_pt", "space_before"), ("space_after_pt", "space_after"), ("left_indent_pt", "left_indent"), ("right_indent_pt", "right_indent"), ("first_line_indent_pt", "first_line_indent")):
            val = getattr(pf, attr, None)
            if val is not None:
                sig[key] = val.pt
        if pf.line_spacing is not None:
            sig["line_spacing"] = pf.line_spacing
    except Exception:
        pass
    try:
        run_sigs = [_extract_run_style_signature(r) for r in paragraph.runs if getattr(r, "text", "").strip()]
        run_sigs = [r for r in run_sigs if r]
        if run_sigs:
            # Prefer the first styled run; question headings often have one run.
            sig["run"] = run_sigs[0]
    except Exception:
        pass
    return sig

def _apply_run_style_signature(target_run, signature: Dict[str, Any]) -> None:
    """Apply a conservative run signature without disturbing text/math/numbering."""
    if not signature:
        return
    try:
        font = target_run.font
        if signature.get("font_name"):
            font.name = signature.get("font_name")
        if signature.get("font_size_pt") is not None:
            font.size = Pt(float(signature.get("font_size_pt")))
        for attr in ("bold", "italic", "underline"):
            if attr in signature:
                setattr(font, attr, bool(signature[attr]))
        if signature.get("color_rgb"):
            from docx.shared import RGBColor
            rgb = str(signature.get("color_rgb")).replace("#", "").strip()
            if len(rgb) == 6:
                font.color.rgb = RGBColor(int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16))
    except Exception as e:
        logger.debug(f"Could not apply run signature: {e}")

def _apply_paragraph_style_signature(target_paragraph, signature: Dict[str, Any]) -> None:
    """Apply paragraph/run style signature defensively."""
    if not signature:
        return
    try:
        style_name = signature.get("style_name")
        if style_name and style_name in [s.name for s in target_paragraph.part.document.styles]:
            target_paragraph.style = style_name
    except Exception:
        pass
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_name = str(signature.get("alignment") or "").upper()
        if align_name and hasattr(WD_ALIGN_PARAGRAPH, align_name):
            target_paragraph.alignment = getattr(WD_ALIGN_PARAGRAPH, align_name)
    except Exception:
        pass
    try:
        pf = target_paragraph.paragraph_format
        if signature.get("space_before_pt") is not None:
            pf.space_before = Pt(float(signature["space_before_pt"]))
        if signature.get("space_after_pt") is not None:
            pf.space_after = Pt(float(signature["space_after_pt"]))
        if signature.get("left_indent_pt") is not None:
            pf.left_indent = Pt(float(signature["left_indent_pt"]))
        if signature.get("first_line_indent_pt") is not None:
            pf.first_line_indent = Pt(float(signature["first_line_indent_pt"]))
        if signature.get("line_spacing") is not None:
            pf.line_spacing = signature.get("line_spacing")
    except Exception:
        pass
    run_sig = signature.get("run") or {}
    for run in getattr(target_paragraph, "runs", []) or []:
        _apply_run_style_signature(run, run_sig)

def _container_text(container) -> str:
    try:
        return _norm_spaces(" ".join([p.text for p in container.paragraphs if (p.text or "").strip()]))
    except Exception:
        return ""

def _container_has_images(container) -> bool:
    try:
        xml = container._element.xml
        return ("<w:drawing" in xml) or ("<w:pict" in xml) or ("graphic" in xml and "r:embed" in xml)
    except Exception:
        return False

def _extract_header_footer_contract(doc: Document) -> Dict[str, Any]:
    contract: Dict[str, Any] = {"different_first_page_header_footer": False}
    try:
        section = doc.sections[0]
    except Exception:
        return contract
    try:
        contract["different_first_page_header_footer"] = bool(section.different_first_page_header_footer)
    except Exception:
        pass
    zones = {
        "first_page_header": getattr(section, "first_page_header", None),
        "default_header": getattr(section, "header", None),
        "first_page_footer": getattr(section, "first_page_footer", None),
        "default_footer": getattr(section, "footer", None),
    }
    for name, container in zones.items():
        if container is None:
            continue
        contract[f"{name}_text"] = _container_text(container)
        contract[f"{name}_has_images"] = _container_has_images(container)
        try:
            paras = [p for p in container.paragraphs if (p.text or "").strip()]
            if paras:
                contract[f"{name}_style_signature"] = _extract_paragraph_style_signature(paras[0])
        except Exception:
            pass
    footer_text = " ".join([contract.get("first_page_footer_text", ""), contract.get("default_footer_text", "")])
    footer_xml = ""
    try:
        footer_xml = " ".join([
            getattr(zones.get("first_page_footer"), "_element", None).xml if zones.get("first_page_footer") is not None else "",
            getattr(zones.get("default_footer"), "_element", None).xml if zones.get("default_footer") is not None else "",
        ])
    except Exception:
        footer_xml = ""
    contract["footer_page_numbering_detected"] = bool(
        re.search(r"page\s*\d|page\s*\{?PAGE|NUMPAGES|of\s+\d", footer_text, re.I)
        or re.search(r"\b(PAGE|NUMPAGES)\b", footer_xml, re.I)
    )
    contract["footer_has_word_fields"] = bool(re.search(r"w:fldChar|w:instrText|PAGE|NUMPAGES", footer_xml, re.I))
    return contract

def _extract_first_page_layout_contract(docx_path: str) -> Dict[str, Any]:
    """Extract a defensive dictionary-based BS5 first-page layout contract."""
    contract: Dict[str, Any] = {"available": False, "source_path": docx_path}
    if not DOCX_AVAILABLE or not docx_path or not os.path.exists(docx_path):
        return contract
    try:
        doc = Document(docx_path)
        contract.update(_extract_header_footer_contract(doc))
        contract["available"] = True
        section = doc.sections[0] if doc.sections else None
        if section is not None:
            contract["page_margins"] = {
                "left_margin_inches": _safe_len_value(section.left_margin),
                "right_margin_inches": _safe_len_value(section.right_margin),
                "top_margin_inches": _safe_len_value(section.top_margin),
                "bottom_margin_inches": _safe_len_value(section.bottom_margin),
                "header_distance_inches": _safe_len_value(section.header_distance),
                "footer_distance_inches": _safe_len_value(section.footer_distance),
            }
            try:
                contract["orientation"] = "landscape" if section.orientation and section.orientation.name == "LANDSCAPE" else "portrait"
            except Exception:
                pass
            try:
                cols = section._sectPr.xpath('./w:cols')
                if cols:
                    c = cols[0]
                    contract["columns"] = {"count": int(c.get(qn('w:num'), 1)), "equal_width": c.get(qn('w:eq')) != "0"}
            except Exception:
                pass
        paras = [p for p in doc.paragraphs[:80] if (p.text or "").strip()]
        first_page_paras = paras[:25]
        body_candidates = [p for p in paras if len((p.text or "").strip()) > 40]
        heading_candidates = [p for p in paras if re.match(r"^\s*(#{0,6}\s*)?(Question|Q)\s*\d+", p.text or "", re.I) or (p.style and "heading" in p.style.name.lower())]
        if first_page_paras:
            contract["dominant_first_page_paragraph_style_signature"] = _extract_paragraph_style_signature(first_page_paras[0])
        if heading_candidates:
            contract["dominant_question_heading_style_signature"] = _extract_paragraph_style_signature(heading_candidates[0])
        if body_candidates:
            contract["dominant_body_style_signature"] = _extract_paragraph_style_signature(body_candidates[0])
        hf_sigs = [contract.get(k) for k in ("first_page_header_style_signature", "default_header_style_signature", "first_page_footer_style_signature", "default_footer_style_signature") if contract.get(k)]
        if hf_sigs:
            contract["dominant_header_footer_style_signature"] = hf_sigs[0]
    except Exception as e:
        contract["error"] = str(e)
    return contract


# ======================================================================================
# DOCX ZIP/XML FALLBACK EXTRACTION HELPERS
# ======================================================================================

def _extract_docx_xml_header_footer_fallback(docx_path: str) -> Dict[str, Any]:
    """
    Low-level DOCX ZIP/XML fallback extractor for headers, footers and page-number fields.

    Purpose:
    - catch header/footer text that python-docx may miss
    - reconstruct footer patterns such as "Page {PAGE} of {NUMPAGES}"
    - preserve this as extraction evidence only, not as donor academic content
    """
    result = {
        "available": False,
        "headers": [],
        "footers": [],
        "page_number_patterns": [],
        "xml_files_checked": [],
        "errors": [],
    }

    if not docx_path or not os.path.exists(docx_path):
        result["errors"].append("DOCX path missing or does not exist.")
        return result

    def _xml_unescape(text: str) -> str:
        text = text or ""
        return (
            text.replace("&amp;", "&")
                .replace("&lt;", "<")
                .replace("&gt;", ">")
                .replace("&quot;", '"')
                .replace("&apos;", "'")
        )

    def _clean_join(parts):
        cleaned = []
        for part in parts or []:
            part = _xml_unescape(part)
            part = part.replace("\xa0", " ").strip()
            if part:
                cleaned.append(part)
        return _norm_spaces(" ".join(cleaned))

    def _extract_text_and_field_pattern(xml: str) -> Tuple[str, str]:
        # Pull normal visible text and field instructions in document order.
        tokens = []
        for match in re.finditer(r"<w:(t|instrText|instr)[^>]*>(.*?)</w:\1>", xml, flags=re.DOTALL):
            tag = match.group(1)
            value = re.sub(r"<[^>]+>", "", match.group(2) or "")
            value = _xml_unescape(value).strip()
            if not value:
                continue
            upper = value.upper().strip()
            if tag in {"instrText", "instr"}:
                if "NUMPAGES" in upper:
                    tokens.append("{NUMPAGES}")
                elif re.search(r"\bPAGE\b", upper):
                    tokens.append("{PAGE}")
            else:
                tokens.append(value)

        visible_text = _clean_join([t for t in tokens if t not in {"{PAGE}", "{NUMPAGES}"}])
        field_pattern = _clean_join(tokens)
        return visible_text, field_pattern

    try:
        import zipfile
        with zipfile.ZipFile(docx_path) as docx_zip:
            xml_files = [
                name for name in docx_zip.namelist()
                if name.startswith("word/")
                and name.endswith(".xml")
                and ("header" in name.lower() or "footer" in name.lower())
            ]

            for xml_file in sorted(xml_files):
                result["xml_files_checked"].append(xml_file)
                try:
                    xml = docx_zip.read(xml_file).decode("utf-8", errors="ignore")
                    visible_text, field_pattern = _extract_text_and_field_pattern(xml)

                    if "header" in xml_file.lower() and visible_text:
                        result["headers"].append(visible_text)
                    elif "footer" in xml_file.lower() and visible_text:
                        result["footers"].append(visible_text)

                    if "footer" in xml_file.lower() and ("{PAGE}" in field_pattern or "{NUMPAGES}" in field_pattern):
                        result["page_number_patterns"].append(field_pattern)
                except Exception as e:
                    result["errors"].append(f"{xml_file}: {e}")

        # Deduplicate while preserving order.
        def _dedupe(items):
            seen = set()
            out = []
            for item in items or []:
                key = str(item).strip().lower()
                if key and key not in seen:
                    seen.add(key)
                    out.append(item)
            return out

        result["headers"] = _dedupe(result["headers"])
        result["footers"] = _dedupe(result["footers"])
        result["page_number_patterns"] = _dedupe(result["page_number_patterns"])
        result["available"] = bool(result["headers"] or result["footers"] or result["page_number_patterns"])
        return result

    except Exception as e:
        result["errors"].append(str(e))
        return result


def _merge_docx_xml_fallback_into_header_footer(header_footer: Dict[str, Any], docx_path: str) -> Dict[str, Any]:
    """
    Merge low-level ZIP/XML header/footer evidence into the python-docx header/footer result.
    Keeps fallback evidence explicit so it remains auditable.
    """
    header_footer = dict(header_footer or {})
    fallback = _extract_docx_xml_header_footer_fallback(docx_path)
    header_footer["xml_fallback"] = fallback

    def _merge_unique(existing, extra, limit=20):
        seen = set()
        out = []
        for item in list(existing or []) + list(extra or []):
            s = (item or "").replace("\xa0", " ").strip()
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out[:limit]

    header_footer["header_texts"] = _merge_unique(
        header_footer.get("header_texts", []), fallback.get("headers", []), limit=20
    )
    header_footer["footer_texts"] = _merge_unique(
        header_footer.get("footer_texts", []), fallback.get("footers", []), limit=20
    )
    header_footer["header_candidates"] = _merge_unique(
        header_footer.get("header_candidates", []), fallback.get("headers", []), limit=20
    )
    header_footer["footer_candidates"] = _merge_unique(
        header_footer.get("footer_candidates", []), fallback.get("footers", []), limit=20
    )
    header_footer["page_number_patterns"] = _merge_unique(
        header_footer.get("page_number_patterns", []), fallback.get("page_number_patterns", []), limit=10
    )

    header_footer["has_header"] = bool(header_footer.get("has_header") or header_footer.get("header_texts"))
    header_footer["has_footer"] = bool(header_footer.get("has_footer") or header_footer.get("footer_texts") or header_footer.get("page_number_patterns"))
    return header_footer


def _extract_docx_media_location_fallback(docx_path: str) -> List[Dict[str, Any]]:
    """
    Low-level DOCX ZIP/XML media scanner.

    Purpose:
    - verify whether images are used in headers, footers or body
    - recover pixel dimensions and file format
    - provide a stable diagnostic record for image extraction debugging
    """
    records = []
    if not docx_path or not os.path.exists(docx_path):
        return records

    try:
        import hashlib
        import zipfile
        import xml.etree.ElementTree as ET
        import io

        rels_by_part = {}
        with zipfile.ZipFile(docx_path) as docx_zip:
            names = set(docx_zip.namelist())
            media_files = [name for name in names if name.startswith("word/media/")]

            for rels_name in [name for name in names if name.startswith("word/") and name.endswith(".rels")]:
                try:
                    xml = docx_zip.read(rels_name)
                    root = ET.fromstring(xml)
                    rel_map = {}
                    for rel in root:
                        rid = rel.attrib.get("Id", "")
                        target = rel.attrib.get("Target", "")
                        if not rid or "media/" not in target:
                            continue
                        media_name = target.split("media/")[-1]
                        rel_map[rid] = "word/media/" + media_name
                    rels_by_part[rels_name] = rel_map
                except Exception:
                    continue

            media_usage = {media: set() for media in media_files}

            xml_parts = [name for name in names if name.startswith("word/") and name.endswith(".xml")]
            for xml_part in xml_parts:
                rels_name = "word/_rels/document.xml.rels" if xml_part == "word/document.xml" else None
                if xml_part.startswith("word/header"):
                    rels_name = f"word/_rels/{xml_part.split('/')[-1]}.rels"
                elif xml_part.startswith("word/footer"):
                    rels_name = f"word/_rels/{xml_part.split('/')[-1]}.rels"

                rel_map = rels_by_part.get(rels_name or "", {})
                if not rel_map:
                    continue

                xml_text = docx_zip.read(xml_part).decode("utf-8", errors="ignore")
                for rid, media in rel_map.items():
                    if rid in xml_text and media in media_usage:
                        if "header" in xml_part.lower():
                            media_usage[media].add("header")
                        elif "footer" in xml_part.lower():
                            media_usage[media].add("footer")
                        else:
                            media_usage[media].add("body")

            for media in sorted(media_files):
                blob = docx_zip.read(media)
                width_px = None
                height_px = None
                image_format = None
                try:
                    if OCR_IMAGE_AVAILABLE:
                        img = Image.open(io.BytesIO(blob))
                        width_px, height_px = img.size
                        image_format = img.format
                except Exception:
                    pass

                usage = sorted(media_usage.get(media) or [])
                position = "header" if "header" in usage else "footer" if "footer" in usage else "body" if "body" in usage else "unknown"
                is_likely_logo = bool(position in {"header", "footer"} or (width_px and height_px and width_px <= 500 and height_px <= 500))

                records.append({
                    "filename": media,
                    "size_bytes": len(blob),
                    "sha1": hashlib.sha1(blob).hexdigest(),
                    "position": position,
                    "usage_zones": usage,
                    "width_px": width_px,
                    "height_px": height_px,
                    "format": image_format,
                    "is_likely_logo": is_likely_logo,
                })

    except Exception as e:
        logger.debug(f"DOCX media location fallback failed: {e}")

    return records


def _merge_docx_media_fallback_into_images(images: List[Dict[str, Any]], docx_path: str) -> List[Dict[str, Any]]:
    """
    Attach ZIP/XML media-location evidence to extracted image records and add missing media diagnostics.
    Does not make educational images inheritable; it only improves extraction evidence.
    """
    images = list(images or [])
    media_records = _extract_docx_media_location_fallback(docx_path)
    if not media_records:
        return images

    by_sha1 = {r.get("sha1"): r for r in media_records if r.get("sha1")}
    attached_sha1 = set()

    try:
        import hashlib
        for img in images:
            blob = img.get("binary_data")
            if blob:
                sha1 = hashlib.sha1(bytes(blob)).hexdigest()
                fallback = by_sha1.get(sha1)
                if fallback:
                    img["media_fallback"] = fallback
                    attached_sha1.add(sha1)
                    if img.get("position") in {None, "", "body", "unknown"} and fallback.get("position") in {"header", "footer"}:
                        img["position"] = fallback.get("position")
                    if fallback.get("is_likely_logo") and img.get("image_role") != "branding":
                        img["image_role"] = "branding"
                        img["classification_reason"] = "media_fallback_likely_logo"
    except Exception:
        pass

    # Preserve records for media not captured by python-docx, without binary payload.
    next_id = len(images) + 1
    for record in media_records:
        if record.get("sha1") in attached_sha1:
            continue
        images.append({
            "image_id": str(next_id),
            "binary_data": b"",
            "width_inches": 1.0,
            "height_inches": 1.0,
            "position": record.get("position", "unknown"),
            "alignment": "left",
            "alt_text": None,
            "relationship_id": None,
            "image_role": "branding" if record.get("is_likely_logo") else "educational",
            "classification_reason": "media_fallback_record_only",
            "paragraph_index": None,
            "media_fallback": record,
        })
        next_id += 1

    return images


def _build_donor_extraction_debug_report(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Human-readable extraction report for debugging donor evidence.
    This mirrors the useful report style of the standalone extractor without treating donor content as authority.
    """
    raw_profile = raw_profile or {}
    hf = raw_profile.get("header_footer", {}) or {}
    images = raw_profile.get("extracted_images", []) or []
    tables = raw_profile.get("extracted_tables", []) or []
    text_stats = raw_profile.get("text_stats", {}) or {}
    style = raw_profile.get("style_preferences", {}) or {}
    first_contract = raw_profile.get("first_page_layout_contract", {}) or {}

    return {
        "summary": {
            "source_filename": raw_profile.get("source_filename", ""),
            "headers_detected": len(hf.get("header_texts", []) or []),
            "footers_detected": len(hf.get("footer_texts", []) or []),
            "page_number_patterns_detected": len(hf.get("page_number_patterns", []) or []),
            "images_detected": len(images),
            "branding_images_detected": len([img for img in images if str(img.get("image_role", "")).lower() == "branding"]),
            "tables_detected": len(tables),
            "paragraphs_detected": text_stats.get("paragraph_count", 0),
            "runs_detected": text_stats.get("run_count", 0),
        },
        "headers_footers": {
            "headers": list(hf.get("header_texts", []) or [])[:10],
            "footers": list(hf.get("footer_texts", []) or [])[:10],
            "page_number_patterns": list(hf.get("page_number_patterns", []) or [])[:10],
            "xml_fallback_available": bool((hf.get("xml_fallback", {}) or {}).get("available")),
        },
        "images": [
            {
                "position": img.get("position"),
                "role": img.get("image_role"),
                "reason": img.get("classification_reason"),
                "width_inches": img.get("width_inches"),
                "height_inches": img.get("height_inches"),
                "media_fallback": img.get("media_fallback"),
            }
            for img in images[:20]
        ],
        "tables": [
            {
                "rows": tbl.get("rows"),
                "cols": tbl.get("cols"),
                "role": tbl.get("table_role"),
                "reason": tbl.get("classification_reason"),
            }
            for tbl in tables[:20]
        ],
        "first_page_layout_contract": {
            "different_first_page_header_footer": bool(first_contract.get("different_first_page_header_footer", False)),
            "first_page_header_text": first_contract.get("first_page_header_text", ""),
            "later_page_header_text": first_contract.get("default_header_text", ""),
            "first_page_footer_text": first_contract.get("first_page_footer_text", ""),
            "later_page_footer_text": first_contract.get("default_footer_text", ""),
            "first_page_header_has_images": bool(first_contract.get("first_page_header_has_images", False)),
            "later_header_has_images": bool(first_contract.get("default_header_has_images", False)),
            "footer_page_numbering_detected": bool(first_contract.get("footer_page_numbering_detected", False)),
            "question_heading_style_detected": bool(first_contract.get("dominant_question_heading_style_signature")),
            "body_style_detected": bool(first_contract.get("dominant_body_style_signature")),
        },
        "formatting": {
            "font_family": style.get("font_family"),
            "font_size_pt": style.get("font_size_pt"),
            "unique_fonts": style.get("unique_fonts"),
            "font_variations": style.get("font_variations"),
            "heading_styles": style.get("heading_styles", []),
        },
    }


# ======================================================================================
# PHASE 1: DONOR STYLE EXTRACTION (ENHANCED - FULL DOCX PROPERTIES)
# ======================================================================================

def _strip_subject_baggage_from_donor(profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Remove donor academic/content baggage so the donor stays a style source,
    not a content master.

    Build Sequence 3 rule:
    - keep institutional style signals
    - keep generic layout/style metadata
    - preserve structural evidence even when uncertain
    - suppress likely academic content conservatively, not destructively
    """

    profile = dict(profile or {})

    # ======================================================================
    # IMPORTANT: These lists are for BAGGAGE REMOVAL ONLY.
    # They are NEVER used by the appropriateness engine to make layout decisions.
    # Layout decisions are based on STRUCTURAL signals (question length,
    # presence of diagrams/tables, document type, etc.), never on subject keywords.
    # ======================================================================
    subject_words = [
        "mathematics", "maths", "math",
        "biology", "science", "chemistry", "physics",
        "english", "history", "geography",
        "economics", "business", "accounting",
        "technology", "engineering", "law", "medicine",
        "exam", "test", "worksheet", "assignment", "memo", "rubric"
    ]

    academic_phrases = [
        "answer all questions",
        "show your working",
        "section a",
        "section b",
        "marks",
        "total marks",
        "question ",
        "calculate",
        "solve",
        "describe",
        "explain",
        "discuss",
        "compare",
        "justify",
        "evaluate",
        "write your answer",
        "time allowed",
    ]

    academic_phrases = [
        "answer all questions",
        "show your working",
        "section a",
        "section b",
        "marks",
        "total marks",
        "question ",
        "calculate",
        "solve",
        "describe",
        "explain",
        "discuss",
        "compare",
        "justify",
        "evaluate",
        "write your answer",
        "time allowed",
    ]

    institution_terms = [
        "school", "college", "academy", "institute", "university",
        "grammar", "campus", "department", "faculty"
    ]

    field_line_terms = [
        "name", "student", "teacher", "class", "date", "term",
        "year", "grade", "subject", "candidate", "candidate name",
        "student name", "examiner", "duration", "time"
    ]

    def _clean_text(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _dedupe_texts(items):
        seen = set()
        out = []
        for item in items or []:
            s = _clean_text(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    def _looks_like_field_line(text: str) -> bool:
        s = _clean_text(text).lower()
        if not s:
            return False

        if re.search(r"^\s*(name|student|teacher|class|date|term|year|grade|subject|candidate)\s*[:_\.]", s):
            return True

        if re.search(r"\b(name|student|teacher|class|date|term|year|grade|subject|candidate)\b", s):
            if any(ch in text for ch in [":", "_", ".", "/"]):
                return True

        if re.search(r"\b(name|student|teacher|class|date|term|year|grade|subject|candidate)\b", s) and len(s) <= 60:
            return True

        return False

    def _looks_like_institutional_identity(text: str) -> bool:
        s = _clean_text(text)
        low = s.lower()
        if not s:
            return False

        if any(term in low for term in institution_terms):
            return True

        if re.search(r"\b(state school|high school|primary school|secondary school)\b", low):
            return True

        if len(s) <= 80 and s.isupper() and len(s.split()) <= 8:
            return True

        return False

    def _is_question_like(text: str) -> bool:
        s = _clean_text(text).lower()
        if not s:
            return False

        if re.search(r"^\d+[\.\)]\s", s):
            return True

        if re.search(r"^\([a-zA-Zivx]+\)\s", s):
            return True

        if re.search(r"\b\d+\s*marks?\b", s):
            return True

        return False

    def _contains_subject_baggage(text: str) -> bool:
        s = _clean_text(text).lower()
        if not s:
            return False

        if any(word in s for word in subject_words):
            return True

        if any(phrase in s for phrase in academic_phrases):
            return True

        if _is_question_like(s):
            return True

        return False

    def _should_preserve_text(text: str, context: str = "") -> bool:
        s = _clean_text(text)
        if not s:
            return False

        low = s.lower()

        if _looks_like_institutional_identity(s):
            return True

        if _looks_like_field_line(s):
            return True

        if context in {"header", "footer", "title_block"}:
            if len(s) <= 120 and not _is_question_like(s):
                return True

        if context == "paragraph":
            if len(s) <= 80 and any(term in low for term in field_line_terms):
                return True

        return False

    cleanup_audit = {
        "suppressed_paragraphs": [],
        "suppressed_header_texts": [],
        "suppressed_footer_texts": [],
        "suppressed_images": [],
        "suppressed_tables": [],
    }

    # -----------------------------
    # 1. Clean template_title
    # -----------------------------
    raw_title = _clean_text(profile.get("template_title") or "")
    if raw_title:
        if _contains_subject_baggage(raw_title) and not _looks_like_institutional_identity(raw_title):
            profile["template_title"] = "Institution Template"
        else:
            profile["template_title"] = raw_title

    # -----------------------------
    # 2. Clean paragraphs, but preserve structural evidence
    # -----------------------------
    cleaned_paragraphs = []
    preserved_paragraph_candidates = []
    suppressed_paragraphs = []

    for para in profile.get("paragraphs", []) or []:
        row = dict(para or {})
        text = _clean_text(row.get("text") or "")
        if not text:
            continue

        if _should_preserve_text(text, context="paragraph"):
            row["preservation_reason"] = "structural_or_identity_signal"
            cleaned_paragraphs.append(row)
            preserved_paragraph_candidates.append(text)
            continue

        if _contains_subject_baggage(text):
            row["suppression_reason"] = "likely_academic_content"
            suppressed_paragraphs.append(row)
            cleanup_audit["suppressed_paragraphs"].append(text)
            continue

        cleaned_paragraphs.append(row)

    profile["paragraphs"] = cleaned_paragraphs
    profile["suppressed_paragraphs"] = suppressed_paragraphs

    # -----------------------------
    # 3. Clean header/footer structures conservatively
    # -----------------------------
    header_footer = dict(profile.get("header_footer", {}) or {})

    def _clean_text_list(items, context_name):
        cleaned = []
        suppressed = []
        seen = set()

        for item in items or []:
            s = _clean_text(item)
            if not s:
                continue

            key = s.lower()
            if key in seen:
                continue
            seen.add(key)

            if _should_preserve_text(s, context=context_name):
                cleaned.append(s)
                continue

            if _contains_subject_baggage(s):
                suppressed.append(s)
                continue

            cleaned.append(s)

        return cleaned, suppressed

    def _clean_paragraph_list(items, context_name):
        cleaned = []
        suppressed = []
        seen = set()

        for item in items or []:
            row = dict(item or {})
            text = _clean_text(row.get("text") or "")
            if not text:
                continue

            key = text.lower()
            if key in seen:
                continue
            seen.add(key)

            if _should_preserve_text(text, context=context_name):
                row["preservation_reason"] = "header_footer_identity_signal"
                cleaned.append(row)
                continue

            if _contains_subject_baggage(text):
                row["suppression_reason"] = "likely_academic_header_footer_content"
                suppressed.append(row)
                continue

            cleaned.append(row)

        return cleaned, suppressed

    header_texts, suppressed_header_texts = _clean_text_list(header_footer.get("header_texts", []), "header")
    footer_texts, suppressed_footer_texts = _clean_text_list(header_footer.get("footer_texts", []), "footer")
    header_candidates, _ = _clean_text_list(header_footer.get("header_candidates", []), "header")
    footer_candidates, _ = _clean_text_list(header_footer.get("footer_candidates", []), "footer")
    header_paragraphs, suppressed_header_paragraphs = _clean_paragraph_list(header_footer.get("header_paragraphs", []), "header")
    footer_paragraphs, suppressed_footer_paragraphs = _clean_paragraph_list(header_footer.get("footer_paragraphs", []), "footer")

    header_footer["header_texts"] = header_texts
    header_footer["footer_texts"] = footer_texts
    header_footer["header_candidates"] = header_candidates
    header_footer["footer_candidates"] = footer_candidates
    header_footer["header_paragraphs"] = header_paragraphs
    header_footer["footer_paragraphs"] = footer_paragraphs
    header_footer["suppressed_header_texts"] = suppressed_header_texts
    header_footer["suppressed_footer_texts"] = suppressed_footer_texts
    header_footer["suppressed_header_paragraphs"] = suppressed_header_paragraphs
    header_footer["suppressed_footer_paragraphs"] = suppressed_footer_paragraphs

    cleanup_audit["suppressed_header_texts"] = suppressed_header_texts
    cleanup_audit["suppressed_footer_texts"] = suppressed_footer_texts

    profile["header_footer"] = header_footer

    # -----------------------------
    # 4. Clean style notes conservatively
    # -----------------------------
    cleaned_style_notes = []
    suppressed_style_notes = []

    for note in profile.get("style_notes", []) or []:
        s = _clean_text(note)
        if not s:
            continue

        if _contains_subject_baggage(s) and not _looks_like_institutional_identity(s):
            suppressed_style_notes.append(s)
            continue

        cleaned_style_notes.append(s)

    profile["style_notes"] = cleaned_style_notes
    profile["suppressed_style_notes"] = suppressed_style_notes

    # -----------------------------
    # 5. Classify extracted images conservatively
    # Keep all evidence, but separate likely branding from likely educational
    # -----------------------------
    kept_images = []
    suppressed_images = []

    for img in profile.get("extracted_images", []) or []:
        row = dict(img or {})
        position = str(row.get("position") or "").strip().lower()
        existing_role = str(row.get("image_role") or "").strip().lower()
        para_idx = row.get("paragraph_index", None)
        has_binary = bool(row.get("binary_data"))
        width_inches = float(row.get("width_inches") or 0) if str(row.get("width_inches") or "").strip() else 0.0
        height_inches = float(row.get("height_inches") or 0) if str(row.get("height_inches") or "").strip() else 0.0

        likely_branding = False
        reason = "uncertain"

        if existing_role == "branding":
            likely_branding = True
            reason = "upstream_branding_role"
        elif position in {"header", "footer", "first_page"}:
            likely_branding = True
            reason = "header_footer_or_first_page_position"
        elif has_binary and position == "body" and isinstance(para_idx, int) and para_idx <= 8:
            likely_branding = True
            reason = "early_body_image"
        elif width_inches and height_inches and width_inches <= 2.5 and height_inches <= 2.5:
            likely_branding = True
            reason = "small_logo_sized_image"

        row["image_role"] = "branding" if likely_branding else "educational"
        row["classification_reason"] = reason

        if likely_branding:
            kept_images.append(row)
        else:
            suppressed_images.append(row)
            cleanup_audit["suppressed_images"].append({
                "position": position,
                "reason": reason,
            })

    profile["extracted_images"] = kept_images
    profile["suppressed_images"] = suppressed_images

    # -----------------------------
    # 6. Classify extracted tables conservatively
    # Keep likely structural tables, but preserve everything suppressed
    # -----------------------------
    kept_tables = []
    suppressed_tables = []

    for tbl in profile.get("extracted_tables", []) or []:
        row = dict(tbl or {})
        rows = int(row.get("rows", 0) or 0)
        cols = int(row.get("cols", 0) or 0)
        has_header_row = bool(row.get("has_header_row", False))

        sample_text = " ".join(
            [
                _clean_text(x)
                for x in row.get("sample_texts", []) or []
                if _clean_text(x)
            ]
        ).strip()

        looks_like_field_table = bool(sample_text) and _looks_like_field_line(sample_text)
        very_small_table = rows <= 3 and cols <= 6
        moderate_front_matter_table = rows <= 5 and cols <= 4

        likely_structural = (
            looks_like_field_table
            or (very_small_table and not has_header_row)
            or moderate_front_matter_table
        )

        row["table_role"] = "structural" if likely_structural else "academic"
        row["classification_reason"] = (
            "field_line_table" if looks_like_field_table
            else "small_or_front_matter_table" if likely_structural
            else "likely_academic_table"
        )

        if likely_structural:
            kept_tables.append(row)
        else:
            suppressed_tables.append(row)
            cleanup_audit["suppressed_tables"].append({
                "rows": rows,
                "cols": cols,
                "reason": row["classification_reason"],
            })

    profile["extracted_tables"] = kept_tables
    profile["suppressed_tables"] = suppressed_tables

    # -----------------------------
    # 7. Preserve useful identity candidates
    # -----------------------------
    institution_identity = dict(profile.get("institution_identity", {}) or {})

    raw_identity_candidates = []
    raw_identity_candidates.extend(header_footer.get("header_candidates", []) or [])
    raw_identity_candidates.extend(header_footer.get("footer_candidates", []) or [])
    raw_identity_candidates.extend(preserved_paragraph_candidates)

    title_block_signals = dict(profile.get("title_block_signals", {}) or {})
    for key in ["school_candidates", "identity_lines", "top_lines", "field_lines"]:
        for item in title_block_signals.get(key, []) or []:
            text = _clean_text(item if isinstance(item, str) else item.get("text", ""))
            if text:
                raw_identity_candidates.append(text)

    raw_identity_candidates = _dedupe_texts(raw_identity_candidates)
    likely_identity_candidates = [x for x in raw_identity_candidates if _looks_like_institutional_identity(x)]

    institution_identity["raw_identity_candidates"] = raw_identity_candidates[:25]
    institution_identity["likely_identity_candidates"] = likely_identity_candidates[:10]

    if not institution_identity.get("likely_institution_name") and likely_identity_candidates:
        institution_identity["likely_institution_name"] = likely_identity_candidates[0]

    profile["institution_identity"] = institution_identity

    # -----------------------------
    # 8. Cleanup audit + completion flag
    # -----------------------------
    profile["cleanup_audit"] = cleanup_audit
    profile["subject_baggage_stripped"] = True

    return profile


def _extract_donor_with_styles(docx_path: str) -> Dict[str, Any]:
    """
    Extract ACTUAL styles from donor DOCX using python-docx.
    
    ENHANCED EXTRACTION:
        - Full font properties (name, size, bold, italic, color, underline)
        - Paragraph spacing (before, after, line spacing)
        - Table styles (borders, shading, cell merging, column widths)
        - Images with positioning and dimensions
        - Section properties (margins, columns, headers, footers)
        - List styles and numbering
    """
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available, using fallback extraction")
        return _extract_donor_fallback(docx_path)
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        import io
        
        doc = Document(docx_path)
        
        # Initialize result structure
        result = {
            "source_type": "docx",
            "source_filename": os.path.basename(docx_path),
            "template_title": os.path.splitext(os.path.basename(docx_path))[0],
            "institution_identity": {},
            "title_block_signals": {},
            "style_preferences": {},
            "layout_features": [],
            "page_setup": {},
            "header_footer": {},
            "text_stats": {},
            "style_notes": [],
            "font_samples": [],
            "paragraphs": [],
            "extracted_images": [],
            "extracted_tables": [],
            "list_styles": []
        }
        
        # Extract page setup from first section
        if doc.sections:
            section = doc.sections[0]
            
            # Margins
            result["page_setup"] = {
                "left_margin_inches": section.left_margin.inches if section.left_margin else 1.0,
                "right_margin_inches": section.right_margin.inches if section.right_margin else 1.0,
                "top_margin_inches": section.top_margin.inches if section.top_margin else 1.0,
                "bottom_margin_inches": section.bottom_margin.inches if section.bottom_margin else 1.0,
                "orientation": "landscape" if section.orientation and section.orientation.name == "LANDSCAPE" else "portrait",
                "page_width_inches": section.page_width.inches if section.page_width else 8.5,
                "page_height_inches": section.page_height.inches if section.page_height else 11.0,
                "header_distance_inches": section.header_distance.inches if section.header_distance else 0.5,
                "footer_distance_inches": section.footer_distance.inches if section.footer_distance else 0.5,
            }
            
            # Detect columns
            cols = section._sectPr.xpath('./w:cols')
            if cols:
                col_elem = cols[0]
                col_count = int(col_elem.get(qn('w:num'), 1))
                col_space = float(col_elem.get(qn('w:space'), 10)) / 20.0 if col_elem.get(qn('w:space')) else 0.5
                
                result["page_setup"]["columns"] = {
                    "count": col_count,
                    "space_inches": col_space,
                    "equal_width": col_elem.get(qn('w:eq')) != "0"
                }
                
                if col_count > 1:
                    # Check if this is math two-column layout using enhanced detection
                    math_detection = _detect_math_columns(doc)
                    is_math_columns = math_detection["is_math_columns"]
                    feature_category = FeatureCategory.NEVER_INHERIT if is_math_columns else FeatureCategory.CONTEXTUAL_INHERIT
                    
                    result["layout_features"].append({
                        "feature_name": "two_column_math" if is_math_columns else "columns",
                        "category": feature_category,
                        "extracted_value": {
                            "count": col_count, 
                            "space_inches": col_space,
                            "equal_width": result["page_setup"]["columns"]["equal_width"]
                        }
                    })
            else:
                result["page_setup"]["columns"] = {"count": 1, "space_inches": 0, "equal_width": True}
        
        # Extract fonts, paragraph styles, and runs
        font_samples = []
        paragraph_styles_seen = set()
        run_properties_seen = set()
        
        for para in doc.paragraphs[:200]:  # Increased sample size
            # Store paragraph text and style
            if para.text.strip():
                para_info = {
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                    "alignment": para.alignment.name if para.alignment else None,
                }
                
                # Extract paragraph spacing
                if para.paragraph_format:
                    para_info["space_before"] = para.paragraph_format.space_before.pt if para.paragraph_format.space_before else None
                    para_info["space_after"] = para.paragraph_format.space_after.pt if para.paragraph_format.space_after else None
                    para_info["line_spacing"] = para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else None
                    para_info["first_line_indent"] = para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else None
                    para_info["left_indent"] = para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else None
                
                result["paragraphs"].append(para_info)
            
            # Extract run properties (fonts)
            if para.runs:
                for run in para.runs:
                    font_info = _extract_run_properties_enhanced(run)
                    if font_info:
                        font_samples.append(font_info)
                        # Track unique run property combinations
                        run_key = f"{font_info.get('name')}_{font_info.get('size_pt')}_{font_info.get('bold')}_{font_info.get('italic')}"
                        run_properties_seen.add(run_key)
            
            # Track paragraph styles
            if para.style and para.style.name:
                style_name = para.style.name
                if style_name not in paragraph_styles_seen:
                    paragraph_styles_seen.add(style_name)
        
        result["font_samples"] = font_samples
        result["run_property_variations"] = len(run_properties_seen)
        
        # Extract style preferences from font samples
        if font_samples:
            font_counter = Counter([f.get("name") for f in font_samples if f.get("name")])
            if font_counter:
                result["style_preferences"]["font_family"] = font_counter.most_common(1)[0][0]
            
            size_counter = Counter([f.get("size_pt") for f in font_samples if f.get("size_pt")])
            if size_counter:
                result["style_preferences"]["font_size_pt"] = size_counter.most_common(1)[0][0]
            
            # Detect if document uses multiple fonts
            result["style_preferences"]["unique_fonts"] = len(font_counter)
            result["style_preferences"]["font_variations"] = len(run_properties_seen)
        
        result["style_preferences"]["paragraph_styles_detected"] = list(paragraph_styles_seen)[:15]
        
        # Detect heading styles
        heading_styles = [s for s in paragraph_styles_seen if "Heading" in s or "Title" in s or "Head" in s]
        result["style_preferences"]["heading_styles"] = heading_styles[:5]
        
        # Extract front-page/title-block body signals
        result["title_block_signals"] = _extract_title_block_signals(doc)

        # Extract header/footer with full content
        result["header_footer"] = _extract_header_footer_enhanced(doc, docx_path)

        # Extract BS5 first-page layout/style contract (defensive, partial OK)
        result["first_page_layout_contract"] = _extract_first_page_layout_contract(docx_path)

        # Detect layout features
        result["layout_features"].extend(_detect_layout_features_enhanced(doc))
        
        # Extract images with full metadata
        result["extracted_images"] = _extract_images_enhanced(doc, docx_path)
        
        # Extract table structures with full styling
        result["extracted_tables"] = _extract_table_structures_enhanced(doc)
        
        # Extract list styles (bullet points, numbering)
        result["list_styles"] = _extract_list_styles(doc)

        # --------------------------------------------------
        # Donor-level integration pass
        # Consolidate enriched BS3 raw evidence before cleanup
        # --------------------------------------------------
        title_block = dict(result.get("title_block_signals", {}) or {})
        header_footer = dict(result.get("header_footer", {}) or {})
        extracted_images = list(result.get("extracted_images", []) or [])
        branding_ocr = _extract_text_candidates_from_branding_image(extracted_images)
        extracted_tables = list(result.get("extracted_tables", []) or [])

        def _clean_text_local(s: str) -> str:
            return (s or "").replace("\xa0", " ").strip()

        def _dedupe_keep_order_local(items):
            seen = set()
            out = []
            for item in items or []:
                s = _clean_text_local(item)
                if not s:
                    continue
                key = s.lower()
                if key in seen:
                    continue
                seen.add(key)
                out.append(s)
            return out

        institution_identity = dict(result.get("institution_identity", {}) or {})

        raw_identity_candidates = []
        raw_identity_candidates.extend(title_block.get("school_candidates", []) or [])
        raw_identity_candidates.extend(title_block.get("identity_lines", []) or [])
        raw_identity_candidates.extend(header_footer.get("header_candidates", []) or [])
        raw_identity_candidates.extend(header_footer.get("footer_candidates", []) or [])

        raw_identity_candidates.extend(branding_ocr.get("school_candidates", []) or [])
        raw_identity_candidates = _dedupe_keep_order_local(raw_identity_candidates)

        likely_identity_candidates = _dedupe_keep_order_local(
            list(title_block.get("identity_lines", []) or []) +
            list(branding_ocr.get("school_candidates", []) or [])
        )[:10]

        likely_institution_name = institution_identity.get("likely_institution_name", "")
        if not likely_institution_name:
            likely_institution_name = branding_ocr.get("likely_institution_name", "") or ""
        if not likely_institution_name and raw_identity_candidates:
            likely_institution_name = raw_identity_candidates[0]

        institution_identity["raw_identity_candidates"] = raw_identity_candidates[:25]
        institution_identity["likely_identity_candidates"] = likely_identity_candidates
        institution_identity["branding_ocr_lines"] = list(branding_ocr.get("ocr_lines", []) or [])[:20]
        institution_identity["branding_ocr_school_candidates"] = list(branding_ocr.get("school_candidates", []) or [])[:10]
        if likely_institution_name:
            institution_identity["likely_institution_name"] = likely_institution_name

        branding_images = [img for img in extracted_images if str(img.get("image_role", "")).lower() == "branding"]
        likely_front_matter_tables = [tbl for tbl in extracted_tables if bool(tbl.get("likely_front_matter_table", False))]
        field_like_tables = [tbl for tbl in extracted_tables if bool(tbl.get("looks_like_field_table", False))]

        institution_identity["branding_image_count"] = len(branding_images)
        institution_identity["likely_front_matter_table_count"] = len(likely_front_matter_tables)

        result["institution_identity"] = institution_identity

        result["donor_signal_summary"] = {
            "title_block_identity_count": len(title_block.get("identity_lines", []) or []),
            "title_block_framing_count": len(title_block.get("framing_lines", []) or []),
            "header_candidate_count": len(header_footer.get("header_candidates", []) or []),
            "footer_candidate_count": len(header_footer.get("footer_candidates", []) or []),
            "branding_image_count": len(branding_images),
            "branding_ocr_line_count": len(branding_ocr.get("ocr_lines", []) or []),
            "branding_ocr_school_candidate_count": len(branding_ocr.get("school_candidates", []) or []),
            "front_matter_table_count": len(likely_front_matter_tables),
            "field_like_table_count": len(field_like_tables),
        }
        
        # Text stats
        full_text = extract_text_from_docx(docx_path) or ""
        result["text_stats"] = {
            "char_count": len(full_text),
            "line_count": len(full_text.split("\n")),
            "table_count": len(doc.tables),
            "paragraph_count": len(doc.paragraphs),
            "run_count": len(font_samples),
            "image_count": len(result["extracted_images"])
        }
        
        result["style_notes"] = [
            "Enhanced DOCX styles extracted successfully.",
            f"Font family: {result['style_preferences'].get('font_family', 'Not detected')}",
            f"Columns detected: {result['page_setup'].get('columns', {}).get('count', 1)}",
            f"Images extracted: {len(result['extracted_images'])}",
            f"Branding images detected: {len([img for img in result.get('extracted_images', []) if str(img.get('image_role', '')).lower() == 'branding'])}",
            f"Branding OCR school candidates: {len(branding_ocr.get('school_candidates', []) or [])}",
            f"Tables extracted: {len(result['extracted_tables'])}",
            f"Likely front-matter tables: {len([tbl for tbl in result.get('extracted_tables', []) if bool(tbl.get('likely_front_matter_table', False))])}",
            f"Title-block identity lines: {len(result.get('title_block_signals', {}).get('identity_lines', []) or [])}",
        ]
        
        result["extraction_debug_report"] = _build_donor_extraction_debug_report(result)

        # STRIP SUBJECT BAGGAGE - Enforce Rule 1
        result = _strip_subject_baggage_from_donor(result)
        
        return result
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return _extract_donor_fallback(docx_path)


def _extract_donor_fallback(docx_path: str) -> Dict[str, Any]:
    """
    Safe fallback donor extractor for DOCX when full python-docx style extraction fails.

    Build Sequence 3 role:
    - never crash the donor pipeline
    - return a donor profile in the same broad structural shape as the main extractor
    - preserve minimal but useful raw donor evidence
    """
    text = ""
    try:
        text = extract_text_from_docx(docx_path) or ""
    except Exception as e:
        logger.warning(f"Fallback DOCX text extraction failed: {e}")
        text = ""

    def _clean_text(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _dedupe_keep_order(items):
        seen = set()
        out = []
        for item in items or []:
            s = _clean_text(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    def _looks_like_field_line(text: str) -> bool:
        s = _clean_text(text).lower()
        if not s:
            return False
        if re.search(r"^\s*(name|student|teacher|class|date|term|year|grade|subject|candidate|duration|examiner)\s*[:_\.]", s):
            return True
        if re.search(r"\b(name|student|teacher|class|date|term|year|grade|subject|candidate|duration|examiner)\b", s):
            if any(ch in text for ch in [":", "_", ".", "/"]):
                return True
        return False

    def _looks_like_identity(text: str) -> bool:
        s = _clean_text(text)
        low = s.lower()
        if not s:
            return False
        if re.search(r"\b(high school|primary school|secondary school|college|grammar|academy|institute|university|school)\b", low):
            return True
        if len(s) <= 80 and s.isupper() and len(s.split()) <= 8 and not re.search(r"\d", s):
            return True
        return False

    def _looks_like_framing(text: str) -> bool:
        s = _clean_text(text).lower()
        if not s:
            return False
        framing_terms = [
            "question/answer booklet", "examination", "exam", "test", "worksheet",
            "assessment", "semester", "calculator assumed", "calculator free",
            "instructions", "time allowed", "reading time"
        ]
        return any(term in s for term in framing_terms)

    lines = [_clean_text(ln) for ln in text.splitlines() if _clean_text(ln)]
    top_lines = lines[:30]
    header_candidates = _dedupe_keep_order(lines[:3])
    footer_candidates = _dedupe_keep_order(lines[-3:] if len(lines) >= 3 else lines[:])

    identity_lines = _dedupe_keep_order([ln for ln in top_lines if _looks_like_identity(ln)])[:15]
    field_lines = _dedupe_keep_order([ln for ln in top_lines if _looks_like_field_line(ln)])[:20]
    framing_lines = _dedupe_keep_order([ln for ln in top_lines if _looks_like_framing(ln)])[:15]

    raw_identity_candidates = _dedupe_keep_order(
        identity_lines + header_candidates + footer_candidates
    )[:25]

    likely_institution_name = raw_identity_candidates[0] if raw_identity_candidates else ""

    title_block_signals = {
        "top_lines": top_lines[:30],
        "school_candidates": identity_lines[:15],
        "field_lines": field_lines[:20],
        "title_candidates": framing_lines[:15],
        "identity_lines": identity_lines[:15],
        "framing_lines": framing_lines[:15],
        "admin_lines": field_lines[:20],
        "acronym_candidates": [],
        "source_tagged_lines": [
            {"text": ln, "source": "fallback_text", "index": idx}
            for idx, ln in enumerate(top_lines[:30])
        ],
    }

    institution_identity = {
        "raw_identity_candidates": raw_identity_candidates,
        "likely_identity_candidates": identity_lines[:10],
        "likely_institution_name": likely_institution_name,
        "branding_image_count": 0,
        "likely_front_matter_table_count": 0,
    }

    raw_profile = {
        "source_type": "docx",
        "source_filename": os.path.basename(docx_path),
        "template_title": os.path.splitext(os.path.basename(docx_path))[0],
        "institution_identity": institution_identity,
        "title_block_signals": title_block_signals,
        "style_preferences": {
            "font_family": "Not extracted (fallback mode)",
            "font_size_pt": 11,
            "heading_style_hint": "Fallback mode",
            "visual_feel": "Fallback extraction only",
            "paragraph_styles_detected": [],
            "has_font_extraction": False,
            "unique_fonts": 0,
            "font_variations": 0,
            "heading_styles": [],
        },
        "layout_features": [],
        "page_setup": {
            "left_margin_inches": 1.0,
            "right_margin_inches": 1.0,
            "top_margin_inches": 1.0,
            "bottom_margin_inches": 1.0,
            "orientation": "portrait",
            "page_width_inches": 8.5,
            "page_height_inches": 11.0,
            "header_distance_inches": 0.5,
            "footer_distance_inches": 0.5,
            "columns": {"count": 1, "space_inches": 0, "equal_width": True},
        },
        "header_footer": {
            "has_header": bool(header_candidates),
            "has_footer": bool(footer_candidates),
            "header_texts": header_candidates,
            "footer_texts": footer_candidates,
            "header_candidates": header_candidates,
            "footer_candidates": footer_candidates,
            "header_paragraphs": [{"text": x, "style": None, "alignment": None} for x in header_candidates],
            "footer_paragraphs": [{"text": x, "style": None, "alignment": None} for x in footer_candidates],
            "header_tables": [],
            "footer_tables": [],
            "suppressed_header_texts": [],
            "suppressed_footer_texts": [],
            "suppressed_header_paragraphs": [],
            "suppressed_footer_paragraphs": [],
        },
        "text_stats": {
            "char_count": len(text),
            "line_count": len(text.splitlines()) if text else 0,
            "table_count": 0,
            "paragraph_count": len(lines),
            "run_count": 0,
            "image_count": 0,
        },
        "style_notes": [
            "Fallback donor extraction used.",
            "Full DOCX style extraction failed, so only minimal text-based signals were captured.",
            f"Fallback identity candidates: {len(raw_identity_candidates)}",
            f"Fallback title-block identity lines: {len(identity_lines)}",
            f"Fallback field lines: {len(field_lines)}",
        ],
        "font_samples": [],
        "run_property_variations": 0,
        "paragraphs": [{"text": ln} for ln in lines[:80]],
        "suppressed_paragraphs": [],
        "extracted_images": [],
        "suppressed_images": [],
        "extracted_tables": [],
        "suppressed_tables": [],
        "list_styles": [],
        "donor_signal_summary": {
            "title_block_identity_count": len(identity_lines),
            "title_block_framing_count": len(framing_lines),
            "header_candidate_count": len(header_candidates),
            "footer_candidate_count": len(footer_candidates),
            "branding_image_count": 0,
            "front_matter_table_count": 0,
            "field_like_table_count": 0,
        },
    }

    raw_profile["header_footer"] = _merge_docx_xml_fallback_into_header_footer(raw_profile.get("header_footer", {}), docx_path)
    raw_profile["first_page_layout_contract"] = _extract_first_page_layout_contract(docx_path)
    raw_profile["extracted_images"] = _merge_docx_media_fallback_into_images(raw_profile.get("extracted_images", []), docx_path)
    raw_profile["extraction_debug_report"] = _build_donor_extraction_debug_report(raw_profile)
    raw_profile = _strip_subject_baggage_from_donor(raw_profile)
    return raw_profile

# ======================================================================================
# BUILD SEQUENCE 4: DONOR INTERPRETATION ENGINE (DIE)
# ======================================================================================

def _interpret_donor_profile(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Build Sequence 4 boundary.

    Convert BS3 raw donor evidence into a structured interpreted donor model.
    This is the donor-side intelligence bridge between extraction and normalization.

    Output shape:
        {
            "schema_version": "interpreted_donor_v1",
            "document_type": "...",
            "identity": {...},
            "framing": {...},
            "structure": {...},
            "discarded_content": {...},
            "inheritance_map": {...},
            "reason_log": [...]
        }
    """
    raw_profile = raw_profile or {}
    if not isinstance(raw_profile, dict):
        raw_profile = {}

    identity = _interpret_donor_identity(raw_profile)
    framing = _interpret_donor_framing(raw_profile)
    structure = _interpret_donor_structure(raw_profile)
    images = _interpret_donor_images(raw_profile)
    discarded = _interpret_donor_discarded_content(raw_profile)

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _contains_any(text: str, terms: List[str]) -> bool:
        low = (text or "").lower()
        for term in terms:
            if term in low:
                return True
        return False

    title_block = dict(raw_profile.get("title_block_signals", {}) or {})
    header_footer = dict(raw_profile.get("header_footer", {}) or {})

    donor_text_pool = []

    donor_text_pool.extend(list(title_block.get("title_candidates", []) or []))
    donor_text_pool.extend(list(title_block.get("framing_lines", []) or []))
    donor_text_pool.extend(list(header_footer.get("header_texts", []) or []))
    donor_text_pool.extend(list(framing.get("document_frame_lines", []) or []))
    donor_text_pool.extend(list(framing.get("generic_instructions", []) or []))
    donor_text_pool.extend(list(framing.get("warning_lines", []) or []))

    donor_text_pool = [_clean(x) for x in donor_text_pool if _clean(x)]
    donor_blob = " | ".join(donor_text_pool).lower()

    exam_terms = [
        "examination", "exam", "question/answer booklet", "question answer booklet",
        "reading time", "working time", "calculator assumed", "calculator free"
    ]
    test_terms = [
        "test", "class test", "quiz"
    ]
    worksheet_terms = [
        "worksheet", "practice", "revision sheet", "activity sheet"
    ]
    memo_terms = [
        "memo", "memorandum", "marking guide", "answer guide", "solutions"
    ]
    rubric_terms = [
        "rubric", "criteria sheet", "assessment criteria"
    ]
    assignment_terms = [
        "assignment", "task sheet", "project brief"
    ]

    document_type = "custom"
    document_type_reason = "No strong donor document-type signal detected."

    if _contains_any(donor_blob, exam_terms):
        document_type = "exam"
        document_type_reason = "Donor framing contains strong examination signals such as booklet/timing/calculator wording."
    elif _contains_any(donor_blob, test_terms):
        document_type = "test"
        document_type_reason = "Donor framing contains test-style wording."
    elif _contains_any(donor_blob, worksheet_terms):
        document_type = "worksheet"
        document_type_reason = "Donor framing contains worksheet/practice wording."
    elif _contains_any(donor_blob, memo_terms):
        document_type = "memo"
        document_type_reason = "Donor framing contains memo/solutions wording."
    elif _contains_any(donor_blob, rubric_terms):
        document_type = "rubric"
        document_type_reason = "Donor framing contains rubric/criteria wording."
    elif _contains_any(donor_blob, assignment_terms):
        document_type = "assignment"
        document_type_reason = "Donor framing contains assignment/task wording."

    inheritance_map = _build_donor_inheritance_map(
        identity=identity,
        framing=framing,
        structure=structure,
        discarded_content=discarded,
    )

    reason_log = []
    reason_log.append({
        "layer": "document_type",
        "decision": "document_type",
        "value": document_type,
        "reason": document_type_reason,
    })
    reason_log.extend(identity.get("reason_log", []))
    reason_log.extend(framing.get("reason_log", []))
    reason_log.extend(structure.get("reason_log", []))
    reason_log.extend(images.get("reason_log", []))
    reason_log.extend(discarded.get("reason_log", []))

    return {
        "schema_version": "interpreted_donor_v1",
        "document_type": document_type,
        "identity": identity,
        "framing": framing,
        "structure": structure,
        "images": images,
        "discarded_content": discarded,
        "inheritance_map": inheritance_map,
        "reason_log": reason_log,
    }


def _interpret_donor_identity(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Interpret donor identity signals into a stable institutional identity block.
    BS4 rule: identify the institution, not the donor subject.
    """
    raw_profile = raw_profile or {}

    institution_identity = dict(raw_profile.get("institution_identity", {}) or {})
    header_footer = dict(raw_profile.get("header_footer", {}) or {})
    title_block = dict(raw_profile.get("title_block_signals", {}) or {})

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _dedupe(items):
        seen = set()
        out = []
        for item in items or []:
            s = _clean(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    def _split_header_identity_candidates(lines):
        """
        Break header/title identity lines into smaller identity candidates.
        Example:
            'Year 8 Mathematics   RSHS   Exam 2'
        should allow 'RSHS' to surface as a candidate.
        """
        out = []
        for line in lines or []:
            s = _clean(line)
            if not s:
                continue

            normalized = s.replace("\t", " ")
            parts = [p.strip() for p in normalized.split("  ") if p.strip()]

            if not parts:
                parts = normalized.split()

            # Keep the full line for fallback inspection
            out.append(s)

            # Add split chunks
            for part in parts:
                part = _clean(part)
                if part:
                    out.append(part)

            # Add token-level uppercase acronym candidates
            for token in normalized.replace("/", " ").replace("-", " ").split():
                token = _clean(token)
                if token.isupper() and token.isalpha() and 2 <= len(token) <= 8:
                    out.append(token)

        return _dedupe(out)

    upstream_likely_name = _clean(institution_identity.get("likely_institution_name", ""))

    raw_header_lines = _dedupe(
        list(header_footer.get("header_identity_lines", []) or [])
        + list(header_footer.get("header_candidates", []) or [])
        + list(header_footer.get("header_texts", []) or [])
    )

    raw_footer_lines = _dedupe(
        list(header_footer.get("footer_identity_lines", []) or [])
        + list(header_footer.get("footer_candidates", []) or [])
        + list(header_footer.get("footer_texts", []) or [])
    )

    split_header_candidates = _split_header_identity_candidates(raw_header_lines)
    split_footer_candidates = _split_header_identity_candidates(raw_footer_lines)

    raw_school_candidates = _dedupe(
        list(institution_identity.get("school_candidates", []) or [])
        + list(institution_identity.get("raw_identity_candidates", []) or [])
        + list(institution_identity.get("likely_identity_candidates", []) or [])
        + list(title_block.get("school_candidates", []) or [])
        + list(title_block.get("identity_lines", []) or [])
        + split_header_candidates
        + split_footer_candidates
    )

    raw_acronym_candidates = _dedupe(
        list(title_block.get("acronym_candidates", []) or [])
        + [
            s for s in raw_school_candidates
            if s.isupper() and len(s) <= 10 and " " not in s and s.isalpha()
        ]
    )

    ocr_identity_lines = _dedupe(
        list(institution_identity.get("branding_ocr_lines", []) or [])
        + list(institution_identity.get("branding_ocr_school_candidates", []) or [])
    )

    SUBJECT_KEYWORDS = {
        "mathematics", "maths", "math", "science", "biology", "chemistry",
        "physics", "english", "history", "geography", "economics",
        "business", "accounting", "technology", "engineering",
        "exam", "test", "worksheet", "assignment", "memo", "rubric",
        "semester", "section"
    }

    FRAMING_KEYWORDS = {
        "time allowed", "reading time", "working time", "important note",
        "question/answer booklet", "question answer booklet", "calculator assumed",
        "calculator free", "marks", "to be provided",
        "material required", "student", "teacher", "name:"
    }

    GENERIC_BAD_IDENTITY = {
        "official", "draft", "confidential"
    }

    INSTITUTION_TERMS = (
        "school", "college", "academy", "institute", "university",
        "grammar", "campus", "faculty", "department", "primary", "secondary"
    )

    def _has_heavy_symbol_noise(s: str) -> bool:
        s = _clean(s)
        if not s:
            return False
        bad_chars = sum(1 for ch in s if not ch.isalnum() and ch not in {" ", "&", "'", "."})
        return bad_chars >= 2

    def _looks_like_ocr_garbage(s: str) -> bool:
        s = _clean(s)
        if not s:
            return True

        letters = sum(1 for ch in s if ch.isalpha())
        digits = sum(1 for ch in s if ch.isdigit())

        # Too many symbols / punctuation
        if _has_heavy_symbol_noise(s):
            return True

        # Tiny meaningless fragments
        if len(s) <= 2:
            return True

        # Mixed noisy OCR fragments like x 0 Senior a / Se % / x | 3
        if digits > 0 and letters > 0 and len(s.split()) <= 4:
            return True

        # Weird one-letter chopped words
        tokens = s.split()
        if len(tokens) >= 2:
            short_tokens = sum(1 for t in tokens if len(t) == 1)
            if short_tokens >= 2:
                return True

        return False

    def _is_bad_identity_line(s: str) -> bool:
        s = _clean(s)
        low = s.lower()

        if not s:
            return True

        if len(s) > 80:
            return True

        if low in GENERIC_BAD_IDENTITY:
            return True

        if re.fullmatch(r"(year|grade)\s*\d{1,2}", low):
            return True

        if any(word in low for word in SUBJECT_KEYWORDS):
            return True

        if any(word in low for word in FRAMING_KEYWORDS):
            return True

        if "intentionally blank" in low:
            return True

        if re.fullmatch(r"[_\-\=\.\s]+", s):
            return True

        if _looks_like_ocr_garbage(s):
            return True

        return False

    def _looks_like_full_institution_name(s: str) -> bool:
        s = _clean(s)
        low = s.lower()
        if not s or _is_bad_identity_line(s):
            return False

        # Strong institution terms always count
        if any(term in low for term in INSTITUTION_TERMS):
            return True

        # Require at least 2 real words, no digits, not all uppercase noise
        words = s.split()
        if len(words) >= 2 and not s.isupper():
            if not any(ch.isdigit() for ch in s):
                if all(any(ch.isalpha() for ch in w) for w in words):
                    return True

        # Uppercase full names only if institution term is present
        if len(words) >= 2 and s.isupper() and any(term in low for term in INSTITUTION_TERMS):
            return True

        return False

    def _looks_like_safe_acronym(s: str) -> bool:
        s = _clean(s)
        if not s:
            return False
        return s.isupper() and s.isalpha() and 2 <= len(s) <= 8

    filtered_school_candidates = [
        s for s in raw_school_candidates
        if not _is_bad_identity_line(s)
    ]

    filtered_running_header_identity = [
        s for s in split_header_candidates
        if not _is_bad_identity_line(s)
    ]

    filtered_running_footer_identity = [
        s for s in split_footer_candidates
        if not _is_bad_identity_line(s)
    ]

    filtered_ocr_identity = [
        s for s in ocr_identity_lines
        if not _is_bad_identity_line(s)
    ]

    filtered_acronym_candidates = [
        s for s in raw_acronym_candidates
        if _looks_like_safe_acronym(s) and not _is_bad_identity_line(s)
    ]

    first_page_identity_lines = _dedupe(
        list(title_block.get("identity_lines", []) or [])
        + list(title_block.get("school_candidates", []) or [])
    )
    first_page_identity_lines = [
        s for s in first_page_identity_lines
        if not _is_bad_identity_line(s)
    ]

    image_interpretation = _interpret_donor_images(raw_profile)
    branding_images = image_interpretation.get("branding_images", []) or []
    structural_images = image_interpretation.get("structural_images", []) or []

    has_logo = bool(
        institution_identity.get("has_logo", False)
        or institution_identity.get("branding_image_count", 0)
        or branding_images
    )

    logo_positions = _dedupe(
        list(institution_identity.get("logo_positions", []) or [])
        + [
            str(img.get("position", "")).strip().lower()
            for img in branding_images
            if str(img.get("position", "")).strip()
        ]
    )

    official_labels = []

    official_source_lines = []
    official_source_lines.extend(list(title_block.get("top_lines", []) or []))
    official_source_lines.extend(list(header_footer.get("header_texts", []) or []))
    official_source_lines.extend(list(header_footer.get("header_candidates", []) or []))
    official_source_lines.extend(list(header_footer.get("footer_texts", []) or []))

    for row in list(title_block.get("source_tagged_lines", []) or []):
        if isinstance(row, dict):
            official_source_lines.append(row.get("text", ""))
        else:
            official_source_lines.append(row)

    for line in official_source_lines:
        s = _clean(line)
        if not s:
            continue

        matches = re.findall(r"\b(official|confidential|draft)\b", s, re.I)
        for match in matches:
            label = match.upper()
            official_labels.append(label)

    official_labels = _dedupe(official_labels)

    # --------------------------------------------------
    # NEW — split first-page vs running header labels
    # --------------------------------------------------

    first_page_header_labels = []
    running_header_labels = []

    def _extract_admin_labels_from_line(line: str):
        labels = []
        s = _clean(line)
        if not s:
            return labels

        for match in re.findall(r"\b(official|confidential|draft)\b", s, re.I):
            labels.append(match.upper())

        return labels

    # FIRST PAGE HEADER labels only
    for row in list(header_footer.get("first_page_header_textboxes", []) or []):
        for label in _extract_admin_labels_from_line(row):
            first_page_header_labels.append(label)

    # RUNNING HEADER labels only
    for row in (
        list(header_footer.get("header_texts", []) or []) +
        list(header_footer.get("header_candidates", []) or [])
    ):
        for label in _extract_admin_labels_from_line(row):
            running_header_labels.append(label)

    first_page_header_labels = _dedupe(first_page_header_labels)
    running_header_labels = _dedupe(running_header_labels)

    candidate_score_rows = []

    combined_candidates = _dedupe(
        filtered_school_candidates
        + filtered_running_header_identity
        + filtered_running_footer_identity
        + first_page_identity_lines
        + filtered_acronym_candidates
        + ([upstream_likely_name] if upstream_likely_name and not _is_bad_identity_line(upstream_likely_name) else [])
        + filtered_ocr_identity
    )

    for cand in combined_candidates:
        s = _clean(cand)
        low = s.lower()
        score = 0
        reasons = []

        if s in first_page_identity_lines:
            score += 4
            reasons.append("first_page_identity")

        if s in filtered_running_header_identity:
            score += 6
            reasons.append("running_header_identity")

        if s in filtered_running_footer_identity:
            score += 3
            reasons.append("running_footer_identity")

        if s in filtered_school_candidates:
            score += 3
            reasons.append("school_candidate_pool")

        if s in filtered_ocr_identity:
            score += 1
            reasons.append("branding_ocr_support")

        if upstream_likely_name and s == upstream_likely_name:
            score += 2
            reasons.append("upstream_likely_name")

        if _looks_like_full_institution_name(s):
            score += 5
            reasons.append("full_institution_name_shape")

        if _looks_like_safe_acronym(s):
            score += 4
            reasons.append("acronym_shape")

            supported = (
                s in filtered_running_header_identity
                or s in first_page_identity_lines
                or s in filtered_school_candidates
            )
            if supported:
                score += 4
                reasons.append("supported_acronym")
            else:
                score -= 4
                reasons.append("unsupported_acronym_penalty")

        # Penalise OCR-only candidates hard unless they have stronger support elsewhere
        if s in filtered_ocr_identity and s not in filtered_running_header_identity and s not in filtered_school_candidates:
            score -= 4
            reasons.append("ocr_only_penalty")

        if low in GENERIC_BAD_IDENTITY:
            score -= 10
            reasons.append("generic_bad_identity_penalty")

        candidate_score_rows.append({
            "candidate": s,
            "score": score,
            "reasons": reasons,
        })

    candidate_score_rows.sort(
        key=lambda row: (row["score"], len(row["candidate"])),
        reverse=True
    )

    likely_name = candidate_score_rows[0]["candidate"] if candidate_score_rows else ""

    reason_log = [
        {
            "layer": "identity",
            "decision": "institution_name",
            "value": likely_name,
            "reason": "Selected by scoring combined donor identity evidence across header/title, school candidates, OCR, and acronym safety rules.",
        },
        {
            "layer": "identity",
            "decision": "logo_presence",
            "value": has_logo,
            "reason": "Determined from preserved branding-image evidence and logo metadata.",
        },
        {
            "layer": "identity",
            "decision": "identity_candidate_scores",
            "value": candidate_score_rows[:10],
            "reason": "Top candidate rankings used to choose the final institution identity.",
        },
    ]

    return {
        "institution_name": likely_name,
        "school_candidates": filtered_school_candidates[:15],
        "acronym_candidates": filtered_acronym_candidates[:15],
        "logo_present": bool(has_logo),
        "logo_positions": logo_positions[:10],
        "first_page_identity_lines": first_page_identity_lines[:15],
        "running_header_identity": filtered_running_header_identity[:12],
        "running_footer_identity": filtered_running_footer_identity[:12],
        "official_labels": official_labels[:10],
        "ocr_identity_lines": filtered_ocr_identity[:15],
        "branding_images": branding_images[:10],
        "structural_images": structural_images[:10],
        "identity_candidate_scores": candidate_score_rows[:15],
        "reason_log": reason_log + list(image_interpretation.get("reason_log", []) or []),
        "first_page_header_labels": first_page_header_labels,
        "running_header_labels": running_header_labels,
    }


def _interpret_donor_framing(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Interpret donor framing signals.
    BS4 rule: preserve institutional framing, not donor academic content.
    """
    raw_profile = raw_profile or {}

    title_block = dict(raw_profile.get("title_block_signals", {}) or {})
    header_footer = dict(raw_profile.get("header_footer", {}) or {})
    paragraphs = list(raw_profile.get("paragraphs", []) or [])
    suppressed_paragraphs = list(raw_profile.get("suppressed_paragraphs", []) or [])

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _dedupe(items):
        seen = set()
        out = []
        for item in items or []:
            s = _clean(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    def _contains_any(text: str, terms: List[str]) -> bool:
        low = (text or "").lower()
        for term in terms:
            if term in low:
                return True
        return False

    framing_terms = [
        "question/answer booklet", "question answer booklet", "instructions",
        "answer all questions", "all questions must be answered",
        "read carefully", "write your answer", "candidate", "student",
        "all working should be shown", "answers should be rounded",
        "clarification questions", "assessment", "booklet"
    ]
    timing_terms = [
        "time allowed", "reading time", "working time", "duration"
    ]
    material_terms = [
        "materials required", "material required", "calculator assumed", "calculator free",
        "equipment", "pen", "pencil", "ruler",
        "to be provided by the supervisor",
        "to be provided by the student",
        "special items", "standard items"
    ]
    warning_terms = [
        "important note", "warning", "do not", "must not",
        "may not be used", "no pen", "no pencils", "cannot be"
    ]
    academic_terms = [
        "question 1", "question 2", "calculate", "solve", "explain", "describe",
        "discuss", "compare", "justify", "evaluate", "label", "prove"
    ]
    rich_instruction_terms = [
        "important note",
        "all questions must be answered",
        "all working should be shown",
        "answers should be rounded",
        "reading time",
        "working time",
        "calculator",
        "clarification questions",
        "materials required",
        "material required",
        "marks to be awarded",
        "supporting reasoning",
        "receive full marks"
    ]

    framing_seed_lines = _dedupe(
        list(title_block.get("framing_lines", []) or [])
        + list(title_block.get("title_candidates", []) or [])
        + list(header_footer.get("header_framing_lines", []) or [])
        + list(header_footer.get("footer_framing_lines", []) or [])
    )

    field_lines = _dedupe(
        list(title_block.get("field_lines", []) or [])
        + list(title_block.get("admin_lines", []) or [])
    )

    document_frame_lines = []
    generic_instructions = []
    timing_lines = []
    material_lines = []
    warning_lines = []
    rich_instruction_paragraphs = []

    framing_pool = []
    framing_pool.extend(framing_seed_lines)

    # --------------------------------------------------
    # 1. Clean paragraph evidence still present in donor
    # --------------------------------------------------
    for para in paragraphs[:120]:
        if not isinstance(para, dict):
            continue

        txt = _clean(para.get("text", ""))
        if not txt:
            continue

        low = txt.lower()

        if _contains_any(low, academic_terms):
            continue

        if len(txt) <= 180:
            if (
                _contains_any(low, framing_terms)
                or _contains_any(low, timing_terms)
                or _contains_any(low, material_terms)
                or _contains_any(low, warning_terms)
            ):
                framing_pool.append(txt)
                continue

        if len(txt) > 120 and _contains_any(low, rich_instruction_terms):
            rich_instruction_paragraphs.append(txt)

    # --------------------------------------------------
    # 2. Recover rich framing paragraphs from suppressed donor
    # This is important because donor cleanup may suppress long
    # generic exam instructions before framing gets to use them.
    # --------------------------------------------------
    for row in suppressed_paragraphs[:60]:
        if isinstance(row, dict):
            txt = _clean(row.get("text", ""))
        else:
            txt = _clean(str(row))

        if not txt:
            continue

        low = txt.lower()

        # Strongly reject donor question content
        if _contains_any(low, academic_terms):
            continue

        if len(txt) > 120 and _contains_any(low, rich_instruction_terms):
            rich_instruction_paragraphs.append(txt)

    # --------------------------------------------------
    # 3. Use improved BS4 table interpretation
    # --------------------------------------------------
    table_interpretation = _interpret_donor_tables(raw_profile)

    def _looks_like_candidate_field(s: str) -> bool:
        s = _clean(s)
        low = s.lower()
        if not s:
            return False

        direct_fields = [
            "name", "teacher", "date", "class", "student", "candidate", "subject",
            "candidate name", "student name", "examiner"
        ]
        if low in direct_fields:
            return True

        if any(low.startswith(f + ":") for f in direct_fields):
            return True

        if any(low.startswith(f + "_") for f in direct_fields):
            return True

        if any(low.startswith(f + ".") for f in direct_fields):
            return True

        if low.startswith("year ") or low.startswith("grade "):
            return False

        if _contains_any(low, [
            "important note",
            "to be provided by the student",
            "to be provided by the supervisor",
            "time allowed",
            "reading time",
            "working time",
            "question/answer booklet",
            "instructions",
        ]):
            return False

        return False

    candidate_fields = []
    for s in field_lines:
        if _looks_like_candidate_field(s):
            candidate_fields.append(s)

    for tbl in table_interpretation.get("front_matter_tables", []) or []:
        for sample in tbl.get("sample_texts", []) or []:
            s = _clean(sample)
            if _looks_like_candidate_field(s):
                candidate_fields.append(s)

    candidate_fields = _dedupe(candidate_fields)

    # --------------------------------------------------
    # 4. Final framing classification
    # --------------------------------------------------
    for txt in _dedupe(framing_pool + rich_instruction_paragraphs):
        low = txt.lower()

        if _contains_any(low, timing_terms):
            timing_lines.append(txt)

        if _contains_any(low, material_terms):
            material_lines.append(txt)

        if _contains_any(low, warning_terms):
            warning_lines.append(txt)

        if _contains_any(low, framing_terms):
            generic_instructions.append(txt)

        if (
            _contains_any(low, framing_terms)
            or _contains_any(low, timing_terms)
            or _contains_any(low, material_terms)
            or _contains_any(low, warning_terms)
        ):
            document_frame_lines.append(txt)

    reason_log = [
        {
            "layer": "framing",
            "decision": "document_frame_lines",
            "value": len(_dedupe(document_frame_lines)),
            "reason": "Collected from title block, header/footer framing clues, live donor paragraphs, suppressed donor instruction recovery, and institutional instruction signals.",
        },
        {
            "layer": "framing",
            "decision": "candidate_fields",
            "value": len(candidate_fields),
            "reason": "Collected from donor field lines and interpreted front-matter tables only.",
        },
        {
            "layer": "framing",
            "decision": "rich_instruction_paragraphs",
            "value": len(_dedupe(rich_instruction_paragraphs)),
            "reason": "Long institutional instruction paragraphs preserved from both live donor paragraphs and suppressed donor text evidence.",
        },
    ]

    return {
        "document_frame_lines": _dedupe(document_frame_lines)[:24],
        "generic_instructions": _dedupe(generic_instructions)[:24],
        "timing_lines": _dedupe(timing_lines)[:12],
        "material_lines": _dedupe(material_lines)[:12],
        "warning_lines": _dedupe(warning_lines)[:12],
        "rich_instruction_paragraphs": _dedupe(rich_instruction_paragraphs)[:10],
        "candidate_fields": candidate_fields[:20],
        "reason_log": reason_log + list(table_interpretation.get("reason_log", []) or []),
    }


def _interpret_donor_tables(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Second-stage BS4 table interpretation.

    Distinguishes:
    - front_matter_tables
    - response_tables
    - academic_tables
    - uncertain_tables

    BS4 goal:
    preserve reusable institutional table grammar without dragging donor subject content forward.
    """
    raw_profile = raw_profile or {}

    extracted_tables = list(raw_profile.get("extracted_tables", []) or [])
    suppressed_tables = list(raw_profile.get("suppressed_tables", []) or [])

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _sample_texts(tbl: Dict[str, Any]) -> List[str]:
        out = []
        for x in tbl.get("sample_texts", []) or []:
            s = _clean(x)
            if s:
                out.append(s)
        return out

    def _blob(tbl: Dict[str, Any]) -> str:
        return " | ".join(_sample_texts(tbl)).strip()

    def _contains_any(text: str, needles: List[str]) -> bool:
        low = (text or "").lower()
        for needle in needles:
            if needle in low:
                return True
        return False

    FIELD_TERMS = [
        "name", "teacher", "date", "class", "student", "candidate",
        "candidate name", "student name", "subject", "grade", "year", "examiner"
    ]

    RESPONSE_TERMS = [
        "working", "answer", "response", "method", "solution", "observation",
        "result", "comment", "reflection", "evidence"
    ]

    ACADEMIC_TERMS = [
        "question", "calculate", "solve", "explain", "describe", "discuss",
        "compare", "justify", "evaluate", "prove", "label", "identify"
    ]

    MARK_TERMS = [
        "mark", "marks", "total"
    ]

    front_matter_tables = []
    response_tables = []
    academic_tables = []
    uncertain_tables = []

    all_tables = []
    for tbl in extracted_tables:
        if isinstance(tbl, dict):
            row = dict(tbl)
            row["_source_bucket"] = "extracted"
            all_tables.append(row)

    for tbl in suppressed_tables:
        if isinstance(tbl, dict):
            row = dict(tbl)
            row["_source_bucket"] = "suppressed"
            all_tables.append(row)

    for tbl in all_tables:
        rows = int(tbl.get("rows", 0) or 0)
        cols = int(tbl.get("cols", 0) or 0)
        has_header_row = bool(tbl.get("has_header_row", False))
        first_pass_role = str(tbl.get("table_role", "") or "").strip().lower()
        first_pass_reason = str(tbl.get("classification_reason", "") or "").strip()
        sample_texts = _sample_texts(tbl)
        blob = _blob(tbl)
        low = blob.lower()

        has_field_terms = _contains_any(low, FIELD_TERMS)
        has_response_terms = _contains_any(low, RESPONSE_TERMS)
        has_marks_terms = _contains_any(low, MARK_TERMS)
        has_academic_terms = _contains_any(low, ACADEMIC_TERMS)

        small_table = rows <= 6 and cols <= 6
        medium_table = rows <= 12 and cols <= 8
        wide_table = cols >= 5
        mostly_blank_like = len(sample_texts) <= max(2, rows)

        interpreted = {
            "rows": rows,
            "cols": cols,
            "has_header_row": has_header_row,
            "sample_texts": sample_texts[:8],
            "first_pass_role": first_pass_role,
            "first_pass_reason": first_pass_reason,
            "source_bucket": tbl.get("_source_bucket", ""),
        }

        # 1. Candidate/admin front matter tables
        if has_field_terms and small_table:
            interpreted["interpreted_table_role"] = "front_matter"
            interpreted["interpretation_reason"] = (
                "contains candidate or admin field language in a small table shape typical of front-matter forms"
            )
            front_matter_tables.append(interpreted)
            continue

        # 2. Reusable response tables
        # Structural tables that look like answer/response scaffolds should survive BS4.
        response_like = False

        if first_pass_role == "structural" and medium_table and not has_academic_terms:
            if has_header_row:
                response_like = True
            elif has_response_terms:
                response_like = True
            elif wide_table and mostly_blank_like:
                response_like = True

        if response_like:
            interpreted["interpreted_table_role"] = "response_table"
            interpreted["interpretation_reason"] = (
                "structural table shape suggests reusable response grammar rather than donor-specific academic content"
            )
            response_tables.append(interpreted)
            continue

        # 3. Academic/data tables
        if first_pass_role == "academic" or has_academic_terms or has_marks_terms:
            interpreted["interpreted_table_role"] = "academic"
            interpreted["interpretation_reason"] = (
                "contains donor academic, question, or marks language and must remain never-inherit content"
            )
            academic_tables.append(interpreted)
            continue

        # 4. Structural-but-uncertain tables
        # We preserve these cautiously instead of killing them too early.
        if first_pass_role == "structural":
            interpreted["interpreted_table_role"] = "uncertain"
            interpreted["interpretation_reason"] = (
                "structural upstream classification preserved, but table meaning is not yet strong enough to promote safely"
            )
            uncertain_tables.append(interpreted)
            continue

        # 5. Everything else remains uncertain
        interpreted["interpreted_table_role"] = "uncertain"
        interpreted["interpretation_reason"] = (
            "table does not clearly fit front-matter, response, or academic categories"
        )
        uncertain_tables.append(interpreted)

    reason_log = [
        {
            "layer": "tables",
            "decision": "front_matter_tables",
            "value": len(front_matter_tables),
            "reason": "Small candidate and admin tables classified as reusable institutional framing.",
        },
        {
            "layer": "tables",
            "decision": "response_tables",
            "value": len(response_tables),
            "reason": "Reusable response grammar tables classified separately from front matter and academic content.",
        },
        {
            "layer": "tables",
            "decision": "academic_tables",
            "value": len(academic_tables),
            "reason": "Academic and donor-content tables classified as never-inherit content.",
        },
        {
            "layer": "tables",
            "decision": "uncertain_tables",
            "value": len(uncertain_tables),
            "reason": "Ambiguous or weakly structural tables preserved for cautious downstream handling instead of early loss.",
        },
    ]

    return {
        "front_matter_tables": front_matter_tables[:12],
        "response_tables": response_tables[:12],
        "academic_tables": academic_tables[:12],
        "uncertain_tables": uncertain_tables[:12],
        "reason_log": reason_log,
    }


def _interpret_donor_images(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Second-stage BS4 image interpretation.

    Distinguishes:
    - branding_images
    - structural_images
    - educational_images
    - uncertain_images

    BS4 goal:
    preserve institutional identity and reusable document-structure visuals
    without dragging donor academic visuals forward.
    """
    raw_profile = raw_profile or {}

    extracted_images = list(raw_profile.get("extracted_images", []) or [])
    suppressed_images = list(raw_profile.get("suppressed_images", []) or [])

    interpreted_branding = []
    interpreted_structural = []
    interpreted_educational = []
    interpreted_uncertain = []

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    all_images = []
    for img in extracted_images:
        if isinstance(img, dict):
            row = dict(img)
            row["_source_bucket"] = "extracted"
            all_images.append(row)

    for img in suppressed_images:
        if isinstance(img, dict):
            row = dict(img)
            row["_source_bucket"] = "suppressed"
            all_images.append(row)

    for img in all_images:
        position = str(img.get("position", "") or "").strip().lower()
        first_pass_role = str(
            img.get("image_role")
            or img.get("interpreted_image_role")
            or img.get("first_pass_role")
            or ""
        ).strip().lower()
        first_pass_reason = str(img.get("classification_reason", "") or "").strip()
        width_inches = float(img.get("width_inches", 0) or 0)
        height_inches = float(img.get("height_inches", 0) or 0)
        para_idx = img.get("paragraph_index", None)
        has_binary = bool(img.get("binary_data"))
        alt_text = _clean(str(img.get("alt_text", "") or ""))

        area = width_inches * height_inches if width_inches and height_inches else 0.0
        is_small = bool(width_inches and height_inches and width_inches <= 3.5 and height_inches <= 3.5)
        is_large = bool(width_inches and height_inches and area >= 16.0)
        is_early_body = bool(isinstance(para_idx, int) and para_idx <= 6)

        interpreted = {
            "position": position,
            "width_inches": width_inches,
            "height_inches": height_inches,
            "first_pass_role": first_pass_role,
            "first_pass_reason": first_pass_reason,
            "source_bucket": img.get("_source_bucket", ""),
            "has_binary": has_binary,
            "paragraph_index": para_idx,
            "binary_data": img.get("binary_data"),
            "alt_text": alt_text,
        }

        # 1. Branding images
        if first_pass_role == "branding":
            interpreted["interpreted_image_role"] = "branding"
            interpreted["interpretation_reason"] = (
                "classified branding upstream and preserved for institutional identity"
            )
            interpreted_branding.append(interpreted)
            continue

        if position in {"header", "footer"} and is_small:
            interpreted["interpreted_image_role"] = "branding"
            interpreted["interpretation_reason"] = (
                "small header/footer image is most likely branding rather than academic content"
            )
            interpreted_branding.append(interpreted)
            continue

        if position == "first_page" and is_small and not is_large:
            interpreted["interpreted_image_role"] = "branding"
            interpreted["interpretation_reason"] = (
                "small first-page image is most likely institution identity or crest branding"
            )
            interpreted_branding.append(interpreted)
            continue

        # 2. Structural images
        if position == "first_page" and is_large:
            interpreted["interpreted_image_role"] = "structural"
            interpreted["interpretation_reason"] = (
                "large first-page image preserved as possible reusable structural front-page visual"
            )
            interpreted_structural.append(interpreted)
            continue

        if position in {"header", "footer", "first_page"}:
            interpreted["interpreted_image_role"] = "structural"
            interpreted["interpretation_reason"] = (
                "position suggests reusable document-structure image rather than donor academic content"
            )
            interpreted_structural.append(interpreted)
            continue

        if is_early_body and is_small and first_pass_role not in {"educational"}:
            interpreted["interpreted_image_role"] = "structural"
            interpreted["interpretation_reason"] = (
                "small early-body image preserved as possible reusable structural visual"
            )
            interpreted_structural.append(interpreted)
            continue

        # 3. Educational images
        if first_pass_role == "educational":
            interpreted["interpreted_image_role"] = "educational"
            interpreted["interpretation_reason"] = (
                "classified educational upstream and treated as never-inherit content"
            )
            interpreted_educational.append(interpreted)
            continue

        if position == "body" and is_large:
            interpreted["interpreted_image_role"] = "educational"
            interpreted["interpretation_reason"] = (
                "large body image is more likely donor academic content than reusable identity or structure"
            )
            interpreted_educational.append(interpreted)
            continue

        if alt_text and any(word in alt_text.lower() for word in [
            "diagram", "graph", "chart", "figure", "map", "illustration"
        ]):
            interpreted["interpreted_image_role"] = "educational"
            interpreted["interpretation_reason"] = (
                "alt text suggests educational content rather than institutional branding"
            )
            interpreted_educational.append(interpreted)
            continue

        # 4. Uncertain
        interpreted["interpreted_image_role"] = "uncertain"
        interpreted["interpretation_reason"] = (
            "image does not clearly fit branding, structural, or educational categories"
        )
        interpreted_uncertain.append(interpreted)

    reason_log = [
        {
            "layer": "images",
            "decision": "branding_images",
            "value": len(interpreted_branding),
            "reason": "Institutional branding images preserved for identity and logo selection.",
        },
        {
            "layer": "images",
            "decision": "structural_images",
            "value": len(interpreted_structural),
            "reason": "Reusable document-structure visuals preserved separately from branding and educational content.",
        },
        {
            "layer": "images",
            "decision": "educational_images",
            "value": len(interpreted_educational),
            "reason": "Educational donor visuals classified as never-inherit content.",
        },
        {
            "layer": "images",
            "decision": "uncertain_images",
            "value": len(interpreted_uncertain),
            "reason": "Ambiguous donor visuals preserved for cautious downstream handling.",
        },
    ]

    return {
        "branding_images": interpreted_branding[:12],
        "structural_images": interpreted_structural[:12],
        "educational_images": interpreted_educational[:12],
        "uncertain_images": interpreted_uncertain[:12],
        "reason_log": reason_log,
    }


def _interpret_donor_structure(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Interpret donor assessment structure and presentation grammar.
    BS4 rule: preserve reusable structure, not donor content.
    """
    raw_profile = raw_profile or {}

    layout_features = list(raw_profile.get("layout_features", []) or [])
    extracted_tables = list(raw_profile.get("extracted_tables", []) or [])
    paragraphs = list(raw_profile.get("paragraphs", []) or [])
    list_styles = list(raw_profile.get("list_styles", []) or [])
    page_setup = dict(raw_profile.get("page_setup", {}) or {})
    style_preferences = dict(raw_profile.get("style_preferences", {}) or {})

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    numbering_style = "unknown"
    marks_style = "unknown"
    answer_space_style = "unknown"
    question_heading_style = "unknown"

    structure_signal_lines = []

    # Paragraph signals
    for para in paragraphs[:100]:
        if not isinstance(para, dict):
            continue
        txt = _clean(para.get("text", ""))
        if txt:
            structure_signal_lines.append(txt)

    # Table signals — crucial for exam-style donors where numbering/marks live in tables
    for tbl in extracted_tables:
        if not isinstance(tbl, dict):
            continue
        for sample in tbl.get("sample_texts", []) or []:
            txt = _clean(sample)
            if txt:
                structure_signal_lines.append(txt)

    # Deduplicate while preserving order
    seen_structure = set()
    deduped_structure_lines = []
    for txt in structure_signal_lines:
        key = txt.lower()
        if key in seen_structure:
            continue
        seen_structure.add(key)
        deduped_structure_lines.append(txt)

    numeric_hits = 0
    alpha_hits = 0
    parenthetical_hits = 0

    bracketed_marks_hits = 0
    inline_marks_hits = 0

    uppercase_heading_hits = 0
    colon_heading_hits = 0

    for txt in deduped_structure_lines:
        # Numbering style
        if txt[:8]:
            stripped = txt.strip()

            if stripped and stripped[0].isdigit():
                if len(stripped) >= 2 and stripped[1] in {".", ")"}:
                    numeric_hits += 1

            if stripped and len(stripped) >= 2:
                if stripped[0].isalpha() and stripped[1] in {".", ")"}:
                    alpha_hits += 1

            if stripped.startswith("("):
                if len(stripped) >= 3 and stripped[2:3] == ")":
                    parenthetical_hits += 1

        # Marks style
        low = txt.lower()
        if "[" in txt and "mark" in low and "]" in txt:
            bracketed_marks_hits += 1
        elif "mark" in low:
            inline_marks_hits += 1

        # Heading style
        if txt.isupper() and len(txt.split()) <= 12:
            uppercase_heading_hits += 1
        elif txt.endswith(":") and len(txt) <= 120:
            colon_heading_hits += 1

    # Final numbering choice
    if numeric_hits > 0 and numeric_hits >= alpha_hits and numeric_hits >= parenthetical_hits:
        numbering_style = "numeric"
    elif parenthetical_hits > 0 and parenthetical_hits >= alpha_hits:
        numbering_style = "subpart_parenthetical"
    elif alpha_hits > 0:
        numbering_style = "subpart_alpha"

    # Final marks choice
    if bracketed_marks_hits > 0 and bracketed_marks_hits >= inline_marks_hits:
        marks_style = "bracketed_marks"
    elif inline_marks_hits > 0:
        marks_style = "inline_marks"

    # Final heading choice
    if uppercase_heading_hits > 0 and uppercase_heading_hits >= colon_heading_hits:
        question_heading_style = "uppercase_heading"
    elif colon_heading_hits > 0:
        question_heading_style = "colon_heading"

    table_interpretation = _interpret_donor_tables(raw_profile)
    front_matter_tables = table_interpretation.get("front_matter_tables", []) or []
    response_tables = table_interpretation.get("response_tables", []) or []
    uncertain_tables = table_interpretation.get("uncertain_tables", []) or []

    has_answer_lines_feature = any(
        isinstance(f, dict) and f.get("feature_name") == "answer_lines"
        for f in layout_features
    )

    has_contextual_columns = any(
        isinstance(f, dict)
        and str(f.get("category", "")).strip() == FeatureCategory.CONTEXTUAL_INHERIT.value
        and f.get("feature_name") == "columns"
        for f in layout_features
    )

    if has_answer_lines_feature:
        answer_space_style = "answer_lines"
    elif response_tables:
        answer_space_style = "response_table"
    elif front_matter_tables:
        answer_space_style = "structured_front_matter"
    elif uncertain_tables:
        answer_space_style = "uncertain_structured_space"
    else:
        answer_space_style = "plain_response_space"

    structural_tables = []
    for tbl in front_matter_tables + response_tables + uncertain_tables:
        structural_tables.append({
            "rows": int(tbl.get("rows", 0) or 0),
            "cols": int(tbl.get("cols", 0) or 0),
            "has_header_row": bool(tbl.get("has_header_row", False)),
            "classification_reason": str(tbl.get("interpretation_reason", "") or "").strip(),
            "sample_texts": list(tbl.get("sample_texts", []) or [])[:8],
            "interpreted_table_role": str(tbl.get("interpreted_table_role", "") or "").strip(),
        })

    column_info = dict(page_setup.get("columns", {}) or {})
    preferred_column_mode = "single"
    if int(column_info.get("count", 1) or 1) > 1:
        preferred_column_mode = "multi"

    # If contextual column evidence exists, keep that as supporting structure signal
    if preferred_column_mode == "single" and has_contextual_columns:
        preferred_column_mode = "multi"

    reason_log = [
        {
            "layer": "structure",
            "decision": "numbering_style",
            "value": numbering_style,
            "reason": "Derived from preserved donor numbering patterns across paragraphs and table evidence.",
        },
        {
            "layer": "structure",
            "decision": "marks_style",
            "value": marks_style,
            "reason": "Derived from preserved donor marks presentation patterns.",
        },
        {
            "layer": "structure",
            "decision": "question_heading_style",
            "value": question_heading_style,
            "reason": "Derived from heading-like donor presentation cues.",
        },
        {
            "layer": "structure",
            "decision": "answer_space_style",
            "value": answer_space_style,
            "reason": "Derived from preserved layout features plus interpreted front-matter/response-table evidence.",
        },
        {
            "layer": "structure",
            "decision": "preferred_column_mode",
            "value": preferred_column_mode,
            "reason": "Derived from page setup and preserved contextual layout feature evidence.",
        },
    ]

    return {
        "numbering_style": numbering_style,
        "marks_style": marks_style,
        "answer_space_style": answer_space_style,
        "question_heading_style": question_heading_style,
        "preferred_column_mode": preferred_column_mode,
        "column_count": int(column_info.get("count", 1) or 1),
        "column_space_inches": float(column_info.get("space_inches", 0.0) or 0.0),
        "list_styles": list_styles[:20],
        "structural_tables": structural_tables[:15],
        "front_matter_tables": front_matter_tables[:12],
        "response_tables": response_tables[:12],
        "uncertain_tables": uncertain_tables[:12],
        "contextual_layout_features": [
            dict(f) for f in layout_features
            if isinstance(f, dict)
            and str(f.get("category", "")).strip() == FeatureCategory.CONTEXTUAL_INHERIT.value
        ],
        "never_inherit_layout_features": [
            dict(f) for f in layout_features
            if isinstance(f, dict)
            and str(f.get("category", "")).strip() == FeatureCategory.NEVER_INHERIT.value
        ],
        "heading_styles": list(style_preferences.get("heading_styles", []) or [])[:10],
        "reason_log": reason_log + list(table_interpretation.get("reason_log", []) or []),
    }


def _interpret_donor_discarded_content(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Capture donor content that must NEVER be inherited.
    This preserves visibility into what BS3/BS4 excluded.
    """
    raw_profile = raw_profile or {}

    suppressed_paragraphs = list(raw_profile.get("suppressed_paragraphs", []) or [])
    suppressed_images = list(raw_profile.get("suppressed_images", []) or [])
    suppressed_tables = list(raw_profile.get("suppressed_tables", []) or [])
    header_footer = dict(raw_profile.get("header_footer", {}) or {})

    table_interpretation = _interpret_donor_tables(raw_profile)
    image_interpretation = _interpret_donor_images(raw_profile)

    academic_tables = list(table_interpretation.get("academic_tables", []) or [])
    uncertain_tables = list(table_interpretation.get("uncertain_tables", []) or [])

    educational_images = list(image_interpretation.get("educational_images", []) or [])
    uncertain_images = list(image_interpretation.get("uncertain_images", []) or [])

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    paragraph_samples = []
    for row in suppressed_paragraphs[:20]:
        if isinstance(row, dict):
            txt = _clean(row.get("text", ""))
            if txt:
                paragraph_samples.append(txt)
        else:
            txt = _clean(str(row))
            if txt:
                paragraph_samples.append(txt)

    suppressed_header_texts = [
        _clean(x) for x in (header_footer.get("suppressed_header_texts", []) or []) if _clean(x)
    ]
    suppressed_footer_texts = [
        _clean(x) for x in (header_footer.get("suppressed_footer_texts", []) or []) if _clean(x)
    ]

    academic_table_samples = []
    for tbl in academic_tables[:10]:
        if not isinstance(tbl, dict):
            continue
        for sample in tbl.get("sample_texts", []) or []:
            s = _clean(sample)
            if s:
                academic_table_samples.append(s)

    educational_image_samples = []
    for img in educational_images[:10]:
        if not isinstance(img, dict):
            continue
        alt = _clean(img.get("alt_text", ""))
        if alt:
            educational_image_samples.append(alt)
        else:
            pos = _clean(str(img.get("position", "") or ""))
            if pos:
                educational_image_samples.append(f"[image at {pos}]")

    reason_log = [
        {
            "layer": "discarded_content",
            "decision": "suppressed_content_summary",
            "value": {
                "paragraphs": len(suppressed_paragraphs),
                "images": len(suppressed_images),
                "tables": len(suppressed_tables),
            },
            "reason": "These donor elements were suppressed during earlier cleanup and must never drive reuse.",
        },
        {
            "layer": "discarded_content",
            "decision": "academic_tables",
            "value": len(academic_tables),
            "reason": "BS4 interpreted these as donor academic tables and excluded them from reusable donor structure.",
        },
        {
            "layer": "discarded_content",
            "decision": "educational_images",
            "value": len(educational_images),
            "reason": "BS4 interpreted these as donor educational visuals and excluded them from reusable donor identity/structure.",
        },
        {
            "layer": "discarded_content",
            "decision": "uncertain_non_promoted",
            "value": {
                "tables": len(uncertain_tables),
                "images": len(uncertain_images),
            },
            "reason": "These donor elements were not promoted into reusable donor meaning and remain excluded evidence unless later logic explicitly handles them.",
        },
    ]

    return {
        "suppressed_paragraph_samples": paragraph_samples[:20],
        "suppressed_header_texts": suppressed_header_texts[:20],
        "suppressed_footer_texts": suppressed_footer_texts[:20],
        "suppressed_image_count": len(suppressed_images),
        "suppressed_table_count": len(suppressed_tables),
        "suppressed_images": suppressed_images[:10],
        "suppressed_tables": suppressed_tables[:10],
        "academic_table_count": len(academic_tables),
        "academic_table_samples": academic_table_samples[:20],
        "educational_image_count": len(educational_images),
        "educational_image_samples": educational_image_samples[:20],
        "uncertain_table_count": len(uncertain_tables),
        "uncertain_image_count": len(uncertain_images),
        "reason_log": reason_log + list(table_interpretation.get("reason_log", []) or []) + list(image_interpretation.get("reason_log", []) or []),
    }


def _build_donor_inheritance_map(
    identity: Dict[str, Any],
    framing: Dict[str, Any],
    structure: Dict[str, Any],
    discarded_content: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Build the BS4 inheritance map required by the DIE contract.
    """
    identity = identity or {}
    framing = framing or {}
    structure = structure or {}
    discarded_content = discarded_content or {}

    return {
        "ALWAYS": {
            "identity": {
                "institution_name": str(identity.get("institution_name") or "").strip(),
                "school_candidates": list(identity.get("school_candidates", []) or []),
                "acronym_candidates": list(identity.get("acronym_candidates", []) or []),
                "logo_present": bool(identity.get("logo_present")),
                "logo_positions": list(identity.get("logo_positions", []) or []),
                "first_page_identity_lines": list(identity.get("first_page_identity_lines", []) or []),
                "running_header_identity": list(identity.get("running_header_identity", []) or []),
                "running_footer_identity": list(identity.get("running_footer_identity", []) or []),
                "official_labels": list(identity.get("official_labels", []) or []),
                "identity_candidate_scores": list(identity.get("identity_candidate_scores", []) or []),
            }
        },

        "CONTEXTUAL": {
            "framing": {
                "document_frame_lines": list(framing.get("document_frame_lines", []) or []),
                "generic_instructions": list(framing.get("generic_instructions", []) or []),
                "timing_lines": list(framing.get("timing_lines", []) or []),
                "material_lines": list(framing.get("material_lines", []) or []),
                "warning_lines": list(framing.get("warning_lines", []) or []),
                "rich_instruction_paragraphs": list(framing.get("rich_instruction_paragraphs", []) or []),
                "candidate_fields": list(framing.get("candidate_fields", []) or []),
            }
        },

        "STRUCTURAL": {
            "structure": {
                "numbering_style": str(structure.get("numbering_style", "unknown") or "unknown"),
                "marks_style": str(structure.get("marks_style", "unknown") or "unknown"),
                "answer_space_style": str(structure.get("answer_space_style", "unknown") or "unknown"),
                "question_heading_style": str(structure.get("question_heading_style", "unknown") or "unknown"),
                "preferred_column_mode": str(structure.get("preferred_column_mode", "single") or "single"),
                "column_count": int(structure.get("column_count", 1) or 1),
                "column_space_inches": float(structure.get("column_space_inches", 0.0) or 0.0),
                "structural_tables": list(structure.get("structural_tables", []) or []),
                "front_matter_tables": list(structure.get("front_matter_tables", []) or []),
                "response_tables": list(structure.get("response_tables", []) or []),
                "uncertain_tables": list(structure.get("uncertain_tables", []) or []),
                "contextual_layout_features": list(structure.get("contextual_layout_features", []) or []),
                "never_inherit_layout_features": list(structure.get("never_inherit_layout_features", []) or []),
                "heading_styles": list(structure.get("heading_styles", []) or []),
                "list_styles": list(structure.get("list_styles", []) or []),
            }
        },

        "NEVER": {
            "discarded_content": {
                "suppressed_paragraph_samples": list(discarded_content.get("suppressed_paragraph_samples", []) or []),
                "suppressed_header_texts": list(discarded_content.get("suppressed_header_texts", []) or []),
                "suppressed_footer_texts": list(discarded_content.get("suppressed_footer_texts", []) or []),
                "suppressed_image_count": int(discarded_content.get("suppressed_image_count", 0) or 0),
                "suppressed_table_count": int(discarded_content.get("suppressed_table_count", 0) or 0),
                "academic_table_count": int(discarded_content.get("academic_table_count", 0) or 0),
                "academic_table_samples": list(discarded_content.get("academic_table_samples", []) or []),
                "educational_image_count": int(discarded_content.get("educational_image_count", 0) or 0),
                "educational_image_samples": list(discarded_content.get("educational_image_samples", []) or []),
                "uncertain_table_count": int(discarded_content.get("uncertain_table_count", 0) or 0),
                "uncertain_image_count": int(discarded_content.get("uncertain_image_count", 0) or 0),
            }
        },
    }


def _extract_run_properties_enhanced(run) -> Dict[str, Any]:
    """
    Extract ENHANCED run-level properties from a python-docx run.
    
    Extracts:
        - font_name (str)
        - font_size (Pt)
        - bold (bool)
        - italic (bool)
        - underline (bool)
        - color (RGBColor or hex string)
        - highlight (bool)
        - subscript/superscript
        - strike (strikethrough)
        - character spacing
    """
    properties = {}
    
    try:
        # Font name
        if run.font.name:
            properties["name"] = run.font.name
        
        # Font size
        if run.font.size:
            properties["size_pt"] = run.font.size.pt
        
        # Bold
        if run.font.bold is not None:
            properties["bold"] = run.font.bold
        
        # Italic
        if run.font.italic is not None:
            properties["italic"] = run.font.italic
        
        # Underline
        if run.font.underline is not None:
            properties["underline"] = run.font.underline
        
        # Color
        if run.font.color and run.font.color.rgb:
            properties["color"] = str(run.font.color.rgb)
        elif run.font.color and run.font.color.theme_color:
            properties["color"] = run.font.color.theme_color
        
        # Highlight
        if run.font.highlight_color:
            properties["highlight"] = run.font.highlight_color
        
        # Subscript / Superscript
        if run.font.subscript:
            properties["subscript"] = True
        if run.font.superscript:
            properties["superscript"] = True
        
        # Strikethrough
        if run.font.strike:
            properties["strike"] = True
        
        # Character spacing (if available via XML)
        try:
            rPr = run._element.get_or_add_rPr()
            spacing = rPr.find(qn('w:spacing'))
            if spacing is not None:
                properties["character_spacing"] = spacing.get(qn('w:val'))
        except:
            pass
        
    except Exception as e:
        logger.debug(f"Enhanced run property extraction failed: {e}")
    
    return properties


def _extract_header_footer_enhanced(doc, docx_path: str = "") -> Dict[str, Any]:
    """
    Zone-aware header/footer extractor.

    Captures:
    - first-page header/footer
    - default/running header/footer
    - even-page header/footer
    - textbox/drawing text such as OFFICIAL
    - donor running-header pattern
    - different-first-page setting
    """
    result = {
        "has_header": False,
        "has_footer": False,
        "different_first_page_header_footer": False,

        "header_texts": [],
        "footer_texts": [],
        "header_candidates": [],
        "footer_candidates": [],
        "header_paragraphs": [],
        "footer_paragraphs": [],
        "header_tables": [],
        "footer_tables": [],

        "first_page_header_texts": [],
        "first_page_footer_texts": [],
        "first_page_header_paragraphs": [],
        "first_page_footer_paragraphs": [],
        "first_page_header_textboxes": [],
        "first_page_footer_textboxes": [],

        "running_header_texts": [],
        "running_footer_texts": [],
        "running_header_paragraphs": [],
        "running_footer_paragraphs": [],
        "running_header_textboxes": [],
        "running_footer_textboxes": [],

        "even_header_texts": [],
        "even_footer_texts": [],
        "even_header_textboxes": [],
        "even_footer_textboxes": [],

        "first_page_header_labels": [],
        "running_header_labels": [],
        "official_labels": [],

        "running_header_pattern": {},
        "header_identity_lines": [],
        "footer_identity_lines": [],
        "header_framing_lines": [],
        "footer_framing_lines": [],

        "has_header_image": False,
        "has_footer_image": False,
    }

    if not getattr(doc, "sections", None):
        return result

    section = doc.sections[0]
    result["different_first_page_header_footer"] = bool(
        getattr(section, "different_first_page_header_footer", False)
    )

    def _clean_text(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _dedupe_keep_order(items):
        seen = set()
        out = []
        for item in items or []:
            s = _clean_text(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    def _extract_textboxes(container):
        found = []
        try:
            for node in container._element.xpath('.//w:txbxContent//w:t'):
                txt = _clean_text(getattr(node, "text", "") or "")
                if txt:
                    found.append(txt)
            for node in container._element.xpath('.//a:t'):
                txt = _clean_text(getattr(node, "text", "") or "")
                if txt:
                    found.append(txt)
        except Exception:
            pass
        return _dedupe_keep_order(found)

    def _extract_paragraphs(container):
        rows = []
        texts = []
        try:
            for para in container.paragraphs:
                text = _clean_text(getattr(para, "text", ""))
                if not text:
                    continue
                texts.append(text)
                rows.append({
                    "text": text,
                    "style": para.style.name if getattr(para, "style", None) else None,
                    "alignment": para.alignment.name if getattr(para, "alignment", None) else None,
                })
        except Exception:
            pass
        return _dedupe_keep_order(texts), rows

    def _extract_table_previews(container, limit_tables=4):
        previews = []
        try:
            tables = container.tables if hasattr(container, "tables") else []
            for t_idx, table in enumerate((tables or [])[:limit_tables]):
                rows_preview = []
                for row in table.rows[:4]:
                    row_vals = []
                    for cell in row.cells[:6]:
                        txt = _clean_text(cell.text)
                        if txt:
                            row_vals.append(txt)
                    if row_vals:
                        rows_preview.append(row_vals)
                if rows_preview:
                    previews.append({
                        "index": t_idx,
                        "rows": len(getattr(table, "rows", []) or []),
                        "cols": len(getattr(table, "columns", []) or []),
                        "preview": rows_preview[:4],
                    })
        except Exception:
            pass
        return previews

    def _has_image(container):
        try:
            for rel in container.part.related_parts.values():
                if hasattr(rel, "image") or hasattr(rel, "blob"):
                    return True
        except Exception:
            pass
        return False

    def _looks_like_identity(text: str) -> bool:
        s = _clean_text(text)
        low = s.lower()
        if not s:
            return False
        if re.search(r"\b(high school|primary school|secondary school|college|grammar|academy|institute|university|school)\b", low):
            return True
        if len(s) <= 80 and s.isupper() and len(s.split()) <= 8 and not re.search(r"\d", s):
            return True
        return False

    def _looks_like_framing(text: str) -> bool:
        s = _clean_text(text).lower()
        if not s:
            return False
        framing_terms = [
            "question/answer booklet", "examination", "exam", "test", "worksheet",
            "assessment", "semester", "calculator assumed", "calculator free",
            "instructions", "time allowed", "reading time"
        ]
        return any(term in s for term in framing_terms)

    def _labels_from_lines(lines):
        labels = []
        for line in lines or []:
            s = _clean_text(line)
            if not s:
                continue
            for match in re.findall(r"\b(official|confidential|draft)\b", s, re.I):
                labels.append(match.upper())
        return _dedupe_keep_order(labels)

    def _split_running_header_pattern(line: str):
        s = _clean_text(line)
        if not s:
            return {}

        tab_parts = [_clean_text(x) for x in s.split("\t") if _clean_text(x)]
        if len(tab_parts) >= 2:
            parts = tab_parts
        else:
            parts = [_clean_text(x) for x in re.split(r"\s{2,}", s) if _clean_text(x)]

        return {
            "raw": s,
            "parts": parts,
            "part_count": len(parts),
        }

    zones = {
        "first_page_header": getattr(section, "first_page_header", None),
        "running_header": getattr(section, "header", None),
        "even_header": getattr(section, "even_page_header", None),
        "first_page_footer": getattr(section, "first_page_footer", None),
        "running_footer": getattr(section, "footer", None),
        "even_footer": getattr(section, "even_page_footer", None),
    }

    for zone_name, container in zones.items():
        if container is None:
            continue

        texts, paragraphs = _extract_paragraphs(container)
        textboxes = _extract_textboxes(container)
        tables = _extract_table_previews(container)

        result[f"{zone_name}_texts"] = texts[:20]
        result[f"{zone_name}_paragraphs"] = paragraphs[:20]
        result[f"{zone_name}_textboxes"] = textboxes[:20]

        if zone_name == "running_header":
            result["header_texts"] = texts[:10]
            result["header_paragraphs"] = paragraphs[:12]
            result["header_tables"] = tables
            result["has_header_image"] = _has_image(container)
            result["has_header"] = bool(texts or textboxes or tables or result["has_header_image"])

        if zone_name == "running_footer":
            result["footer_texts"] = texts[:10]
            result["footer_paragraphs"] = paragraphs[:12]
            result["footer_tables"] = tables
            result["has_footer_image"] = _has_image(container)
            result["has_footer"] = bool(texts or textboxes or tables or result["has_footer_image"])

    result["first_page_header_labels"] = _labels_from_lines(
        result.get("first_page_header_texts", []) + result.get("first_page_header_textboxes", [])
    )

    result["running_header_labels"] = _labels_from_lines(
        result.get("running_header_texts", []) + result.get("running_header_textboxes", [])
    )

    result["official_labels"] = _dedupe_keep_order(
        result["first_page_header_labels"] + result["running_header_labels"]
    )

    result["header_candidates"] = []
    for text in result.get("header_texts", []) or []:
        if len(_clean_text(text)) <= 120:
            result["header_candidates"].append(text)
    result["header_candidates"] = _dedupe_keep_order(result["header_candidates"])[:12]

    result["footer_candidates"] = []
    for text in result.get("footer_texts", []) or []:
        if len(_clean_text(text)) <= 120:
            result["footer_candidates"].append(text)
    result["footer_candidates"] = _dedupe_keep_order(result["footer_candidates"])[:12]

    result["header_identity_lines"] = [
        x for x in result.get("header_texts", []) if _looks_like_identity(x)
    ][:10]

    result["footer_identity_lines"] = [
        x for x in result.get("footer_texts", []) if _looks_like_identity(x)
    ][:10]

    result["header_framing_lines"] = [
        x for x in result.get("header_texts", []) if _looks_like_framing(x)
    ][:10]

    result["footer_framing_lines"] = [
        x for x in result.get("footer_texts", []) if _looks_like_framing(x)
    ][:10]

    if result.get("header_texts"):
        result["running_header_pattern"] = _split_running_header_pattern(result["header_texts"][0])

    result["has_header"] = bool(
        result.get("header_texts")
        or result.get("running_header_textboxes")
        or result.get("first_page_header_textboxes")
        or result.get("has_header_image")
    )

    result["has_footer"] = bool(
        result.get("footer_texts")
        or result.get("running_footer_textboxes")
        or result.get("first_page_footer_textboxes")
        or result.get("has_footer_image")
    )

    if docx_path:
        result = _merge_docx_xml_fallback_into_header_footer(result, docx_path)

    return result


def _detect_layout_features_enhanced(doc) -> List[Dict[str, Any]]:
    """
    Detect ENHANCED layout features from donor DOCX.

    Build Sequence 3 role:
    - preserve generic layout signals as raw donor evidence
    - keep answer-line, table-density, page-border, watermark, and textbox clues
    - do not over-interpret these features yet
    """
    features = []

    def _clean_text(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    # --------------------------------------------------
    # 1. Detect answer-line signals with examples
    # --------------------------------------------------
    answer_line_count = 0
    answer_line_patterns = []
    answer_line_examples = []

    for para in getattr(doc, "paragraphs", []):
        text = _clean_text(getattr(para, "text", ""))
        if not text:
            continue

        matched = False
        if "__________" in text:
            answer_line_count += 1
            answer_line_patterns.append("underscore")
            matched = True
        elif "______" in text:
            answer_line_count += 1
            answer_line_patterns.append("underscore_short")
            matched = True
        elif "________" in text:
            answer_line_count += 1
            answer_line_patterns.append("underscore_long")
            matched = True
        elif re.search(r"\.{4,}", text):
            answer_line_count += 1
            answer_line_patterns.append("dotted")
            matched = True

        if matched and len(answer_line_examples) < 8:
            answer_line_examples.append(text[:120])

    if answer_line_count > 0:
        dominant_pattern = max(set(answer_line_patterns), key=answer_line_patterns.count) if answer_line_patterns else "underscore"

        features.append({
            "feature_name": "answer_lines",
            "category": FeatureCategory.CONTEXTUAL_INHERIT,
            "extracted_value": {
                "count": answer_line_count,
                "style": dominant_pattern,
                "patterns": sorted(list(set(answer_line_patterns))),
                "examples": answer_line_examples,
            }
        })

    # --------------------------------------------------
    # 2. Detect table-layout signals using richer table evidence
    # --------------------------------------------------
    if getattr(doc, "tables", None):
        table_details = []
        dense_tables = 0
        likely_front_matter_tables = 0
        bordered_tables = 0

        for table in doc.tables[:10]:
            rows = len(getattr(table, "rows", []) or [])
            cols = len(getattr(table, "columns", []) or []) if getattr(table, "columns", None) else 0

            non_empty_cells = 0
            sample_texts = []

            try:
                for row in table.rows[:6]:
                    for cell in row.cells[:6]:
                        txt = _clean_text(cell.text)
                        if txt:
                            non_empty_cells += 1
                            if len(sample_texts) < 8:
                                sample_texts.append(txt)
            except Exception:
                pass

            total_cells = max(rows * max(cols, 1), 1)
            fill_ratio = non_empty_cells / total_cells
            density = "dense" if fill_ratio >= 0.7 else "moderate" if fill_ratio >= 0.35 else "sparse"

            looks_like_field_table = any(
                re.search(r"^\s*(name|student|teacher|class|date|term|year|grade|subject|candidate|duration|examiner)\s*[:_\.]", x.lower())
                for x in sample_texts
            )
            likely_front_matter = looks_like_field_table or (rows <= 4 and cols <= 4 and density != "dense")

            has_borders = False
            try:
                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr is not None:
                    borders = tblPr.find(qn('w:tblBorders'))
                    has_borders = borders is not None
            except Exception:
                pass

            if density == "dense":
                dense_tables += 1
            if likely_front_matter:
                likely_front_matter_tables += 1
            if has_borders:
                bordered_tables += 1

            table_details.append({
                "rows": rows,
                "cols": cols,
                "density": density,
                "has_borders": has_borders,
                "looks_like_field_table": looks_like_field_table,
                "likely_front_matter_table": likely_front_matter,
                "sample_texts": sample_texts[:6],
            })

        features.append({
            "feature_name": "tables",
            "category": FeatureCategory.CONTEXTUAL_INHERIT,
            "extracted_value": {
                "count": len(doc.tables),
                "dense_table_count": dense_tables,
                "bordered_table_count": bordered_tables,
                "front_matter_table_count": likely_front_matter_tables,
                "details": table_details[:5],
            }
        })

    # --------------------------------------------------
    # 3. Detect textbox / drawing text presence
    # --------------------------------------------------
    textbox_line_count = 0
    try:
        textbox_nodes = doc.element.xpath('.//w:txbxContent//w:t')
        drawing_text_nodes = doc.element.xpath('.//a:t')
        textbox_line_count = len(textbox_nodes) + len(drawing_text_nodes)
    except Exception:
        textbox_line_count = 0

    if textbox_line_count > 0:
        features.append({
            "feature_name": "textboxes",
            "category": FeatureCategory.CONTEXTUAL_INHERIT,
            "extracted_value": {
                "line_count": textbox_line_count,
                "has_textboxes": True,
            }
        })

    # --------------------------------------------------
    # 4. Detect page borders
    # --------------------------------------------------
    if getattr(doc, "sections", None):
        section = doc.sections[0]
        try:
            sect_pr = section._sectPr
            pg_borders = sect_pr.find(qn('w:pgBorders'))
            if pg_borders is not None:
                features.append({
                    "feature_name": "page_borders",
                    "category": FeatureCategory.ALWAYS_INHERIT,
                    "extracted_value": {"has_borders": True}
                })
        except Exception:
            pass

    # --------------------------------------------------
    # 5. Detect watermarks
    # --------------------------------------------------
    try:
        if getattr(doc, "sections", None):
            section = doc.sections[0]
            sect_pr = section._sectPr
            watermark = sect_pr.find(qn('w:watermark'))
            if watermark is not None:
                features.append({
                    "feature_name": "watermark",
                    "category": FeatureCategory.ALWAYS_INHERIT,
                    "extracted_value": {"has_watermark": True}
                })
    except Exception:
        pass

    return features


def _table_has_borders(table) -> bool:
    """Check if a table has borders enabled."""
    try:
        if table.style and 'grid' in table.style.name.lower():
            return True
        # Check first cell for borders
        if table.rows and table.rows[0].cells:
            cell = table.rows[0].cells[0]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.find(qn('w:tcBorders'))
            if borders is not None:
                return True
    except:
        pass
    return False


def _detect_math_columns(doc) -> Dict[str, Any]:
    r"""
    Detect if columns contain math equations with equation numbering.
    Returns dict with classification and confidence.

    ENHANCED DETECTION:
        - Equation numbering patterns: (1), [1], 1., 1)
        - LaTeX math delimiters: $$, $, \( \), \[ \]
        - Math environments: \begin{equation}, \begin{align}
        - Equation references: "see equation (1)", "as shown in (2)"
        - Sequential numbering detection
        - Left/right column alignment patterns typical of math exams
    """
    math_indicators = {
        "equation_numbers": 0,
        "latex_delimiters": 0,
        "math_environments": 0,
        "equation_references": 0,
        "math_keywords": 0,
        "sequential_numbers": False,
        "column_alignment": False
    }
    
    sample_size = 0
    equation_numbers_found = []
    
    # Pattern 1: Equation numbering patterns
    # Matches: (1), [1], 1., 1), (1.2), [1.2], (i), (a), etc.
    eq_number_patterns = [
        r'\(\s*\d+\s*\)',           # (1)
        r'\[\s*\d+\s*\]',           # [1]
        r'^\s*\d+\.\s*$',           # 1. on its own line
        r'\s\d+\.\s*$',             # 1. at end of line
        r'\(\s*\d+\s*\)\s*$',       # (1) at end of line
        r'\[\s*\d+\s*\]\s*$',       # [1] at end of line
        r'\(\s*\d+\.\d+\s*\)',      # (1.2)
        r'\[\s*\d+\.\d+\s*\]',      # [1.2]
        r'^\s*\([ivx]+\)\s*$',      # (i), (ii), (iii) - roman numerals
        r'^\s*\([a-z]\)\s*$',       # (a), (b), (c)
    ]
    
    # Pattern 2: LaTeX math delimiters
    latex_patterns = [
        r'\$\$[^$]+\$\$',           # $$ equation $$
        r'\$[^$]+\$',               # $equation$
        r'\\\(.*?\\\)',             # \( equation \)
        r'\\\[.*?\\\]',             # \[ equation \]
        r'\\begin\{equation\}',     # \begin{equation}
        r'\\end\{equation\}',       # \end{equation}
        r'\\begin\{align\}',        # \begin{align}
        r'\\end\{align\}',          # \end{align}
        r'\\begin\{eqnarray\}',     # \begin{eqnarray}
        r'\\end\{eqnarray\}',       # \end{eqnarray}
    ]
    
    # Pattern 3: Equation references in text
    ref_patterns = [
        r'equation\s*\(?\s*\d+\s*\)?',      # equation (1)
        r'as shown in \(?\s*\d+\s*\)?',     # as shown in (1)
        r'see equation\s*\d+',              # see equation 1
        r'refer to equation\s*\d+',         # refer to equation 1
        r'from equation\s*\(?\s*\d+',       # from equation (1)
        r'in equation\s*\(?\s*\d+',         # in equation (1)
        r'by equation\s*\(?\s*\d+',         # by equation (1)
    ]
    
    # Pattern 4: Math keywords (even without LaTeX)
    math_keywords = [
        r'\b(solve|calculate|find|determine|evaluate|simplify|expand|factorise|differentiate|integrate)\b',
        r'\b(equation|formula|expression|function|derivative|integral|matrix|vector)\b',
        r'\b(graph|plot|sketch|draw|curve|asymptote|intercept|tangent)\b',
        r'\b(prove|show|verify|demonstrate|derive)\b',
        r'\b(limit|sum|product|series|sequence)\b',
        r'[+\-*/=<>≤≥≠√∫∑∏∂∇]',  # Math symbols
        r'\b(sin|cos|tan|sec|csc|cot|arcsin|arccos|arctan)\b',
        r'\b(ln|log|exp|e\^|π|theta|alpha|beta|gamma|delta)\b',
    ]
    
    # Collect text from all paragraphs in first section (where columns would be)
    all_text = ""
    left_column_text = ""
    right_column_text = ""
    
    # First, detect if this is likely a two-column math exam
    for para in doc.paragraphs[:100]:
        text = para.text.strip()
        if not text:
            continue
        
        sample_size += 1
        all_text += " " + text
        
        # Detect equation numbers
        for pattern in eq_number_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                math_indicators["equation_numbers"] += len(matches)
                for m in matches:
                    # Extract the number for sequential detection
                    num_match = re.search(r'(\d+)', m)
                    if num_match:
                        equation_numbers_found.append(int(num_match.group(1)))
                break
        
        # Detect LaTeX delimiters
        for pattern in latex_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                math_indicators["latex_delimiters"] += 1
                break
        
        # Detect math environments
        if re.search(r'\\begin\{equation\}', text, re.IGNORECASE) or \
           re.search(r'\\begin\{align\}', text, re.IGNORECASE) or \
           re.search(r'\\begin\{eqnarray\}', text, re.IGNORECASE):
            math_indicators["math_environments"] += 1
        
        # Detect equation references
        for pattern in ref_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                math_indicators["equation_references"] += 1
                break
        
        # Detect math keywords
        for pattern in math_keywords:
            if re.search(pattern, text, re.IGNORECASE):
                math_indicators["math_keywords"] += 1
                break
    
    # Detect sequential numbering (e.g., 1, 2, 3 or 1, 3, 5)
    if len(equation_numbers_found) >= 3:
        sorted_nums = sorted(equation_numbers_found)
        # Check if numbers are sequential or have common difference
        differences = [sorted_nums[i+1] - sorted_nums[i] for i in range(len(sorted_nums)-1)]
        if all(d == 1 for d in differences[:5]):  # Sequential: 1,2,3,4,5
            math_indicators["sequential_numbers"] = True
        elif len(set(differences)) == 1 and differences[0] > 0:  # Arithmetic progression: 1,3,5,7
            math_indicators["sequential_numbers"] = True
    
    # Detect column alignment patterns (math exams often have left/right columns)
    # Check if paragraphs are short and aligned in two columns
    short_paragraph_count = 0
    for para in doc.paragraphs[:80]:
        text = para.text.strip()
        if text and len(text) < 80:  # Short lines typical of two-column layout
            short_paragraph_count += 1
    
    # If many short paragraphs, likely two-column layout
    if short_paragraph_count > 20 and sample_size > 30:
        math_indicators["column_alignment"] = True
    
    # Calculate weighted score
    score = 0
    max_score = 0
    
    # Equation numbers: high weight (most reliable indicator)
    eq_score = min(math_indicators["equation_numbers"] * 15, 40)
    score += eq_score
    max_score += 40
    
    # LaTeX delimiters: medium-high weight
    latex_score = min(math_indicators["latex_delimiters"] * 10, 25)
    score += latex_score
    max_score += 25
    
    # Math environments: high weight (very reliable)
    env_score = min(math_indicators["math_environments"] * 20, 20)
    score += env_score
    max_score += 20
    
    # Equation references: medium weight
    ref_score = min(math_indicators["equation_references"] * 5, 10)
    score += ref_score
    max_score += 10
    
    # Math keywords: low weight (can appear in non-math contexts)
    keyword_score = min(math_indicators["math_keywords"] * 2, 10)
    score += keyword_score
    max_score += 10
    
    # Sequential numbers bonus
    if math_indicators["sequential_numbers"]:
        score += 10
        max_score += 10
    
    # Column alignment bonus
    if math_indicators["column_alignment"]:
        score += 5
        max_score += 5
    
    # Calculate confidence (0-1)
    confidence = score / max_score if max_score > 0 else 0
    
    # Determine classification
    is_math_columns = confidence > 0.4  # Threshold for math detection
    
    # Build detailed result
    result = {
        "is_math_columns": is_math_columns,
        "confidence": round(confidence, 2),
        "indicators": {
            "equation_numbers": math_indicators["equation_numbers"],
            "latex_delimiters": math_indicators["latex_delimiters"],
            "math_environments": math_indicators["math_environments"],
            "equation_references": math_indicators["equation_references"],
            "math_keywords": math_indicators["math_keywords"],
            "sequential_numbers": math_indicators["sequential_numbers"],
            "column_alignment": math_indicators["column_alignment"]
        },
        "sample_size": sample_size
    }
    
    # Add reasoning for transparency
    reasons = []
    if math_indicators["equation_numbers"] >= 2:
        reasons.append(f"Found {math_indicators['equation_numbers']} equation numbers")
    if math_indicators["latex_delimiters"] >= 1:
        reasons.append(f"Found {math_indicators['latex_delimiters']} LaTeX math delimiters")
    if math_indicators["math_environments"] >= 1:
        reasons.append(f"Found {math_indicators['math_environments']} math environments")
    if math_indicators["equation_references"] >= 1:
        reasons.append(f"Found {math_indicators['equation_references']} equation references")
    if math_indicators["sequential_numbers"]:
        reasons.append("Sequential equation numbering detected")
    if math_indicators["column_alignment"]:
        reasons.append("Two-column layout pattern detected")
    
    result["reasons"] = reasons
    
    logger.info(f"Math columns detection: {is_math_columns} (confidence: {confidence:.2%}) - {', '.join(reasons[:3])}")
    
    return result


# Update the caller in _extract_donor_with_styles to use the enhanced detection
# Replace this line:
#     is_math_columns = _detect_math_columns(doc)
#
# With:
#     math_detection = _detect_math_columns(doc)
#     is_math_columns = math_detection["is_math_columns"]


def _extract_images_enhanced(doc, docx_path: str = "") -> List[Dict[str, Any]]:
    """
    Extract ENHANCED images/logo from donor document with richer raw metadata.

    Build Sequence 3 role:
    - preserve raw branding-image evidence
    - keep body/header/footer images in a consistent schema
    - classify conservatively, not decisively
    - preserve size / placement clues where possible
    - deduplicate obvious repeats
    """
    images = []
    image_id = 1
    seen_keys = set()

    def _safe_inches_from_emu(val):
        try:
            return round(float(val) / 914400.0, 3)
        except Exception:
            return 1.0

    def _extract_size_from_run(run):
        width_inches = 1.0
        height_inches = 1.0

        try:
            ext_nodes = run._element.xpath('.//*[local-name()="extent"]')
            if ext_nodes:
                ext = ext_nodes[0]
                cx = ext.get("cx") or ext.get(qn("cx"))
                cy = ext.get("cy") or ext.get(qn("cy"))
                if cx:
                    width_inches = _safe_inches_from_emu(cx)
                if cy:
                    height_inches = _safe_inches_from_emu(cy)
        except Exception:
            pass

        return width_inches, height_inches

    def _dedupe_key(rel_id, blob, position, width_inches, height_inches):
        blob_len = len(blob) if blob else 0

        # Dedupe should primarily follow the actual image payload and its rough
        # rendered footprint, not the relationship id, because the same logo can
        # be referenced through different rel_ids in body/header/footer parts.
        blob_signature = None
        if blob:
            try:
                blob_signature = (blob_len, blob[:32], blob[-32:] if blob_len >= 32 else blob)
            except Exception:
                blob_signature = (blob_len, None, None)

        return (
            str(position or ""),
            round(width_inches, 2),
            round(height_inches, 2),
            blob_signature,
        )

    def _classify_image(position, para_idx=None, width_inches=1.0, height_inches=1.0, existing_role=""):
        existing_role = str(existing_role or "").strip().lower()

        if existing_role == "branding":
            return "branding", "upstream_branding_role"

        if position == "header":
            return "branding", "header_image"

        if position == "footer":
            return "branding", "footer_image"

        if position == "first_page":
            return "branding", "first_page_image"

        if position == "body":
            if isinstance(para_idx, int) and para_idx <= 8:
                if width_inches <= 2.75 and height_inches <= 2.75:
                    return "branding", "early_small_body_image"
                return "branding", "early_body_image"
            if width_inches <= 2.25 and height_inches <= 2.25:
                return "branding", "small_logo_sized_body_image"
            return "educational", "body_image_default"

        return "educational", "unknown_position_default"

    def _append_image(blob, rel_id, position, alignment="left", para_idx=None, width_inches=1.0, height_inches=1.0, alt_text=None, existing_role=""):
        nonlocal image_id

        image_role, classification_reason = _classify_image(
            position=position,
            para_idx=para_idx,
            width_inches=width_inches,
            height_inches=height_inches,
            existing_role=existing_role,
        )

        key = _dedupe_key(rel_id, blob, position, width_inches, height_inches)
        if key in seen_keys:
            return

        seen_keys.add(key)

        image = {
            "image_id": str(image_id),
            "binary_data": blob,
            "width_inches": width_inches,
            "height_inches": height_inches,
            "position": position,
            "alignment": alignment or "left",
            "alt_text": alt_text,
            "relationship_id": rel_id,
            "image_role": image_role,
            "classification_reason": classification_reason,
            "paragraph_index": para_idx,
        }
        images.append(image)
        image_id += 1

    # --------------------------------------------------
    # 1. Extract images from body paragraphs / runs
    # --------------------------------------------------
    for para_idx, para in enumerate(getattr(doc, "paragraphs", [])):
        alignment = para.alignment.name if getattr(para, "alignment", None) else "left"

        for run in getattr(para, "runs", []):
            if not run._element.xpath('.//w:drawing'):
                continue

            try:
                rel_id = None
                blob = None
                alt_text = None

                blips = run._element.xpath('.//a:blip')
                if blips:
                    rel_id = blips[0].get(qn('r:embed'))

                if rel_id:
                    related_part = doc.part.related_parts.get(rel_id)
                    if related_part is not None:
                        blob = getattr(related_part, "blob", None)

                # Try to recover basic alt text / descr
                try:
                    docpr_nodes = run._element.xpath('.//*[local-name()="docPr"]')
                    if docpr_nodes:
                        alt_text = docpr_nodes[0].get("descr") or docpr_nodes[0].get("title")
                except Exception:
                    pass

                width_inches, height_inches = _extract_size_from_run(run)

                _append_image(
                    blob=blob,
                    rel_id=rel_id,
                    position="body",
                    alignment=alignment,
                    para_idx=para_idx,
                    width_inches=width_inches,
                    height_inches=height_inches,
                    alt_text=alt_text,
                )
            except Exception as e:
                logger.debug(f"Body image extraction failed for run: {e}")

    # --------------------------------------------------
    # 2. Extract images from first-section header/footer
    # --------------------------------------------------
    if getattr(doc, "sections", None):
        section = doc.sections[0]

        if getattr(section, "header", None):
            try:
                for rel_id, rel in section.header.part.related_parts.items():
                    blob = getattr(rel, "blob", None)
                    if not blob:
                        continue

                    _append_image(
                        blob=blob,
                        rel_id=rel_id,
                        position="header",
                        alignment="left",
                        para_idx=None,
                        width_inches=1.0,
                        height_inches=1.0,
                        alt_text=None,
                    )
            except Exception as e:
                logger.debug(f"Header image extraction failed: {e}")

        if getattr(section, "footer", None):
            try:
                for rel_id, rel in section.footer.part.related_parts.items():
                    blob = getattr(rel, "blob", None)
                    if not blob:
                        continue

                    _append_image(
                        blob=blob,
                        rel_id=rel_id,
                        position="footer",
                        alignment="left",
                        para_idx=None,
                        width_inches=1.0,
                        height_inches=1.0,
                        alt_text=None,
                    )
            except Exception as e:
                logger.debug(f"Footer image extraction failed: {e}")

    if docx_path:
        images = _merge_docx_media_fallback_into_images(images, docx_path)

    return images


def _extract_table_structures_enhanced(doc) -> List[Dict[str, Any]]:
    """
    Extract ENHANCED complete table structures from donor document.

    Build Sequence 3 role:
    - preserve raw table evidence without over-interpreting it
    - capture structural clues, sample text, border clues, and front-matter hints
    - provide later cleanup / DIE with enough raw donor evidence to reason safely
    """
    tables = []

    def _clean_text(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _dedupe_keep_order(items):
        seen = set()
        out = []
        for item in items or []:
            s = _clean_text(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    def _cell_text(cell) -> str:
        try:
            return _clean_text(cell.text)
        except Exception:
            return ""

    def _table_border_flags(table) -> Dict[str, bool]:
        flags = {
            "top": False,
            "bottom": False,
            "left": False,
            "right": False,
            "inside_h": False,
            "inside_v": False,
        }

        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is not None:
                borders = tblPr.find(qn('w:tblBorders'))
                if borders is not None:
                    for key, tag in [
                        ("top", 'w:top'),
                        ("bottom", 'w:bottom'),
                        ("left", 'w:left'),
                        ("right", 'w:right'),
                        ("inside_h", 'w:insideH'),
                        ("inside_v", 'w:insideV'),
                    ]:
                        node = borders.find(qn(tag))
                        if node is not None:
                            flags[key] = True
        except Exception:
            pass

        # Fallback to cell-border heuristic if table-level borders absent
        if not any(flags.values()):
            try:
                if table.rows and table.rows[0].cells:
                    cell = table.rows[0].cells[0]
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    borders = tcPr.find(qn('w:tcBorders'))
                    if borders is not None:
                        for key, tag in [
                            ("top", 'w:top'),
                            ("bottom", 'w:bottom'),
                            ("left", 'w:left'),
                            ("right", 'w:right'),
                        ]:
                            node = borders.find(qn(tag))
                            if node is not None:
                                flags[key] = True
            except Exception:
                pass

        return flags

    def _looks_like_field_line(text: str) -> bool:
        s = _clean_text(text).lower()
        if not s:
            return False

        if re.search(r"^\s*(name|student|teacher|class|date|term|year|grade|subject|candidate|duration|examiner)\s*[:_\.]", s):
            return True

        if re.search(r"\b(name|student|teacher|class|date|term|year|grade|subject|candidate|duration|examiner)\b", s):
            if any(ch in text for ch in [":", "_", ".", "/"]):
                return True

        return False

    for table_idx, table in enumerate(getattr(doc, "tables", [])[:20]):
        try:
            row_count = len(table.rows)
            col_count = len(table.columns) if getattr(table, "columns", None) else 0

            table_structure = {
                "index": table_idx,
                "rows": row_count,
                "cols": col_count,
                "has_header_row": False,
                "borders": _table_border_flags(table),
                "column_widths": [],
                "cell_merges": [],
                "cell_styles": {},
                "alignment": None,
                "first_row_style": {},
                "sample_texts": [],
                "first_row_texts": [],
                "first_cells_preview": [],
                "non_empty_cell_count": 0,
                "text_density": "sparse",
                "looks_like_field_table": False,
                "likely_front_matter_table": False,
            }

            # --------------------------------------------------
            # 1. Collect text previews / sample texts
            # --------------------------------------------------
            sample_texts = []
            first_row_texts = []
            first_cells_preview = []
            non_empty_cell_count = 0

            for row_idx, row in enumerate(table.rows[:8]):
                row_preview = []
                for col_idx, cell in enumerate(row.cells[:8]):
                    txt = _cell_text(cell)
                    if txt:
                        non_empty_cell_count += 1
                        sample_texts.append(txt)
                        row_preview.append(txt)

                        if len(first_cells_preview) < 12:
                            first_cells_preview.append(txt)

                if row_idx == 0:
                    first_row_texts = row_preview[:10]

            sample_texts = _dedupe_keep_order(sample_texts)[:20]
            table_structure["sample_texts"] = sample_texts
            table_structure["first_row_texts"] = _dedupe_keep_order(first_row_texts)[:10]
            table_structure["first_cells_preview"] = _dedupe_keep_order(first_cells_preview)[:12]
            table_structure["non_empty_cell_count"] = non_empty_cell_count

            # --------------------------------------------------
            # 2. Text density clue
            # --------------------------------------------------
            total_visible_cells = max(row_count * max(col_count, 1), 1)
            fill_ratio = non_empty_cell_count / total_visible_cells

            if fill_ratio >= 0.7:
                table_structure["text_density"] = "dense"
            elif fill_ratio >= 0.35:
                table_structure["text_density"] = "moderate"
            else:
                table_structure["text_density"] = "sparse"

            # --------------------------------------------------
            # 3. Detect likely field/front-matter table
            # --------------------------------------------------
            field_like_hits = 0
            for txt in sample_texts[:12]:
                if _looks_like_field_line(txt):
                    field_like_hits += 1

            looks_like_field_table = field_like_hits >= 1
            likely_front_matter_table = (
                looks_like_field_table
                or (row_count <= 4 and col_count <= 4 and table_structure["text_density"] != "dense")
            )

            table_structure["looks_like_field_table"] = looks_like_field_table
            table_structure["likely_front_matter_table"] = likely_front_matter_table

            # --------------------------------------------------
            # 4. Detect if first row is likely a header row
            # --------------------------------------------------
            if table.rows and len(table.rows) > 0:
                first_row = table.rows[0]
                first_row_cells = [_cell_text(cell) for cell in first_row.cells]
                non_empty_first_row = [cell for cell in first_row_cells if cell]

                if non_empty_first_row:
                    short_first_row = all(len(cell) < 30 for cell in non_empty_first_row)
                    fieldish_first_row = any(_looks_like_field_line(cell) for cell in non_empty_first_row)

                    # Only call it a header row if it looks compact but NOT like
                    # a candidate-info / field-line row.
                    if short_first_row and not fieldish_first_row:
                        table_structure["has_header_row"] = True

                    # Extract first-row style hints
                    try:
                        first_cell = first_row.cells[0]
                        for para in first_cell.paragraphs:
                            for run in para.runs:
                                if run.font.bold:
                                    table_structure["first_row_style"]["bold"] = True
                                if run.font.italic:
                                    table_structure["first_row_style"]["italic"] = True
                    except Exception:
                        pass

            # --------------------------------------------------
            # 5. Extract column widths
            # --------------------------------------------------
            if getattr(table, "columns", None):
                for col in table.columns:
                    try:
                        width_inches = col.width.inches if col.width else 1.0
                        table_structure["column_widths"].append(round(width_inches, 2))
                    except Exception:
                        table_structure["column_widths"].append(1.0)

            # --------------------------------------------------
            # 6. Detect merged cells
            # --------------------------------------------------
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    try:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()

                        grid_span = tcPr.find(qn('w:gridSpan'))
                        if grid_span is not None:
                            span_val = int(grid_span.get(qn('w:val'), 1))
                            if span_val > 1:
                                table_structure["cell_merges"].append({
                                    "type": "horizontal",
                                    "row": row_idx,
                                    "col": col_idx,
                                    "span": span_val
                                })

                        v_merge = tcPr.find(qn('w:vMerge'))
                        if v_merge is not None:
                            merge_val = v_merge.get(qn('w:val'), 'continue')
                            table_structure["cell_merges"].append({
                                "type": "vertical",
                                "row": row_idx,
                                "col": col_idx,
                                "value": merge_val
                            })
                    except Exception:
                        pass

            # --------------------------------------------------
            # 7. Detect table alignment
            # --------------------------------------------------
            try:
                if table.alignment:
                    table_structure["alignment"] = table.alignment.name
            except Exception:
                pass

            # --------------------------------------------------
            # 8. Add light cell-style hints
            # --------------------------------------------------
            try:
                if table.rows and table.rows[0].cells:
                    first_cell = table.rows[0].cells[0]
                    tc = first_cell._tc
                    tcPr = tc.get_or_add_tcPr()

                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill:
                            table_structure["cell_styles"]["first_cell_fill"] = fill
            except Exception:
                pass

            tables.append(table_structure)

        except Exception as e:
            logger.warning(f"Table {table_idx} extraction failed: {e}")

    return tables


def _extract_list_styles(doc) -> List[Dict[str, Any]]:
    """
    Extract list styles (bullet points, numbering) from document.
    
    Returns list of list style dicts with:
        - type (bullet/numbered)
        - level
        - format (decimal, lowerLetter, upperRoman, etc.)
        - start_at (starting number)
    """
    list_styles = []
    
    for para in doc.paragraphs:
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            
            if 'list' in style_name or 'bullet' in style_name:
                list_info = {
                    "type": "bullet" if 'bullet' in style_name else "numbered",
                    "level": 0,
                    "format": None,
                    "start_at": None,
                    "style_name": para.style.name
                }
                
                # Try to extract numbering format from paragraph properties
                try:
                    pPr = para._element.get_or_add_pPr()
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        ilvl = numPr.find(qn('w:ilvl'))
                        if ilvl is not None:
                            list_info["level"] = int(ilvl.get(qn('w:val'), 0))
                        
                        numId = numPr.find(qn('w:numId'))
                        if numId is not None:
                            list_info["num_id"] = numId.get(qn('w:val'))
                except:
                    pass
                
                # Avoid duplicates
                if list_info not in list_styles:
                    list_styles.append(list_info)
    
    return list_styles[:10]  # Limit to 10 list styles


# ======================================================================================
# PHASE 2: FEATURE CLASSIFICATION & APPROPRIATENESS RULES
# ======================================================================================

# Appropriateness rules - defines when layout features should be inherited
# These rules are SUBJECT-AGNOSTIC - they analyze question structure, not subject names
LAYOUT_APPROPRIATENESS_RULES = {
    "columns": {
        "category": FeatureCategory.CONTEXTUAL_INHERIT,
        "decision_factors": [
            "short_question_count",
            "avg_question_length",
            "has_diagrams",
            "has_tables",
            "has_matching_exercises"
        ],
        "weights": {
            "short_question_count": 0.4,
            "avg_question_length": 0.3,
            "has_diagrams": 0.5,
            "has_tables": 0.4,
            "has_matching": 0.3
        },
        "default": False,
        "uncertainty_threshold": 0.6
    },
    "tables": {
        "category": FeatureCategory.CONTEXTUAL_INHERIT,
        "decision_factors": ["has_data_tables_in_draft", "table_density"],
        "default": True,
        "uncertainty_threshold": 0.5
    },
    "answer_lines": {
        "category": FeatureCategory.ALWAYS_INHERIT,
        "default": True
    },
    "margins": {
        "category": FeatureCategory.ALWAYS_INHERIT,
        "default": True
    },
    "two_column_math": {
        "category": FeatureCategory.NEVER_INHERIT,
        "default": False
    },
    "logo": {
        "category": FeatureCategory.ALWAYS_INHERIT,
        "default": True
    },
    "header_footer": {
        "category": FeatureCategory.ALWAYS_INHERIT,
        "default": True
    }
}


def _analyze_question_structure_rules(questions: List[Dict]) -> Dict[str, Any]:
    """
    Rule-based question structure analysis (fast path).
    
    Returns:
        {
            "recommends_columns": bool,
            "confidence": float,
            "reasons": List[str],
            "metrics": {...},
            "uncertainty_factors": List[str]
        }
    """
    if not questions:
        return {
            "recommends_columns": False,
            "recommends_single_column": True,
            "confidence": 0.5,
            "reasons": ["No questions to analyze"],
            "metrics": {},
            "uncertainty_factors": ["empty_question_list"]
        }
    
    short_count = 0
    long_count = 0
    has_diagrams = False
    has_tables = False
    has_matching = False
    has_calculations = False
    total_chars = 0
    
    # Regex patterns for detection
    diagram_patterns = [r"\b(diagram|figure|image|picture|graph|chart|illustration|photo)\b", r"\b(see below|shown above|refer to)\b"]
    table_patterns = [r"\b(table|matrix|grid|spreadsheet|data set)\b"]
    matching_patterns = [r"\b(match|pair|connect|link|correspond|associate|relate)\b"]
    calculation_patterns = [r"\b(calculate|compute|determine|solve|find|evaluate|derive)\b", r"\$.*\$", r"\\frac|\\sqrt|\\sum"]
    
    for q in questions:
        q_text = " ".join(q.get("question_texts", []))
        char_len = len(q_text)
        total_chars += char_len
        
        if char_len < 60:
            short_count += 1
        elif char_len > 200:
            long_count += 1
        
        # Detect patterns
        for pattern in diagram_patterns:
            if re.search(pattern, q_text, re.I):
                has_diagrams = True
                break
        
        for pattern in table_patterns:
            if re.search(pattern, q_text, re.I):
                has_tables = True
                break
        
        for pattern in matching_patterns:
            if re.search(pattern, q_text, re.I):
                has_matching = True
                break
        
        for pattern in calculation_patterns:
            if re.search(pattern, q_text, re.I):
                has_calculations = True
                break
    
    avg_length = total_chars / max(len(questions), 1)
    
    # Decision logic
    reasons = []
    uncertainty_factors = []
    recommend_columns = False
    confidence = 0.5
    
    # Case 1: Many short questions + no visual obstacles
    if short_count >= 8 and not has_diagrams and not has_tables:
        recommend_columns = True
        confidence = 0.85
        reasons.append(f"{short_count} short questions (<60 chars each)")
        
        if has_matching:
            confidence = 0.95
            reasons.append("Matching/pairing exercises detected - columns beneficial")
        
        if short_count >= 15:
            confidence = 0.95
            reasons.append("High volume of short questions - columns strongly recommended")
    
    # Case 2: Diagrams - must have full width
    if has_diagrams:
        recommend_columns = False
        confidence = 0.95
        reasons.append("Diagrams/figures detected - need full width")
    
    # Case 3: Tables - need full width
    if has_tables:
        recommend_columns = False
        confidence = 0.90
        reasons.append("Tables detected - need full width")
    
    # Case 4: Mixed signals - flag uncertainty
    if short_count >= 5 and has_diagrams:
        uncertainty_factors.append("mixed: short questions + diagrams")
        confidence = 0.55
    
    if short_count >= 5 and long_count >= 5:
        uncertainty_factors.append("mixed: equal short and long questions")
        confidence = 0.50
    
    if avg_length > 150:
        recommend_columns = False
        confidence = 0.80
        reasons.append(f"Long average question length ({avg_length:.0f} chars)")
    
    # Default case
    if not reasons and not recommend_columns:
        reasons.append("Default to single column (conservative)")
    
    return {
        "recommends_columns": recommend_columns,
        "recommends_single_column": not recommend_columns,
        "confidence": confidence,
        "reasons": reasons,
        "metrics": {
            "total_questions": len(questions),
            "short_questions": short_count,
            "long_questions": long_count,
            "avg_question_length_chars": round(avg_length, 1),
            "has_diagrams": has_diagrams,
            "has_tables": has_tables,
            "has_matching": has_matching,
            "has_calculations": has_calculations
        },
        "uncertainty_factors": uncertainty_factors
    }


def _analyze_question_structure(questions: List[Dict]) -> Dict[str, Any]:
    """
    Analyze question structure using rules first, LLM for ambiguous cases.
    
    This is the MAIN entry point for question analysis.
    """
    # 1. Run rule-based analysis first (fast path)
    rule_based_result = _analyze_question_structure_rules(questions)
    
    # 2. Check if LLM is needed
    confidence = rule_based_result.get("confidence", 0.5)
    uncertainty_factors = rule_based_result.get("uncertainty_factors", [])
    has_mixed_signals = len(uncertainty_factors) > 0
    
    if _should_use_llm(confidence, has_mixed_signals):
        logger.info(f"Rule-based confidence low ({confidence}), invoking LLM for question analysis")
        
        # 3. Run LLM analysis (deep path)
        llm_result = _analyze_question_structure_with_llm(questions)
        
        # 4. Combine results (LLM overrides when confidence is higher)
        if llm_result.get("confidence", 0) > confidence:
            llm_result["llm_used"] = True
            llm_result["rule_based_fallback"] = rule_based_result
            return llm_result
    
    # 5. Return rule-based result with LLM flag false
    rule_based_result["llm_used"] = False
    return rule_based_result


def _detect_document_type_rules(draft_md: str) -> Tuple[DocumentType, float]:
    """
    Rule-based document type detection.

    BS5 RULE:
    - classify from the draft's opening/front matter first
    - do not let later answer-key or memo sections redefine the whole document
    - worksheet/test/exam identity should come from the visible primary document,
      not an appendix or trailing memo block
    """
    text = (draft_md or "").strip()
    low = text.lower()

    if not low:
        return (DocumentType.CUSTOM, 0.3)

    # --------------------------------------------------
    # Front slice is the authoritative classification zone
    # --------------------------------------------------
    front_slice = low[:2500]

    def _has(pattern: str, blob: str) -> bool:
        return bool(re.search(pattern, blob, re.I | re.M))

    # --------------------------------------------------
    # 1. Strong front-matter signals first
    # --------------------------------------------------
    if _has(r"\bworksheet\b", front_slice):
        return (DocumentType.WORKSHEET, 0.95)

    if _has(r"\brubric\b", front_slice):
        return (DocumentType.RUBRIC, 0.95)

    if _has(r"\bassignment\b|\btask sheet\b", front_slice):
        return (DocumentType.ASSIGNMENT, 0.90)

    if _has(r"\binvestigation\b|\binquiry\b|\bpractical investigation\b|\bscientific investigation\b", front_slice):
        return (DocumentType.INVESTIGATION, 0.92)

    if _has(r"\blesson\b|\bactivity\b", front_slice):
        return (DocumentType.LESSON, 0.85)

    if _has(r"\bexam\b|\bexamination\b", front_slice):
        return (DocumentType.EXAM, 0.90)

    if _has(r"\btest\b|\bquiz\b", front_slice):
        return (DocumentType.TEST, 0.85)

    # Memo only if it is truly the visible front document
    if _has(r"^\s*#*\s*(memo|answer key|marking key)\b", front_slice):
        return (DocumentType.MEMO, 0.85)

    # --------------------------------------------------
    # 2. Whole-document memo detection only if document is mostly memo-like
    # --------------------------------------------------
    whole_has_memo = _has(r"\bmemo\b|\bmarking key\b|\banswer key\b", low)
    whole_has_worksheet = _has(r"\bworksheet\b", low)
    whole_has_exam = _has(r"\bexam\b|\bexamination\b", low)
    whole_has_test = _has(r"\btest\b|\bquiz\b", low)

    # If worksheet/exam/test is clearly present in the front, don't let a later memo section win.
    if whole_has_memo and not (whole_has_worksheet or whole_has_exam or whole_has_test):
        return (DocumentType.MEMO, 0.75)

    # --------------------------------------------------
    # 3. Heuristic signals
    # --------------------------------------------------
    has_questions = bool(re.search(r"^\s*1\.", text, flags=re.M))
    has_marks = bool(re.search(r"\b\d+\s*marks?\b|\(\s*\d+\s*marks?\s*\)", text, re.I))

    if has_questions and has_marks:
        return (DocumentType.TEST, 0.70)
    if has_questions:
        return (DocumentType.WORKSHEET, 0.65)

    return (DocumentType.CUSTOM, 0.40)


def _detect_document_type(draft_md: str) -> Tuple[DocumentType, float, bool]:
    """
    Detect document type using rules first, LLM for ambiguous cases.
    
    Returns: (document_type, confidence, llm_used)
    """
    # Rule-based
    doc_type, confidence = _detect_document_type_rules(draft_md)
    
    if _should_use_llm(confidence, False):
        logger.info(f"Document type confidence low ({confidence}), invoking LLM")
        llm_result = _detect_document_type_with_llm(draft_md)
        
        if llm_result.get("confidence", 0) > confidence:
            type_str = llm_result.get("document_type", "custom")
            try:
                doc_type = DocumentType(type_str)
            except ValueError:
                doc_type = DocumentType.CUSTOM
            confidence = llm_result.get("confidence", confidence)
            return (doc_type, confidence, True)
    
    return (doc_type, confidence, False)


def _detect_education_level_rules(draft_md: str) -> Tuple[EducationLevel, float]:
    """
    Rule-based education level detection.
    
    Returns: (education_level, confidence)
    """
    text = (draft_md or "").strip().lower()
    
    if not text:
        return (EducationLevel.UNKNOWN, 0.2)
    
    # Year level patterns (school)
    year_match = re.search(r"\byear\s+(\d+)\b", text, re.I)
    if year_match:
        year = int(year_match.group(1))
        if year <= 6:
            return (EducationLevel.PRIMARY, 0.90)
        elif year <= 12:
            return (EducationLevel.SECONDARY, 0.90)
    
    # Course code patterns (university)
    course_patterns = [
        r"\b[A-Z]{2,4}\d{3,4}\b",  # MATH101, PSYC2001
        r"\b[A-Z]{2,4}\s+\d{3,4}\b",  # MATH 101
        r"\b(undergraduate|postgraduate|graduate|doctoral|phd|masters?)\b"
    ]
    
    for pattern in course_patterns:
        if re.search(pattern, text, re.I):
            if re.search(r"\b(phd|doctoral|postgraduate|masters?)\b", text, re.I):
                return (EducationLevel.POSTGRADUATE, 0.85)
            return (EducationLevel.UNDERGRADUATE, 0.80)
    
    # Higher-order verb patterns (suggest university)
    university_verbs = [r"\b(critique|synthesize|evaluate|analyze|theorize|conceptualize)\b"]
    for pattern in university_verbs:
        if re.search(pattern, text, re.I):
            return (EducationLevel.UNDERGRADUATE, 0.65)
    
    return (EducationLevel.UNKNOWN, 0.40)


def _detect_education_level(draft_md: str) -> Tuple[EducationLevel, float, bool]:
    """
    Detect education level using rules first, LLM for ambiguous cases.
    
    Returns: (education_level, confidence, llm_used)
    """
    # Rule-based
    edu_level, confidence = _detect_education_level_rules(draft_md)
    
    if _should_use_llm(confidence, False):
        logger.info(f"Education level confidence low ({confidence}), invoking LLM")
        llm_result = _detect_education_level_with_llm(draft_md)
        
        if llm_result.get("confidence", 0) > confidence:
            level_str = llm_result.get("education_level", "unknown")
            try:
                edu_level = EducationLevel(level_str)
            except ValueError:
                edu_level = EducationLevel.UNKNOWN
            confidence = llm_result.get("confidence", confidence)
            return (edu_level, confidence, True)
    
    return (edu_level, confidence, False)


def _detect_university_subject_patterns(text: str) -> Dict[str, Any]:
    """
    Detect university-level subject patterns.
    """
    text = (text or "").strip().lower()
    
    detected_patterns = []
    confidence = 0.0
    
    # Course code patterns
    if re.search(r"\b[A-Z]{2,4}\d{3,4}\b", text, re.I):
        detected_patterns.append("course_codes")
        confidence += 0.3
    
    # Higher-order terminology
    higher_order_terms = ["critique", "synthesize", "evaluate", "theorize", "conceptualize", "methodology", "theoretical", "framework", "literature review"]
    for term in higher_order_terms:
        if re.search(rf"\b{term}\b", text, re.I):
            detected_patterns.append(term)
            confidence += 0.1
    
    # Reference formats
    if re.search(r"\bet al\.\b|\bibid\.\b|harvard|apa|mla|chicago", text, re.I):
        detected_patterns.append("academic_referencing")
        confidence += 0.2
    
    confidence = min(confidence, 1.0)
    
    suggested_level = EducationLevel.UNDERGRADUATE
    if confidence >= 0.6:
        if re.search(r"\b(phd|doctoral|postgraduate|thesis|dissertation)\b", text, re.I):
            suggested_level = EducationLevel.POSTGRADUATE
    
    return {
        "is_university_level": confidence >= 0.4,
        "confidence": confidence,
        "detected_patterns": detected_patterns,
        "suggested_education_level": suggested_level
    }


def _extract_user_layout_hints(draft_md: str) -> Dict[str, bool]:
    """
    Extract teacher layout hints from markdown comments.
    """
    hints = {}
    
    if not draft_md:
        return hints
    
    # Pattern: <!-- layout: key=value -->
    pattern1 = r"<!--\s*layout:\s*(\w+)\s*=\s*(true|false)\s*-->"
    for match in re.finditer(pattern1, draft_md, re.I):
        key = match.group(1).lower()
        value = match.group(2).lower() == "true"
        hints[key] = value
        logger.info(f"Extracted layout hint: {key}={value}")
    
    # Pattern: <!-- layout: two-column -->
    if re.search(r"<!--\s*layout:\s*two-column\s*-->", draft_md, re.I):
        hints["columns"] = True
        logger.info("Extracted layout hint: columns=True (two-column)")
    
    # Pattern: <!-- layout: single-column -->
    if re.search(r"<!--\s*layout:\s*single-column\s*-->", draft_md, re.I):
        hints["columns"] = False
        logger.info("Extracted layout hint: columns=False (single-column)")
    
    return hints


def _resolve_conflicting_hints(hints: Dict[str, List[Tuple[int, bool]]]) -> Dict[str, Dict[str, Any]]:
    """
    Resolve conflicting user hints.
    """
    resolved = {}
    
    for key, occurrences in hints.items():
        if len(occurrences) == 1:
            line_num, value = occurrences[0]
            resolved[key] = {
                "value": value,
                "source_line": line_num,
                "warning": None
            }
        else:
            # Multiple occurrences - last one wins
            last_line, last_value = occurrences[-1]
            resolved[key] = {
                "value": last_value,
                "source_line": last_line,
                "warning": f"Multiple conflicting hints for '{key}', using last one (line {last_line})"
            }
            logger.warning(resolved[key]["warning"])
    
    return resolved


def _calculate_confidence_from_factors(
    factors: Dict[str, Any],
    weights: Dict[str, float],
    default_confidence: float = 0.5
) -> float:
    """
    Calculate confidence score based on weighted factors.
    """
    total_weight = 0
    weighted_sum = 0
    
    for factor_name, weight in weights.items():
        if factor_name in factors:
            factor_value = factors[factor_name]
            total_weight += weight
            
            # Normalize factor value to 0-1 range
            if isinstance(factor_value, bool):
                normalized = 1.0 if factor_value else 0.0
            elif isinstance(factor_value, (int, float)):
                # For counts, normalize by expected max (e.g., 20 short questions)
                if factor_name == "short_question_count":
                    normalized = min(factor_value / 20.0, 1.0)
                elif factor_name == "avg_question_length":
                    normalized = max(1.0 - (factor_value / 200.0), 0.0)
                else:
                    normalized = min(factor_value / 10.0, 1.0)
            else:
                normalized = 0.5
            
            weighted_sum += normalized * weight
    
    if total_weight == 0:
        return default_confidence
    
    confidence = weighted_sum / total_weight
    return min(max(confidence, 0.0), 1.0)


def _flag_uncertain_decisions(decisions: Dict[str, Decision]) -> Dict[str, Decision]:
    """
    Mark decisions that need teacher confirmation.

    Phase 2 upgrade:
    - low confidence detection
    - conflicting reasoning detection
    - weak/default logic detection
    - provide meaningful alternatives
    """
    for key, decision in decisions.items():

        # -----------------------------
        # Rule 1: Low confidence
        # -----------------------------
        if decision.confidence < 0.6:
            decision.requires_teacher_confirmation = True
            decision.alternative = f"Confirm whether '{key}' should be applied"
            logger.info(f"[UNCERTAIN] {key}: low confidence ({decision.confidence})")
            continue

        # -----------------------------
        # Rule 2: Default / weak logic
        # -----------------------------
        if decision.source == "default":
            decision.requires_teacher_confirmation = True
            decision.alternative = f"No rule defined for '{key}' — teacher input recommended"
            logger.info(f"[UNCERTAIN] {key}: default rule used")
            continue

        # -----------------------------
        # Rule 3: Contradictory signals
        # -----------------------------
        if "suppressed" in decision.reason.lower() and "preferred" in decision.reason.lower():
            decision.requires_teacher_confirmation = True
            decision.alternative = f"Conflicting signals for '{key}' — confirm preference"
            logger.info(f"[UNCERTAIN] {key}: conflicting reasoning detected")
            continue

        # -----------------------------
        # Rule 4: Medium confidence gray zone
        # -----------------------------
        if 0.6 <= decision.confidence <= 0.75:
            decision.requires_teacher_confirmation = True
            decision.alternative = f"Moderate confidence for '{key}' — optional confirmation"
            logger.info(f"[UNCERTAIN] {key}: moderate confidence ({decision.confidence})")

    return decisions


def _build_teacher_confirmation_prompt(uncertain_decisions: Dict[str, Decision]) -> str:
    """
    Build a prompt for teacher when engine is uncertain.
    """
    if not uncertain_decisions:
        return ""
    
    prompt_lines = [
        "The engine is uncertain about the following layout decisions:",
        ""
    ]
    
    for feature, decision in uncertain_decisions.items():
        prompt_lines.append(f"**{feature}**:")
        prompt_lines.append(f"  - Recommended: {'Apply' if decision.apply else 'Do not apply'}")
        prompt_lines.append(f"  - Confidence: {decision.confidence * 100:.0f}%")
        prompt_lines.append(f"  - Reason: {decision.reason}")
        if decision.alternative:
            prompt_lines.append(f"  - Suggested action: {decision.alternative}")
        prompt_lines.append("")
    
    prompt_lines.extend([
        "Please choose:",
        "  (A) Accept the engine's recommendation",
        "  (B) Override - Apply the feature",
        "  (C) Override - Do not apply the feature",
        "  (D) Keep donor's original layout"
    ])
    
    return "\n".join(prompt_lines)


def _decide_contextual_feature_from_layout_needs(
    feature_name: str,
    layout_needs: Dict[str, Any],
    document_type: DocumentType
) -> Decision:
    """
    Decide contextual donor-feature inheritance from draft structure only.

    Design Contract rule:
    - donor contextual features are not inherited by default
    - draft structure must justify them
    - when uncertain, prefer clean single-column usability over donor leakage
    """
    layout_needs = layout_needs or {}

    prefers_columns = bool(layout_needs.get("prefers_columns", False))
    prefers_single_column = bool(layout_needs.get("prefers_single_column", False))
    has_tables = bool(layout_needs.get("has_tables", False))
    has_diagrams = bool(layout_needs.get("has_diagrams", False))
    has_subparts = bool(layout_needs.get("has_subparts", False))
    requires_written_explanations = bool(layout_needs.get("requires_written_explanations", False))
    requires_calculations = bool(layout_needs.get("requires_calculations", False))
    vertical_flow_required = bool(layout_needs.get("vertical_flow_required", False))
    requires_wide_layout = bool(layout_needs.get("requires_wide_layout", False))

    compact_question_flow = bool(layout_needs.get("compact_question_flow", False))
    dense_short_questions = bool(layout_needs.get("dense_short_questions", False))
    long_response_heavy = bool(layout_needs.get("long_response_heavy", False))

    confidence = float(layout_needs.get("confidence", 0.75) or 0.75)

    blocking_wide_or_vertical_signals = any([
        prefers_single_column,
        has_diagrams,
        has_tables,
        vertical_flow_required,
        requires_wide_layout,
        requires_written_explanations,
        long_response_heavy,
    ])

    # --------------------------------------------------
    # COLUMNS
    # --------------------------------------------------
    if feature_name == "columns":
        # ========= NEW: Document‑type hard block =========
        if document_type in {DocumentType.WORKSHEET, DocumentType.MEMO, DocumentType.RUBRIC}:
            return Decision(
                apply=False,
                confidence=0.95,
                reason=f"{document_type.value.title()} document type: columns suppressed by default for readability and vertical flow.",
                source="rule"
            )
        # =================================================
        if blocking_wide_or_vertical_signals:
            return Decision(
                apply=False,
                confidence=max(confidence, 0.90),
                reason=(
                    "Columns suppressed: draft requires readability, vertical flow, "
                    "wide layout, diagrams, tables, or longer written responses."
                ),
                source="rule"
            )

        # ========= TIGHTENED COLUMN THRESHOLD (F4) =========
        # Columns are allowed ONLY for EXAM/TEST with very dense short questions
        # and no conflicting layout needs.
        allowed_doc_types = {DocumentType.EXAM, DocumentType.TEST}
        is_exam_or_test = document_type in allowed_doc_types

        # Extract additional metrics from layout_needs (already available)
        total_qs = layout_needs.get("total_questions", 0)
        short_qs = layout_needs.get("short_question_count", 0)
        long_qs = layout_needs.get("long_question_count", 0)
        has_conflicting = layout_needs.get("has_diagrams", False) or \
                          layout_needs.get("has_tables", False) or \
                          layout_needs.get("has_subparts", False) or \
                          layout_needs.get("requires_written_explanations", False)

        # Strict conditions:
        # - Must be exam/test
        # - At least 12 questions
        # - At least 10 short questions (<60 chars)
        # - No long questions (>200 chars)
        # - No conflicting elements (diagrams, tables, subparts, written explanations)
        # - Compact question flow flag (already computed) must be true
        meets_strict_criteria = (
            is_exam_or_test and
            total_qs >= 12 and
            short_qs >= 10 and
            long_qs == 0 and
            not has_conflicting and
            compact_question_flow
        )

        if meets_strict_criteria:
            return Decision(
                apply=True,
                confidence=0.95,
                reason=(
                    f"Columns allowed: {document_type.value.title()} with {total_qs} total questions, "
                    f"{short_qs} short questions, no long questions, no diagrams/tables/subparts, "
                    "and compact question flow."
                ),
                source="rule"
            )
        # ===================================================

        return Decision(
            apply=False,
            confidence=max(confidence, 0.85),
            reason=(
                "Columns suppressed: no strong draft-structural reason exists to inherit donor columns."
            ),
            source="rule"
        )

    # --------------------------------------------------
    # TABLES
    # --------------------------------------------------
    if feature_name == "tables":
        if has_tables:
            return Decision(
                apply=True,
                confidence=max(confidence, 0.90),
                reason="Tables allowed: draft explicitly requires tabular structure.",
                source="rule"
            )

        return Decision(
            apply=False,
            confidence=max(confidence, 0.90),
            reason="Tables suppressed: donor tables are not inherited unless the draft requires tables.",
            source="rule"
        )

    # --------------------------------------------------
    # ANSWER LINES
    # --------------------------------------------------
    if feature_name == "answer_lines":
        if has_tables or has_diagrams or requires_wide_layout:
            return Decision(
                apply=False,
                confidence=max(confidence, 0.90),
                reason="Answer lines suppressed: draft needs table, diagram, or wide-layout response areas.",
                source="rule"
            )

        if requires_calculations or has_subparts or dense_short_questions:
            return Decision(
                apply=True,
                confidence=max(confidence, 0.85),
                reason="Answer lines allowed: draft needs ordinary response space for calculations, subparts, or short answers.",
                source="rule"
            )

        if requires_written_explanations or long_response_heavy:
            return Decision(
                apply=False,
                confidence=max(confidence, 0.85),
                reason="Generic donor answer lines suppressed: draft needs longer written response space instead.",
                source="rule"
            )

        return Decision(
            apply=False,
            confidence=max(confidence, 0.80),
            reason="Answer lines suppressed: no strong draft-structural reason exists to inherit donor answer lines.",
            source="rule"
        )

    # --------------------------------------------------
    # TEXTBOXES / RESPONSE CONTAINERS
    # --------------------------------------------------
    if feature_name == "textboxes":
        if requires_written_explanations and not has_tables and not has_diagrams:
            return Decision(
                apply=True,
                confidence=max(confidence, 0.85),
                reason="Textboxes allowed: draft needs sustained written responses without conflicting visual/table needs.",
                source="rule"
            )

        return Decision(
            apply=False,
            confidence=max(confidence, 0.85),
            reason="Textboxes suppressed: draft does not strongly require donor-style response boxes.",
            source="rule"
        )

    # --------------------------------------------------
    # DEFAULT CONTEXTUAL FEATURE RULE
    # --------------------------------------------------
    return Decision(
        apply=False,
        confidence=max(confidence, 0.80),
        reason=f"Contextual feature '{feature_name}' suppressed: no draft-structural rule supports inheritance.",
        source="rule"
    )


def _decide_textboxes_from_draft(
    draft_questions: List[Dict[str, Any]],
    layout_needs: Dict[str, Any],
    document_type: DocumentType,
) -> Decision:
    """
    Decide textbox / structured response inheritance from draft truth only.

    BS5 RULE:
    - donor textbox presence alone must never trigger teacher confirmation
    - apply structured response layout only when the draft clearly needs it
    """
    draft_questions = draft_questions or []
    layout_needs = layout_needs or {}

    has_tables = bool(layout_needs.get("has_tables", False))
    has_diagrams = bool(layout_needs.get("has_diagrams", False))
    vertical_flow_required = bool(layout_needs.get("vertical_flow_required", False))

    show_working_count = 0
    table_response_count = 0
    long_response_count = 0

    for q in draft_questions:
        if not isinstance(q, dict):
            continue

        answer_style = str(q.get("answer_style", "") or "").strip().lower()
        marks_value = _try_int(q.get("marks_value", 0), 0)

        if bool(q.get("requires_table", False)) or answer_style == "table_response":
            table_response_count += 1

        if answer_style == "show_working":
            show_working_count += 1

        if marks_value >= 4 or answer_style in {"paragraph_response", "show_working"}:
            long_response_count += 1

    # Deterministic rule:
    # apply structured boxes only where the draft strongly suggests structured response areas
    should_apply = any([
        has_tables,
        has_diagrams,
        show_working_count > 0,
        table_response_count > 0,
        long_response_count >= 2 and vertical_flow_required,
        document_type in {DocumentType.EXAM, DocumentType.TEST} and long_response_count >= 2,
    ])

    if should_apply:
        return Decision(
            apply=True,
            confidence=0.90,
            reason="Draft requires structured response space based on table/workings/long-response signals.",
            alternative=None,
            requires_teacher_confirmation=False,
            source="rule",
        )

    return Decision(
        apply=False,
        confidence=0.90,
        reason="Draft does not require donor textbox-style structured response areas.",
        alternative=None,
        requires_teacher_confirmation=False,
        source="rule",
    )


def _apply_appropriateness_rules(
    layout_features: List[Dict[str, Any]],
    draft_questions: List[Dict],
    education_level: EducationLevel,
    document_type: DocumentType,
    user_hints: Dict[str, bool] = None,
    layout_needs: Dict[str, Any] = None
) -> AppropriatenessResult:
    """
    Decide which layout features to apply.
    """
    user_hints = user_hints or {}
    layout_needs = layout_needs or {}

    # --------------------------------------------------
    # SAFETY NORMALISATION
    # --------------------------------------------------
    # If draft layout signals are weak or missing, default to a clean,
    # readable, single-column document rather than inheriting donor layout.
    layout_needs = {
        "total_questions": _try_int(layout_needs.get("total_questions", 0), 0),
        "avg_question_length": _try_int(layout_needs.get("avg_question_length", 0), 0),
        "short_question_count": _try_int(layout_needs.get("short_question_count", 0), 0),
        "medium_question_count": _try_int(layout_needs.get("medium_question_count", 0), 0),
        "long_question_count": _try_int(layout_needs.get("long_question_count", 0), 0),
        "question_density": str(layout_needs.get("question_density", "low") or "low"),

        "has_tables": bool(layout_needs.get("has_tables", False)),
        "has_diagrams": bool(layout_needs.get("has_diagrams", False)),
        "has_subparts": bool(layout_needs.get("has_subparts", False)),
        "requires_written_explanations": bool(layout_needs.get("requires_written_explanations", False)),
        "requires_calculations": bool(layout_needs.get("requires_calculations", False)),

        "paragraph_response_count": _try_int(layout_needs.get("paragraph_response_count", 0), 0),
        "show_working_count": _try_int(layout_needs.get("show_working_count", 0), 0),
        "table_response_count": _try_int(layout_needs.get("table_response_count", 0), 0),
        "visual_response_count": _try_int(layout_needs.get("visual_response_count", 0), 0),

        "dense_short_questions": bool(layout_needs.get("dense_short_questions", False)),
        "compact_question_flow": bool(layout_needs.get("compact_question_flow", False)),
        "long_response_heavy": bool(layout_needs.get("long_response_heavy", False)),

        "prefers_columns": bool(layout_needs.get("prefers_columns", False)),
        "prefers_single_column": bool(layout_needs.get("prefers_single_column", True)),
        "requires_wide_layout": bool(layout_needs.get("requires_wide_layout", True)),
        "vertical_flow_required": bool(layout_needs.get("vertical_flow_required", True)),
        "confidence": float(layout_needs.get("confidence", 0.6) or 0.6),
    }

    decisions = {}
    
    # Analyze question structure once (handles LLM internally)
    structure_analysis = _analyze_question_structure(draft_questions)
    
    for feature in layout_features:
        feature_name = feature.get("feature_name")
        category = feature.get("category")
        
        # Priority 1: NEVER_INHERIT features are ALWAYS blocked (cannot be overridden)
        if category == FeatureCategory.NEVER_INHERIT:
            decisions[feature_name] = Decision(
                apply=False,
                confidence=1.0,
                reason=f"NEVER_INHERIT category: {feature_name} is subject-specific and cannot be inherited under any circumstances.",
                source="rule"
            )
            continue

        # Priority 2: User hints (only for ALWAYS and CONTEXTUAL)
        if feature_name in user_hints:
            decisions[feature_name] = Decision(
                apply=user_hints[feature_name],
                confidence=1.0,
                reason=f"Teacher override: {feature_name}={user_hints[feature_name]}",
                source="teacher_hint"
            )
            continue
        
        if category == FeatureCategory.ALWAYS_INHERIT:
            decisions[feature_name] = Decision(
                apply=True,
                confidence=1.0,
                reason=f"ALWAYS_INHERIT category: {feature_name} is institutional branding",
                source="rule"
            )
            continue
        
        # At this point, the feature is CONTEXTUAL_INHERIT
        # Check for a learned rule from aggregated teacher corrections
        learned = _get_learned_rule(document_type, feature_name)
        if learned is not None:
            decisions[feature_name] = Decision(
                apply=learned["apply"],
                confidence=learned["confidence"],
                reason=learned["reason"],
                source="learned"
            )
            continue
        
        # Fall back to draft-driven decisions
        if feature_name == "textboxes":
            decisions[feature_name] = _decide_textboxes_from_draft(
                draft_questions=draft_questions,
                layout_needs=layout_needs,
                document_type=document_type
            )
        else:
            decisions[feature_name] = _decide_contextual_feature_from_layout_needs(
                feature_name=feature_name,
                layout_needs=layout_needs,
                document_type=document_type
            )
    
    # Flag uncertain decisions
    decisions = _flag_uncertain_decisions(decisions)
    
    # Apply document type rules
    decisions = _apply_document_type_rules(document_type, decisions)
    
    return AppropriatenessResult(
        decisions=decisions,
        structure_analysis=structure_analysis,
        user_hints_applied=user_hints,
        education_level=education_level,
        document_type=document_type,
        uncertainty_flags=[d.reason for d in decisions.values() if d.requires_teacher_confirmation],
        llm_used=structure_analysis.get("llm_used", False)
    )


def _apply_document_type_rules(
    document_type: DocumentType,
    decisions: Dict[str, Decision]
) -> Dict[str, Decision]:
    """
    Apply document-type adjustments without overriding draft truth blindly.

    Contract rule:
    - document type may strengthen safe decisions
    - document type may suppress risky layout
    - document type must not force donor structures where the draft did not require them
    """
    modified = decisions.copy()

    # --------------------------------------------------
    # EXAM / TEST
    # --------------------------------------------------
    # Exams and tests may tolerate dense layouts, but only when the
    # draft-structure decision already allowed them.
    if document_type in {DocumentType.EXAM, DocumentType.TEST}:
        if "columns" in modified and modified["columns"].apply is True:
            modified["columns"].reason += "; Exam/test type supports dense layout only because draft structure already allowed columns."
            modified["columns"].confidence = min(1.0, modified["columns"].confidence + 0.05)

    # --------------------------------------------------
    # MEMO
    # --------------------------------------------------
    # Memos should remain simple and linear.
    elif document_type == DocumentType.MEMO:
        if "columns" in modified:
            modified["columns"].apply = False
            modified["columns"].confidence = 1.0
            modified["columns"].reason = "Memo document type: columns suppressed because memo layouts should remain linear."
            modified["columns"].source = "rule"

    # --------------------------------------------------
    # RUBRIC
    # --------------------------------------------------
    # Rubrics often use tables, but do not force donor tables unless
    # draft/table logic already supports table use.
    elif document_type == DocumentType.RUBRIC:
        if "columns" in modified:
            modified["columns"].apply = False
            modified["columns"].confidence = 1.0
            modified["columns"].reason = "Rubric document type: columns suppressed because rubric layouts need clear criteria flow."
            modified["columns"].source = "rule"

        if "tables" in modified:
            if modified["tables"].apply is True:
                modified["tables"].confidence = max(modified["tables"].confidence, 0.95)
                modified["tables"].reason += "; Rubric document type supports table use because draft/table logic already allowed it."
            else:
                modified["tables"].reason += "; Rubric type noted, but donor tables remain suppressed because draft did not require tables."

    # --------------------------------------------------
    # WORKSHEET
    # --------------------------------------------------
    # Worksheets favour readability, space, diagrams, and vertical flow.
    elif document_type == DocumentType.WORKSHEET:
        if "columns" in modified:
            modified["columns"].apply = False
            modified["columns"].confidence = max(modified["columns"].confidence, 0.90)
            modified["columns"].reason = "Worksheet document type: columns suppressed for readability, whitespace, and vertical flow."
            modified["columns"].source = "rule"

    return modified


def _modify_layout_feature(
    feature: Dict[str, Any],
    modification: str,
    parameters: Dict
) -> Dict[str, Any]:
    """
    Modify a layout feature rather than just applying or suppressing.
    """
    modified = feature.copy()
    
    if modification == "column_count" and feature.get("feature_name") == "columns":
        modified["extracted_value"]["count"] = parameters.get("count", 2)
        modified["modified"] = True
        modified["modification_reason"] = f"Column count adjusted from {feature.get('extracted_value', {}).get('count', 2)} to {parameters.get('count', 2)}"
    
    elif modification == "table_density" and feature.get("feature_name") == "tables":
        modified["extracted_value"]["density"] = parameters.get("density", "simple")
        modified["modified"] = True
        modified["modification_reason"] = f"Table density set to {parameters.get('density', 'simple')}"
    
    elif modification == "margin_width" and feature.get("feature_name") == "margins":
        modified["extracted_value"]["left"] = parameters.get("left", modified["extracted_value"].get("left", 1.0))
        modified["extracted_value"]["right"] = parameters.get("right", modified["extracted_value"].get("right", 1.0))
        modified["modified"] = True
        modified["modification_reason"] = "Margin width adjusted"
    
    return modified


# ======================================================================================
# PHASE 2.5: LLM-POWERED INTELLIGENCE (when rules are uncertain)
# ======================================================================================

def _should_use_llm(confidence: float, has_mixed_signals: bool) -> bool:
    """Determine if LLM should be invoked for better analysis."""
    return confidence < 0.6 or has_mixed_signals


def _get_llm_confidence_threshold() -> float:
    """Configurable threshold for LLM invocation."""
    return 0.6


def _analyze_question_structure_with_llm(questions: List[Dict]) -> Dict[str, Any]:
    """
    Use LLM to intelligently analyze question structure when rule-based
    confidence is low (<0.6) or signals are mixed.
    """
    if not questions:
        return {
            "recommends_columns": False,
            "recommends_single_column": True,
            "confidence": 0.5,
            "reasons": ["No questions to analyze"],
            "metrics": {},
            "uncertainty_factors": ["empty_question_list"],
            "llm_used": True
        }
    
    # Build prompt with question texts
    questions_text = []
    for i, q in enumerate(questions[:20], 1):  # Limit to 20 questions for token efficiency
        q_text = " ".join(q.get("question_texts", []))
        marks = q.get("marks_label", "")
        if marks:
            questions_text.append(f"{i}. {q_text} {marks}")
        else:
            questions_text.append(f"{i}. {q_text}")
    
    questions_prompt = "\n".join(questions_text)
    
    prompt = f"""You are an expert educational designer analyzing question structure to determine optimal layout.

Analyze these questions and determine if they would benefit from a TWO-COLUMN layout.

CONSIDER:
- Short questions (<60 characters) → columns beneficial
- Long questions (>200 characters) → single column better
- Diagrams, figures, images → need full width (NO columns)
- Tables → need full width (NO columns)
- Matching/pairing exercises → columns beneficial
- Calculations with working space → single column better
- Subparts (a), (b), (c) → usually single column

QUESTIONS:
{questions_prompt}

Return ONLY valid JSON with this exact structure:
{{
    "recommends_columns": true/false,
    "confidence": 0.0-1.0,
    "reasoning": "brief explanation",
    "identified_patterns": ["pattern1", "pattern2"]
}}"""

    try:
        response = call_llm(prompt, model_name="gpt-4o-mini")
        
        # Extract JSON from response
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            return {
                "recommends_columns": result.get("recommends_columns", False),
                "recommends_single_column": not result.get("recommends_columns", False),
                "confidence": min(max(result.get("confidence", 0.7), 0.0), 1.0),
                "reasons": [result.get("reasoning", "LLM analysis")],
                "metrics": {
                    "total_questions": len(questions),
                    "identified_patterns": result.get("identified_patterns", [])
                },
                "uncertainty_factors": [],
                "llm_used": True
            }
    except Exception as e:
        logger.error(f"LLM question structure analysis failed: {e}")
    
    # Fallback to conservative default
    return {
        "recommends_columns": False,
        "recommends_single_column": True,
        "confidence": 0.5,
        "reasons": ["LLM analysis failed, using conservative default"],
        "metrics": {"total_questions": len(questions)},
        "uncertainty_factors": ["llm_failed"],
        "llm_used": True
    }


def _detect_document_type_with_llm(draft_md: str) -> Dict[str, Any]:
    """
    Use LLM to detect document type when pattern matching is uncertain.
    
    Handles:
        - Novel document types not in DocumentType enum
        - Hybrid documents (exam with worksheet sections)
        - Professional/industry documents
    """
    # Take first 3000 chars for context
    preview = draft_md[:3000] if draft_md else ""
    
    prompt = f"""What type of educational document is this?

Choose from: worksheet, exam, test, memo, rubric, assignment, lesson, custom

Document content:
{preview}

Return ONLY valid JSON with this exact structure:
{{
    "document_type": "one of the types above",
    "confidence": 0.0-1.0,
    "reasoning": "brief explanation",
    "alternative_types": ["other possible types"]
}}"""

    try:
        response = call_llm(prompt, model_name="gpt-4o-mini")
        
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            return {
                "document_type": result.get("document_type", "custom"),
                "confidence": min(max(result.get("confidence", 0.7), 0.0), 1.0),
                "reasoning": result.get("reasoning", "LLM classification"),
                "alternative_types": result.get("alternative_types", [])
            }
    except Exception as e:
        logger.error(f"LLM document type detection failed: {e}")
    
    return {
        "document_type": "custom",
        "confidence": 0.5,
        "reasoning": "LLM analysis failed, defaulting to custom",
        "alternative_types": []
    }


def _detect_education_level_with_llm(draft_md: str) -> Dict[str, Any]:
    """
    Use LLM to detect education level when pattern matching fails.
    
    Detects:
        - Graduate vs undergraduate from language complexity
        - Professional certification documents
        - Industry training materials
    """
    preview = draft_md[:3000] if draft_md else ""
    
    prompt = f"""What education level is this document designed for?

Choose from: primary, secondary, undergraduate, postgraduate, professional

Document content:
{preview}

Return ONLY valid JSON with this exact structure:
{{
    "education_level": "one of the levels above",
    "confidence": 0.0-1.0,
    "reasoning": "brief explanation of why",
    "indicators": ["indicator1", "indicator2"]
}}"""

    try:
        response = call_llm(prompt, model_name="gpt-4o-mini")
        
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            return {
                "education_level": result.get("education_level", "unknown"),
                "confidence": min(max(result.get("confidence", 0.7), 0.0), 1.0),
                "reasoning": result.get("reasoning", "LLM classification"),
                "indicators": result.get("indicators", [])
            }
    except Exception as e:
        logger.error(f"LLM education level detection failed: {e}")
    
    return {
        "education_level": "unknown",
        "confidence": 0.5,
        "reasoning": "LLM analysis failed, defaulting to unknown",
        "indicators": []
    }


def _classify_layout_feature_with_llm(
    feature_name: str,
    feature_context: str
) -> FeatureCategory:
    """
    Use LLM to classify ambiguous layout features.
    
    Example: Are these columns from a math exam (NEVER_INHERIT) 
             or general columns (CONTEXTUAL_INHERIT)?
    """
    prompt = f"""Classify this layout feature for a document template system.

Feature: {feature_name}
Context: {feature_context[:500]}

Categories:
- ALWAYS_INHERIT: Institutional branding (logo, margins, fonts, header/footer)
- CONTEXTUAL_INHERIT: Layout that depends on content (columns, tables, answer lines)
- NEVER_INHERIT: Subject-specific structures (math two-column, equation formatting)

Return ONLY valid JSON:
{{
    "category": "ALWAYS_INHERIT" or "CONTEXTUAL_INHERIT" or "NEVER_INHERIT",
    "confidence": 0.0-1.0,
    "reasoning": "brief explanation"
}}"""

    try:
        response = call_llm(prompt, model_name="gpt-4o-mini")
        
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            category_str = result.get("category", "CONTEXTUAL_INHERIT")
            for cat in FeatureCategory:
                if cat.value == category_str or cat.name == category_str:
                    return cat
    except Exception as e:
        logger.error(f"LLM feature classification failed: {e}")
    
    return FeatureCategory.CONTEXTUAL_INHERIT


def _infer_answer_space_needs_with_llm(
    question_text: str,
    marks: int,
    document_type: DocumentType
) -> Dict[str, Any]:
    """
    Use LLM to infer answer space needs for complex questions.
    
    Example: "Draw a diagram showing the water cycle and label each stage"
    → Rule-based might miss that this needs a LABEL_DIAGRAM style.
    """
    doc_type_str = document_type.value if hasattr(document_type, 'value') else str(document_type)
    
    prompt = f"""What answer space style is appropriate for this question?

Question: {question_text}
Marks: {marks}
Document type: {doc_type_str}

Choose from:
- short_response: 2-3 lines for short answers
- paragraph_response: 6-10 lines for explanations
- show_working: 6-8 lines for calculations with working space
- label_diagram: diagram labeling area (no lines)
- table_response: table completion area
- subparts: multiple sub-questions (a, b, c)

Return ONLY valid JSON:
{{
    "answer_style": "one of the styles above",
    "suggested_lines": number,
    "confidence": 0.0-1.0,
    "reasoning": "brief explanation"
}}"""

    try:
        response = call_llm(prompt, model_name="gpt-4o-mini")
        
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            style_str = result.get("answer_style", "short_response")
            
            # Convert to AnswerSpaceStyle enum
            for style in AnswerSpaceStyle:
                if style.value == style_str or style.name == style_str.upper():
                    return {
                        "answer_style": style,
                        "suggested_lines": result.get("suggested_lines", 3),
                        "confidence": min(max(result.get("confidence", 0.7), 0.0), 1.0),
                        "reasoning": result.get("reasoning", ""),
                        "llm_used": True
                    }
    except Exception as e:
        logger.error(f"LLM answer space inference failed: {e}")
    
    return {
        "answer_style": AnswerSpaceStyle.SHORT_RESPONSE,
        "suggested_lines": 3,
        "confidence": 0.5,
        "reasoning": "LLM analysis failed, using default",
        "llm_used": False
    }


def _build_teacher_confirmation_prompt_with_llm(
    uncertain_decisions: Dict[str, Decision],
    draft_context: str
) -> str:
    """
    Use LLM to generate a more natural, context-aware confirmation prompt.
    
    Instead of:
        "The engine is uncertain about columns. Recommended: single column (60%)"
    
    Generate:
        "I notice your biology worksheet has 12 short identification questions
         but also 3 diagrams that need full width. In similar worksheets,
         teachers usually prefer single column for diagrams, but columns work
         well for the short ID questions. What would you prefer?"
    """
    if not uncertain_decisions:
        return ""
    
    # Build context for LLM
    decisions_text = []
    for feature, decision in uncertain_decisions.items():
        decisions_text.append(
            f"- {feature}: recommends {decision.apply} (confidence {decision.confidence:.0%})\n"
            f"  Reason: {decision.reason}"
        )
    
    draft_preview = draft_context[:1500] if draft_context else ""
    
    prompt = f"""Generate a helpful, natural-language confirmation prompt for a teacher.

The document analysis system is uncertain about some layout decisions.

UNCERTAIN DECISIONS:
{chr(10).join(decisions_text)}

DRAFT PREVIEW:
{draft_preview}

Generate a friendly, clear prompt that:
1. Explains what the system is uncertain about
2. Provides the analysis findings in plain English
3. Asks the teacher for their preference with clear options:
   (A) Accept the system's recommendation
   (B) Apply the feature
   (C) Do not apply the feature
   (D) Keep the donor's original layout

Return ONLY the prompt text, no JSON, no extra formatting."""

    try:
        response = call_llm(prompt, model_name="gpt-4o-mini")
        if response and len(response) > 50:
            return response.strip()
    except Exception as e:
        logger.error(f"LLM confirmation prompt generation failed: {e}")
    
    # Fallback to rule-based prompt
    return _build_teacher_confirmation_prompt(uncertain_decisions)


def _enhance_analysis_with_llm(
    analysis_type: str,
    input_data: Any,
    rule_based_result: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Route to appropriate LLM analyzer based on type.
    """
    if analysis_type == "question_structure":
        return _analyze_question_structure_with_llm(input_data)
    elif analysis_type == "document_type":
        return _detect_document_type_with_llm(input_data)
    elif analysis_type == "education_level":
        return _detect_education_level_with_llm(input_data)
    elif analysis_type == "layout_feature":
        result = _classify_layout_feature_with_llm(
            input_data.get("feature_name", ""),
            input_data.get("feature_context", "")
        )
        return {"category": result}
    elif analysis_type == "answer_space":
        return _infer_answer_space_needs_with_llm(
            input_data.get("question_text", ""),
            input_data.get("marks", 0),
            input_data.get("document_type", DocumentType.CUSTOM)
        )
    elif analysis_type == "confirmation_prompt":
        return {"prompt": _build_teacher_confirmation_prompt_with_llm(
            input_data.get("uncertain_decisions", {}),
            input_data.get("draft_context", "")
        )}
    
    return rule_based_result


# ======================================================================================
# PHASE 3: PROFILE NORMALIZATION
# ======================================================================================

def _normalize_template_profile(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Build Sequence 1 normalization boundary.

    Rule:
        - The strict Phase-1 core lives in normalized_profile
        - Legacy outer keys remain available as compatibility mirrors
        - Outer mirrors must be derived from the strict core, not treated as a second truth system
    """
    raw_profile = raw_profile or {}
    if not isinstance(raw_profile, dict):
        raw_profile = {}

    # --------------------------------------
    # SOURCE METADATA
    # --------------------------------------
    source_type = str(raw_profile.get("source_type") or "unknown").strip().lower()
    source_filename = str(raw_profile.get("source_filename") or "").strip()
    template_title = str(raw_profile.get("template_title") or "").strip()

    if not template_title and source_filename:
        template_title = os.path.splitext(source_filename)[0]
    elif not template_title:
        template_title = "School Template"

    # --------------------------------------
    # EXTRACT RAW DONOR SIGNALS
    # Sanitize malformed nested structures BEFORE helper calls.
    # --------------------------------------
    helper_input = dict(raw_profile)

    if not isinstance(helper_input.get("header_footer"), dict):
        helper_input["header_footer"] = {}

    if not isinstance(helper_input.get("title_block_signals"), dict):
        helper_input["title_block_signals"] = {}

    if not isinstance(helper_input.get("style_preferences"), dict):
        helper_input["style_preferences"] = {}

    if not isinstance(helper_input.get("style"), dict):
        helper_input["style"] = {}

    if not isinstance(helper_input.get("page_setup"), dict):
        helper_input["page_setup"] = {}

    if not isinstance(helper_input.get("text_stats"), dict):
        helper_input["text_stats"] = {}

    if not isinstance(helper_input.get("layout_features"), list):
        helper_input["layout_features"] = []

    if not isinstance(helper_input.get("font_samples"), list):
        helper_input["font_samples"] = []

    if not isinstance(helper_input.get("paragraphs"), list):
        helper_input["paragraphs"] = []

    if not isinstance(helper_input.get("extracted_tables"), list):
        nested_raw = {}
        if isinstance(helper_input.get("normalized_profile"), dict):
            nested_raw = helper_input.get("normalized_profile", {}).get("raw", {}) or {}

        if isinstance(nested_raw.get("extracted_tables"), list):
            helper_input["extracted_tables"] = nested_raw.get("extracted_tables", [])
        else:
            helper_input["extracted_tables"] = []

    if not isinstance(helper_input.get("extracted_images"), list) or not helper_input.get("extracted_images"):
        nested_raw = {}
        if isinstance(helper_input.get("normalized_profile"), dict):
            nested_raw = helper_input.get("normalized_profile", {}).get("raw", {}) or {}

        existing_interpreted = helper_input.get("interpreted_donor", {}) or {}
        existing_identity = existing_interpreted.get("identity", {}) if isinstance(existing_interpreted, dict) else {}

        if isinstance(nested_raw.get("extracted_images"), list) and nested_raw.get("extracted_images"):
            helper_input["extracted_images"] = nested_raw.get("extracted_images", [])
        elif isinstance(existing_identity.get("branding_images"), list) and existing_identity.get("branding_images"):
            helper_input["extracted_images"] = existing_identity.get("branding_images", [])
        else:
            helper_input["extracted_images"] = []

    if not isinstance(helper_input.get("style_notes"), list):
        helper_input["style_notes"] = []

    if not isinstance(helper_input.get("first_page_layout_contract"), dict):
        nested_raw = {}
        if isinstance(helper_input.get("normalized_profile"), dict):
            nested_raw = helper_input.get("normalized_profile", {}).get("raw", {}) or {}
        helper_input["first_page_layout_contract"] = nested_raw.get("first_page_layout_contract", {}) if isinstance(nested_raw.get("first_page_layout_contract"), dict) else {}

    # Filter helper-facing list contents so downstream helper functions
    # never receive mixed scalar garbage where they expect dict-like rows.
    helper_input["font_samples"] = [
        x for x in helper_input.get("font_samples", [])
        if isinstance(x, dict)
    ]

    helper_input["paragraphs"] = [
        x for x in helper_input.get("paragraphs", [])
        if isinstance(x, dict)
    ]

    helper_input["extracted_tables"] = [
        x for x in helper_input.get("extracted_tables", [])
        if isinstance(x, dict)
    ]

    helper_input["extracted_images"] = [
        x for x in helper_input.get("extracted_images", [])
        if isinstance(x, dict)
    ]

    helper_input["style_notes"] = [
        str(x).strip()
        for x in helper_input.get("style_notes", [])
        if str(x).strip()
    ]

    interpreted_donor = _interpret_donor_profile(helper_input)
    interpreted_identity = dict(interpreted_donor.get("identity", {}) or {})
    interpreted_framing = dict(interpreted_donor.get("framing", {}) or {})
    interpreted_structure = dict(interpreted_donor.get("structure", {}) or {})
    interpreted_discarded = dict(interpreted_donor.get("discarded_content", {}) or {})
    interpreted_inheritance_map = dict(interpreted_donor.get("inheritance_map", {}) or {})

    # Keep raw style extraction for typography/style normalization.
    # BS4 replaces donor meaning inference, not donor typography extraction.
    institution_identity = _extract_institution_identity(helper_input)
    style_preferences = _extract_style_preferences(helper_input)

    # --------------------------------------
    # NORMALIZE / JSON-SAFE LAYOUT FEATURES
    # --------------------------------------
    raw_layout_features = helper_input.get("layout_features", []) or []
    safe_layout_features = []

    for feature in raw_layout_features:
        if not isinstance(feature, dict):
            feature_copy = {"feature_name": str(feature or "unknown").strip()}
        else:
            feature_copy = dict(feature)

        feature_name = str(feature_copy.get("feature_name") or "unknown").strip()

        cat = feature_copy.get("category")
        if isinstance(cat, FeatureCategory):
            cat = cat.value
        elif isinstance(cat, str):
            cat = cat.strip()
        else:
            cat = None

        if not cat:
            if feature_name in {"margins", "answer_lines", "logo", "header_footer", "page_borders"}:
                cat = FeatureCategory.ALWAYS_INHERIT.value
            elif feature_name in {"two_column_math", "equation_formatting"}:
                cat = FeatureCategory.NEVER_INHERIT.value
            else:
                cat = FeatureCategory.CONTEXTUAL_INHERIT.value

        feature_copy["category"] = cat
        safe_layout_features.append(feature_copy)

    layout_features = safe_layout_features

    # --------------------------------------
    # PAGE SETUP DEFAULTS + COERCION
    # --------------------------------------
    page_setup = dict(helper_input.get("page_setup", {}) or {})
    raw_columns = dict(page_setup.get("columns", {}) or {})

    normalized_columns = {
        "count": max(1, _try_int(raw_columns.get("count"), 1)),
        "space_inches": _try_float(raw_columns.get("space_inches"), 0.0),
        "equal_width": bool(raw_columns.get("equal_width", True)),
    }

    normalized_page_setup = {
        "left_margin_inches": _try_float(page_setup.get("left_margin_inches"), 1.0),
        "right_margin_inches": _try_float(page_setup.get("right_margin_inches"), 1.0),
        "top_margin_inches": _try_float(page_setup.get("top_margin_inches"), 1.0),
        "bottom_margin_inches": _try_float(page_setup.get("bottom_margin_inches"), 1.0),
        "orientation": str(page_setup.get("orientation") or "portrait").strip().lower() or "portrait",
        "page_width_inches": _try_float(page_setup.get("page_width_inches"), 8.5),
        "page_height_inches": _try_float(page_setup.get("page_height_inches"), 11.0),
        "header_distance_inches": _try_float(page_setup.get("header_distance_inches"), 0.5),
        "footer_distance_inches": _try_float(page_setup.get("footer_distance_inches"), 0.5),
        "columns": normalized_columns,
    }

    # --------------------------------------
    # HEADER / FOOTER DEFAULTS
    # --------------------------------------
    raw_header_footer = dict(helper_input.get("header_footer", {}) or {})
    header_footer = {
        "has_header": bool(raw_header_footer.get("has_header", False)),
        "has_footer": bool(raw_header_footer.get("has_footer", False)),
        "header_texts": list(raw_header_footer.get("header_texts", []) or []),
        "footer_texts": list(raw_header_footer.get("footer_texts", []) or []),
        "header_candidates": list(raw_header_footer.get("header_candidates", []) or []),
        "footer_candidates": list(raw_header_footer.get("footer_candidates", []) or []),
        "header_paragraphs": list(raw_header_footer.get("header_paragraphs", []) or []),
        "footer_paragraphs": list(raw_header_footer.get("footer_paragraphs", []) or []),
        "header_tables": list(raw_header_footer.get("header_tables", []) or []),
        "footer_tables": list(raw_header_footer.get("footer_tables", []) or []),
        "has_header_image": bool(raw_header_footer.get("has_header_image", False)),
    }

    # --------------------------------------
    # TEXT / NOTES DEFAULTS
    # --------------------------------------
    raw_text_stats = dict(helper_input.get("text_stats", {}) or {})
    text_stats = {
        "char_count": _try_int(raw_text_stats.get("char_count"), 0),
        "line_count": _try_int(raw_text_stats.get("line_count"), 0),
        "table_count": _try_int(raw_text_stats.get("table_count"), 0),
        "paragraph_count": _try_int(raw_text_stats.get("paragraph_count"), 0),
        "run_count": _try_int(raw_text_stats.get("run_count"), 0),
        "image_count": _try_int(raw_text_stats.get("image_count"), 0),
    }

    style_notes = [str(x).strip() for x in (helper_input.get("style_notes", []) or []) if str(x).strip()]
    first_page_layout_contract = dict(helper_input.get("first_page_layout_contract", {}) or {})

    # --------------------------------------
    # RAW TABLE / IMAGE / PARAGRAPH DEFAULTS
    # --------------------------------------
    extracted_tables = [t for t in (helper_input.get("extracted_tables", []) or []) if isinstance(t, dict)]
    extracted_images = [i for i in (helper_input.get("extracted_images", []) or []) if isinstance(i, dict)]
    paragraphs = [p for p in (helper_input.get("paragraphs", []) or []) if isinstance(p, dict)]

    # --------------------------------------
    # PARAGRAPH SPACING NORMALIZATION
    # --------------------------------------
    before_vals = []
    after_vals = []
    line_vals = []

    for para in paragraphs:
        sb = para.get("space_before")
        sa = para.get("space_after")
        ls = para.get("line_spacing")

        if sb is not None:
            before_vals.append(_try_float(sb, 0.0))
        if sa is not None:
            after_vals.append(_try_float(sa, 0.0))
        if ls is not None:
            try:
                line_vals.append(float(ls))
            except Exception:
                pass

    def _dominant_numeric(values, default):
        if not values:
            return default
        rounded = [round(v, 2) for v in values]
        counter = Counter(rounded)
        return counter.most_common(1)[0][0]

    paragraph_spacing = {
        "before": _dominant_numeric(before_vals, 0.0),
        "after": _dominant_numeric(after_vals, 6.0),
        "line": _dominant_numeric(line_vals, 1.15),
    }

    # --------------------------------------
    # STRICT PHASE-1 CORE (AUTHORITATIVE)
    # --------------------------------------
    strict_core = {
        "schema_version": PROFILE_SCHEMA_VERSION,

        "identity": {
            "institution_name": str(interpreted_identity.get("institution_name") or "").strip(),
            "school_candidates": list(interpreted_identity.get("school_candidates", []) or []),
            "logo_present": bool(interpreted_identity.get("logo_present", False)),
            "logo_positions": list(interpreted_identity.get("logo_positions", []) or []),
            "header_lines": list(header_footer.get("header_texts", []) or []),
            "footer_lines": list(header_footer.get("footer_texts", []) or []),
            "header_identity_lines": list(interpreted_identity.get("running_header_identity", []) or []),
            "footer_identity_lines": list(interpreted_identity.get("running_footer_identity", []) or []),
            "top_block_lines": list(interpreted_identity.get("first_page_identity_lines", []) or []),
            "field_lines": list(interpreted_framing.get("candidate_fields", []) or []),
            "title_candidates": list(interpreted_framing.get("document_frame_lines", []) or []),
        },

        "page": {
            "orientation": normalized_page_setup["orientation"],
            "margins": {
                "top": normalized_page_setup["top_margin_inches"],
                "bottom": normalized_page_setup["bottom_margin_inches"],
                "left": normalized_page_setup["left_margin_inches"],
                "right": normalized_page_setup["right_margin_inches"],
            },
            "page_width_inches": normalized_page_setup["page_width_inches"],
            "page_height_inches": normalized_page_setup["page_height_inches"],
            "header_distance_inches": normalized_page_setup["header_distance_inches"],
            "footer_distance_inches": normalized_page_setup["footer_distance_inches"],
            "columns": {
                "count": normalized_columns["count"],
                "space_inches": normalized_columns["space_inches"],
                "equal_width": normalized_columns["equal_width"],
            },
        },

        "style": {
            "font_family": style_preferences.get("font_family", "Not detected (default: Calibri)"),
            "font_size": _try_float(style_preferences.get("font_size_pt"), 11.0),
            "heading_styles": list(style_preferences.get("heading_styles", []) or []),
            "heading_style_hint": style_preferences.get("heading_style_hint", "Not detected"),
            "visual_feel": style_preferences.get("visual_feel", "Standard document layout"),
            "paragraph_styles_detected": list(style_preferences.get("paragraph_styles_detected", []) or []),
            "has_font_extraction": bool(style_preferences.get("has_font_extraction", False)),
            "paragraph_spacing": paragraph_spacing,
        },

        "layout": {
            "has_columns": normalized_columns["count"] > 1,
            "column_count": normalized_columns["count"],
            "column_space_inches": normalized_columns["space_inches"],
            "columns_equal_width": normalized_columns["equal_width"],
            "has_tables": (
                len(extracted_tables) > 0
                or any(
                    f.get("feature_name") == "tables"
                    and isinstance(f.get("extracted_value"), dict)
                    and int(f.get("extracted_value", {}).get("count", 0) or 0) > 0
                    for f in layout_features
                )
            ),
            "table_count": (
                len(extracted_tables)
                if len(extracted_tables) > 0
                else max(
                    [
                        int(f.get("extracted_value", {}).get("count", 0) or 0)
                        for f in layout_features
                        if f.get("feature_name") == "tables"
                        and isinstance(f.get("extracted_value"), dict)
                    ] or [0]
                )
            ),
            "has_answer_lines": any(f.get("feature_name") == "answer_lines" for f in layout_features),
            "has_page_borders": any(f.get("feature_name") == "page_borders" for f in layout_features),
            "contextual_features_present": [
                f.get("feature_name")
                for f in layout_features
                if f.get("category") == FeatureCategory.CONTEXTUAL_INHERIT.value
            ],
            "never_inherit_features_present": [
                f.get("feature_name")
                for f in layout_features
                if f.get("category") == FeatureCategory.NEVER_INHERIT.value
            ],
        },

        "tables": {
            "count": (
                len(extracted_tables)
                if len(extracted_tables) > 0
                else max(
                    [
                        int(f.get("extracted_value", {}).get("count", 0) or 0)
                        for f in layout_features
                        if f.get("feature_name") == "tables"
                        and isinstance(f.get("extracted_value"), dict)
                    ] or [0]
                )
            ),
            "has_headers": (
                any(bool(t.get("has_header_row")) for t in extracted_tables)
                or any(
                    isinstance(f.get("extracted_value"), dict)
                    and bool(f.get("extracted_value", {}).get("details"))
                    and any(bool(d.get("has_header")) for d in (f.get("extracted_value", {}).get("details") or []))
                    for f in layout_features
                    if f.get("feature_name") == "tables"
                )
            ),
            "dominant_cols": (
                max([_try_int(t.get("cols"), 0) for t in extracted_tables], default=0)
                if len(extracted_tables) > 0
                else max(
                    [
                        max(
                            [_try_int(d.get("cols"), 0) for d in (f.get("extracted_value", {}).get("details") or [])],
                            default=0
                        )
                        for f in layout_features
                        if f.get("feature_name") == "tables"
                        and isinstance(f.get("extracted_value"), dict)
                    ] or [0]
                )
            ),
        },

        "images": {
            "count": len(extracted_images),
            "has_header_logo": any(img.get("position") == "header" for img in extracted_images),
        },

        "features": layout_features,

        "raw": {
            **raw_profile,
            "layout_features": layout_features,
            "page_setup": normalized_page_setup,
            "header_footer": header_footer,
            "first_page_layout_contract": first_page_layout_contract,
            "text_stats": text_stats,
            "interpreted_donor": interpreted_donor,
        },
    }

    # --------------------------------------------------
    # Layout Zone Detection (V1)
    # --------------------------------------------------
    try:
        layout_zones = _detect_layout_zones_from_donor({
            **raw_profile,
            "normalized_profile": strict_core,
            "interpreted_donor": interpreted_donor,
            "institution_identity": {
                "top_block_lines": strict_core["identity"]["top_block_lines"],
                "field_lines": strict_core["identity"]["field_lines"],
                "title_candidates": strict_core["identity"]["title_candidates"],
            },
            "header_footer": header_footer,
        })
    except Exception as e:
        layout_zones = {
            "schema_version": "layout_zones_v1",
            "zones": {},
            "error": str(e)
        }

    # Attach to strict core
    if "layout" in strict_core and isinstance(strict_core["layout"], dict):
        strict_core["layout"]["layout_zones"] = layout_zones

    # --------------------------------------
    # COMPATIBILITY MIRRORS DERIVED FROM CORE
    # --------------------------------------
    compatibility_institution_identity = {
        "likely_institution_name": strict_core["identity"]["institution_name"],
        "school_candidates": strict_core["identity"]["school_candidates"],
        "header_identity_lines": strict_core["identity"]["header_identity_lines"],
        "footer_identity_lines": strict_core["identity"]["footer_identity_lines"],
        "top_block_lines": strict_core["identity"]["top_block_lines"],
        "field_lines": strict_core["identity"]["field_lines"],
        "title_candidates": strict_core["identity"]["title_candidates"],
        "has_logo": strict_core["identity"]["logo_present"],
        "logo_positions": strict_core["identity"]["logo_positions"],
    }

    compatibility_style_preferences = {
        "font_family": strict_core["style"]["font_family"],
        "font_size_pt": strict_core["style"]["font_size"],
        "heading_styles": strict_core["style"]["heading_styles"],
        "heading_style_hint": strict_core["style"]["heading_style_hint"],
        "visual_feel": strict_core["style"]["visual_feel"],
        "paragraph_styles_detected": strict_core["style"]["paragraph_styles_detected"],
        "has_font_extraction": strict_core["style"]["has_font_extraction"],
        "paragraph_spacing": strict_core["style"]["paragraph_spacing"],
    }

    compatibility_page_setup = {
        "left_margin_inches": strict_core["page"]["margins"]["left"],
        "right_margin_inches": strict_core["page"]["margins"]["right"],
        "top_margin_inches": strict_core["page"]["margins"]["top"],
        "bottom_margin_inches": strict_core["page"]["margins"]["bottom"],
        "orientation": strict_core["page"]["orientation"],
        "page_width_inches": strict_core["page"]["page_width_inches"],
        "page_height_inches": strict_core["page"]["page_height_inches"],
        "header_distance_inches": strict_core["page"]["header_distance_inches"],
        "footer_distance_inches": strict_core["page"]["footer_distance_inches"],
        "columns": dict(strict_core["page"]["columns"]),
    }

    normalized = {
        "schema_version": PROFILE_SCHEMA_VERSION,

        "source_meta": {
            "source_type": source_type,
            "source_filename": source_filename,
            "template_title": template_title,
        },

        "institution_identity": compatibility_institution_identity,
        "style_preferences": compatibility_style_preferences,
        "layout_features": strict_core["features"],
        "page_setup": compatibility_page_setup,
        "header_footer": header_footer,
        "first_page_layout_contract": first_page_layout_contract,
        "text_stats": text_stats,
        "preview_notes": style_notes,
        "interpreted_donor": interpreted_donor,
        "layout_zones": layout_zones,

        # Authoritative Phase-1 core
        "normalized_profile": strict_core,

        "engine_metadata": {
            "profile_schema_version": PROFILE_SCHEMA_VERSION,
            "normalization_mode": "phase1_strict_compatibility_bridge",
            "docx_available": DOCX_AVAILABLE,
            "normalized_at": datetime.now().isoformat(),
        },
    }

    # Backward compatibility convenience fields
    normalized["source_type"] = source_type
    normalized["source_filename"] = source_filename
    normalized["template_title"] = template_title

    return normalized


def _extract_title_block_signals(doc) -> Dict[str, Any]:
    """
    Extract richer first-page institution / title-block raw signals from the donor.

    Build Sequence 3 role:
    - preserve first-page evidence from multiple source zones
    - separate likely identity lines from academic/framing lines
    - preserve admin / field-line structures
    - preserve acronym candidates and source context
    - do NOT interpret inheritance yet
    """

    result = {
        "top_lines": [],
        "school_candidates": [],
        "field_lines": [],
        "title_candidates": [],
        "identity_lines": [],
        "framing_lines": [],
        "admin_lines": [],
        "acronym_candidates": [],
        "source_tagged_lines": [],
    }

    if not doc:
        return result

    question_start_re = re.compile(r"^\s*\d+[\.\)]\s*$|^\s*\d+[\.\)]\s+\S")
    subject_words = [
        "mathematics", "maths", "math",
        "biology", "chemistry", "physics",
        "english", "history", "geography",
        "science", "economics", "business", "accounting"
    ]

    institution_terms = [
        "high school", "primary school", "secondary school",
        "college", "grammar", "academy", "institute", "university", "school"
    ]

    admin_terms = [
        "name", "student", "teacher", "date", "class", "term", "year",
        "grade", "subject", "candidate", "candidate name", "student name",
        "duration", "examiner"
    ]

    framing_terms = [
        "question/answer booklet", "examination", "exam", "test", "worksheet",
        "assessment", "semester", "calculator assumed", "calculator free",
        "instructions", "time allowed", "reading time"
    ]

    SCHOOL_TOKEN_STOPWORDS = {
        "OFFICIAL", "YEAR", "EXAM", "TEST", "SECTION", "CALCULATOR",
        "ASSUMED", "FREE", "QUESTION", "ANSWER", "BOOKLET",
        "SEMESTER", "MATHEMATICS", "MATHS", "MATH", "SCIENCE",
        "ENGLISH", "HISTORY", "GEOGRAPHY", "BIOLOGY", "CHEMISTRY",
        "PHYSICS", "BUSINESS", "ACCOUNTING", "TWO", "ONE", "THREE",
        "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE", "TEN"
    }

    def clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def dedupe_keep_order(items):
        seen = set()
        out = []
        for item in items or []:
            key_source = item
            if isinstance(item, dict):
                key_source = clean(item.get("text", ""))
            s = clean(key_source)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(item if isinstance(item, dict) else s)
        return out

    def dedupe_tagged_rows(rows):
        seen = set()
        out = []
        for row in rows or []:
            if not isinstance(row, dict):
                continue
            text = clean(row.get("text", ""))
            source = clean(row.get("source", ""))
            if not text:
                continue
            key = (text.lower(), source.lower())
            if key in seen:
                continue
            seen.add(key)
            out.append({
                "text": text,
                "source": source,
                "index": row.get("index"),
            })
        return out

    def looks_like_sentence(text: str) -> bool:
        text = clean(text)
        if not text:
            return False
        if len(text) > 60:
            return True
        if re.search(r"[,:;]\s", text):
            return True
        if len(text.split()) >= 5:
            return True
        return False

    def is_field_line(text: str) -> bool:
        s = clean(text)
        low = s.lower()
        if not s:
            return False

        if re.search(r"^\s*(name|student|teacher|class|date|term|year|grade|subject|candidate|duration|examiner)\s*[:_\.]", low):
            return True

        if any(term in low for term in admin_terms):
            if any(ch in s for ch in [":", "_", ".", "/"]):
                return True

        compact_matches = re.findall(r"\b(Name|Teacher|Date|Class|Student|Subject|Grade|Candidate|Duration|Examiner)\s*:", s, re.I)
        if compact_matches:
            return True

        return False

    def is_identity_line(text: str) -> bool:
        s = clean(text)
        low = s.lower()
        if not s:
            return False

        if any(term in low for term in institution_terms):
            return True

        if len(s) <= 80 and s.isupper() and len(s.split()) <= 8 and not re.search(r"\d", s):
            return True

        return False

    def is_framing_line(text: str) -> bool:
        s = clean(text)
        low = s.lower()
        if not s:
            return False

        if any(term in low for term in framing_terms):
            return True

        if "year " in low and any(sw in low for sw in subject_words):
            return True

        return False

    def add_collected_line(bucket, text, source, index=None):
        text = clean(text)
        if not text:
            return False
        if len(text) > 220:
            return False
        if question_start_re.match(text):
            return False
        bucket.append({
            "text": text,
            "source": source,
            "index": index,
        })
        return True

    # --------------------------------------------------
    # 1. Gather early front-page text with source context
    # --------------------------------------------------
    collected_rows = []

    def _collect_textbox_lines_from_element(element, source_prefix, limit=40):
        rows = []
        if element is None:
            return rows

        try:
            for idx, node in enumerate(element.xpath('.//w:txbxContent//w:t')):
                txt = clean(getattr(node, "text", "") or "")
                if not txt:
                    continue
                if question_start_re.match(txt):
                    break
                rows.append({"text": txt, "source": source_prefix, "index": idx})
        except Exception:
            pass

        try:
            base_idx = len(rows)
            for idx, node in enumerate(element.xpath('.//a:t')):
                txt = clean(getattr(node, "text", "") or "")
                if not txt:
                    continue
                if question_start_re.match(txt):
                    break
                rows.append({"text": txt, "source": source_prefix, "index": base_idx + idx})
        except Exception:
            pass

        rows = dedupe_tagged_rows(rows)
        return rows[:limit]

    # Early body paragraphs
    for idx, para in enumerate(getattr(doc, "paragraphs", [])[:80]):
        text = clean(getattr(para, "text", ""))
        if not text:
            continue
        if question_start_re.match(text):
            break
        add_collected_line(collected_rows, text, "body_paragraph", idx)

    # Textboxes / grouped drawing text in main body
    try:
        collected_rows.extend(_collect_textbox_lines_from_element(doc.element, "body_textbox", limit=30))
    except Exception:
        pass

    # First header/footer textbox text
    if getattr(doc, "sections", None):
        section = doc.sections[0]

        try:
            if getattr(section, "header", None) is not None:
                collected_rows.extend(_collect_textbox_lines_from_element(section.header._element, "header_textbox", limit=20))
        except Exception:
            pass

        try:
            if getattr(section, "footer", None) is not None:
                collected_rows.extend(_collect_textbox_lines_from_element(section.footer._element, "footer_textbox", limit=10))
        except Exception:
            pass

    # Early table cell text
    for t_idx, table in enumerate(getattr(doc, "tables", [])[:8]):
        stop = False
        for r_idx, row in enumerate(table.rows[:12]):
            for c_idx, cell in enumerate(row.cells[:10]):
                text = clean(cell.text)
                if not text:
                    continue
                if question_start_re.match(text):
                    stop = True
                    break
                add_collected_line(
                    collected_rows,
                    text,
                    "table_cell",
                    f"{t_idx}:{r_idx}:{c_idx}"
                )
            if stop:
                break
        if stop:
            break

    collected_rows = dedupe_tagged_rows(collected_rows)

    result["source_tagged_lines"] = collected_rows[:60]
    result["top_lines"] = [row["text"] for row in collected_rows[:30]]

    # --------------------------------------------------
    # 2. Field lines and admin lines
    # --------------------------------------------------
    field_lines = []
    admin_lines = []

    for row in collected_rows:
        text = row["text"]
        low = text.lower()

        if is_field_line(text):
            field_lines.append(text)
            admin_lines.append(text)
            continue

        if any(term in low for term in admin_terms):
            if len(text) <= 120:
                admin_lines.append(text)

        compact_matches = re.findall(
            r"\b(Name|Teacher|Date|Class|Student|Subject|Grade|Candidate|Duration|Examiner)\s*:",
            text,
            re.I
        )
        for m in compact_matches:
            field_lines.append(f"{m.title()}:")

    result["field_lines"] = dedupe_keep_order(field_lines)[:20]
    result["admin_lines"] = dedupe_keep_order(admin_lines)[:20]

    # --------------------------------------------------
    # 3. Title / framing candidates
    # --------------------------------------------------
    title_candidates = []
    framing_lines = []

    for row in collected_rows:
        text = row["text"]
        low = text.lower()

        if len(text) > 160:
            continue

        if is_framing_line(text):
            title_candidates.append(text)
            framing_lines.append(text)

    result["title_candidates"] = dedupe_keep_order(title_candidates)[:15]
    result["framing_lines"] = dedupe_keep_order(framing_lines)[:15]

    # --------------------------------------------------
    # 4. School / institution candidates + identity lines
    # --------------------------------------------------
    school_candidates = []
    identity_lines = []
    acronym_candidates = []

    for row in collected_rows:
        s = row["text"]
        low = s.lower()
        source = row.get("source", "")

        if is_identity_line(s):
            identity_lines.append(s)
            school_candidates.append(s)

        # Acronym tokens only count when the line shape looks front-matter-ish
        line_is_front_matterish = (
            len(s) <= 60
            and not looks_like_sentence(s)
        ) or ("  " in s) or ("\t" in s) or source in {"header_textbox", "body_textbox", "table_cell"}

        if line_is_front_matterish:
            acronym_tokens = re.findall(r"\b[A-Z]{3,6}\b", s)
            for token in acronym_tokens:
                token_up = token.upper()
                token_low = token.lower()

                if token_up in SCHOOL_TOKEN_STOPWORDS:
                    continue
                # Do not treat generic institution words as acronyms
                if token_up in {"HIGH", "SCHOOL", "PRIMARY", "SECONDARY", "COLLEGE", "ACADEMY", "INSTITUTE", "UNIVERSITY", "GRAMMAR"}:
                    continue
                if token_up.isdigit():
                    continue
                if any(word in token_low for word in subject_words):
                    continue

                acronym_candidates.append(token_up)
                school_candidates.append(token_up)

        # Short uppercase whole line may be institutional, but filter hard
        if len(s) <= 20 and s.isupper():
            s_up = s.upper()
            if (
                s_up not in SCHOOL_TOKEN_STOPWORDS
                and not re.fullmatch(r"YEAR\s+\d+", s_up)
                and not re.search(r"\d", s_up)
            ):
                identity_lines.append(s_up)
                school_candidates.append(s_up)

    result["identity_lines"] = dedupe_keep_order(identity_lines)[:15]
    result["acronym_candidates"] = dedupe_keep_order(acronym_candidates)[:15]
    result["school_candidates"] = dedupe_keep_order(school_candidates)[:15]

    return result


def _extract_institution_identity(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract institution branding signals.

    STRICT RULE:
        - NO subject information (maths, biology, etc.)
        - NO year level information
        - ONLY branding / institutional framing / reusable front-matter identity
    """
    raw_profile = raw_profile or {}

    header_footer = raw_profile.get("header_footer", {}) or {}
    title_block = raw_profile.get("title_block_signals", {}) or {}
    extracted_images = raw_profile.get("extracted_images", []) or []
    branding_ocr = _extract_text_candidates_from_branding_image(extracted_images)

    # --------------------------------------------------
    # Logo clues
    # --------------------------------------------------
    logo_positions = []
    for img in extracted_images:
        if not isinstance(img, dict):
            continue
        if img.get("binary_data"):
            pos = str(img.get("position", "") or "").strip().lower()
            if pos and pos not in logo_positions:
                logo_positions.append(pos)

    has_logo = len(logo_positions) > 0

    # --------------------------------------------------
    # Gather raw candidate pools
    # --------------------------------------------------
    header_candidates = list(header_footer.get("header_candidates", []) or [])
    footer_candidates = list(header_footer.get("footer_candidates", []) or [])
    header_texts = list(header_footer.get("header_texts", []) or [])

    title_school_candidates = list(title_block.get("school_candidates", []) or [])
    top_block_lines = list(title_block.get("top_lines", []) or [])
    title_candidates = list(title_block.get("title_candidates", []) or [])
    field_lines = list(title_block.get("field_lines", []) or [])

    school_candidates = []
    school_candidates.extend(header_candidates)
    school_candidates.extend(footer_candidates)
    school_candidates.extend(title_school_candidates)
    school_candidates.extend(branding_ocr.get("school_candidates", []) or [])

    for text in header_texts:
        if re.search(r"\b(high|primary|secondary|grammar|college|academy|institute|university|school)\b", str(text), re.I):
            school_candidates.append(text)

    # --------------------------------------------------
    # Strict cleaning / ranking helpers
    # --------------------------------------------------
    subject_keywords = [
        "math", "biology", "chemistry", "physics", "english", "history", "geography",
        "science", "art", "music", "drama", "economics", "accounting", "business"
    ]

    banned_exact = {
        "OFFICIAL", "YEAR", "EXAM", "TEST", "SECTION", "QUESTION", "BOOKLET",
        "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE", "TEN",
        "MATHEMATICS", "BIOLOGY", "CHEMISTRY", "PHYSICS", "ENGLISH", "HISTORY", "GEOGRAPHY",
        "CALCULATOR", "ASSUMED", "FREE", "NAME", "TEACHER", "DATE", "CLASS", "STUDENT"
    }

    banned_fragments = [
        "semester", "question/answer booklet", "calculator assumed", "calculator free",
        "time allowed", "reading time", "working time", "material required",
        "important note", "marks", "section two", "section one"
    ]

    def _clean(s: str) -> str:
        return (s or "").replace("\xa0", " ").strip()

    def _is_subject_or_noise(s: str) -> bool:
        low = s.lower()
        if not low:
            return True
        if low.upper() in banned_exact:
            return True
        if any(keyword in low for keyword in subject_keywords):
            return True
        if any(fragment in low for fragment in banned_fragments):
            return True
        if re.fullmatch(r"\d+", s):
            return True
        if re.fullmatch(r"year\s+\d+", low):
            return True
        return False

    # --------------------------------------------------
    # Clean school candidates strictly
    # --------------------------------------------------
    cleaned_school_candidates = []
    seen = set()

    for item in school_candidates:
        s = _clean(str(item))
        if not s:
            continue

        # Split mixed lines into acronym tokens too
        tokens = re.findall(r"\b[A-Z]{2,10}\b", s)
        token_candidates = [s] + tokens

        for cand in token_candidates:
            c = _clean(cand)
            if not c:
                continue
            if _is_subject_or_noise(c):
                continue

            # Accept institution phrases
            accept = False
            if re.search(r"\b(high school|primary school|secondary school|college|grammar|academy|institute|university|school)\b", c, re.I):
                accept = True

            # Only accept acronym-style candidates if they came from genuine
            # text-layer sources, not noisy OCR logo reconstruction.
            elif re.fullmatch(r"[A-Z]{2,10}", c):
                source_text = _clean(str(item))
                if source_text in (header_candidates + footer_candidates + title_school_candidates):
                    accept = True

            if not accept:
                continue

            key = c.lower()
            if key in seen:
                continue
            seen.add(key)
            cleaned_school_candidates.append(c)

    # --------------------------------------------------
    # Clean field lines
    # --------------------------------------------------
    cleaned_fields = []
    seen_fields = set()

    for field in field_lines:
        s = _clean(str(field))
        if not s:
            continue

        matches = re.findall(r"\b(Name|Teacher|Date|Class|Student)\s*:", s, re.I)
        if matches:
            for m in matches:
                label = f"{m.title()}:"
                key = label.lower()
                if key not in seen_fields:
                    seen_fields.add(key)
                    cleaned_fields.append(label)
            continue

        if ":" in s and len(s) < 80 and not any(keyword in s.lower() for keyword in subject_keywords):
            key = s.lower()
            if key not in seen_fields:
                seen_fields.add(key)
                cleaned_fields.append(s)

    # --------------------------------------------------
    # Choose likely institution name STRICTLY from real
    # institution evidence only.
    #
    # IMPORTANT:
    # - Do NOT let exam titles become institution names.
    # - If no school/header/logo evidence exists, leave blank.
    # --------------------------------------------------
    likely_name = ""

    # Prefer explicit institution candidates first
    if cleaned_school_candidates:
        likely_name = cleaned_school_candidates[0]

    else:
        # Try short clean header/footer identity lines only
        identity_line_candidates = []

        for item in header_candidates + footer_candidates:
            s = _clean(str(item))
            if not s:
                continue
            if _is_subject_or_noise(s):
                continue

            if re.search(r"\b(high school|primary school|secondary school|college|grammar|academy|institute|university|school)\b", s, re.I):
                identity_line_candidates.append(s)
            elif re.fullmatch(r"[A-Z]{2,10}", s):
                identity_line_candidates.append(s)

        deduped_identity_line_candidates = []
        seen_identity = set()
        for item in identity_line_candidates:
            key = item.lower()
            if key in seen_identity:
                continue
            seen_identity.add(key)
            deduped_identity_line_candidates.append(item)

        if deduped_identity_line_candidates:
            likely_name = deduped_identity_line_candidates[0]
        else:
            # --------------------------------------------------
            # OCR reconstruction fallback
            # Use noisy OCR fragments conservatively when logo text
            # is clearly pointing to a school identity.
            # --------------------------------------------------
            ocr_lines = list(branding_ocr.get("ocr_lines", []) or [])
            ocr_blob = " | ".join([_clean(x).lower() for x in ocr_lines if _clean(x)])

            has_senior = "senior" in ocr_blob
            has_high = "high" in ocr_blob
            has_school = "school" in ocr_blob

            # Try to recover a leading place/name token from OCR lines
            leading_name = ""
            for line in ocr_lines:
                s = _clean(line)
                if not s:
                    continue

                # Remove obvious noise
                s2 = re.sub(r"[^A-Za-z\s]", " ", s)
                s2 = re.sub(r"\s+", " ", s2).strip()
                if not s2:
                    continue

                # Example target: "Rockingham Senior High School"
                m = re.search(r"([A-Z][a-zA-Z]+)\s+Senior", s2)
                if m:
                    leading_name = m.group(1)
                    break

                # Fallback: first title-case token that is not just "Senior"
                parts = s2.split()
                for part in parts:
                    if part.lower() in {"senior", "high", "school", "independent", "public"}:
                        continue
                    if re.fullmatch(r"[A-Z][a-zA-Z]+", part):
                        leading_name = part
                        break
                if leading_name:
                    break

            if leading_name and has_senior and (has_high or has_school):
                likely_name = f"{leading_name} Senior High School"
            else:
                # OCR fragments alone are not enough unless they form
                # a strong institution phrase.
                likely_name = ""

    return {
        "likely_institution_name": likely_name or branding_ocr.get("likely_institution_name", ""),
        "school_candidates": (
            cleaned_school_candidates[:10]
            if cleaned_school_candidates
            else ([likely_name] if likely_name else [])
        ),
        "header_identity_lines": header_candidates[:5],
        "footer_identity_lines": footer_candidates[:5],
        "top_block_lines": top_block_lines[:10],
        "field_lines": cleaned_fields[:10],
        "title_candidates": title_candidates[:5],
        "has_logo": has_logo,
        "logo_positions": logo_positions[:5],
        "ocr_identity_lines": (branding_ocr.get("ocr_lines", []) or [])[:10],
    }


def _extract_style_preferences(raw_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract typography and visual style.

    Extracts:
        - Font family (most common)
        - Font size (most common)
        - Heading style detection
        - Paragraph spacing hints
    """
    raw_profile = raw_profile or {}

    style = raw_profile.get("style", {}) or {}
    style_preferences = raw_profile.get("style_preferences", {}) or {}

    # --------------------------------------------------
    # Prefer explicitly supplied style_preferences first
    # --------------------------------------------------
    font_family = (
        style_preferences.get("font_family")
        or style.get("dominant_font_family")
        or "Unknown"
    )

    font_size = (
        style_preferences.get("font_size_pt")
        or style.get("dominant_font_size_pt")
        or "Unknown"
    )

    heading_hint = (
        style_preferences.get("heading_style_hint")
        or style.get("heading_style_hint")
        or ""
    )

    visual_feel = (
        style_preferences.get("visual_feel")
        or style.get("visual_feel")
        or ""
    )

    paragraph_styles = (
        style_preferences.get("paragraph_styles_detected")
        or style.get("paragraph_styles")
        or []
    )

    has_font_extraction = bool(style_preferences.get("has_font_extraction", False))

    # --------------------------------------------------
    # Improve from font samples if explicit values missing
    # --------------------------------------------------
    font_samples = raw_profile.get("font_samples", []) or []

    if font_samples and font_family == "Unknown":
        font_counter = Counter([f.get("name") for f in font_samples if f.get("name")])
        if font_counter:
            font_family = font_counter.most_common(1)[0][0]

    if font_samples and font_size == "Unknown":
        size_counter = Counter([f.get("size_pt") for f in font_samples if f.get("size_pt")])
        if size_counter:
            font_size = size_counter.most_common(1)[0][0]

    if font_samples:
        has_font_extraction = True

    # --------------------------------------------------
    # Detect heading hint from donor paragraphs if still weak
    # --------------------------------------------------
    paragraphs = raw_profile.get("paragraphs", []) or []
    if not heading_hint or heading_hint == "Not detected":
        for para in paragraphs[:20]:
            text = (para.get("text") or "").strip()
            if len(text) < 100 and (text.isupper() or text.endswith(":")):
                heading_hint = "All caps or colon-terminated lines detected"
                break

    # --------------------------------------------------
    # Clean paragraph styles
    # --------------------------------------------------
    cleaned_styles = []
    seen = set()
    for ps in paragraph_styles:
        s = str(ps).strip() if ps else ""
        if not s:
            continue
        key = s.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned_styles.append(s)

    return {
        "font_family": font_family if font_family != "Unknown" else "Not detected (default: Calibri)",
        "font_size_pt": _try_float(font_size, 11),
        "heading_style_hint": heading_hint or "Not detected",
        "visual_feel": visual_feel or "Standard document layout",
        "paragraph_styles_detected": cleaned_styles[:10],
        "has_font_extraction": has_font_extraction,
    }


def _build_clean_template_md_from_profile(profile: Dict[str, Any]) -> str:
    """
    Build human-readable preview markdown from profile.
    
    INTELLIGENT ASPECTS:
        - Explains what each feature means
        - Shows classification (always/contextual/never)
        - Previews how decisions will be made
        - Warns about missing or uncertain data
    """
    profile = _normalize_template_profile(profile)
    
    source_meta = profile.get("source_meta", {})
    institution = profile.get("institution_identity", {})
    style = profile.get("style_preferences", {})
    layout_features = profile.get("layout_features", [])
    page_setup = profile.get("page_setup", {})
    first_contract = profile.get("first_page_layout_contract", {}) or {}
    engine_meta = profile.get("engine_metadata", {})
    
    lines = []
    
    # Header
    title = source_meta.get("template_title", "Institution Style Profile")
    lines.append(f"# {title}")
    lines.append("")
    
    # What this is
    lines.append("## What this is")
    lines.append("This is a normalized **institution style profile** extracted from your uploaded donor document.")
    lines.append("")
    lines.append("**The donor teaches the engine how the institution presents documents.**")
    lines.append("**Your draft decides what the new document actually is.**")
    lines.append("")
    
    # Source metadata
    lines.append("## Source Metadata")
    lines.append(f"- **Source type:** {source_meta.get('source_type', 'Unknown')}")
    lines.append(f"- **Source filename:** {source_meta.get('source_filename', 'Unknown')}")
    lines.append(f"- **Profile version:** {engine_meta.get('profile_schema_version', 'unknown')}")
    lines.append("")
    
    # Institution identity (NO subject)
    lines.append("## Institution Identity")
    institution_name = institution.get("likely_institution_name") or "Not clearly detected"
    lines.append(f"- **Likely institution:** {institution_name}")
    
    school_candidates = institution.get("school_candidates", [])
    if school_candidates:
        lines.append("- **Identity signals:**")
        for item in school_candidates[:3]:
            lines.append(f"  - {item}")
    
    field_lines = institution.get("field_lines", [])
    if field_lines:
        lines.append("- **Official field labels:**")
        for item in field_lines[:3]:
            lines.append(f"  - {item}")
    lines.append("")
    
    # Style preferences
    lines.append("## Style Preferences")
    lines.append(f"- **Font family:** {style.get('font_family', 'Not detected')}")
    lines.append(f"- **Font size:** {style.get('font_size_pt', 'Not detected')} pt")
    if style.get("heading_style_hint"):
        lines.append(f"- **Heading hint:** {style.get('heading_style_hint')}")
    lines.append("")
    
    # Page setup
    lines.append("## Page Setup")
    lines.append(f"- **Margins:** L:{page_setup.get('left_margin_inches', 1.0)}\" R:{page_setup.get('right_margin_inches', 1.0)}\" T:{page_setup.get('top_margin_inches', 1.0)}\" B:{page_setup.get('bottom_margin_inches', 1.0)}\"")
    lines.append(f"- **Orientation:** {page_setup.get('orientation', 'portrait').title()}")
    
    columns = page_setup.get("columns", {})
    col_count = columns.get("count", 1)
    lines.append(f"- **Columns:** {col_count} column{'s' if col_count > 1 else ''}")
    lines.append("")

    # First-page layout contract
    lines.append("## First Page Layout Contract")
    lines.append(f"- **Different first-page header/footer:** {str(bool(first_contract.get('different_first_page_header_footer', False))).lower()}")
    lines.append(f"- **First-page header text:** {_short(first_contract.get('first_page_header_text', ''), 120) or 'Not detected'}")
    lines.append(f"- **Later-page header text:** {_short(first_contract.get('default_header_text', ''), 120) or 'Not detected'}")
    lines.append(f"- **First-page footer text:** {_short(first_contract.get('first_page_footer_text', ''), 120) or 'Not detected'}")
    lines.append(f"- **Later-page footer text:** {_short(first_contract.get('default_footer_text', ''), 120) or 'Not detected'}")
    lines.append(f"- **First-page header has images:** {str(bool(first_contract.get('first_page_header_has_images', False))).lower()}")
    lines.append(f"- **Later header has images:** {str(bool(first_contract.get('default_header_has_images', False))).lower()}")
    lines.append(f"- **Footer page numbering detected:** {str(bool(first_contract.get('footer_page_numbering_detected', False))).lower()}")
    lines.append(f"- **Question heading style detected:** {str(bool(first_contract.get('dominant_question_heading_style_signature'))).lower()}")
    lines.append(f"- **Body style detected:** {str(bool(first_contract.get('dominant_body_style_signature'))).lower()}")
    lines.append("")
    
    # Layout features with classification
    lines.append("## Detected Layout Features")
    
    always_features = [f for f in layout_features if f.get("category") == FeatureCategory.ALWAYS_INHERIT]
    contextual_features = [f for f in layout_features if f.get("category") == FeatureCategory.CONTEXTUAL_INHERIT]
    never_features = [f for f in layout_features if f.get("category") == FeatureCategory.NEVER_INHERIT]
    
    if always_features:
        lines.append("### Always Inherit (Institutional Branding)")
        for f in always_features[:5]:
            name = f.get("feature_name", "unknown")
            lines.append(f"- ✅ **{name}** - will always be applied")
        lines.append("")
    
    if contextual_features:
        lines.append("### Contextual Inherit (Layout Decisions)")
        for f in contextual_features[:5]:
            name = f.get("feature_name", "unknown")
            lines.append(f"- 🔄 **{name}** - will be decided based on your draft's structure")
        lines.append("")
    
    if never_features:
        lines.append("### Never Inherit (Subject-Specific)")
        for f in never_features[:5]:
            name = f.get("feature_name", "unknown")
            lines.append(f"- ❌ **{name}** - will NOT be transferred to other subjects")
        lines.append("")
    
    # How decisions are made
    lines.append("## How Layout Decisions Are Made")
    lines.append("")
    lines.append("The engine analyzes your **draft's question structure** to decide which contextual features to apply:")
    lines.append("")
    lines.append("- **Columns** → Applied if many short questions, suppressed if diagrams or tables present")
    lines.append("- **Tables** → Preserved if your draft references data tables")
    lines.append("- **Answer lines** → Always preserved (institutional style)")
    lines.append("- **Margins** → Always preserved (institutional style)")
    lines.append("")
    lines.append("### Teacher Overrides")
    lines.append("You can override any decision by adding hints to your draft:")
    lines.append("```markdown")
    lines.append("<!-- layout: columns=true -->")
    lines.append("<!-- layout: columns=false -->")
    lines.append("<!-- layout: tables=false -->")
    lines.append("```")
    lines.append("")
    
    # Confidence and uncertainty
    lines.append("## Confidence & Uncertainty")
    lines.append("The engine provides confidence scores for every decision.")
    lines.append("- **Confidence ≥ 60%** → Engine decides automatically")
    lines.append("- **Confidence < 60%** → Engine asks for teacher confirmation")
    lines.append("")
    
    # Engine status
    lines.append("## Engine Status")
    lines.append(f"- **DOCX extraction:** {'✅ Available' if engine_meta.get('docx_available') else '⚠️ Fallback mode'}")
    lines.append(f"- **Normalization mode:** {engine_meta.get('normalization_mode', 'unknown')}")
    lines.append(f"- **LLM enhancement:** Ready for ambiguous cases")
    lines.append("")
    
    return "\n".join(lines)


def _create_fallback_template_from_draft(draft_model: Dict) -> Dict[str, Any]:
    """
    Create a fallback style profile when donor extraction fails.

    This function is called when:
    - The uploaded donor file cannot be parsed (corrupted, unsupported format)
    - The donor extraction pipeline (DOCX/PDF/PPTX) raises an exception
    - Required libraries (python-docx, etc.) are missing

    The fallback profile uses safe, neutral defaults:
    - Font: Calibri, 11pt
    - Margins: 1 inch all sides
    - Orientation: Portrait
    - Columns: Single column
    - Header/Footer: None (blank)
    - Logo: None
    - Answer lines: Underscore style (if draft requires them)

    The resulting document will be minimally branded and may lack institutional
    identity. Teachers are strongly encouraged to re-upload a valid donor document
    for full institution styling and layout fidelity.

    Returns a normalized profile dictionary safe for downstream blueprint
    construction, matching the schema expected by _normalize_template_profile().
    """
    draft_model = draft_model or {}
    
    document_type = draft_model.get("document_type", DocumentType.CUSTOM)
    document_type_str = document_type.value if hasattr(document_type, 'value') else str(document_type)
    
    # Create a minimal but functional profile
    fallback_profile = {
        "schema_version": PROFILE_SCHEMA_VERSION,
        "source_type": "fallback",
        "source_filename": "generated_from_draft",
        "template_title": f"{document_type_str.title()} Template",
        "institution_identity": {
            "likely_institution_name": "",
            "school_candidates": [],
            "header_identity_lines": [],
            "footer_identity_lines": [],
            "top_block_lines": [],
            "field_lines": [],
            "title_candidates": [document_type_str.title()]
        },
        "style_preferences": {
            "font_family": "Calibri",
            "font_size_pt": 11,
            "heading_style_hint": "Bold headings detected from draft structure",
            "visual_feel": "Standard document layout",
            "paragraph_styles_detected": ["Normal"],
            "has_font_extraction": False
        },
        "layout_features": [
            {
                "feature_name": "margins",
                "category": FeatureCategory.ALWAYS_INHERIT,
                "extracted_value": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0}
            },
            {
                "feature_name": "answer_lines",
                "category": FeatureCategory.ALWAYS_INHERIT,
                "extracted_value": {"style": "underscore", "count": 0}
            },
            {
                "feature_name": "columns",
                "category": FeatureCategory.CONTEXTUAL_INHERIT,
                "extracted_value": {"count": 1}
            }
        ],
        "page_setup": {
            "left_margin_inches": 1.0,
            "right_margin_inches": 1.0,
            "top_margin_inches": 1.0,
            "bottom_margin_inches": 1.0,
            "orientation": "portrait",
            "page_width_inches": 8.5,
            "page_height_inches": 11.0,
            "columns": {"count": 1, "space_inches": 0}
        },
        "header_footer": {
            "has_header": False,
            "has_footer": False,
            "header_texts": [],
            "footer_texts": []
        },
        "text_stats": {
            "char_count": 0,
            "line_count": 0,
            "table_count": 0
        },
        "preview_notes": [
            "This is a fallback template created because donor extraction failed.",
            "Using default settings: Calibri font, 1-inch margins, single column.",
            "For better results, upload a valid DOCX donor document."
        ],
        "engine_metadata": {
            "profile_schema_version": PROFILE_SCHEMA_VERSION,
            "normalization_mode": "fallback_generation",
            "docx_available": DOCX_AVAILABLE,
            "is_fallback": True,
            "source_document_type": document_type_str
        }
    }
    
    # Add source_type for backward compatibility
    fallback_profile["source_type"] = "fallback"
    fallback_profile["source_filename"] = "generated_from_draft"
    fallback_profile["template_title"] = f"{document_type_str.title()} Template"
    
    return fallback_profile




# ======================================================================================
# PHASE 4: DRAFT INTERPRETATION (CONTENT AUTHORITY)
# ======================================================================================

def _derive_layout_needs_from_question_map(question_map: Dict[str, Any]) -> Dict[str, Any]:
    """
    Derive draft layout needs from question structure only.

    Design Contract rule:
    - layout decisions must come from draft structure
    - no subject-based assumptions
    - uncertain drafts default toward clean single-column usability
    """
    questions = (question_map or {}).get("questions", []) or []

    short_question_count = 0
    medium_question_count = 0
    long_question_count = 0

    has_diagrams = False
    has_tables = False
    has_subparts = False
    requires_written_explanations = False
    requires_calculations = False

    paragraph_response_count = 0
    show_working_count = 0
    table_response_count = 0
    visual_response_count = 0

    total_chars = 0

    diagram_patterns = [r"\bdiagram\b", r"\bfigure\b", r"\bdraw\b", r"\blabel\b", r"\bmap\b", r"\bgraph\b", r"\bchart\b"]
    explanation_patterns = [r"\bexplain\b", r"\bdescribe\b", r"\bdiscuss\b", r"\bwhy\b", r"\bcompare\b", r"\bjustify\b", r"\bevaluate\b"]
    calculation_patterns = [r"\bcalculate\b", r"\bsolve\b", r"\bwork out\b", r"\bfind\b", r"\bdetermine\b", r"\bsimplify\b", r"\bexpand\b", r"\bfactorise\b"]

    for q in questions:
        if not isinstance(q, dict):
            continue

        q_texts = q.get("question_texts", []) or []
        joined = " ".join(str(x) for x in q_texts if x is not None).strip()
        low = joined.lower()
        total_chars += len(joined)

        if joined and len(joined) < 80:
            short_question_count += 1
        elif len(joined) <= 180:
            medium_question_count += 1
        elif len(joined) > 180:
            long_question_count += 1

        if q.get("subparts", []):
            has_subparts = True

        answer_style = str(q.get("answer_style", "") or "").strip().lower()

        if q.get("requires_diagram", False):
            has_diagrams = True
            visual_response_count += 1
        elif any(re.search(p, low) for p in diagram_patterns):
            has_diagrams = True
            visual_response_count += 1

        if q.get("requires_table", False):
            has_tables = True
            table_response_count += 1

        if answer_style == "paragraph_response":
            requires_written_explanations = True
            paragraph_response_count += 1
        elif any(re.search(p, low) for p in explanation_patterns):
            requires_written_explanations = True
            paragraph_response_count += 1

        if answer_style == "show_working":
            requires_calculations = True
            show_working_count += 1
        elif any(re.search(p, low) for p in calculation_patterns):
            requires_calculations = True
            show_working_count += 1

        if answer_style == "table_response":
            has_tables = True
            table_response_count += 1

        if answer_style == "label_diagram":
            has_diagrams = True
            visual_response_count += 1

    total_questions = len(questions)
    avg_question_length = int(total_chars / total_questions) if total_questions else 0

    if total_questions > 15:
        question_density = "high"
    elif total_questions >= 6:
        question_density = "medium"
    else:
        question_density = "low"

    dense_short_questions = (
        total_questions >= 8
        and short_question_count >= max(6, int(total_questions * 0.75))
        and long_question_count == 0
    )

    compact_question_flow = (
        dense_short_questions
        and not has_diagrams
        and not has_tables
        and not has_subparts
        and not requires_written_explanations
    )

    long_response_heavy = (
        long_question_count >= 2
        or paragraph_response_count >= 2
        or avg_question_length >= 160
    )

    requires_wide_layout = bool(has_diagrams or has_tables)

    vertical_flow_required = bool(
        long_response_heavy
        or has_diagrams
        or has_tables
        or has_subparts
        or requires_written_explanations
    )

    prefers_columns = bool(
        compact_question_flow
        and not requires_wide_layout
        and not vertical_flow_required
    )

    prefers_single_column = bool(
        vertical_flow_required
        or requires_wide_layout
        or long_response_heavy
        or total_questions == 0
    )

    confidence = 0.85
    if total_questions == 0:
        confidence = 0.4
    elif prefers_columns and prefers_single_column:
        confidence = 0.6
    elif compact_question_flow or vertical_flow_required:
        confidence = 0.9

    return {
        "total_questions": total_questions,
        "avg_question_length": avg_question_length,
        "short_question_count": short_question_count,
        "medium_question_count": medium_question_count,
        "long_question_count": long_question_count,
        "question_density": question_density,

        "has_diagrams": has_diagrams,
        "has_tables": has_tables,
        "has_subparts": has_subparts,
        "requires_written_explanations": requires_written_explanations,
        "requires_calculations": requires_calculations,

        "paragraph_response_count": paragraph_response_count,
        "show_working_count": show_working_count,
        "table_response_count": table_response_count,
        "visual_response_count": visual_response_count,

        "dense_short_questions": dense_short_questions,
        "compact_question_flow": compact_question_flow,
        "long_response_heavy": long_response_heavy,

        "prefers_columns": prefers_columns,
        "prefers_single_column": prefers_single_column,
        "requires_wide_layout": requires_wide_layout,
        "vertical_flow_required": vertical_flow_required,
        "confidence": confidence,
    }

def analyze_draft_content(
    draft_md: str, 
    year_level: str = "", 
    subject: str = "",
    education_level_hint: str = ""
) -> Dict[str, Any]:
    """
    Parse draft markdown into structured model.
    
    INTELLIGENT ASPECTS:
        - Auto-detects document type (worksheet, exam, memo, rubric, etc.)
        - Extracts education level from content if not provided
        - Parses teacher layout hints from markdown comments
        - Identifies visual slot requirements
        - Builds question map with subpart detection
        - Detects cross-references
    """
    draft_md = (draft_md or "").strip()
    raw_lines = [ln.rstrip() for ln in draft_md.splitlines()]
    nonempty_lines = [ln.strip() for ln in raw_lines if ln.strip()]
    
    # Extract user layout hints first (they can appear anywhere)
    user_hints = _extract_user_layout_hints(draft_md)
    
    # Detect document type (handles LLM fallback internally)
    doc_type, doc_confidence, doc_llm_used = _detect_document_type(draft_md)
    
    # Extract title
    title = _extract_draft_title(draft_md)
    
    # Detect education level (handles LLM fallback internally)
    edu_level, edu_confidence, edu_llm_used = _detect_education_level(draft_md)
    
    # Override with explicit hints if provided
    if year_level:
        edu_level = EducationLevel.SECONDARY if "year" in year_level.lower() else edu_level
    if education_level_hint:
        try:
            edu_level = EducationLevel(education_level_hint.lower())
        except ValueError:
            pass
    
    # Extract instruction block
    instruction_block = _extract_instruction_block(draft_md)
    
    # Extract visual slots
    visual_slots = _extract_visual_slots(draft_md)
    
    # Extract question map
    question_map = extract_draft_question_map(draft_md)
    
    # Detect cross-references
    cross_references = _detect_cross_references(draft_md)
    
    # Detect page breaks
    page_breaks = _detect_draft_page_breaks(draft_md)
    
    # Build draft model
    draft_model = {
        # --------------------------------------------------
        # BACKWARD-COMPATIBLE TOP LEVEL
        # --------------------------------------------------
        "document_type": doc_type,
        "document_type_confidence": doc_confidence,
        "document_type_llm_used": doc_llm_used,
        "title": title,
        "year_level": year_level or "",
        "subject": subject or "",
        "education_level": edu_level,
        "education_level_confidence": edu_confidence,
        "education_level_llm_used": edu_llm_used,
        "user_layout_hints": user_hints,
        "instruction_block": instruction_block,
        "visual_slots": visual_slots,
        "question_map": question_map,
        "cross_references": cross_references,
        "page_breaks": page_breaks,
        "raw_text": draft_md,
        "raw_lines": nonempty_lines,
        "question_count": len(question_map.get("questions", [])),
        "uncertainties": [],
        "layout_needs": _derive_layout_needs_from_question_map(question_map),

        # --------------------------------------------------
        # NEW PHASE-1 CORE TRUTH LAYER
        # --------------------------------------------------
        "normalized_draft_model": {
            "document_type": doc_type.value if isinstance(doc_type, DocumentType) else str(doc_type),
            "title": title,
            "subject": subject or "",
            "year_level": year_level or "",
            "question_map": question_map,
            "instruction_block": instruction_block,
            "layout_needs": _derive_layout_needs_from_question_map(question_map),
            "visual_slots": visual_slots,
            "raw_text": draft_md,
            "raw_lines": nonempty_lines,
        },

        # --------------------------------------------------
        # ENGINE / DIAGNOSTIC METADATA
        # --------------------------------------------------
        "analysis_metadata": {
            "education_level": edu_level.value if isinstance(edu_level, EducationLevel) else str(edu_level),
            "education_level_confidence": edu_confidence,
            "education_level_llm_used": edu_llm_used,
            "document_type_confidence": doc_confidence,
            "document_type_llm_used": doc_llm_used,
            "cross_reference_count": len(cross_references),
            "page_break_count": len(page_breaks),
            "question_count": len(question_map.get("questions", [])),
        },
    }
    
    # Infer answer space needs (after question map is built)
    draft_model["answer_space_needs"] = _infer_answer_space_needs(draft_model)
    
    # Collect uncertainties
    if doc_confidence < 0.7:
        draft_model["uncertainties"].append(f"Document type detection confidence low: {doc_confidence}")
    if edu_confidence < 0.7:
        draft_model["uncertainties"].append(f"Education level detection confidence low: {edu_confidence}")
    
    return draft_model


def _detect_document_type(draft_md: str) -> Tuple[DocumentType, float, bool]:
    """
    Detect document type using the BS5-safe rules-first path.

    IMPORTANT:
    This definition exists later in the file and therefore becomes the
    effective runtime definition. It must stay aligned with the earlier
    BS5-safe detector logic.
    """
    # Rule-based
    doc_type, confidence = _detect_document_type_rules(draft_md)

    if _should_use_llm(confidence, False):
        logger.info(f"Document type confidence low ({confidence}), invoking LLM")
        llm_result = _detect_document_type_with_llm(draft_md)

        if llm_result.get("confidence", 0) > confidence:
            type_str = llm_result.get("document_type", "custom")
            try:
                doc_type = DocumentType(type_str)
            except ValueError:
                doc_type = DocumentType.CUSTOM
            confidence = llm_result.get("confidence", confidence)
            return (doc_type, confidence, True)

    return (doc_type, confidence, False)


def _extract_draft_title(draft_md: str) -> str:
    """
    Extract a clean draft title.

    Rules:
    - Prefer genuine markdown titles/headings
    - Reject instructions, marks-only lines, visual placeholders, and raw question lines
    - Clean "Question 1 (4 marks): Motion" down to "Motion"
    - Fall back safely if no true title exists
    """
    lines = [ln.strip() for ln in (draft_md or "").splitlines() if ln.strip()]

    if not lines:
        return "Untitled Document"

    def _clean_title_candidate(s: str) -> str:
        s = (s or "").strip()

        # Remove markdown heading markers
        s = re.sub(r"^#{1,6}\s*", "", s).strip()

        # If this is a heading like "Question 1 (4 marks): Motion"
        # reduce it to the trailing human title part if present.
        m = re.match(r"^Question\s+\d+\s*(?:\([^)]+\))?\s*[:\-–—]\s*(.+)$", s, re.I)
        if m:
            s = m.group(1).strip()

        # Strip trailing marks if the candidate itself contains them
        s = re.sub(r"\s*(\[[^\]]*marks?[^\]]*\]|\(\s*\d+\s*marks?\s*\)|\b\d+\s*marks?\b)\s*$", "", s, flags=re.I).strip()

        # Remove surrounding punctuation noise
        s = re.sub(r"^[\s:\-–—]+", "", s).strip()
        s = re.sub(r"[\s:\-–—]+$", "", s).strip()

        return s

    def _is_instruction_like_line(s: str) -> bool:
        low = (s or "").strip().lower()

        if not low:
            return True

        instruction_patterns = [
            r"^instructions?\s*:?\s*$",
            r"^answer all questions\b",
            r"^show all working\b",
            r"^read (the|each|all)\b",
            r"^write your answers\b",
            r"^calculator\b",
            r"^materials required\b",
            r"^time allowed\b",
            r"^do not use\b",
        ]
        return any(re.search(p, low, re.I) for p in instruction_patterns)

    def _is_visual_placeholder_line(s: str) -> bool:
        s = (s or "").strip()
        if re.match(r"^\[\[VISUAL\s*:", s, re.I):
            return True
        if re.match(r"^v\d+$", s, re.I):
            return True
        return False

    def _is_marks_only_line(s: str) -> bool:
        s = (s or "").strip()
        return bool(re.fullmatch(r"(\[[^\]]*marks?[^\]]*\]|\(\s*\d+\s*marks?\s*\)|\b\d+\s*marks?\b)", s, re.I))

    def _is_question_like_line(s: str) -> bool:
        s = (s or "").strip()

        # Standard numbered question starts
        if re.match(r"^\d+\.\s*", s):
            return True

        # Weak-format numbered starts, e.g. "1 what is photosynthesis"
        if re.match(r"^\d+\s+\S+", s):
            return True

        # Subpart-style starts
        if re.match(r"^\(?[a-zA-Z]\)\s*", s):
            return True

        # Weak subpart formatting, e.g. "a name the gas absorbed"
        if re.match(r"^[a-zA-Z]\s+\S+", s):
            return True

        # Heading-style question markers
        if re.match(r"^#{1,6}\s*Question\s+\d+", s, re.I):
            return True

        return False

    def _is_probable_title_line(s: str) -> bool:
        raw = (s or "").strip()
        cleaned = _clean_title_candidate(raw)

        if not cleaned:
            return False
        if len(cleaned) <= 3:
            return False
        if _is_instruction_like_line(raw):
            return False
        if _is_visual_placeholder_line(raw):
            return False
        if _is_marks_only_line(raw):
            return False

        # Accept markdown heading lines if they clean into something human
        if raw.startswith("#"):
            # But reject bare "Question 1" style headings with no meaningful suffix
            if re.match(r"^#{1,6}\s*Question\s+\d+\s*$", raw, re.I):
                return False
            return True

        # For non-heading lines, reject obvious question-like content
        if _is_question_like_line(raw):
            return False

        # Also reject plain body/prose lines.
        # A non-heading title candidate should usually look title-like,
        # not like a full sentence from question content.
        if re.search(r"[.!?]$", raw):
            return False

        # Reject long prose-like lines unless they look like a document label
        if len(cleaned) > 60:
            return False

        titleish_patterns = [
            r"\b(test|exam|worksheet|assignment|rubric|memo|lesson|quiz|activity)\b",
            r"\byear\s+\d+\b",
            r"\bterm\s+\d+\b",
            r"\bsemester\s+\d+\b",
            r"\bscience\b|\bmathematics\b|\bmaths\b|\benglish\b|\bhistory\b|\bgeography\b",
        ]

        if any(re.search(p, cleaned, re.I) for p in titleish_patterns):
            return True

        # Otherwise, for non-heading lines, be conservative
        return False

    # --------------------------------------------------
    # Priority 1: genuine markdown title / heading
    # --------------------------------------------------
    for line in lines[:12]:
        if line.startswith("#"):
            cleaned = _clean_title_candidate(line)
            if _is_probable_title_line(line) and cleaned:
                return cleaned[:150]

    # --------------------------------------------------
    # Priority 2: first probable non-question, non-instruction title line
    # --------------------------------------------------
    for line in lines[:12]:
        if _is_probable_title_line(line):
            cleaned = _clean_title_candidate(line)
            if cleaned:
                return cleaned[:150]

    # --------------------------------------------------
    # Priority 3: recover only from heading-style question headings
    # Example: "## Question 1 (4 marks): Motion" -> "Motion"
    # Do NOT promote ordinary question body lines into titles.
    # --------------------------------------------------
    for line in lines[:12]:
        raw = (line or "").strip()

        if not raw.startswith("#"):
            continue

        if not re.match(r"^#{1,6}\s*Question\s+\d+", raw, re.I):
            continue

        cleaned = _clean_title_candidate(raw)
        if cleaned and cleaned.lower() not in {"question 1", "question 2", "question 3"}:
            if not _is_instruction_like_line(cleaned) and not _is_marks_only_line(cleaned):
                if len(cleaned) > 3:
                    return cleaned[:150]

    return "Untitled Document"


def _extract_instruction_block(draft_md: str) -> Dict[str, Any]:
    """Extract instruction section from draft."""
    lines = (draft_md or "").splitlines()
    
    instruction_lines = []
    in_instruction_block = False
    instruction_heading_found = False
    
    # Patterns that indicate instruction content
    instruction_patterns = [
        r"^\s*instructions?\s*:?\s*$",
        r"^\s*special instructions?\s*:?\s*$",
        r"^\s*time allowed\b",
        r"^\s*answer all questions\b",
        r"^\s*show all working\b",
        r"^\s*read (the|each|all)\b",
        r"^\s*use the .* image below\b",
        r"^\s*write your answers\b",
        r"^\s*do not use\b",
        r"^\s*calculator\b",
        r"^\s*materials required\b",
    ]
    
    for line in lines:
        s = line.strip()
        if not s:
            if in_instruction_block:
                # Empty line might end instruction block
                in_instruction_block = False
            continue
        
        # Check for instruction heading
        if re.match(r"^\s*instructions?\s*:?\s*$", s, re.I):
            in_instruction_block = True
            instruction_heading_found = True
            continue
        
        # Check for numbered question - ends instruction block
        if re.match(r"^\s*\d+\.", s):
            in_instruction_block = False
            continue
        
        # If we're in instruction block, capture everything
        if in_instruction_block:
            instruction_lines.append(s)
            continue
        
        # Check for standalone instruction patterns
        if any(re.search(pat, s, re.I) for pat in instruction_patterns):
            instruction_lines.append(s)
    
    # Deduplicate while preserving order
    deduped = []
    seen = set()
    for item in instruction_lines:
        clean = item.strip()
        if not clean:
            continue
        key = clean.lower()[:100]  # Limit key length
        if key in seen:
            continue
        seen.add(key)
        deduped.append(clean)
    
    return {
        "has_explicit_instructions": bool(deduped),
        "instruction_lines": deduped[:20],
        "instruction_count": len(deduped),
        "has_instruction_heading": instruction_heading_found,
        "mode": "explicit" if deduped else "none"
    }


def _extract_visual_slots(draft_md: str) -> List[Dict[str, Any]]:
    """
    Extract visual placeholder markers.

    DETECTS:
        - [[VISUAL: description]]
        - [[VISUAL ...attributes...]]
        - v1, v2, v3 tokens (on their own line)
        - prose references such as "See diagram below"
    """
    lines = (draft_md or "").splitlines()
    visual_slots = []

    def _clean(value) -> str:
        return str(value or "").strip()

    def _extract_attr(raw: str, attr_name: str) -> str:
        pattern = rf'{attr_name}\s*=\s*"([^"]*)"'
        m = re.search(pattern, raw, re.I)
        return m.group(1).strip() if m else ""

    for idx, line in enumerate(lines, start=1):
        s = line.strip()
        if not s:
            continue

        slot_type = None
        description = None
        slot = None

        # --------------------------------------------------
        # Pattern 1A: legacy [[VISUAL: description]]
        # --------------------------------------------------
        legacy_match = re.search(r"\[\[VISUAL\s*:\s*([^\]]*)\]\]", s, re.I)
        if legacy_match:
            desc = legacy_match.group(1).strip() or "Unnamed visual"
            slot_type = "visual_placeholder"
            description = desc
            slot = {
                "line_no": idx,
                "slot_type": slot_type,
                "description": description,
                "raw": s[:500],
            }

        # --------------------------------------------------
        # Pattern 1B: structured placeholder
        # e.g. [[VISUAL id="v1" kind="image" where="Q1" prompt="..." notes="..." data=""]]
        # --------------------------------------------------
        if slot is None:
            structured_match = re.search(r"\[\[VISUAL\b(.*?)\]\]", s, re.I)
            if structured_match:
                raw_attrs = structured_match.group(1) or ""

                visual_id = _extract_attr(raw_attrs, "id")
                kind = _extract_attr(raw_attrs, "kind")
                where = _extract_attr(raw_attrs, "where")
                prompt = _extract_attr(raw_attrs, "prompt")
                notes = _extract_attr(raw_attrs, "notes")
                data = _extract_attr(raw_attrs, "data")

                desc_parts = []
                if where:
                    desc_parts.append(where)
                if kind:
                    desc_parts.append(kind)
                if prompt:
                    desc_parts.append(prompt)

                description = " | ".join([p for p in desc_parts if p]) or visual_id or "Unnamed visual"

                slot = {
                    "line_no": idx,
                    "slot_type": "visual_placeholder",
                    "description": description,
                    "raw": s[:500],
                    "visual_id": visual_id,
                    "kind": kind,
                    "where": where,
                    "prompt": prompt,
                    "notes": notes,
                    "data": data,
                }

        # --------------------------------------------------
        # Pattern 2: v1, v2, v3 tokens (alone on line)
        # --------------------------------------------------
        if slot is None and re.match(r"^v\d+$", s.lower()):
            slot = {
                "line_no": idx,
                "slot_type": "diagram_token",
                "description": s,
                "raw": s[:500],
                "visual_id": s.lower(),
            }

        # --------------------------------------------------
        # Pattern 3: "See diagram below" references
        # --------------------------------------------------
        if slot is None and re.search(r"\b(see|refer to|as shown in)\s+(the\s+)?(diagram|figure|image|picture|graph|chart|illustration|map)(\s+below)?\b", s, re.I):
            ref_match = re.search(r"(diagram|figure|image|picture|graph|chart|illustration|map)", s, re.I)
            description = f"Reference to {ref_match.group(1) if ref_match else 'visual'}"
            slot = {
                "line_no": idx,
                "slot_type": "visual_reference",
                "description": description,
                "raw": s[:500],
            }

        # --------------------------------------------------
        # Pattern 4: "Label the diagram" instructions
        # --------------------------------------------------
        if slot is None and re.search(r"\b(label|complete|fill in|identify on)\b.*\b(diagram|figure|image|picture|graph)\b", s, re.I):
            slot = {
                "line_no": idx,
                "slot_type": "visual_instruction",
                "description": s[:200],
                "raw": s[:500],
            }

        if slot:
            visual_slots.append(slot)

    # Deduplicate by line number + strongest available identity
    seen = set()
    deduped = []
    for slot in visual_slots:
        key = (
            slot.get("line_no"),
            slot.get("visual_id", ""),
            slot.get("where", ""),
            slot.get("description", "")
        )
        if key not in seen:
            seen.add(key)
            deduped.append(slot)

    return deduped


def extract_draft_question_map(draft_md: str) -> Dict[str, Any]:
    """
    Parse draft markdown into question structure.
    
    INTELLIGENT ASPECTS:
        - Detects parent questions (1., 2., etc.)
        - Detects subparts (a), (b), (i), (ii)
        - Preserves markdown formatting within questions
        - Handles multi-line questions
        - Extracts marks labels [4 marks] or (4 marks)
    """
    if not draft_md:
        return {"questions": [], "counts": {}}
    
    draft_md = draft_md.replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.rstrip() for ln in draft_md.split("\n")]
    
    # Regex patterns
    q_start_re = re.compile(r"^\s*(\d+)\.\s*$")
    q_inline_re = re.compile(r"^\s*(\d+)\.\s+(.*)$")

    # Weak-format fallback:
    # supports lines like:
    #   "1 what is photosynthesis"
    #   "2 complete the table below"
    # but avoids matching proper dotted numbering already handled above.
    q_inline_weak_re = re.compile(r"^\s*(\d+)\s+(?!marks?\b)(.+)$", re.I)

    q_heading_re = re.compile(r"^\s*#{1,6}\s*(?:Question\s+|Q)(\d+)\s*[\.:]?\s*(.*)$", re.I)
    q_bold_heading_re = re.compile(r"^\s*\*\*\s*Question\s+(\d+)\s*[\.:]?\s*\*\*\s*(.*)$", re.I)
    q_plain_heading_re = re.compile(r"^\s*Question\s+(\d+)\s*[\.:]?\s*(.*)$", re.I)
    memo_heading_re = re.compile(r"^\s*#{1,6}\s*(answer key|memo|marking key)\b", re.I)
    subpart_re = re.compile(r"^\s*\(?([a-zA-Z])\)\s*(.*)$")
    bold_subpart_re = re.compile(r"^\s*\*\*\(?([a-zA-Z])\)\*\*\s*(.*)$")

    # Weak-format subparts:
    # supports lines like:
    #   "a name the gas absorbed"
    #   "b explain your answer"
    # but avoids capturing marks-only lines like "2 marks"
    weak_subpart_re = re.compile(r"^\s*([a-zA-Z])\s+(?!\d+\s*marks?\b)(.+)$", re.I)

    marks_re = re.compile(r"(\[[^\]]*marks?[^\]]*\]|\(\s*\d+\s*marks?\s*\)|\b\d+\s*marks?\b)", re.I)
    visual_token_re = re.compile(r"^\s*(v\d+|\[\[VISUAL.*\]\])\s*$", re.I)
    total_heading_re = re.compile(r"^\s*#{1,6}\s*Total\s*:\s*.*$", re.I)
    step_heading_re = re.compile(r"^\s*#{0,6}\s*Step\s+(\d+)\s*[:.\-–—]?\s*(.*)$", re.I)
    hr_re = re.compile(r"^\s*---+\s*$")
    
    def clean(text: str) -> str:
        return (text or "").strip()
    
    def _extract_marks_value(marks_label: str) -> Optional[int]:
        """
        Convert a marks label like:
            "[4 marks]" / "(10 marks)" / "6 marks"
        into an integer marks value.
        """
        s = clean(marks_label)
        if not s:
            return None

        m = re.search(r"(\d+)", s)
        if not m:
            return None

        try:
            return int(m.group(1))
        except Exception:
            return None

    def _infer_question_flags(question_obj: Dict[str, Any]) -> Dict[str, Any]:
        """
        Infer structural truth flags from the current question content.

        This is still rule-based for now, but it gives downstream functions
        explicit question-level truth instead of forcing them to rediscover it.
        """
        q_text = " ".join(question_obj.get("question_texts", []) or []).strip().lower()
        raw_text = " ".join(question_obj.get("raw_lines", []) or []).strip().lower()
        combined = f"{q_text} {raw_text}".strip()

        has_visual_slot = bool(question_obj.get("has_visual_slot"))
        visual_tokens = question_obj.get("visual_tokens", []) or []
        has_subparts = bool(question_obj.get("subparts"))

        diagram_reference = bool(
            re.search(r"\b(diagram|figure|picture|graph|chart|illustration|map)\b", combined)
        )

        image_reference = bool(
            re.search(r"\bimage\b", combined)
        )

        explicit_visual_action = bool(
            re.search(r"\b(label|identify|name|indicate|sketch)\b", combined)
        )

        explicit_draw_diagram = bool(
            re.search(r"\bdraw\b.*\b(diagram|graph|chart|figure|picture|map)\b", combined)
        )

        # Guard against false positives like:
        # - "draw a line to match..."
        # - "calculate the magnification of the image"
        matching_draw_line = bool(
            re.search(r"\bdraw a line to match\b", combined)
        )

        image_only_reference = image_reference and not (
            explicit_visual_action
            or diagram_reference
            or explicit_draw_diagram
        )

        requires_diagram = (
            has_visual_slot
            or bool(visual_tokens)
            or explicit_draw_diagram
            or ((diagram_reference or image_reference) and explicit_visual_action)
        )

        if matching_draw_line:
            requires_diagram = False

        if image_only_reference:
            requires_diagram = False

        # Distinguish table reference from actual table completion
        table_completion = bool(
            re.search(r"\b(complete the table|fill in the table|table below|complete the grid|fill in the grid|use the table above to complete)\b", combined)
        )
        table_reference = bool(
            re.search(r"\b(table|grid|data table)\b", combined)
        )
        requires_table = table_completion

        # Detect matching / pairing style questions
        matching_question = bool(
            re.search(r"\b(match|match the following|draw a line to match|pair|connect each|link each)\b", combined)
        )

        explanation_question = bool(
            re.search(r"\b(explain|describe|discuss|compare|contrast|justify|why|how)\b", combined)
        )

        calculation_question = bool(
            re.search(r"\b(calculate|compute|solve|find|determine|work out|simplify|expand|factorise)\b", combined)
        )

        evaluative_question = bool(
            re.search(r"\b(evaluate)\b", combined)
        )

        if has_subparts:
            answer_style = "subparts"
        elif matching_question:
            # Keep this as short_response for now, but do not misclassify as table_response
            answer_style = "short_response"
        elif requires_diagram and re.search(r"\b(label|identify|name|indicate)\b", combined):
            answer_style = "label_diagram"
        elif table_completion:
            answer_style = "table_response"
        elif calculation_question:
            answer_style = "show_working"
        elif explanation_question or (table_reference and evaluative_question):
            # e.g. "Use Table 1 to compare/explain..."
            answer_style = "paragraph_response"
        elif table_reference:
            # Mere table reference does not automatically mean table completion
            answer_style = "short_response"
        elif evaluative_question:
            answer_style = "paragraph_response"
        else:
            answer_style = "short_response"

        return {
            "requires_diagram": requires_diagram,
            "requires_table": requires_table,
            "answer_style": answer_style,
        }

    def new_question(q_no: str) -> Dict:
        return {
            "q_no": str(q_no),

            # Backward-compatible field still used elsewhere
            "marks_label": "",

            # New Phase-1 truth field
            "marks": None,

            "subparts": [],
            "question_texts": [],
            "raw_lines": [],
            "visual_tokens": [],
            "has_visual_slot": False,

            # New explicit structural truth fields
            "requires_diagram": False,
            "requires_table": False,
            "answer_style": "",

            # Existing compatibility field
            "layout_type": "paragraph"
        }
    
    questions = []
    current_q = None
    
    for line in lines:
        s = clean(line)
        if not s:
            continue

        # Stop parsing when memo / answer key begins
        if memo_heading_re.match(s):
            break

        # Stop parsing when total/footer heading begins
        if total_heading_re.match(s):
            break
        
        # Skip section headers
        if re.match(r"^\s*section\b", s, re.I):
            continue

        # Investigation / step-section heading
        m_step = step_heading_re.match(s)
        if m_step:
            if current_q:
                questions.append(current_q)

            step_no = clean(m_step.group(1))
            step_title = clean(m_step.group(2))

            current_q = new_question(f"Step {step_no}")
            current_q["raw_lines"].append(s)

            if step_title:
                current_q["question_texts"].append(step_title)

            current_q["structure_type"] = "step_section"
            current_q["section_label"] = f"Step {step_no}"
            current_q["section_title"] = step_title

            continue
        
        # Parent question number alone
        m_start = q_start_re.match(s)
        if m_start:
            # If already inside a heading-based question, treat bare numbered lines
            # as content/list items, not as brand-new parent questions.
            if current_q and current_q.get("raw_lines"):
                first_raw = (current_q["raw_lines"][0] or "").strip()
                if q_heading_re.match(first_raw):
                    current_q["question_texts"].append(s)
                    current_q["raw_lines"].append(s)
                    continue

            if current_q:
                questions.append(current_q)
            current_q = new_question(m_start.group(1))
            current_q["raw_lines"].append(s)
            continue
        
        # Inline parent question
        m_inline = q_inline_re.match(s)
        weak_mode = False

        if not m_inline:
            m_inline = q_inline_weak_re.match(s)
            weak_mode = bool(m_inline)

        if m_inline:
            # If already inside a heading-based question, treat numbered lines
            # like "1. ____" as content, not as new top-level questions.
            if current_q and current_q.get("raw_lines"):
                first_raw = (current_q["raw_lines"][0] or "").strip()
                if q_heading_re.match(first_raw):
                    current_q["raw_lines"].append(s)
                    current_q["question_texts"].append(s)
                    continue

            if current_q and current_q.get("structure_type") == "step_section":
                current_q["raw_lines"].append(s)
                current_q["question_texts"].append(s)
                continue

            if current_q:
                questions.append(current_q)

            q_no = m_inline.group(1)
            rest = clean(m_inline.group(2))
            current_q = new_question(q_no)
            current_q["raw_lines"].append(s)

            # Extract marks
            m_marks = marks_re.search(rest)
            if m_marks:
                current_q["marks_label"] = clean(m_marks.group(1))
                rest = rest.replace(m_marks.group(1), "").strip()

            if rest:
                # Check if rest is a subpart
                m_sub = subpart_re.match(rest)
                if not m_sub:
                    m_sub = bold_subpart_re.match(rest)
                if not m_sub:
                    m_sub = weak_subpart_re.match(rest)

                if m_sub:
                    label = clean(m_sub.group(1)).lower()
                    body = clean(m_sub.group(2))
                    current_q["subparts"].append(label)
                    current_q["question_texts"].append(f"({label}) {body}" if body else f"({label})")
                else:
                    current_q["question_texts"].append(rest)

            continue
        
        # Bold markdown question heading, e.g. "**Question 1.** Refer to the table above..."
        m_bold_heading = q_bold_heading_re.match(s)
        if m_bold_heading:
            if current_q:
                questions.append(current_q)

            q_no = m_bold_heading.group(1)
            rest = clean(m_bold_heading.group(2))
            current_q = new_question(q_no)
            current_q["raw_lines"].append(s)

            m_marks = marks_re.search(rest)
            if m_marks:
                current_q["marks_label"] = clean(m_marks.group(1))
                rest = rest.replace(m_marks.group(1), "").strip()

            rest = re.sub(r"^[\s:\-–—\.]+", "", rest).strip()

            if rest:
                current_q["question_texts"].append(rest)

            continue
        
        # Plain question heading, e.g. "Question 1." or "Question 1. Refer to the table above..."
        m_plain_heading = q_plain_heading_re.match(s)
        if m_plain_heading:
            if current_q:
                questions.append(current_q)

            q_no = m_plain_heading.group(1)
            rest = clean(m_plain_heading.group(2))
            current_q = new_question(q_no)
            current_q["raw_lines"].append(s)

            m_marks = marks_re.search(rest)
            if m_marks:
                current_q["marks_label"] = clean(m_marks.group(1))
                rest = rest.replace(m_marks.group(1), "").strip()

            rest = re.sub(r"^[\s:\-–—\.]+", "", rest).strip()

            if rest:
                current_q["question_texts"].append(rest)

            continue
        
        # Markdown heading question, e.g. "## Question 3 (10 marks): Title"
        m_heading = q_heading_re.match(s)
        if m_heading:
            if current_q:
                questions.append(current_q)

            q_no = m_heading.group(1)
            rest = clean(m_heading.group(2))
            current_q = new_question(q_no)
            current_q["raw_lines"].append(s)

            m_marks = marks_re.search(rest)
            if m_marks:
                current_q["marks_label"] = clean(m_marks.group(1))
                rest = rest.replace(m_marks.group(1), "").strip()

            # Remove leading punctuation like ":" or "-" after the heading
            rest = re.sub(r"^[\s:\-–—]+", "", rest).strip()

            if rest:
                current_q["question_texts"].append(rest)

            continue

        # Before first question, skip
        if current_q is None:
            continue

        # Ignore horizontal rules inside question parsing
        if hr_re.match(s):
            continue
        
        # Store raw line
        current_q["raw_lines"].append(s)
        
        # Visual token
        if visual_token_re.match(s):
            current_q["has_visual_slot"] = True
            current_q["visual_tokens"].append(s)
            continue
        
        # Marks-only line
        m_marks_only = marks_re.match(s)
        if m_marks_only:
            marks_text = clean(m_marks_only.group(1))

            # If parent question has no marks yet, use this as parent marks
            if not current_q.get("marks_label"):
                current_q["marks_label"] = marks_text

            # In all cases, do NOT treat standalone marks-only lines as question text
            continue
        
        # Subpart line
        m_sub = subpart_re.match(s)
        if not m_sub:
            m_sub = bold_subpart_re.match(s)
        if not m_sub:
            m_sub = weak_subpart_re.match(s)

        if m_sub:
            label = clean(m_sub.group(1)).lower()
            body = clean(m_sub.group(2))
            current_q["subparts"].append(label)
            current_q["question_texts"].append(f"({label}) {body}" if body else f"({label})")
            continue
        
        # Embedded marks
        if not current_q.get("marks_label"):
            m_marks_any = marks_re.search(s)
            if m_marks_any:
                current_q["marks_label"] = clean(m_marks_any.group(1))
        
        # Regular question text
        current_q["question_texts"].append(s)
    
    # Add last question
    if current_q:
        questions.append(current_q)
    
    # Clean up each question
    for q in questions:
        # Deduplicate text arrays
        for key in ["question_texts", "raw_lines", "visual_tokens", "subparts"]:
            if key in q:
                seen = set()
                deduped = []
                for item in q[key]:
                    if item not in seen:
                        seen.add(item)
                        deduped.append(item)
                q[key] = deduped

        # Finalize numeric marks field from marks_label
        q["marks"] = _extract_marks_value(q.get("marks_label", ""))

        # Finalize explicit structural truth flags
        inferred = _infer_question_flags(q)
        q["requires_diagram"] = inferred["requires_diagram"]
        q["requires_table"] = inferred["requires_table"]
        q["answer_style"] = inferred["answer_style"]

        # Keep existing compatibility layout_type logic, but strengthen it
        if q["subparts"]:
            q["layout_type"] = "subparts"
        elif q["requires_table"]:
            q["layout_type"] = "table"
        elif q["requires_diagram"] or q["has_visual_slot"]:
            q["layout_type"] = "visual"
        else:
            q["layout_type"] = "paragraph"
    
    return {
        "questions": questions,
        "counts": {
            "questions": len(questions),
            "with_marks": sum(1 for q in questions if q.get("marks_label")),
            "with_subparts": sum(1 for q in questions if q.get("subparts")),
            "with_visuals": sum(1 for q in questions if q.get("has_visual_slot"))
        }
    }


def _infer_answer_space_needs(draft_model: Dict[str, Any]) -> Dict[str, Any]:
    """
    Infer answer space style and length per question.
    
    INTELLIGENT ASPECTS:
        - Analyzes question verbs (explain → paragraph, calculate → working)
        - Considers marks allocation
        - Adapts to document type (exam vs worksheet)
    """
    question_map = draft_model.get("question_map", {})
    questions = question_map.get("questions", [])
    document_type = draft_model.get("document_type", DocumentType.CUSTOM)
    
    def infer_style(text: str, marks: int, has_subparts: bool) -> Tuple[str, int]:
        """Infer answer style and suggested lines from question text."""
        if has_subparts:
            return ("subparts", 0)
        
        lower_text = text.lower()
        
        # Label diagram questions
        if re.search(r"\b(label|identify|name the part|indicate)\b", lower_text):
            if re.search(r"\b(diagram|figure|image|picture|drawing)\b", lower_text):
                return ("label_diagram", 0)
        
        # Table response
        if re.search(r"\b(complete the table|fill in the table|table below)\b", lower_text):
            return ("table_response", 0)
        
        # Show working (calculations)
        if re.search(r"\b(calculate|compute|solve|find|determine|work out|simplify|expand|factorise)\b", lower_text):
            # More marks = more lines
            lines = 6 if marks < 4 else 8
            return ("show_working", lines)
        
        # Paragraph response (explain, describe, discuss)
        if re.search(r"\b(explain|describe|discuss|compare|contrast|justify|evaluate|why|how)\b", lower_text):
            lines = 6 if marks < 4 else 10
            return ("paragraph_response", lines)
        
        # Default short response
        lines = 2 if marks < 2 else (3 if marks < 4 else 4)
        return ("short_response", lines)
    
    per_question = []
    for q in questions:
        q_no = q.get("q_no", "")
        marks_label = q.get("marks_label", "")
        marks_num = q.get("marks", None)

        if marks_num is None:
            marks_num = 0

        # Combine question text
        q_text = " ".join(q.get("question_texts", []))
        has_subparts = bool(q.get("subparts"))

        # Prefer explicit parser truth first
        existing_style = (q.get("answer_style") or "").strip()
        if existing_style:
            style = existing_style

            if style == "subparts":
                lines = 0
            elif style == "label_diagram":
                lines = 0
            elif style == "table_response":
                lines = 0
            elif style == "show_working":
                lines = 6 if marks_num < 4 else 8
            elif style == "paragraph_response":
                lines = 6 if marks_num < 4 else 10
            else:
                lines = 2 if marks_num < 2 else (3 if marks_num < 4 else 4)
        else:
            # Rule‑based inference first
            style, lines = infer_style(q_text, marks_num, has_subparts)
            
            # Check for ambiguous questions that may need LLM refinement
            ambiguous_patterns = [
                r"\b(draw and label|sketch and label|diagram and label)\b",
                r"\b(evaluate|critique|assess)\b",
                r"\b(compare and contrast|discuss the similarities and differences)\b",
                r"\b(justify|explain why|explain how)\b.*\b(marks?.*[4-9]|10|12)\b",
                r"\b(design|plan|propose)\b",
            ]
            ambiguous = any(re.search(p, q_text, re.I) for p in ambiguous_patterns)
            marks_worth = marks_num >= 4
            
            if ambiguous or (marks_worth and len(q_text) > 150):
                try:
                    llm_result = _infer_answer_space_needs_with_llm(q_text, marks_num, document_type)
                    if llm_result.get("confidence", 0) > 0.7:
                        style = llm_result.get("answer_style", style)
                        lines = llm_result.get("suggested_lines", lines)
                        logger.info(f"Question {q_no}: LLM overrode answer style to {style} ({lines} lines) reason: {llm_result.get('reasoning', '')}")
                except Exception as e:
                    logger.warning(f"LLM answer space inference failed for Q{q_no}: {e}")

        per_question.append({
            "q_no": q_no,
            "marks_label": marks_label,
            "marks_value": marks_num,
            "answer_style": style,
            "suggested_lines": lines,
            "has_subparts": has_subparts,
            "requires_diagram": bool(q.get("requires_diagram", False)),
            "requires_table": bool(q.get("requires_table", False)),
        })
    
    return {
        "per_question": per_question,
        "summary": {
            "total_questions": len(per_question),
            "with_subparts": sum(1 for q in per_question if q["has_subparts"]),
            "paragraph_responses": sum(1 for q in per_question if q["answer_style"] == "paragraph_response"),
            "show_working": sum(1 for q in per_question if q["answer_style"] == "show_working"),
            "label_diagram": sum(1 for q in per_question if q["answer_style"] == "label_diagram"),
            "short_responses": sum(1 for q in per_question if q["answer_style"] == "short_response")
        }
    }


def _detect_cross_references(draft_md: str) -> List[Dict[str, Any]]:
    """
    Detect cross-references in draft that must be preserved.
    
    PATTERNS:
        - "See Question 3" → question reference
        - "Refer to diagram 2" → visual reference
        - "as shown in Table 1" → table reference
        - "see Appendix A" → section reference
    """
    if not draft_md:
        return []
    
    references = []
    
    # Question references
    q_patterns = [
        (r"\b(see|refer to|in)\s+[Qq]uestion\s+(\d+)\b", "question"),
        (r"\b[Qq]uestion\s+(\d+)\b", "question"),
        (r"\b(see|refer to)\s+Q(\d+)\b", "question"),
    ]
    
    for pattern, ref_type in q_patterns:
        for match in re.finditer(pattern, draft_md, re.I):
            ref_num = match.group(2) if len(match.groups()) > 1 else match.group(1)
            references.append({
                "type": ref_type,
                "target": str(ref_num),
                "text": match.group(0),
                "position": match.start()
            })
    
    # Diagram/figure references
    diagram_patterns = [
        (r"\b(see|refer to|in|as shown in)\s+(the\s+)?(diagram|figure|image|picture|graph|chart)\s+(\d+)\b", "visual"),
        (r"\b(diagram|figure|image|picture|graph|chart)\s+(\d+)\b", "visual"),
    ]
    
    for pattern, ref_type in diagram_patterns:
        for match in re.finditer(pattern, draft_md, re.I):
            ref_num = match.group(4) if len(match.groups()) > 3 else match.group(2)
            references.append({
                "type": ref_type,
                "target": str(ref_num),
                "text": match.group(0),
                "position": match.start()
            })
    
    # Table references
    table_patterns = [
        (r"\b(see|refer to|in|as shown in)\s+(the\s+)?table\s+(\d+)\b", "table"),
        (r"\btable\s+(\d+)\b", "table"),
    ]
    
    for pattern, ref_type in table_patterns:
        for match in re.finditer(pattern, draft_md, re.I):
            ref_num = match.group(3) if len(match.groups()) > 2 else match.group(1)
            references.append({
                "type": ref_type,
                "target": str(ref_num),
                "text": match.group(0),
                "position": match.start()
            })
    
    # Appendix references
    appendix_patterns = [
        (r"\b(see|refer to)\s+[Aa]ppendix\s+([A-Z])\b", "appendix"),
        (r"\b[Aa]ppendix\s+([A-Z])\b", "appendix"),
    ]
    
    for pattern, ref_type in appendix_patterns:
        for match in re.finditer(pattern, draft_md, re.I):
            ref_target = match.group(2) if len(match.groups()) > 1 else match.group(1)
            references.append({
                "type": ref_type,
                "target": ref_target,
                "text": match.group(0),
                "position": match.start()
            })
    
    # Deduplicate by position
    seen = set()
    unique_refs = []
    for ref in references:
        key = (ref["type"], ref["target"], ref["position"])
        if key not in seen:
            seen.add(key)
            unique_refs.append(ref)
    
    return unique_refs


def _detect_draft_page_breaks(draft_md: str) -> List[int]:
    """
    Detect explicit page break requests in draft.
    
    SYNTAX:
        - "---pagebreak---"
        - "<!-- pagebreak -->"
        - "\\newpage"
        - "===page==="
    """
    if not draft_md:
        return []
    
    page_breaks = []
    lines = draft_md.splitlines()
    
    patterns = [
        r"^---pagebreak---$",
        r"^<!--\s*pagebreak\s*-->$",
        r"^\\newpage$",
        r"^===page===$",
        r"^<div style=\"page-break-before: always\">$",
    ]
    
    for idx, line in enumerate(lines):
        line = line.strip().lower()
        for pattern in patterns:
            if re.match(pattern, line, re.I):
                page_breaks.append(idx + 1)  # 1-indexed line number
                break
    
    return page_breaks


def _detect_layout_zones_from_donor(template_profile: Dict[str, Any]) -> Dict[str, Any]:
    """
    Layout Zone Detection V1.

    Purpose:
    - convert donor evidence into logical layout zones
    - preserve donor layout relationships for later spatial reconstruction
    - do not copy donor academic content
    """

    template_profile = template_profile or {}

    normalized_profile = template_profile.get("normalized_profile", {}) or {}
    raw_profile = normalized_profile.get("raw", {}) if isinstance(normalized_profile, dict) else {}

    header_footer = template_profile.get("header_footer", {}) or raw_profile.get("header_footer", {}) or {}
    institution_identity = template_profile.get("institution_identity", {}) or {}
    interpreted_donor = template_profile.get("interpreted_donor", {}) or {}
    interpreted_identity = interpreted_donor.get("identity", {}) or {}
    interpreted_framing = interpreted_donor.get("framing", {}) or {}

    interpreted_images = interpreted_donor.get("images", {}) or {}

    extracted_images = []
    extracted_images.extend(interpreted_images.get("branding_images", []) or [])
    extracted_images.extend(interpreted_images.get("structural_images", []) or [])
    extracted_images.extend(raw_profile.get("extracted_images", []) or [])
    extracted_images.extend(template_profile.get("extracted_images", []) or [])

    # Deduplicate image evidence while preserving order.
    deduped_images = []
    seen_image_keys = set()

    for img in extracted_images:
        if not isinstance(img, dict):
            continue

        key = (
            str(img.get("paragraph_index")),
            str(img.get("position")),
            str(img.get("width_inches")),
            str(img.get("height_inches")),
            str(img.get("relationship_id")),
            str(img.get("source_bucket")),
        )

        if key in seen_image_keys:
            continue

        seen_image_keys.add(key)
        deduped_images.append(img)

    extracted_images = deduped_images

    extracted_tables = (
        raw_profile.get("extracted_tables", [])
        or template_profile.get("extracted_tables", [])
        or []
    )

    layout_zones = {
        "top_banner": [],
        "identity_block": [],
        "admin_fields": [],
        "title_block": [],
        "instruction_block": [],
        "content_body": [],
        "footer": [],
    }

    order_index = 1

    def _add(zone, element_type, role, text="", position="", alignment="", confidence=0.5, reason="", source="donor", extra=None):
        nonlocal order_index

        item = {
            "element_type": element_type,
            "role": role,
            "text": text or "",
            "position": position or "",
            "alignment": alignment or "",
            "order_index": order_index,
            "source": source,
            "confidence": float(confidence),
            "reason": reason or "",
        }

        if isinstance(extra, dict):
            item.update(extra)

        layout_zones[zone].append(item)
        order_index += 1

    # --------------------------------------------------
    # Header / top banner evidence
    # --------------------------------------------------
    for line in header_footer.get("header_texts", []) or []:
        line = str(line or "").strip()
        if not line:
            continue

        _add(
            zone="top_banner",
            element_type="header_text",
            role="running_header_text",
            text=line,
            position="header",
            alignment="unknown",
            confidence=0.75,
            reason="Extracted from donor header text."
        )

    # --------------------------------------------------
    # Branding / identity images
    # --------------------------------------------------
    for img in extracted_images:
        if not isinstance(img, dict):
            continue

        role = str(img.get("image_role") or img.get("interpreted_image_role") or "").strip().lower()
        position = str(img.get("position") or "").strip().lower()
        reason = str(img.get("classification_reason") or img.get("interpretation_reason") or "").strip()

        is_branding = (
            role == "branding"
            or position in {"header", "footer", "first_page"}
            or "logo" in reason.lower()
            or "branding" in reason.lower()
        )

        if not is_branding:
            continue

        zone = "identity_block"
        if position == "header":
            zone = "top_banner"
        elif position == "footer":
            zone = "footer"

        _add(
            zone=zone,
            element_type="image",
            role="primary_logo_candidate",
            text="",
            position=position or "unknown",
            alignment=str(img.get("alignment") or "unknown"),
            confidence=0.85,
            reason=reason or "Image classified as likely institutional branding.",
            extra={
                "width_inches": img.get("width_inches"),
                "height_inches": img.get("height_inches"),
                "image_role": role or "branding",
                "has_binary": bool(img.get("binary_data")),
            }
        )

    # --------------------------------------------------
    # Institution identity text
    # --------------------------------------------------
    identity_lines = []
    identity_lines.extend(institution_identity.get("top_block_lines", []) or [])
    identity_lines.extend(interpreted_identity.get("first_page_identity_lines", []) or [])
    identity_lines.extend(interpreted_identity.get("running_header_identity", []) or [])

    seen_identity = set()
    for line in identity_lines:
        line = str(line or "").strip()
        key = line.lower()
        if not line or key in seen_identity:
            continue
        seen_identity.add(key)

        _add(
            zone="identity_block",
            element_type="text",
            role="institution_identity_text",
            text=line,
            position="first_page",
            alignment="unknown",
            confidence=0.75,
            reason="Line identified as institutional identity evidence."
        )

    # --------------------------------------------------
    # Administrative fields
    # --------------------------------------------------
    field_lines = []
    field_lines.extend(institution_identity.get("field_lines", []) or [])
    field_lines.extend(interpreted_framing.get("candidate_fields", []) or [])

    seen_fields = set()
    for line in field_lines:
        line = str(line or "").strip()
        key = line.lower()
        if not line or key in seen_fields:
            continue
        seen_fields.add(key)

        _add(
            zone="admin_fields",
            element_type="text",
            role="admin_field",
            text=line,
            position="body",
            alignment="unknown",
            confidence=0.8,
            reason="Line identified as candidate/student/admin metadata field."
        )

    # --------------------------------------------------
    # Title / frame candidates
    # --------------------------------------------------
    title_lines = []
    title_lines.extend(institution_identity.get("title_candidates", []) or [])
    title_lines.extend(interpreted_framing.get("document_frame_lines", []) or [])

    seen_titles = set()
    for line in title_lines:
        line = str(line or "").strip()
        key = line.lower()
        if not line or key in seen_titles:
            continue
        seen_titles.add(key)

        _add(
            zone="title_block",
            element_type="text",
            role="donor_title_style_candidate",
            text=line,
            position="body",
            alignment="unknown",
            confidence=0.65,
            reason="Line identified as document framing/title-style evidence."
        )

    # --------------------------------------------------
    # Instruction evidence
    # --------------------------------------------------
    instruction_sources = []
    for key in ["generic_instructions", "timing_lines", "material_lines", "warning_lines"]:
        instruction_sources.extend(interpreted_framing.get(key, []) or [])

    seen_instructions = set()
    for line in instruction_sources:
        line = str(line or "").strip()
        key = line.lower()
        if not line or key in seen_instructions:
            continue
        seen_instructions.add(key)

        _add(
            zone="instruction_block",
            element_type="text",
            role="instruction_style_candidate",
            text=line,
            position="body",
            alignment="unknown",
            confidence=0.65,
            reason="Line identified as donor instruction/framing evidence."
        )

    # --------------------------------------------------
    # Structural tables
    # --------------------------------------------------
    for tbl in extracted_tables:
        if not isinstance(tbl, dict):
            continue

        table_role = str(tbl.get("table_role") or "").strip().lower()
        looks_like_field_table = bool(tbl.get("looks_like_field_table", False))
        likely_front_matter_table = bool(tbl.get("likely_front_matter_table", False))

        if table_role == "structural" or looks_like_field_table or likely_front_matter_table:
            _add(
                zone="admin_fields",
                element_type="table",
                role="admin_or_front_matter_table",
                text="",
                position="body",
                alignment=str(tbl.get("alignment") or "unknown"),
                confidence=0.7,
                reason="Table classified as structural/front-matter evidence.",
                extra={
                    "rows": tbl.get("rows"),
                    "cols": tbl.get("cols"),
                    "table_role": table_role or "structural",
                }
            )

    # --------------------------------------------------
    # Footer evidence
    # --------------------------------------------------
    for line in header_footer.get("footer_texts", []) or []:
        line = str(line or "").strip()
        if not line:
            continue

        _add(
            zone="footer",
            element_type="footer_text",
            role="running_footer_text",
            text=line,
            position="footer",
            alignment="unknown",
            confidence=0.75,
            reason="Extracted from donor footer text."
        )

    return {
        "schema_version": "layout_zones_v1",
        "zones": layout_zones,
        "zone_order": [
            "top_banner",
            "identity_block",
            "admin_fields",
            "title_block",
            "instruction_block",
            "content_body",
            "footer",
        ],
        "notes": [
            "Layout zones detected from donor evidence.",
            "This is V1 evidence capture for later spatial reconstruction.",
            "No donor academic content is treated as draft truth.",
        ],
    }


# ======================================================================================
# PHASE 5: BLUEPRINT COMPOSITION
# ======================================================================================

def build_document_blueprint(
    template_profile: Dict[str, Any],
    draft_model: Dict[str, Any],
    year_level: str = "",
    subject: str = "",
    appropriateness_result: Optional[AppropriatenessResult] = None,
) -> Dict[str, Any]:
    """
    Combine donor style + draft content + appropriateness decisions.

    BS5 RULE:
    - preserve backward compatibility for existing app / renderer consumers
    - make the Phase 1 contract blueprint authoritative at top level
    - keep donor style separate from draft content truth
    """
    template_profile = template_profile or {}
    draft_model = draft_model or {}

    # --------------------------------------------------
    # Core inputs
    # --------------------------------------------------
    layout_features = template_profile.get("layout_features", []) or []

    question_map = draft_model.get("question_map", {}) or {}
    questions = question_map.get("questions", []) or []

    user_hints = draft_model.get("user_layout_hints", {}) or {}

    education_level = draft_model.get("education_level", EducationLevel.UNKNOWN)
    document_type = draft_model.get("document_type", DocumentType.CUSTOM)

    if appropriateness_result is None:
        appropriateness_result = _apply_appropriateness_rules(
            layout_features=layout_features,
            draft_questions=questions,
            education_level=education_level,
            document_type=document_type,
            user_hints=user_hints,
            layout_needs=draft_model.get("layout_needs", {})
        )

    # --------------------------------------------------
    # Build BS5 blocks
    # --------------------------------------------------
    header_block = _build_header_block(template_profile, draft_model)
    title_block = _build_title_block(template_profile, draft_model, year_level, subject)
    instruction_block = _build_instruction_block(template_profile, draft_model)
    question_blocks = _build_question_blocks(template_profile, draft_model, appropriateness_result)
    optional_sections = _build_optional_teacher_sections(template_profile, draft_model)

    cross_references = draft_model.get("cross_references", []) or []
    visual_slots = draft_model.get("visual_slots", []) or []
    page_breaks = draft_model.get("page_breaks", []) or []
    layout_needs = draft_model.get("layout_needs", {}) or {}

    layout_zones = (
        template_profile.get("layout_zones", {})
        or (template_profile.get("normalized_profile", {}) or {}).get("layout", {}).get("layout_zones", {})
        or _detect_layout_zones_from_donor(template_profile)
    )

    normalized_profile = template_profile.get("normalized_profile", {}) or {}
    normalized_draft_model = draft_model.get("normalized_draft_model", {}) or {}

    # --------------------------------------------------
    # Safe normalized profile access
    # --------------------------------------------------
    normalized_page = normalized_profile.get("page", {}) if isinstance(normalized_profile, dict) else {}
    normalized_style = normalized_profile.get("style", {}) if isinstance(normalized_profile, dict) else {}
    raw_profile_core = normalized_profile.get("raw", {}) if isinstance(normalized_profile, dict) else {}

    # --------------------------------------------------
    # Select best logo candidate
    # --------------------------------------------------
    interpreted_donor = template_profile.get("interpreted_donor", {}) or {}
    interpreted_identity = interpreted_donor.get("identity", {}) or {}
    interpreted_images = interpreted_donor.get("images", {}) or {}

    # 🔴 FIX: Support BOTH locations
    branding_images = (
        interpreted_images.get("branding_images", [])
        or interpreted_identity.get("branding_images", [])
        or template_profile.get("interpreted_donor", {}).get("identity", {}).get("branding_images", [])
        or []
    )
    structural_images = (
        interpreted_images.get("structural_images", [])
        or interpreted_identity.get("structural_images", [])
        or []
    )
    extracted_images = (
        raw_profile_core.get("extracted_images", [])
        or template_profile.get("extracted_images", [])
        or []
    )

    logo_candidate = None

    def _ensure_binary(img):
        if not isinstance(img, dict):
            return None

        blob = img.get("binary_data")

        if isinstance(blob, bytes):
            return img

        if isinstance(blob, dict) and blob.get("__binary_base64__") and blob.get("data"):
            try:
                import base64
                img["binary_data"] = base64.b64decode(blob.get("data"))
                return img
            except Exception:
                pass

        if blob:
            return img

        path = img.get("path") or img.get("source_path")
        if path:
            try:
                with open(path, "rb") as f:
                    img["binary_data"] = f.read()
                    return img
            except Exception:
                pass

        return None

    logo_candidate = None

    for img in branding_images:
        if img.get("position") == "header":
            loaded = _ensure_binary(img)
            if loaded:
                logo_candidate = loaded
                break

    if logo_candidate is None:
        for img in branding_images:
            loaded = _ensure_binary(img)
            if loaded:
                logo_candidate = loaded
                break

    if logo_candidate is None:
        for img in structural_images:
            loaded = _ensure_binary(img)
            if loaded:
                logo_candidate = loaded
                break

    if logo_candidate is None:
        for img in extracted_images:
            loaded = _ensure_binary(img)
            if loaded:
                logo_candidate = loaded
                break

    # --------------------------------------------------
    # Core contract blocks (authoritative BS5 surface)
    # --------------------------------------------------
    document_type_value = document_type.value if hasattr(document_type, "value") else str(document_type)
    education_level_value = education_level.value if hasattr(education_level, "value") else str(education_level)

    identity_block = {
        "institution_name": header_block.get("institution_name", ""),
        "header_lines": header_block.get("header_lines", []),
        "field_lines": header_block.get("field_lines", []),
        "show_header": header_block.get("show_header", True),
        "logo_present": bool(logo_candidate),
    }

    layout_plan = {
        "layout_needs": layout_needs,
        "layout_zones": layout_zones,
        "appropriateness_decisions": {
            feature: {
                "apply": decision.apply,
                "confidence": decision.confidence,
                "reason": decision.reason,
                "source": decision.source,
                "requires_confirmation": decision.requires_teacher_confirmation,
            }
            for feature, decision in appropriateness_result.decisions.items()
        },
        "structure_analysis": appropriateness_result.structure_analysis,
        "visual_slots": visual_slots,
        "page_breaks": page_breaks,
        "cross_references": cross_references,
    }

    notes = [
        "Blueprint built from donor style and draft truth.",
        f"Questions: {len(question_blocks)}",
        f"LLM used in appropriateness: {appropriateness_result.llm_used}",
        f"Uncertainty flags: {len(appropriateness_result.uncertainty_flags)}",
    ]

    normalized_blueprint = {
        "schema_version": BLUEPRINT_SCHEMA_VERSION,
        "title": title_block.get("title", "Untitled Document"),
        "document_type": document_type_value,
        "identity_block": identity_block,
        "page_setup": normalized_page,
        "style_preferences": normalized_style,
        "instruction_block": instruction_block,
        "question_blocks": question_blocks,
        "layout_plan": layout_plan,
        "notes": notes,
    }

    # --------------------------------------------------
    # Final blueprint
    # Top level now satisfies the BS5 contract,
    # while legacy fields are preserved for compatibility.
    # --------------------------------------------------
    blueprint = {
        # ===== Phase 1 / BS5 authoritative contract surface =====
        "schema_version": BLUEPRINT_SCHEMA_VERSION,
        "title": normalized_blueprint["title"],
        "document_type": normalized_blueprint["document_type"],
        "identity_block": identity_block,
        "page_setup": normalized_page,
        "style_preferences": normalized_style,
        "instruction_block": instruction_block,
        "question_blocks": question_blocks,
        "layout_plan": layout_plan,
        "notes": notes,

        # ===== Existing runtime metadata =====
        "year_level": title_block.get("year_level", ""),
        "subject": title_block.get("subject", ""),
        "education_level": education_level_value,
        "uncertainty_flags": appropriateness_result.uncertainty_flags,
        "llm_used": appropriateness_result.llm_used,
        "created_at": datetime.now().isoformat(),

        # ===== Backward-compatible legacy fields =====
        "header_block": header_block,
        "title_block": title_block,
        "optional_teacher_sections": optional_sections,
        "appropriateness_decisions": {
            feature: {
                "apply": decision.apply,
                "confidence": decision.confidence,
                "reason": decision.reason,
                "source": decision.source,
                "requires_confirmation": decision.requires_teacher_confirmation
            }
            for feature, decision in appropriateness_result.decisions.items()
        },
        "structure_analysis": appropriateness_result.structure_analysis,
        "user_hints_applied": user_hints,
        "cross_references": cross_references,
        "visual_slots": visual_slots,
        "logo_candidate": logo_candidate,
        "page_breaks": page_breaks,

        # ===== Truth layers =====
        "normalized_profile": normalized_profile,
        "normalized_draft_model": normalized_draft_model,
        "layout_needs": layout_needs,

        # ===== Nested compatibility mirror =====
        "normalized_blueprint": normalized_blueprint,
    }

    logger.info(
        f"Blueprint built: {len(question_blocks)} questions, "
        f"LLM used: {appropriateness_result.llm_used}, "
        f"Uncertainties: {len(appropriateness_result.uncertainty_flags)}"
    )

    return blueprint


def debug_template_engine_snapshot(
    current_draft_md: str,
    template_bundle: str,
    year_level: str = "",
    subject: str = ""
) -> str:
    """
    Build a live BS5 debug snapshot without rendering.

    This is for diagnostics only.
    It captures the exact upstream values that drive the final output so the app
    can show what template_engine is actually doing in real time.
    """
    try:
        # --------------------------------------------------
        # 1. Unpack template bundle and normalize profile
        # --------------------------------------------------
        clean_md, template_profile = unpack_template_bundle(template_bundle)

        if (not template_profile) or (not isinstance(template_profile, dict)):
            try:
                direct = json.loads(template_bundle or "")
                if isinstance(direct, dict):
                    template_profile = direct
            except Exception:
                template_profile = {}

        normalized_profile = _normalize_template_profile(template_profile or {})

        # --------------------------------------------------
        # 2. Analyze draft + detector diagnostics
        # --------------------------------------------------
        raw_draft_text = current_draft_md or ""

        detected_doc_type, detected_doc_confidence, detected_doc_llm_used = _detect_document_type(raw_draft_text)
        rules_doc_type, rules_doc_confidence = _detect_document_type_rules(raw_draft_text)
        detector_front_slice = raw_draft_text[:2500]

        draft_model = analyze_draft_content(
            draft_md=raw_draft_text,
            year_level=year_level or "",
            subject=subject or "",
            education_level_hint=""
        )

        # --------------------------------------------------
        # 3. Build blueprint (authoritative BS5 surface)
        # --------------------------------------------------
        blueprint = build_document_blueprint(
            template_profile=normalized_profile,
            draft_model=draft_model,
            year_level=year_level or "",
            subject=subject or "",
            appropriateness_result=None
        )

        # --------------------------------------------------
        # 4. Pull useful live values
        # --------------------------------------------------
        title_block = blueprint.get("title_block", {}) or {}
        header_block = blueprint.get("header_block", {}) or {}
        instruction_block = blueprint.get("instruction_block", {}) or {}
        identity_block = blueprint.get("identity_block", {}) or {}
        layout_plan = blueprint.get("layout_plan", {}) or {}
        question_blocks = blueprint.get("question_blocks", []) or []

        interpreted_donor = normalized_profile.get("interpreted_donor", {}) or {}
        interpreted_identity = interpreted_donor.get("identity", {}) or {}
        interpreted_framing = interpreted_donor.get("framing", {}) or {}
        interpreted_images = interpreted_donor.get("images", {}) or {}

        normalized_core = normalized_profile.get("normalized_profile", {}) or {}
        raw_core = normalized_core.get("raw", {}) or {}

        def _logo_blob_status(img):
            if not isinstance(img, dict):
                return "not_dict"

            blob = img.get("binary_data")

            if isinstance(blob, bytes):
                return f"bytes:{len(blob)}"

            if isinstance(blob, str):
                return f"string:{blob[:40]}"

            if blob is None:
                return "none"

            return type(blob).__name__

        def _summarise_logo_images(items):
            out = []
            for img in (items or [])[:8]:
                if not isinstance(img, dict):
                    continue
                out.append({
                    "position": img.get("position"),
                    "image_role": img.get("image_role") or img.get("interpreted_image_role"),
                    "classification_reason": img.get("classification_reason") or img.get("interpretation_reason"),
                    "has_binary": bool(img.get("binary_data")),
                    "binary_status": _logo_blob_status(img),
                    "has_path": bool(img.get("path") or img.get("source_path")),
                    "width_inches": img.get("width_inches"),
                    "height_inches": img.get("height_inches"),
                    "paragraph_index": img.get("paragraph_index"),
                })
            return out

        draft_doc_type = draft_model.get("document_type", "")
        if hasattr(draft_doc_type, "value"):
            draft_doc_type = draft_doc_type.value

        meta_parts = []
        if title_block.get("year_level"):
            meta_parts.append(str(title_block.get("year_level")))
        if title_block.get("subject"):
            meta_parts.append(str(title_block.get("subject")))
        if title_block.get("document_type_display"):
            meta_parts.append(str(title_block.get("document_type_display")))

        meta_preview = " | ".join(meta_parts)

        question_preview = []
        for q in question_blocks[:5]:
            question_preview.append({
                "q_no": q.get("q_no"),
                "marks_label": q.get("marks_label"),
                "answer_style": q.get("answer_style"),
                "has_table_response": q.get("has_table_response"),
                "table_prompt": q.get("table_prompt"),
                "answer_space_plan": q.get("answer_space_plan", {}),
            })

        snapshot = {
            "status": "ok",
            "draft_model": {
                "document_type": draft_doc_type,
                "title": draft_model.get("title", ""),
                "year_level": draft_model.get("year_level", ""),
                "subject": draft_model.get("subject", ""),
                "education_level": (
                    draft_model.get("education_level").value
                    if hasattr(draft_model.get("education_level"), "value")
                    else str(draft_model.get("education_level", ""))
                ),
                "instruction_block": draft_model.get("instruction_block", {}),
                "question_count": draft_model.get("question_count", 0),
                "uncertainties": draft_model.get("uncertainties", []),
            },
            "document_type_diagnostics": {
                "rules_only_type": (
                    rules_doc_type.value if hasattr(rules_doc_type, "value") else str(rules_doc_type)
                ),
                "rules_only_confidence": rules_doc_confidence,
                "final_type": (
                    detected_doc_type.value if hasattr(detected_doc_type, "value") else str(detected_doc_type)
                ),
                "final_confidence": detected_doc_confidence,
                "llm_used": detected_doc_llm_used,
                "front_slice_preview": detector_front_slice[:1200],
            },
            "title_block": title_block,
            "header_block": header_block,
            "instruction_block": instruction_block,
            "identity_block": identity_block,
            "blueprint_top_level": {
                "title": blueprint.get("title", ""),
                "document_type": blueprint.get("document_type", ""),
                "year_level": blueprint.get("year_level", ""),
                "subject": blueprint.get("subject", ""),
                "education_level": blueprint.get("education_level", ""),
                "meta_preview": meta_preview,
                "notes": blueprint.get("notes", []),
                "uncertainty_flags": blueprint.get("uncertainty_flags", []),
            },
            "donor_summary": {
                "source_filename": normalized_profile.get("source_meta", {}).get("source_filename", ""),
                "likely_institution_name": normalized_profile.get("institution_identity", {}).get("likely_institution_name", ""),
                "interpreted_document_type": interpreted_donor.get("document_type", ""),
                "institution_name": interpreted_identity.get("institution_name", ""),
                "running_header_identity": interpreted_identity.get("running_header_identity", []),
                "generic_instructions": interpreted_framing.get("generic_instructions", []),
                "timing_lines": interpreted_framing.get("timing_lines", []),
                "material_lines": interpreted_framing.get("material_lines", []),
                "warning_lines": interpreted_framing.get("warning_lines", []),
            },
            "logo_diagnostics": {
                "identity_logo_present": interpreted_identity.get("logo_present", False),
                "identity_logo_positions": interpreted_identity.get("logo_positions", []),
                "blueprint_logo_present": bool(blueprint.get("logo_candidate")),
                "branding_images_count": len(interpreted_images.get("branding_images", []) or []),
                "structural_images_count": len(interpreted_images.get("structural_images", []) or []),
                "educational_images_count": len(interpreted_images.get("educational_images", []) or []),
                "uncertain_images_count": len(interpreted_images.get("uncertain_images", []) or []),
                "raw_extracted_images_count": len(raw_core.get("extracted_images", []) or []),
                "raw_suppressed_images_count": len(raw_core.get("suppressed_images", []) or []),
                "branding_images": _summarise_logo_images(interpreted_images.get("branding_images", []) or []),
                "structural_images": _summarise_logo_images(interpreted_images.get("structural_images", []) or []),
                "raw_extracted_images": _summarise_logo_images(raw_core.get("extracted_images", []) or []),
                "logo_candidate_status": _logo_blob_status(blueprint.get("logo_candidate")),
            },
            "layout_plan": layout_plan,
            "question_preview": question_preview,
        }

        return json.dumps(snapshot, ensure_ascii=False, indent=2, default=str)

    except Exception as e:
        return json.dumps({
            "status": "error",
            "error_type": type(e).__name__,
            "error_message": str(e),
        }, ensure_ascii=False, indent=2, default=str)


def _build_header_block(template_profile: Dict, draft_model: Dict) -> Dict:
    """
    Build header/front-matter block from donor institution signals.

    BS5 RULE:
    - running header should remain donor-style, but draft-truth aware
    - institution identity comes from donor
    - year/subject/document type come from draft truth
    - first-page front matter is separate from running header
    """
    template_profile = template_profile or {}
    draft_model = draft_model or {}

    interpreted_donor = template_profile.get("interpreted_donor", {}) or {}
    interpreted_identity = interpreted_donor.get("identity", {}) or {}
    interpreted_framing = interpreted_donor.get("framing", {}) or {}

    institution_identity = template_profile.get("institution_identity", {}) or {}
    header_footer = template_profile.get("header_footer", {}) or {}

    normalized_profile = template_profile.get("normalized_profile", {}) or {}
    normalized_style = normalized_profile.get("style", {}) if isinstance(normalized_profile, dict) else {}
    legacy_style_prefs = template_profile.get("style_preferences", {}) or {}

    def _clean_text(value) -> str:
        return str(value or "").replace("\xa0", " ").strip()

    institution_name = (
        interpreted_identity.get("institution_name")
        or institution_identity.get("likely_institution_name", "")
        or ""
    )

    school_candidates = (
        interpreted_identity.get("school_candidates", [])
        or institution_identity.get("school_candidates", [])
        or []
    )

    first_page_identity_lines = (
        interpreted_identity.get("first_page_identity_lines", [])
        or institution_identity.get("top_block_lines", [])
        or []
    )

    running_header_identity = (
        interpreted_identity.get("running_header_identity", [])
        or institution_identity.get("header_identity_lines", [])
        or []
    )

    running_footer_identity = (
        interpreted_identity.get("running_footer_identity", [])
        or institution_identity.get("footer_identity_lines", [])
        or []
    )

    field_lines = (
        interpreted_framing.get("candidate_fields", [])
        or institution_identity.get("field_lines", [])
        or []
    )

    first_page_header_labels = interpreted_identity.get("first_page_header_labels", []) or []
    running_header_labels = interpreted_identity.get("running_header_labels", []) or []
    official_labels = interpreted_identity.get("official_labels", []) or []

    if not running_header_labels and official_labels:
        running_header_labels = official_labels

    if not first_page_header_labels and official_labels:
        first_page_header_labels = official_labels

    header_texts = header_footer.get("header_texts", []) or []
    header_candidates = header_footer.get("header_candidates", []) or []
    header_paragraphs = header_footer.get("header_paragraphs", []) or []
    footer_paragraphs = header_footer.get("footer_paragraphs", []) or []

    def _infer_dominant_alignment(paragraphs, default="CENTER") -> str:
        counts = {}
        for para in paragraphs:
            if not isinstance(para, dict):
                continue
            alignment = str(para.get("alignment", "") or "").strip().upper()
            if not alignment or alignment == "NONE":
                continue
            counts[alignment] = counts.get(alignment, 0) + 1

        if not counts:
            return default

        ranked = sorted(counts.items(), key=lambda x: (-x[1], x[0]))
        return ranked[0][0]

    running_header_alignment = _infer_dominant_alignment(header_paragraphs, default="CENTER")
    running_footer_alignment = _infer_dominant_alignment(footer_paragraphs, default="CENTER")

    def _is_safe_institution_name(value: str) -> bool:
        s = _clean_text(value)
        low = s.lower()

        if not s:
            return False

        if "centertop" in low:
            return False

        if "official" in low and s.upper() != "OFFICIAL":
            return False

        if len(s) > 40:
            return False

        if "\t" in s:
            return False

        if "|" in s:
            return False

        return True

    safe_institution_name = _clean_text(institution_name)
    if not _is_safe_institution_name(safe_institution_name):
        safe_institution_name = ""
    
    ordered_lines = []
    seen = set()

    def _is_bad_header_line(line: str) -> bool:
        """
        Header output filter.

        This is intentionally local to header rendering:
        - it blocks mixed donor header strings;
        - it blocks obvious XML/OCR/textbox artefacts;
        - it does not hardcode any school name;
        - it allows clean institution identity and draft-built metadata.
        """
        s = _clean_text(line)
        if not s:
            return True

        low = s.lower()

        if "centertop" in low:
            return True

        if "\t" in s:
            return True

        if "  " in s:
            return True

        if len(s) > 60:
            return True

        if "|" in s:
            return False

        if re.search(r"\b(year|grade)\s+\d{1,2}\b", low):
            return True

        if re.search(r"\b(exam|test|worksheet|assignment|semester|section)\b", low):
            return True

        return False

    def _normalise_header_line(line: str) -> str:
        s = _clean_text(line)

        if s.upper() == "OFFICIAL OFFICIAL":
            return "OFFICIAL"

        return s

    def push(line):
        s = _normalise_header_line(line)

        if _is_bad_header_line(s):
            return

        key = s.lower()
        if key in seen:
            return

        seen.add(key)
        ordered_lines.append(s)

    # Donor institution identity for running header
    push(safe_institution_name)

    for line in school_candidates[:3]:
        push(line)

    for line in running_header_identity[:3]:
        push(line)

    for line in running_footer_identity[:2]:
        push(line)

    # Running-header administrative labels such as OFFICIAL
    for line in running_header_labels[:2]:
        push(line)

    # Draft-truth running metadata line
    year_level = _clean_text(draft_model.get("year_level", ""))
    subject = _clean_text(draft_model.get("subject", ""))

    if not year_level:
        title_text = _clean_text(draft_model.get("title", ""))
        year_match = re.search(r"\b(year|grade|yr)\s+(\d+)\b", title_text, re.I)
        if year_match:
            year_level = f"Year {year_match.group(2)}"

    doc_type = draft_model.get("document_type")
    if hasattr(doc_type, "value"):
        document_type_display = str(doc_type.value).title()
    else:
        document_type_display = (
            _clean_text(draft_model.get("document_type_display", ""))
            or _clean_text(doc_type).title()
        )

    meta_parts = []
    if year_level:
        meta_parts.append(year_level)
    if subject:
        meta_parts.append(subject)
    if document_type_display and document_type_display != "Custom":
        meta_parts.append(document_type_display)

    running_meta_line = " | ".join(meta_parts).strip()
    if running_meta_line:
        push(running_meta_line)

    if not ordered_lines:
        for line in header_candidates[:3]:
            push(line)
        for line in header_texts[:3]:
            push(line)

    cleaned_field_lines = []
    seen_fields = set()
    for line in field_lines[:8]:
        s = _clean_text(line)
        if not s:
            continue
        key = s.lower()
        if key in seen_fields:
            continue
        seen_fields.add(key)
        cleaned_field_lines.append(s)

    header_pattern = {
        "pattern_version": "header_pattern_v1",
        "confidence": 0.0,
        "first_page": {
            "admin_labels": [
                _clean_text(x) for x in first_page_header_labels[:8] if _clean_text(x)
            ],
            "has_distinct_first_page_header": bool(first_page_header_labels),
        },
        "running": {
            "admin_labels": [
                _clean_text(x) for x in running_header_labels[:8] if _clean_text(x)
            ],
            "layout": "three_part" if len(meta_parts) >= 2 and safe_institution_name else "simple",
            "left": " ".join(meta_parts[:2]).strip() if len(meta_parts) >= 2 else running_meta_line,
            "center": safe_institution_name,
            "right": document_type_display if document_type_display and document_type_display != "Custom" else "",
            "has_underline": True,
            "source": "draft_truth_mapped_to_donor_running_header",
        },
    }

    if header_pattern["running"]["layout"] == "three_part":
        header_pattern["confidence"] = 0.85
    elif header_pattern["running"]["left"] or header_pattern["running"]["center"]:
        header_pattern["confidence"] = 0.65

    show_header = bool(
        ordered_lines
        or cleaned_field_lines
        or header_footer.get("has_header", False)
    )

    return {
        "show_header": show_header,
        "institution_name": safe_institution_name,
        "header_lines": ordered_lines[:12],
        "header_pattern": header_pattern,
        "field_lines": cleaned_field_lines,
        "first_page_header_labels": [
            _clean_text(x) for x in first_page_header_labels[:8] if _clean_text(x)
        ],
        "running_header_labels": [
            _clean_text(x) for x in running_header_labels[:8] if _clean_text(x)
        ],
        "first_page_identity_lines": [
            _clean_text(x) for x in first_page_identity_lines[:8] if _clean_text(x)
        ],
        "running_header_alignment": running_header_alignment,
        "running_footer_alignment": running_footer_alignment,
        "font_family": (
            normalized_style.get("font_family")
            or legacy_style_prefs.get("font_family")
            or "Calibri"
        ),
        "font_size_pt": (
            normalized_style.get("font_size_pt")
            or legacy_style_prefs.get("font_size_pt")
            or 11
        ),
    }


def _build_title_block(
    template_profile: Dict,
    draft_model: Dict,
    year_level: str = "",
    subject: str = ""
) -> Dict:
    """
    Build title block from draft truth.

    BS5 RULE:
    - explicit teacher/user inputs win
    - draft model truth comes next
    - conservative inference is allowed for year level only
    - subject must not be guessed from hardcoded keyword lists
    """
    draft_model = draft_model or {}

    def _clean_text(value) -> str:
        return str(value or "").strip()

    title = _clean_text(draft_model.get("title", "")) or "Untitled Document"
    doc_type = draft_model.get("document_type", DocumentType.CUSTOM)

    # --------------------------------------------------
    # 1. Explicit overrides win
    # --------------------------------------------------
    final_year = _clean_text(year_level) or _clean_text(draft_model.get("year_level", ""))
    final_subject = _clean_text(subject) or _clean_text(draft_model.get("subject", ""))

    # --------------------------------------------------
    # 2. Conservative year inference from title only
    # --------------------------------------------------
    if not final_year:
        title_lower = title.lower()

        year_patterns = [
            r"\byear\s+(\d+)\b",
            r"\bgrade\s+(\d+)\b",
            r"\byr\s+(\d+)\b",
        ]

        for pattern in year_patterns:
            match = re.search(pattern, title_lower)
            if match:
                final_year = f"Year {match.group(1)}"
                break

    # --------------------------------------------------
    # 3. Conservative subject fallback from existing structured draft metadata only
    #    Do NOT guess from hardcoded title keywords.
    # --------------------------------------------------
    if not final_subject:
        metadata_candidates = []

        draft_metadata = draft_model.get("metadata", {}) or {}
        if isinstance(draft_metadata, dict):
            metadata_candidates.extend([
                draft_metadata.get("subject", ""),
                draft_metadata.get("learning_area", ""),
                draft_metadata.get("course", ""),
            ])

        question_map = draft_model.get("question_map", {}) or {}
        if isinstance(question_map, dict):
            metadata_candidates.append(question_map.get("subject", ""))

        for candidate in metadata_candidates:
            cleaned = _clean_text(candidate)
            if cleaned:
                final_subject = cleaned
                break

    # --------------------------------------------------
    # 4. Display formatting
    # --------------------------------------------------
    doc_type_display = doc_type.value.title() if hasattr(doc_type, "value") else str(doc_type).title()

    return {
        "title": title,
        "document_type": doc_type,
        "document_type_display": doc_type_display,
        "year_level": final_year,
        "subject": final_subject,
        "has_year": bool(final_year),
        "has_subject": bool(final_subject)
    }


def _build_instruction_block(template_profile: Dict, draft_model: Dict) -> Dict:
    """
    Build instruction block from draft truth first, then safe donor framing fallback.

    BS5 RULE:
    - draft instructions always win
    - donor framing may only survive if it is document-type compatible with the draft
    - exam/test-specific donor framing must not leak into worksheet/custom outputs
    - when in doubt, prefer placeholder over donor contamination
    """
    template_profile = template_profile or {}
    draft_model = draft_model or {}

    interpreted_donor = template_profile.get("interpreted_donor", {}) or {}
    interpreted_framing = interpreted_donor.get("framing", {}) or {}
    donor_document_type = str(interpreted_donor.get("document_type", "") or "").strip().lower()

    draft_document_type = draft_model.get("document_type", DocumentType.CUSTOM)
    draft_document_type_value = (
        draft_document_type.value.lower()
        if hasattr(draft_document_type, "value")
        else str(draft_document_type).strip().lower()
    )

    def _clean_line(text) -> str:
        return str(text or "").replace("\xa0", " ").strip()

    def _dedupe_keep_order(items):
        seen = set()
        out = []
        for item in items or []:
            s = _clean_line(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    def _is_exam_specific_line(text: str) -> bool:
        s = _clean_line(text).lower()
        if not s:
            return False

        exam_terms = [
            "question/answer booklet",
            "question answer booklet",
            "calculator assumed",
            "calculator free",
            "time allowed",
            "reading time",
            "working time",
            "important note to students",
            "to be provided by the supervisor",
            "to be provided by the student",
            "material required/recommended",
            "all questions must be answered in both sections",
            "section a",
            "section b",
            "use black pen",
            "total marks",
            "exam comprises",
            "during reading time",
            "supervisor",
        ]
        return any(term in s for term in exam_terms)

    def _is_generic_reusable_line(text: str) -> bool:
        s = _clean_line(text).lower()
        if not s:
            return False

        generic_terms = [
            "write your answers",
            "read the following",
            "show all working",
            "answer all questions",
            "round your answers",
            "show your working",
        ]
        return any(term in s for term in generic_terms)

    # --------------------------------------------------
    # 1. Draft explicit instructions always win
    # --------------------------------------------------
    draft_instruction_block = draft_model.get("instruction_block", {}) or {}
    explicit_instructions = draft_instruction_block.get("instruction_lines", []) or []
    cleaned_explicit = _dedupe_keep_order(explicit_instructions)

    if cleaned_explicit:
        return {
            "mode": "explicit",
            "instruction_lines": cleaned_explicit,
            "has_instructions": True,
            "source": "draft"
        }

    # --------------------------------------------------
    # 2. Draft implicit signals come next
    # --------------------------------------------------
    raw_lines = draft_model.get("raw_lines", []) or []
    lower_blob = " ".join([str(x) for x in raw_lines]).lower()

    implicit_signals = [
        "answer all questions",
        "show all working",
        "time allowed",
        "do not use",
        "calculator",
        "read the following",
        "write your answers",
        "use black pen",
        "total marks",
        "section a"
    ]

    has_implicit = any(signal in lower_blob for signal in implicit_signals)

    if has_implicit:
        return {
            "mode": "implicit",
            "instruction_lines": [],
            "has_instructions": True,
            "source": "implicit",
            "note": "Instructions implied by draft language"
        }

    # --------------------------------------------------
    # 3. Safe donor framing fallback
    # Document-type compatibility filter
    # --------------------------------------------------
    donor_instruction_lines = []
    donor_instruction_lines.extend(interpreted_framing.get("generic_instructions", []) or [])
    donor_instruction_lines.extend(interpreted_framing.get("timing_lines", []) or [])
    donor_instruction_lines.extend(interpreted_framing.get("material_lines", []) or [])
    donor_instruction_lines.extend(interpreted_framing.get("warning_lines", []) or [])

    donor_instruction_lines = _dedupe_keep_order(donor_instruction_lines)

    compatible_lines = []

    # Case A: same or compatible assessment type -> allow more donor framing
    same_type = bool(donor_document_type) and donor_document_type == draft_document_type_value
    donor_is_exam_family = donor_document_type in {"exam", "test"}
    draft_is_exam_family = draft_document_type_value in {"exam", "test", "memo", "rubric"}

    for line in donor_instruction_lines:
        if same_type:
            compatible_lines.append(line)
            continue

        if donor_is_exam_family and not draft_is_exam_family:
            # Strong suppression for worksheet/custom/assignment-style drafts
            if _is_generic_reusable_line(line):
                compatible_lines.append(line)
            continue

        # Default conservative path:
        # only keep clearly generic reusable instruction lines
        if _is_generic_reusable_line(line):
            compatible_lines.append(line)

    compatible_lines = _dedupe_keep_order(compatible_lines)

    if compatible_lines:
        return {
            "mode": "donor_framing_fallback",
            "instruction_lines": compatible_lines[:6],
            "has_instructions": True,
            "source": "interpreted_donor",
            "note": "Recovered from donor framing after document-type compatibility filtering; teacher may keep, edit, or remove",
            "removable": True
        }

    # --------------------------------------------------
    # 4. No safe instructions found - provide teacher placeholder
    # --------------------------------------------------
    return {
        "mode": "placeholder",
        "instruction_lines": [
            "[Optional instructions for teacher to keep, edit, or remove]"
        ],
        "has_instructions": False,
        "source": "placeholder",
        "removable": True
    }


def _build_question_blocks(
    template_profile: Dict,
    draft_model: Dict,
    appropriateness_result: AppropriatenessResult
) -> List[Dict]:
    """
    Build question blocks as blueprint structures, not rendered output.

    BS5 RULE:
    - draft question content is authoritative
    - donor may influence presentation grammar only
    - answer-space planning must be explicit and deterministic
    """
    template_profile = template_profile or {}
    draft_model = draft_model or {}

    question_map = draft_model.get("question_map", {}) or {}
    questions = question_map.get("questions", []) or []

    answer_needs = draft_model.get("answer_space_needs", {}) or {}
    per_question_answer_needs = answer_needs.get("per_question", []) or []
    answer_by_q = {
        str(item.get("q_no", "")).strip(): item
        for item in per_question_answer_needs
        if isinstance(item, dict)
    }

    interpreted_donor = template_profile.get("interpreted_donor", {}) or {}
    interpreted_structure = interpreted_donor.get("structure", {}) or {}

    interpreted_answer_space_style = interpreted_structure.get("answer_space_style", "plain_response_space")
    donor_question_heading_style = interpreted_structure.get("question_heading_style", "unknown")
    donor_numbering_style = interpreted_structure.get("numbering_style", "unknown")

    # --------------------------------------------------
    # Determine donor answer-line style
    # --------------------------------------------------
    answer_line_style = "underscore"

    has_donor_answer_lines_feature = False
    for feature in template_profile.get("layout_features", []) or []:
        if not isinstance(feature, dict):
            continue
        if feature.get("feature_name") == "answer_lines":
            has_donor_answer_lines_feature = True
            extracted_value = feature.get("extracted_value", {}) or {}
            answer_line_style = extracted_value.get("style", answer_line_style)
            break

    if interpreted_answer_space_style == "answer_lines":
        has_donor_answer_lines_feature = True

    # --------------------------------------------------
    # Determine whether answer lines should actually apply
    # Conservative rule:
    # - explicit appropriateness decision wins
    # - otherwise only allow if donor actually has answer lines
    # - default must not blindly leak donor styling
    # --------------------------------------------------
    decisions = getattr(appropriateness_result, "decisions", {}) or {}

    if "answer_lines" in decisions:
        apply_answer_lines = bool(decisions["answer_lines"].apply)
    else:
        # Design Contract safety rule:
        # donor answer-line styling must not apply merely because the donor had it.
        # If the appropriateness engine did not explicitly approve answer lines,
        # suppress them and let the draft answer-space plan control the output.
        apply_answer_lines = False

    def _clean_line(text) -> str:
        s = str(text or "").strip()
        if not s:
            return ""
        if re.match(r"^\s*---+\s*$", s):
            return ""
        return s

    def _build_answer_space_plan(base_answer_style: str, base_suggested_lines: int) -> Dict[str, Any]:
        safe_lines = max(0, _try_int(base_suggested_lines, 0))

        return {
            "apply_answer_lines": apply_answer_lines,
            "answer_line_style": answer_line_style if apply_answer_lines else "none",
            "suggested_lines": safe_lines,
            "answer_style": base_answer_style,
            "answer_space_style": interpreted_answer_space_style,
        }

    def _infer_subpart_answer_plan(parent_answer_style: str) -> Tuple[str, int]:
        """
        Derive subpart answer planning conservatively from parent draft truth.
        We do NOT have independent per-subpart answer inference yet,
        so this is a structure-safe bridge, not a hardcoded donor assumption.
        """
        parent_answer_style = str(parent_answer_style or "").strip().lower()

        if parent_answer_style == "show_working":
            return ("show_working", 4)

        if parent_answer_style == "paragraph_response":
            return ("paragraph_response", 4)

        if parent_answer_style == "label_diagram":
            return ("label_diagram", 0)

        if parent_answer_style == "table_response":
            return ("table_response", 0)

        if parent_answer_style == "subparts":
            return ("short_response", 2)

        return ("short_response", 2)

    blocks = []

    for q in questions:
        if not isinstance(q, dict):
            continue

        q_no = str(q.get("q_no", "?")).strip() or "?"
        marks_label = str(q.get("marks_label", "")).strip()

        raw_question_texts = q.get("question_texts", []) or []
        question_texts = [_clean_line(text) for text in raw_question_texts]
        question_texts = [text for text in question_texts if text]

        subparts = q.get("subparts", []) or []
        has_visual_slot = bool(q.get("has_visual_slot", False))
        visual_tokens = q.get("visual_tokens", []) or []

        # Draft-side answer truth for this question
        answer_info = answer_by_q.get(q_no, {}) or {}
        parent_answer_style = str(answer_info.get("answer_style", "short_response")).strip() or "short_response"
        suggested_lines = max(0, _try_int(answer_info.get("suggested_lines", 3), 3))
        marks_value = max(0, _try_int(answer_info.get("marks_value", 0), 0))

        requires_diagram = bool(
            answer_info.get("requires_diagram", False)
            or q.get("requires_diagram", False)
        )
        requires_table = bool(
            answer_info.get("requires_table", False)
            or q.get("requires_table", False)
        )

        subpart_blocks = []
        main_question_lines = []

        if subparts:
            subpart_pattern = re.compile(r"^\(([a-zA-Z])\)\s*(.*)$")
            current_subpart = None

            for text in question_texts:
                match = subpart_pattern.match(text)
                if match:
                    label = match.group(1).lower()
                    body = match.group(2).strip()

                    subpart_answer_style, subpart_suggested_lines = _infer_subpart_answer_plan(parent_answer_style)

                    current_subpart = {
                        "label": label,
                        "body": body,
                        "answer_style": subpart_answer_style,
                        "suggested_lines": subpart_suggested_lines,
                        "numbering_style": donor_numbering_style,
                        "question_heading_style": donor_question_heading_style,
                        "answer_space_plan": _build_answer_space_plan(
                            subpart_answer_style,
                            subpart_suggested_lines
                        ),
                    }
                    subpart_blocks.append(current_subpart)
                    continue

                if current_subpart is None:
                    main_question_lines.append(text)
                else:
                    existing = (current_subpart.get("body") or "").strip()
                    current_subpart["body"] = f"{existing}\n{text}".strip() if existing else text

            # Fallback: structural subparts exist but explicit subpart lines did not survive parsing
            if not subpart_blocks and subparts:
                for label in subparts:
                    label_text = str(label or "").strip().lower()
                    if not label_text:
                        continue

                    subpart_answer_style, subpart_suggested_lines = _infer_subpart_answer_plan(parent_answer_style)

                    subpart_blocks.append({
                        "label": label_text,
                        "body": "",
                        "answer_style": subpart_answer_style,
                        "suggested_lines": subpart_suggested_lines,
                        "numbering_style": donor_numbering_style,
                        "question_heading_style": donor_question_heading_style,
                        "answer_space_plan": _build_answer_space_plan(
                            subpart_answer_style,
                            subpart_suggested_lines
                        ),
                    })
        else:
            main_question_lines = question_texts

        final_answer_style = "subparts" if subpart_blocks else parent_answer_style
        final_suggested_lines = 0 if subpart_blocks else suggested_lines
        has_table_response = final_answer_style == "table_response"

        block = {
            "q_no": q_no,
            "marks_label": marks_label,
            "marks_value": marks_value,
            "question_lines": main_question_lines,
            "has_subparts": bool(subpart_blocks),
            "subparts": subpart_blocks,
            "has_visual_slot": has_visual_slot,
            "visual_tokens": visual_tokens,
            "requires_diagram": requires_diagram,
            "requires_table": requires_table,
            "draft_answer_style": parent_answer_style,
            "answer_style": final_answer_style,
            "suggested_lines": final_suggested_lines,
            "answer_line_style": answer_line_style if apply_answer_lines else "none",
            "apply_answer_lines": apply_answer_lines,
            "has_table_response": has_table_response,
            "table_prompt": "[Complete the table]" if has_table_response else "",
            "numbering_style": donor_numbering_style,
            "question_heading_style": donor_question_heading_style,
            "answer_space_style": interpreted_answer_space_style,
            "answer_space_plan": _build_answer_space_plan(
                final_answer_style,
                final_suggested_lines
            ),
        }

        blocks.append(block)

    return blocks


def _build_optional_teacher_sections(template_profile: Dict, draft_model: Dict) -> List[Dict]:
    """
    Build optional teacher-only sections.
    
    Goal:
        - Keep the generated document practical
        - Avoid always injecting unnecessary teacher sections
        - Only provide a removable teacher note area when it may genuinely help
    """
    document_type = draft_model.get("document_type", DocumentType.CUSTOM)
    doc_type_str = document_type.value if hasattr(document_type, 'value') else str(document_type)
    
    raw_lines = draft_model.get("raw_lines", [])
    lower_blob = " ".join(raw_lines).lower()
    
    # Don't add teacher sections for these document types
    if doc_type_str.lower() in ["memo", "rubric"]:
        return []
    
    # Check if draft already seems complete
    completion_signals = [
        "answer all questions",
        "show all working",
        "time allowed",
        "read the following",
        "total marks",
        "section a",
        "section b"
    ]
    
    has_completion_signal = any(signal in lower_blob for signal in completion_signals)
    
    # If draft has clear structure and enough content, don't add clutter
    if has_completion_signal and len(raw_lines) >= 10:
        return []
    
    # Check if this is a worksheet or test (often benefit from teacher notes)
    if doc_type_str.lower() in ["worksheet", "test", "exam"]:
        return [
            {
                "section_type": "teacher_notes",
                "label": "Teacher Notes",
                "content": "[Optional: Add marking guide, differentiation notes, or answer key here]",
                "removable": True,
                "position": "end"
            }
        ]
    
    # Default: no optional sections
    return []


# ======================================================================================
# PHASE 6: RENDERING
# ======================================================================================

# ======================================================================================
# PHASE 6: RENDERING
# ======================================================================================

# Answer space format configurations
ANSWER_SPACE_FORMATS = {
    AnswerSpaceStyle.SHORT_RESPONSE: {
        "lines": 3,
        "pattern": "_________________________",
        "space_before": 6,
        "space_after": 12
    },
    AnswerSpaceStyle.PARAGRAPH_RESPONSE: {
        "lines": 8,
        "pattern": "_________________________",
        "space_before": 6,
        "space_after": 6
    },
    AnswerSpaceStyle.SHOW_WORKING: {
        "lines": 6,
        "pattern": "_________________________",
        "space_before": 6,
        "space_after": 12,
        "extra": "[Working space]"
    },
    AnswerSpaceStyle.LABEL_DIAGRAM: {
        "lines": 0,
        "pattern": "[Label the diagram below]",
        "space_before": 12,
        "space_after": 6
    },
    AnswerSpaceStyle.TABLE_RESPONSE: {
        "lines": 0,
        "pattern": "[Complete the table]",
        "space_before": 6,
        "space_after": 12
    },
    AnswerSpaceStyle.SUBPARTS: {
        "lines": 0,
        "pattern": "",
        "space_before": 0,
        "space_after": 0
    }
}


def render_blueprint_to_markdown(blueprint: Dict[str, Any]) -> str:
    """Render blueprint to markdown for preview."""
    if not blueprint:
        return "# No blueprint provided\n\nUnable to render document."
    
    lines = []
    
    # Header
    header_block = blueprint.get("header_block", {})
    if header_block.get("show_header"):
        institution_name = header_block.get("institution_name", "")
        if institution_name:
            lines.append(f"**{institution_name}**")
            lines.append("")
    
    # Title
    title = blueprint.get("title", "Untitled Document")
    lines.append(f"# {title}")
    lines.append("")
    
    # Metadata
    year_level = blueprint.get("year_level", "")
    subject = blueprint.get("subject", "")
    
    # TEMP: suppress metadata line to avoid title/metadata conflicts
    if False and (year_level or subject):
        meta_parts = []
        if year_level:
            meta_parts.append(f"**Year Level:** {year_level}")
        if subject:
            meta_parts.append(f"**Subject:** {subject}")
        lines.append(" | ".join(meta_parts))
        lines.append("")
    
    # Instructions
    instruction_block = blueprint.get("instruction_block", {})
    instruction_lines = instruction_block.get("instruction_lines", [])
    if instruction_lines:
        lines.append("## Instructions")
        lines.append("")
        for item in instruction_lines:
            lines.append(f"- {item}")
        lines.append("")

    # Layout decisions confidence summary (F9)
    layout_plan = blueprint.get("layout_plan", {})
    decisions = layout_plan.get("appropriateness_decisions", {})
    if decisions:
        lines.append("## Layout Decisions & Confidence")
        lines.append("")
        lines.append("The engine made the following decisions based on your draft's structure:")
        lines.append("")
        for feature_name, decision in decisions.items():
            apply_text = "✅ Apply" if decision.get("apply") else "❌ Do not apply"
            confidence = decision.get("confidence", 0.0) * 100
            reason = decision.get("reason", "No reason provided")
            source = decision.get("source", "engine")
            lines.append(f"**{feature_name.title()}**: {apply_text}  \n")
            lines.append(f"- **Confidence:** {confidence:.0f}%  \n")
            lines.append(f"- **Reason:** {reason}  \n")
            lines.append(f"- **Source:** {source}  \n")
            lines.append("")
        lines.append("---")
        lines.append("")
    
    # Appropriateness decisions (confidence transparency)
    appropriateness_decisions = blueprint.get("appropriateness_decisions", {})
    if appropriateness_decisions:
        lines.append("## Layout Decisions (with Confidence)")
        lines.append("")
        lines.append("The engine made the following decisions based on your draft's structure:")
        lines.append("")
        lines.append("| Feature | Decision | Confidence | Reason |")
        lines.append("|---------|----------|------------|--------|")
        for feature, decision in appropriateness_decisions.items():
            apply_text = "✅ Apply" if decision.get("apply") else "❌ Suppress"
            confidence = f"{decision.get('confidence', 0)*100:.0f}%"
            reason = decision.get("reason", "")[:60]
            lines.append(f"| {feature} | {apply_text} | {confidence} | {reason} |")
        lines.append("")
    
    # Questions
    visual_slots = blueprint.get("visual_slots", []) or []

    def _visual_slots_for_question(q_no: str) -> List[Dict[str, Any]]:
        raw_q_no = str(q_no or "").strip()
        raw_upper = raw_q_no.upper()

        candidate_keys = set()

        if raw_upper:
            candidate_keys.add(raw_upper)
            candidate_keys.add(f"Q{raw_upper}")

        if raw_upper.startswith("STEP "):
            step_number = raw_upper.replace("STEP ", "", 1).strip()
            if step_number:
                candidate_keys.add(f"STEP {step_number}")
                candidate_keys.add(f"Q{step_number}")

        matched = []

        for slot in visual_slots:
            if not isinstance(slot, dict):
                continue

            where = str(slot.get("where", "") or "").strip().upper()

            if where in candidate_keys:
                matched.append(slot)

        return matched

    def _visual_slot_to_marker(slot: Dict[str, Any]) -> str:
        raw = slot.get("raw")

        if isinstance(raw, str) and "[[VISUAL" in raw:
            return raw.strip()

        visual_id = str(slot.get("visual_id", "") or "").strip()
        kind = str(slot.get("kind", "") or "image").strip()
        where = str(slot.get("where", "") or "").strip()
        prompt = str(slot.get("prompt", "") or slot.get("description", "") or "").strip()
        notes = str(slot.get("notes", "") or "").strip()
        data = str(slot.get("data", "") or "").strip()

        attrs = []
        if visual_id:
            attrs.append(f'id="{visual_id}"')
        if kind:
            attrs.append(f'kind="{kind}"')
        if where:
            attrs.append(f'where="{where}"')
        if prompt:
            attrs.append(f'prompt="{prompt}"')
        if notes:
            attrs.append(f'notes="{notes}"')
        attrs.append(f'data="{data}"')

        return "[[VISUAL " + " ".join(attrs) + "]]"

    question_blocks = blueprint.get("question_blocks", [])

    # ✅ THIS IS THE ONLY FIX (LOOP RESTORED)
    for q in question_blocks:
        q_no = q.get("q_no", "?")
        marks_label = q.get("marks_label", "")
        
        heading = f"## Question {q_no}"
        if marks_label:
            heading += f" {marks_label}"
        lines.append(heading)
        lines.append("")
        
        # Question text
        for line in q.get("question_lines", []):
            if line.strip():
                lines.append(line)
        
        if q.get("question_lines"):
            lines.append("")

        q_visual_slots = _visual_slots_for_question(q_no)
        for slot in q_visual_slots:
            lines.append(_visual_slot_to_marker(slot))
            lines.append("")
        
        # Subparts
        subparts = q.get("subparts", [])
        if subparts:
            for sp in subparts:
                label = sp.get("label", "")
                body = (sp.get("body", "") or "").strip()

                if body:
                    body_lines = body.splitlines()
                    first = body_lines[0].strip()
                    rest = [ln.rstrip() for ln in body_lines[1:]]

                    lines.append(f"**({label})** {first}")
                    lines.append("")
                    for ln in rest:
                        lines.append(ln)

                    lines.append("_______________________________________________________________________________________________")
                    lines.append("")

                else:
                    lines.append(f"**({label})**")

                lines.append("")
        else:
            # Answer space preview must obey the blueprint answer-space plan.
            answer_style = q.get("answer_style", "short_response")
            suggested_lines = q.get("suggested_lines", 3)
            answer_space_plan = q.get("answer_space_plan", {}) or {}

            apply_answer_lines = bool(
                answer_space_plan.get("apply_answer_lines", False)
                or q.get("requires_written_response", False)
                or q.get("answer_style") in ["short_response", "long_response"]
            )
            planned_style = str(answer_space_plan.get("answer_line_style", "") or "").strip().lower()
            planned_lines = max(0, _try_int(answer_space_plan.get("suggested_lines", suggested_lines), suggested_lines))

            pattern = "_________________________"
            if planned_style == "none":
                pattern = ""
            elif planned_style in {"underscore", "underscores"}:
                pattern = "_________________________"
            elif planned_style in {"dotted", "dots"}:
                pattern = ". . . . . . . . . . . . . . . . . . . . ."
            elif planned_style in {"solid", "rule"}:
                pattern = "________________________________________"

            if apply_answer_lines and planned_lines > 0 and pattern:
                for _ in range(planned_lines):
                    lines.append(pattern)

            lines.append("")
    
    return "\n".join(lines)


def _render_bs5_institutional_front_matter(doc, blueprint: Dict[str, Any], first_page_contract: Optional[Dict[str, Any]] = None) -> None:
    """
    Render BS5 institutional front matter from the composed blueprint only.

    GOAL:
    - make the output feel institution-authored, not Workspace-authored
    - keep draft content authority intact
    - use donor institutional identity/style, not donor academic content
    - do NOT dump running-header identity into page-1 body text
    - when donor identity is logo-led, suppress duplicate plain-text identity
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    identity_block = blueprint.get("identity_block", {}) or {}
    header_block = blueprint.get("header_block", {}) or {}
    title_block = blueprint.get("title_block", {}) or {}
    instruction_block = blueprint.get("instruction_block", {}) or {}
    logo_candidate = blueprint.get("logo_candidate") or {}
    first_page_contract = first_page_contract or {}
    logo_candidate = logo_candidate or {}

    institution_name = str(identity_block.get("institution_name", "") or "").strip()
    field_lines = list(identity_block.get("field_lines", []) or [])

    # Prefer first-page-specific identity lines from header_block, not running header lines
    first_page_identity_lines = list(header_block.get("first_page_identity_lines", []) or [])

    title = str(blueprint.get("title", "") or "Untitled Document").strip()
    year_level = str(blueprint.get("year_level", "") or "").strip()
    subject = str(blueprint.get("subject", "") or "").strip()
    document_type_display = str(title_block.get("document_type_display", "") or "").strip()

    logo_present = bool(
        identity_block.get("logo_present", False)
        or (isinstance(logo_candidate, dict) and bool(logo_candidate.get("binary_data")))
    )

    # --------------------------------------------------
    # 0. Page-one institutional logo
    # --------------------------------------------------
    # If the donor stores its page-one logo in the first-page header, the
    # header/footer renderer owns it. Do not also centre it in body front matter.
    contract_first_header_has_logo = bool(first_page_contract.get("first_page_header_has_images", False))
    logo_position_hint = str(logo_candidate.get("position", "") if isinstance(logo_candidate, dict) else "").strip().lower()
    render_body_logo = bool(
        isinstance(logo_candidate, dict)
        and logo_candidate.get("binary_data")
        and not contract_first_header_has_logo
        and logo_position_hint not in {"header", "first_page_header"}
    )

    if render_body_logo:
        try:
            from docx.shared import Inches
            from io import BytesIO

            logo_para = doc.add_paragraph()
            logo_alignment_name = str(
                (first_page_contract.get("dominant_first_page_paragraph_style_signature", {}) or {}).get("alignment")
                or logo_candidate.get("alignment")
                or "CENTER"
            ).upper()
            logo_para.alignment = getattr(WD_ALIGN_PARAGRAPH, logo_alignment_name, WD_ALIGN_PARAGRAPH.CENTER)

            logo_width = float(logo_candidate.get("width_inches", 1.25) or 1.25)
            logo_width = max(0.5, min(logo_width, 1.5))

            run = logo_para.add_run()
            run.add_picture(BytesIO(logo_candidate.get("binary_data")), width=Inches(logo_width))

            doc.add_paragraph()
        except Exception as e:
            logger.warning(f"Failed to render BS5 front-matter logo: {e}")

    def _clean_line(text: str) -> str:
        return str(text or "").replace("\xa0", " ").strip()

    def _dedupe_keep_order(items):
        seen = set()
        out = []
        for item in items:
            s = _clean_line(item)
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    # --------------------------------------------------
    # 1. Institutional identity block (page 1 only)
    #
    # IMPORTANT:
    # - If donor is logo-led, do NOT duplicate plain-text institution name.
    # - Only use first-page-specific identity lines here.
    # - Running header lines belong to later pages, not front matter body text.
    # --------------------------------------------------
    identity_lines = []

    if not logo_present:
        if institution_name:
            identity_lines.append(institution_name)

        identity_lines.extend(first_page_identity_lines)
    else:
        # Logo-led donor: only keep first-page identity text if it adds something
        # beyond the institution name itself.
        for line in first_page_identity_lines:
            s = _clean_line(line)
            if not s:
                continue
            if s.lower() == institution_name.lower():
                continue
            identity_lines.append(s)

    identity_lines = _dedupe_keep_order(identity_lines)

    for idx, line in enumerate(identity_lines[:6]):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)

        if idx == 0:
            run.bold = True
            if getattr(run.font, "size", None) is None:
                try:
                    from docx.shared import Pt
                    run.font.size = Pt(14)
                except Exception:
                    pass
        else:
            if getattr(run.font, "size", None) is None:
                try:
                    from docx.shared import Pt
                    run.font.size = Pt(11)
                except Exception:
                    pass

    if identity_lines:
        doc.add_paragraph()

    # --------------------------------------------------
    # 2. Teacher/admin field lines
    # --------------------------------------------------
    cleaned_fields = _dedupe_keep_order(field_lines)
    for line in cleaned_fields[:6]:
        doc.add_paragraph(line)

    if cleaned_fields:
        doc.add_paragraph()

    # --------------------------------------------------
    # 3. Draft-led title and document framing
    # --------------------------------------------------
    if title:
        title_para = doc.add_paragraph()
        title_signature = first_page_contract.get("dominant_first_page_paragraph_style_signature") or {}
        title_alignment_name = str(title_signature.get("alignment") or "CENTER").upper()
        title_para.alignment = getattr(WD_ALIGN_PARAGRAPH, title_alignment_name, WD_ALIGN_PARAGRAPH.CENTER)
        title_run = title_para.add_run(title)
        if title_signature:
            _apply_paragraph_style_signature(title_para, title_signature)
        else:
            title_run.bold = True
            try:
                from docx.shared import Pt
                title_run.font.size = Pt(16)
            except Exception:
                pass

    meta_parts = []
    if year_level:
        meta_parts.append(year_level)
    if subject:
        meta_parts.append(subject)
    if document_type_display:
        meta_parts.append(document_type_display)

    if meta_parts:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(" | ".join(meta_parts))
        meta_run.italic = True

    doc.add_paragraph()

    # --------------------------------------------------
    # 4. Instruction framing
    # --------------------------------------------------
    instruction_lines = list(instruction_block.get("instruction_lines", []) or [])
    instruction_mode = str(instruction_block.get("mode", "") or "").strip()

    if instruction_lines and instruction_mode != "implicit":
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Instructions")
        heading_run.bold = True

        for item in instruction_lines[:10]:
            s = _clean_line(item)
            if s:
                doc.add_paragraph(f"• {s}")

        doc.add_paragraph()


def render_blueprint_to_docx(blueprint: Dict[str, Any], template_profile: Dict[str, Any]) -> str:
    """
    Render blueprint to DOCX with actual styles.

    BS5 RULE:
    - the composed blueprint is the rendering authority
    - institutional opening/front matter must come from BS5 blueprint blocks only
    - do not apply a second legacy header/front-matter path before BS5 front matter
    - answer-space rendering must obey explicit blueprint answer-space plans
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available, cannot render DOCX")
        raise RuntimeError("python-docx is required for DOCX rendering. Install with: pip install python-docx")

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --------------------------------------------------
    # Apply dominant font from template_profile to Normal style
    # --------------------------------------------------
    try:
        from docx.shared import Pt
        from collections import Counter

        style_prefs = template_profile.get("style_preferences", {}) or {}
        font_samples = template_profile.get("font_samples", []) or []

        font_name = style_prefs.get("font_name") or style_prefs.get("font_family")
        font_size = style_prefs.get("font_size_pt") or style_prefs.get("font_size")

        if not font_name and font_samples:
            names = [s.get("name") for s in font_samples if s.get("name")]
            if names:
                font_name = Counter(names).most_common(1)[0][0]

        if not font_size and font_samples:
            sizes = [s.get("size_pt") for s in font_samples if s.get("size_pt")]
            if sizes:
                font_size = Counter(sizes).most_common(1)[0][0]

        normal_style = doc.styles["Normal"]
        if font_name:
            normal_style.font.name = font_name
            logger.info(f"RENDERER: Applied Normal font name: {font_name}")
        if font_size:
            normal_style.font.size = Pt(float(font_size))
            logger.info(f"RENDERER: Applied Normal font size: {font_size}pt")
    except Exception as e:
        logger.warning(f"RENDERER: Could not apply Normal style font (non-fatal): {e}")

    # --------------------------------------------------
    # Page setup from blueprint contract surface first
    # --------------------------------------------------
    appropriateness_decisions = blueprint.get("appropriateness_decisions", {}) or {}

    columns_decision = appropriateness_decisions.get("columns", {})
    two_column_math_decision = appropriateness_decisions.get("two_column_math", {})

    columns_allowed = bool(columns_decision.get("apply", False))
    math_columns_blocked = ("two_column_math" in appropriateness_decisions) and (
        not bool(two_column_math_decision.get("apply", False))
    )

    final_columns_allowed = columns_allowed and not math_columns_blocked

    # Log the column decision for debugging
    if final_columns_allowed:
        logger.info("RENDERER: Applying columns – appropriateness decision allowed and math columns not blocked.")
    else:
        reason = "columns decision suppressed"
        if not columns_allowed:
            reason = "appropriateness decision set columns.apply=False"
        elif math_columns_blocked:
            reason = "math columns detected – blocked by NEVER_INHERIT rule"
        logger.info(f"RENDERER: Columns suppressed – {reason}")

    page_setup = blueprint.get("page_setup", {}) or template_profile.get("page_setup", {}) or {}
    _apply_page_setup_to_doc(doc, page_setup, final_columns_allowed)

    # --------------------------------------------------
    # BS5 running header/footer
    # - page 1 remains owned by institutional front matter
    # - later pages get safe institutional running header/footer
    # --------------------------------------------------
    header_block = blueprint.get("header_block", {}) or {}
    _preserve_header_footer_structure(
        doc,
        header_block,
        blueprint.get("title", ""),
        datetime.now().strftime("%Y-%m-%d"),
        template_profile.get("first_page_layout_contract", {}) or {},
        blueprint.get("logo_candidate") or {}
    )

    # --------------------------------------------------
    # BS5 institutional front matter
    # --------------------------------------------------
    _render_bs5_institutional_front_matter(doc, blueprint, template_profile.get("first_page_layout_contract", {}) or {})

    title = blueprint.get("title", "Untitled Document")

    # --------------------------------------------------
    # Visual slots (question-aware placement)
    # --------------------------------------------------
    visual_slots = blueprint.get("visual_slots", []) or []

    def _visual_slots_for_question(q_no: str):
        raw_q_no = str(q_no or "").strip()
        raw_upper = raw_q_no.upper()

        candidate_keys = set()
        candidate_prefixes = set()

        if raw_upper:
            candidate_keys.add(raw_upper)
            candidate_keys.add(f"Q{raw_upper}")

        if raw_upper.startswith("STEP "):
            step_number = raw_upper.replace("STEP ", "", 1).strip()
            if step_number:
                candidate_keys.add(f"STEP {step_number}")
                candidate_keys.add(f"Q{step_number}")
                candidate_prefixes.add(f"STEP {step_number} ")

        matched = []
        leftovers = []

        for slot in visual_slots:
            if not isinstance(slot, dict):
                leftovers.append(slot)
                continue

            where = str(slot.get("where", "") or "").strip().upper()

            exact_match = where in candidate_keys
            prefix_match = any(where.startswith(prefix) for prefix in candidate_prefixes)

            if exact_match or prefix_match:
                matched.append(slot)
            else:
                leftovers.append(slot)

        return matched, leftovers

    def _is_document_level_visual_slot(slot: Dict[str, Any]) -> bool:
        where = str(slot.get("where", "") or "").strip().lower()
        return where in {
            "document",
            "top of document",
            "top",
            "before q1",
            "before question 1",
            "document top",
        }

    document_level_visual_slots = [
        slot for slot in visual_slots
        if isinstance(slot, dict) and _is_document_level_visual_slot(slot)
    ]

    if document_level_visual_slots:
        _preserve_visual_slots_in_output(doc, document_level_visual_slots)

    # --------------------------------------------------
    # Questions
    # --------------------------------------------------
    first_page_contract = template_profile.get("first_page_layout_contract", {}) or {}
    donor_question_heading_signature = first_page_contract.get("dominant_question_heading_style_signature") or {}
    donor_body_signature = first_page_contract.get("dominant_body_style_signature") or {}

    question_blocks = blueprint.get("question_blocks", []) or []
    for q in question_blocks:
        q_no = str(q.get("q_no", "?")).strip() or "?"
        marks_label = str(q.get("marks_label", "")).strip()

        if q_no.lower().startswith("step "):
            heading = q_no
        else:
            heading = f"Question {q_no}"

        if marks_label:
            heading += f" {marks_label}"

        q_para = doc.add_paragraph(heading)
        if donor_question_heading_signature:
            _apply_paragraph_style_signature(q_para, donor_question_heading_signature)
        else:
            q_para.style = doc.styles["Heading 2"]

        for line in q.get("question_lines", []) or []:
            clean_line = str(line).strip()
            if not clean_line:
                continue

            if q_no.lower().startswith("step ") and re.match(r"^\d+\.\s+", clean_line):
                doc.add_paragraph(clean_line, style="List Number")
            else:
                body_para = doc.add_paragraph(clean_line)
                if donor_body_signature:
                    _apply_paragraph_style_signature(body_para, donor_body_signature)

        if q.get("question_lines"):
            doc.add_paragraph()

        # --------------------------------------------------
        # Question-specific visual slots
        # --------------------------------------------------
        q_visual_slots, _unused_visuals = _visual_slots_for_question(q_no)
        if q_visual_slots:
            _preserve_visual_slots_in_output(doc, q_visual_slots)

        subparts = q.get("subparts", []) or []
        if subparts:
            for sp in subparts:
                label = sp.get("label", "")
                body = (sp.get("body", "") or "").strip()

                if body:
                    body_lines = body.splitlines()
                    first = body_lines[0].strip()
                    rest = [ln.rstrip() for ln in body_lines[1:]]

                    body_para = doc.add_paragraph(f"({label}) {first}")
                    if donor_body_signature:
                        _apply_paragraph_style_signature(body_para, donor_body_signature)
                    for ln in rest:
                        if ln.strip():
                            body_para = doc.add_paragraph(ln.strip())
                            if donor_body_signature:
                                _apply_paragraph_style_signature(body_para, donor_body_signature)
                else:
                    body_para = doc.add_paragraph(f"({label})")
                    if donor_body_signature:
                        _apply_paragraph_style_signature(body_para, donor_body_signature)

                subpart_plan = sp.get("answer_space_plan", {}) or {}
                if subpart_plan:
                    _format_answer_space(
                        doc,
                        style=sp.get("answer_style", "short_response"),
                        marks=0,
                        has_subparts=False,
                        answer_space_plan=subpart_plan,
                    )
        else:
            marks_value = q.get("marks_value", 0)
            has_table_response = bool(q.get("has_table_response", False))
            table_prompt = (q.get("table_prompt") or "").strip()
            answer_space_plan = q.get("answer_space_plan", {}) or {}

            if has_table_response:
                _render_basic_table_response(doc, marks_value, table_prompt)
            else:
                _format_answer_space(
                    doc,
                    style=q.get("answer_style", "short_response"),
                    marks=marks_value,
                    has_subparts=False,
                    answer_space_plan=answer_space_plan,
                )

        doc.add_paragraph()

    # --------------------------------------------------
    # Visual slots
    # --------------------------------------------------
    # Question-linked visual slots are rendered inside the question loop.
    # No global end-of-document dump here.

    # --------------------------------------------------
    # Page breaks
    # --------------------------------------------------
    page_breaks = blueprint.get("page_breaks", []) or []
    if page_breaks:
        _apply_intelligent_page_breaks(doc, question_blocks, page_breaks)

    # --------------------------------------------------
    # Save
    # --------------------------------------------------
    out_dir = tempfile.mkdtemp(prefix="template_engine_docx_")
    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "_", title)[:50] or "rendered_document"
    out_path = os.path.join(out_dir, f"{safe_title}.docx")

    doc.save(out_path)
    logger.info(f"DOCX rendered: {out_path}")

    return out_path


def _apply_page_setup_to_doc(doc, page_setup: Dict, columns_allowed: bool) -> None:
    """Apply margins, orientation, and columns (if allowed)."""
    if not doc.sections:
        return

    section = doc.sections[0]

    # Apply margins (always)
    left = page_setup.get("left_margin_inches", 1.0)
    right = page_setup.get("right_margin_inches", 1.0)
    top = page_setup.get("top_margin_inches", 1.0)
    bottom = page_setup.get("bottom_margin_inches", 1.0)

    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)

    # Apply orientation
    if page_setup.get("orientation") == "landscape":
        section.orientation = WD_SECTION.LANDSCAPE
        section.page_width = Inches(page_setup.get("page_height_inches", 11.0))
        section.page_height = Inches(page_setup.get("page_width_inches", 8.5))

    # Apply columns ONLY if allowed
    try:
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml

        sect_pr = section._sectPr
        cols_elem = sect_pr.find(qn('w:cols'))
        if cols_elem is None:
            cols_elem = parse_xml(
                r'<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            sect_pr.append(cols_elem)

        if columns_allowed:
            columns = page_setup.get("columns", {})
            col_count = columns.get("count", 1)

            if col_count > 1:
                cols_elem.set(qn('w:num'), str(col_count))
                logger.info(f"Applied {col_count} columns to document")
            else:
                cols_elem.set(qn('w:num'), "1")
                logger.info("Columns allowed, but donor column count <= 1; enforced single column")
        else:
            # Explicitly force single-column output when columns are suppressed
            cols_elem.set(qn('w:num'), "1")
            logger.info("Columns suppressed by appropriateness rules; enforced single column")

    except Exception as e:
        logger.warning(f"Failed to apply column setup: {e}")


def _apply_logo_to_output(output_doc, logo_data: ExtractedImage) -> None:
    """Apply extracted logo to output document."""
    if not logo_data or not logo_data.binary_data:
        logger.warning("No logo data to apply")
        return

    try:
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO

        temp_logo = BytesIO(logo_data.binary_data)

        raw_width = logo_data.width_inches if logo_data.width_inches else 1.5
        logo_width = max(0.5, min(2.5, raw_width))

        in_header_table = getattr(logo_data, "in_header_table", False)
        institution_name = getattr(logo_data, "institution_name", None)

        if in_header_table and institution_name and output_doc.sections:
            header = output_doc.sections[0].header
            # Clear existing header paragraphs
            for para in header.paragraphs:
                para.clear()

            table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            table.style = "Table Grid"
            # Remove borders
            tbl = table._tbl
            tbl_pr = tbl.find(qn("w:tblPr"))
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                tbl.insert(0, tbl_pr)
            tbl_borders = OxmlElement("w:tblBorders")
            for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border_el = OxmlElement(f"w:{border_name}")
                border_el.set(qn("w:val"), "none")
                tbl_borders.append(border_el)
            tbl_pr.append(tbl_borders)

            left_cell = table.cell(0, 0)
            temp_logo.seek(0)
            left_run = left_cell.paragraphs[0].add_run()
            left_run.add_picture(temp_logo, width=Inches(logo_width))

            right_cell = table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_run = right_para.add_run(institution_name)
            right_run.bold = True
            right_run.font.size = Pt(11)

            logger.info("Applied logo + institution name as borderless header table")

        elif logo_data.position == "header" and output_doc.sections:
            header = output_doc.sections[0].header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = paragraph.add_run()
            temp_logo.seek(0)
            run.add_picture(temp_logo, width=Inches(logo_width))
            logger.info("Applied logo to header paragraph")

        else:
            paragraph = output_doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if logo_data.alignment == "center" else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run()
            temp_logo.seek(0)
            run.add_picture(temp_logo, width=Inches(logo_width))
            logger.info("Applied logo to body paragraph")

    except Exception as e:
        logger.warning(f"Failed to apply logo: {e}")


def _preserve_header_footer_structure(
    output_doc,
    header_block: Dict[str, Any],
    draft_title: str,
    draft_date: str,
    first_page_contract: Optional[Dict[str, Any]] = None,
    logo_candidate: Optional[Dict[str, Any]] = None
) -> None:
    """
    Apply BS5-composed running header/footer structure.

    BS5 RULE:
    - do NOT preserve donor academic header/footer text directly
    - use only the composed header block surface
    - keep page 1 front matter separate from running header/footer
    - restore safe institutional running header + page-number footer
    """
    if not output_doc.sections:
        return

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = output_doc.sections[0]
    first_page_contract = first_page_contract or {}
    logo_candidate = logo_candidate or {}

    # --------------------------------------------------
    # Always separate and clear header/footer containers first
    # so stale content can never survive when show_header=False
    # --------------------------------------------------
    section.different_first_page_header_footer = bool(first_page_contract.get("different_first_page_header_footer", True))

    header = section.header
    first_header = section.first_page_header
    footer = section.footer
    first_footer = section.first_page_footer

    def _clear_container(container):
        try:
            for para in container.paragraphs:
                para.clear()
        except Exception:
            pass

    _clear_container(header)
    _clear_container(first_header)
    _clear_container(footer)
    _clear_container(first_footer)

    def _add_logo_to_container(container, style_sig=None) -> bool:
        if not isinstance(logo_candidate, dict) or not logo_candidate.get("binary_data"):
            return False
        try:
            from docx.shared import Inches
            from io import BytesIO

            para = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
            para.clear()
            _apply_paragraph_style_signature(para, style_sig or {})

            width = float(logo_candidate.get("width_inches", 1.25) or 1.25)
            width = max(0.5, min(width, 2.0))
            para.add_run().add_picture(BytesIO(logo_candidate.get("binary_data")), width=Inches(width))
            return True
        except Exception as e:
            logger.warning(f"Could not restore contract header/footer logo: {e}")
            return False

    # Do not rewrite donor header/footer text as plain text: that can flatten fields,
    # tables, drawings, and old pagination. Use the contract as layout evidence only.
    wrote_donor_default_header = False
    wrote_donor_first_header = False
    if bool(first_page_contract.get("default_header_has_images", False)):
        wrote_donor_default_header = _add_logo_to_container(header, first_page_contract.get("default_header_style_signature"))
    if section.different_first_page_header_footer and bool(first_page_contract.get("first_page_header_has_images", False)):
        wrote_donor_first_header = _add_logo_to_container(first_header, first_page_contract.get("first_page_header_style_signature"))

    if not isinstance(header_block, dict):
        return

    show_header = bool(header_block.get("show_header", False))
    if not show_header:
        logger.info("BS5 header/footer cleared; no running header applied because show_header=False")
        return

    institution_name = str(header_block.get("institution_name", "") or "").strip()
    header_lines = header_block.get("header_lines", []) or []
    field_lines = header_block.get("field_lines", []) or []

    running_header_alignment_name = str(header_block.get("running_header_alignment", "CENTER") or "CENTER").strip().upper()
    running_footer_alignment_name = str(header_block.get("running_footer_alignment", "CENTER") or "CENTER").strip().upper()

    alignment_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    running_header_alignment = alignment_map.get(running_header_alignment_name, WD_ALIGN_PARAGRAPH.CENTER)
    running_footer_alignment = alignment_map.get(running_footer_alignment_name, WD_ALIGN_PARAGRAPH.CENTER)

    cleaned_header_lines = []
    seen = set()

    def push(text):
        s = str(text or "").strip()
        if not s:
            return
        key = s.lower()
        if key in seen:
            return
        seen.add(key)
        cleaned_header_lines.append(s)

    # --------------------------------------------------
    # BS5 RULE:
    # Do NOT repeat institution identity in running header
    # Front matter already owns identity on page 1
    # --------------------------------------------------

    for line in header_lines[:10]:
        push(line)

    # Remove obvious field/admin lines from running header
    safe_running_lines = []
    for line in cleaned_header_lines:
        low = line.lower()
        if low.startswith("name:") or low.startswith("teacher:"):
            continue
        safe_running_lines.append(line)

    wrote_any = bool(wrote_donor_default_header)

    header_pattern = header_block.get("header_pattern", {}) or {}
    pattern_confidence = float(header_pattern.get("confidence", 0.0) or 0.0)
    running_pattern = header_pattern.get("running", {}) or {}

    use_header_pattern = (
        isinstance(header_pattern, dict)
        and pattern_confidence >= 0.75
        and isinstance(running_pattern, dict)
    )

    def _add_bottom_border(paragraph):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    if use_header_pattern:
        for text in (running_pattern.get("admin_labels", []) or [])[:1]:
            s = str(text or "").strip()
            if not s:
                continue
            admin_para = header.add_paragraph()
            admin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            admin_para.add_run(s.upper())
            wrote_any = True

        left_text = str(running_pattern.get("left", "") or "").strip()
        centre_text = str(running_pattern.get("center", "") or "").strip()
        right_text = str(running_pattern.get("right", "") or "").strip()

        if left_text or centre_text or right_text:
            running_para = header.add_paragraph()
            running_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            try:
                from docx.enum.text import WD_TAB_ALIGNMENT
                from docx.shared import Inches

                tab_stops = running_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(3.15), WD_TAB_ALIGNMENT.CENTER)
                tab_stops.add_tab_stop(Inches(6.30), WD_TAB_ALIGNMENT.RIGHT)
            except Exception:
                pass

            if left_text:
                running_para.add_run(left_text)

            running_para.add_run("\t")

            if centre_text:
                centre_run = running_para.add_run(centre_text)
                centre_run.bold = True

            running_para.add_run("\t")

            if right_text:
                running_para.add_run(right_text)

            if bool(running_pattern.get("has_underline", False)):
                _add_bottom_border(running_para)

            wrote_any = True

    if not wrote_any:
        admin_label = ""
        for line in safe_running_lines:
            if str(line).strip().lower() in {"official", "draft", "confidential"}:
                admin_label = str(line).strip().upper()
                break

        if admin_label:
            admin_para = header.add_paragraph()
            admin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            admin_para.add_run(admin_label)
            wrote_any = True

        year_subject_line = ""
        doc_type_line = ""

        for line in safe_running_lines:
            s = str(line or "").strip()
            if not s:
                continue
            if s.lower() in {"official", "draft", "confidential"}:
                continue
            if "|" in s:
                parts = [p.strip() for p in s.split("|") if p.strip()]
                if len(parts) >= 2:
                    year_subject_line = " ".join(parts[:2])
                    if len(parts) >= 3:
                        doc_type_line = parts[2]
                    break

        if not year_subject_line:
            year_subject_line = draft_title or institution_name

        running_para = header.add_paragraph()
        running_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        try:
            from docx.enum.text import WD_TAB_ALIGNMENT
            from docx.shared import Inches

            tab_stops = running_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3.15), WD_TAB_ALIGNMENT.CENTER)
            tab_stops.add_tab_stop(Inches(6.30), WD_TAB_ALIGNMENT.RIGHT)
        except Exception:
            pass

        if year_subject_line:
            running_para.add_run(year_subject_line)

        running_para.add_run("\t")

        if institution_name:
            centre_run = running_para.add_run(institution_name)
            centre_run.bold = True

        running_para.add_run("\t")

        if doc_type_line:
            running_para.add_run(doc_type_line)

        _add_bottom_border(running_para)

        wrote_any = True

    # --------------------------------------------------
    # FIRST PAGE HEADER (separate from running header)
    # --------------------------------------------------

    first_page_labels = header_block.get("first_page_header_labels", []) or []

    wrote_first = bool(wrote_donor_first_header)

    for idx, text in enumerate(first_page_labels[:2]):
        s = str(text or "").strip()
        if not s:
            continue

        para = first_header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(s)

        if idx == 0:
            run.bold = True

        wrote_first = True

    # Fallback: reuse OFFICIAL if present anywhere
    if not wrote_first:
        for line in header_lines:
            if str(line).strip().lower() == "official":
                para = first_header.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run("OFFICIAL")
                break

    # --------------------------------------------------
    # Running footer (pages after first) - page numbering
    # --------------------------------------------------
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()
    footer_para.alignment = running_footer_alignment

    def _append_complex_field(paragraph, field_name: str, placeholder_text: str = "1"):
        """
        Insert a Word complex field so PAGE / NUMPAGES render correctly in DOCX.
        """
        # Begin field
        run_begin = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_char_begin)

        # Instruction text
        run_instr = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = field_name
        run_instr._r.append(instr_text)

        # Separate
        run_sep = paragraph.add_run()
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_char_sep)

        # Placeholder display text (Word updates this when fields refresh)
        run_text = paragraph.add_run(placeholder_text)

        # End field
        run_end = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_char_end)

    # Always render dynamic fields for page numbering. If the donor had page
    # numbering, this preserves dynamic PAGE/NUMPAGES semantics; if not, this
    # keeps the existing BS5 running-footer behaviour.
    footer_para.add_run("Page ")
    _append_complex_field(footer_para, " PAGE ", "1")
    footer_para.add_run(" of ")
    _append_complex_field(footer_para, " NUMPAGES ", "1")

    # Keep first-page footer blank so page 1 opening stays clean
    if not first_footer.paragraphs:
        first_footer.add_paragraph("")

    logger.info("BS5 running header/footer applied: first page clean, later pages branded")


def _recreate_table_in_output(output_doc, table_structure: TableStructure, questions: List) -> None:
    """Recreate table in output document, optionally replacing content."""
    try:
        rows = table_structure.rows
        cols = table_structure.cols
        
        if rows == 0 or cols == 0:
            return
        
        table = output_doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Apply column widths
        if table_structure.column_widths:
            for col_idx, width in enumerate(table_structure.column_widths[:cols]):
                if col_idx < len(table.columns):
                    table.columns[col_idx].width = Inches(width)
        
        # Mark header row if present
        if table_structure.has_header_row and table.rows:
            for cell in table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True if cell.paragraphs[0].runs else None
        
        logger.info(f"Recreated table: {rows}x{cols}")
    except Exception as e:
        logger.warning(f"Failed to recreate table: {e}")


def _render_basic_table_response(
    output_doc,
    marks: int = 0,
    prompt: str = "[Complete the table]"
) -> None:
    """
    Render a neutral blank response table for draft questions that require
    a table-style answer, without reusing donor table structure.

    This is an interim draft-driven scaffold:
    - no donor table injection
    - no guessed academic content
    - only a simple blank response grid
    """
    # Only show prompt if it is NOT the generic placeholder
    clean_prompt = (prompt or "").strip()

    if clean_prompt and clean_prompt.lower() not in ["[complete the table]", "complete the table"]:
        output_doc.add_paragraph(clean_prompt)

    # Keep sizing conservative and predictable
    cols = 3
    rows = 4 if marks < 4 else 6

    table = output_doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    # Leave cells blank for student completion
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""

    output_doc.add_paragraph()


def _format_answer_space(
    output_doc,
    style: Union[AnswerSpaceStyle, str],
    marks: int,
    has_subparts: bool = False,
    answer_space_plan: Optional[Dict[str, Any]] = None,
) -> None:
    """
    Format answer space according to the BS5 blueprint answer-space plan.

    Priority:
    1. explicit answer_space_plan from blueprint
    2. fallback to enum/default answer-space formats
    """
    answer_space_plan = answer_space_plan or {}

    # --------------------------------------------------
    # Normalize style
    # --------------------------------------------------
    resolved_style = AnswerSpaceStyle.SHORT_RESPONSE

    if isinstance(style, AnswerSpaceStyle):
        resolved_style = style
    elif isinstance(style, str):
        style_str = style.strip()
        for s in AnswerSpaceStyle:
            if s.value == style_str or s.name == style_str.upper():
                resolved_style = s
                break

    format_config = ANSWER_SPACE_FORMATS.get(
        resolved_style,
        ANSWER_SPACE_FORMATS[AnswerSpaceStyle.SHORT_RESPONSE]
    )

    # --------------------------------------------------
    # Read blueprint answer-space plan
    # --------------------------------------------------
    # Design Contract safety rule:
    # if no explicit blueprint plan approves answer lines, do not create them by default.
    apply_answer_lines = bool(answer_space_plan.get("apply_answer_lines", False))
    planned_style = str(answer_space_plan.get("answer_line_style", "") or "").strip().lower()
    planned_lines = max(0, _try_int(answer_space_plan.get("suggested_lines", format_config.get("lines", 3)), format_config.get("lines", 3)))

    pattern = format_config.get("pattern", "_________________________")
    space_before = format_config.get("space_before", 6)
    space_after = format_config.get("space_after", 12)

    # --------------------------------------------------
    # Respect donor answer-line suppression without removing draft answer space
    # --------------------------------------------------
    lines = planned_lines

    if not apply_answer_lines:
        # Donor-specific answer-line styling is suppressed, but the draft still
        # needs ordinary usable answer space. Use blank paragraph spacing instead.
        pattern = ""
    else:
        if planned_style == "none":
            pattern = ""
        elif planned_style in {"underscore", "underscores"}:
            pattern = "_________________________"
        elif planned_style in {"dotted", "dots"}:
            pattern = ". . . . . . . . . . . . . . . . . . . . ."
        elif planned_style in {"solid", "rule"}:
            pattern = "________________________________________"
        elif planned_style:
            # Unknown style -> keep default pattern conservatively
            pass

    # --------------------------------------------------
    # Conservative marks-based fallback only when plan is sparse
    # --------------------------------------------------
    if apply_answer_lines and lines <= 0:
        if resolved_style == AnswerSpaceStyle.SHORT_RESPONSE:
            lines = 4 if marks >= 4 else max(2, format_config.get("lines", 3))
        elif resolved_style == AnswerSpaceStyle.PARAGRAPH_RESPONSE:
            lines = 10 if marks >= 6 else max(6, format_config.get("lines", 8))
        elif resolved_style == AnswerSpaceStyle.SHOW_WORKING:
            lines = max(6, format_config.get("lines", 6))

    # --------------------------------------------------
    # Spacing before
    # --------------------------------------------------
    if space_before > 0:
        for _ in range(space_before // 6):
            output_doc.add_paragraph()

    # --------------------------------------------------
    # Main answer space
    # --------------------------------------------------
    if lines > 0 and not has_subparts:
        for _ in range(lines):
            output_doc.add_paragraph(pattern if pattern else "")
    elif pattern and not has_subparts and resolved_style in {AnswerSpaceStyle.LABEL_DIAGRAM, AnswerSpaceStyle.TABLE_RESPONSE}:
        output_doc.add_paragraph(pattern)

    # --------------------------------------------------
    # Extra show-working cue
    # --------------------------------------------------
    extra = format_config.get("extra")
    if extra and resolved_style == AnswerSpaceStyle.SHOW_WORKING and not has_subparts:
        output_doc.add_paragraph(extra)

    # --------------------------------------------------
    # Spacing after
    # --------------------------------------------------
    if space_after > 0:
        for _ in range(space_after // 6):
            output_doc.add_paragraph()


def _apply_intelligent_page_breaks(
    output_doc,
    question_blocks: List[Dict],
    explicit_breaks: List[int]
) -> None:
    """Apply page breaks intelligently."""
    if not output_doc.sections:
        return
    
    section = output_doc.sections[0]
    
    # Honor explicit breaks from draft (simplified - would need paragraph mapping)
    for break_line in explicit_breaks:
        try:
            # Add page break at specified position
            if len(output_doc.paragraphs) > break_line:
                output_doc.paragraphs[break_line].insert_paragraph_before().add_run().add_break()
        except Exception as e:
            logger.warning(f"Failed to add explicit page break: {e}")
    
    # Add break before each major question (question 1, question 6, question 11, etc.)
    for idx, q in enumerate(question_blocks):
        if idx > 0 and idx % 5 == 0:  # Every 5 questions
            try:
                # Find the paragraph with the question heading
                for para in output_doc.paragraphs:
                    if f"Question {q.get('q_no')}" in para.text:
                        para.insert_paragraph_before().add_run().add_break()
                        break
            except Exception:
                pass


def _preserve_cross_references(output_doc, references: List[Dict]) -> None:
    """Ensure cross-references remain valid after layout changes."""
    if not references:
        return
    
    # Note: Full cross-reference preservation requires complex tracking
    # This is a placeholder that logs the references found
    logger.info(f"Preserving {len(references)} cross-references")
    
    for ref in references:
        ref_type = ref.get("type", "unknown")
        target = ref.get("target", "?")
        logger.debug(f"Cross-reference: {ref_type} → {target}")


def _preserve_visual_slots_in_output(
    output_doc, 
    visual_slots: List[Dict[str, Any]]
) -> None:
    """Preserve visual slot placeholders in output without rendering."""
    if not visual_slots:
        return
    
    for slot in visual_slots:
        slot_type = slot.get("slot_type", "unknown")
        description = slot.get("description", "Visual")
        
        if slot_type == "visual_placeholder":
            output_doc.add_paragraph(f"[VISUAL: {description}]")
        elif slot_type == "diagram_token":
            output_doc.add_paragraph(f"[Diagram: {description}]")
        else:
            output_doc.add_paragraph(f"[Visual: {description}]")
        
        output_doc.add_paragraph()  # Add spacing


def _apply_run_properties_to_output(run, properties: Dict) -> None:
    """Apply exact run properties to output document run."""
    try:
        from docx.shared import Pt, RGBColor

        if "name" in properties and properties["name"]:
            run.font.name = properties["name"]

        if "size_pt" in properties and properties["size_pt"]:
            run.font.size = Pt(properties["size_pt"])

        run.font.bold = bool(properties.get("bold"))
        run.font.italic = bool(properties.get("italic"))
        run.font.underline = bool(properties.get("underline"))

        color = properties.get("color")
        if color and isinstance(color, (list, tuple)) and len(color) == 3:
            run.font.color.rgb = RGBColor(int(color[0]), int(color[1]), int(color[2]))

        logger.debug("Applied run properties to output")
    except Exception as e:
        logger.warning(f"Failed to apply run properties: {e}")


def _apply_paragraph_properties_to_output(paragraph, style_preferences: Dict) -> None:
    """Apply donor paragraph spacing and indent from style_preferences to an output paragraph."""
    try:
        from docx.shared import Pt, Inches

        fmt = paragraph.paragraph_format

        space_before = style_preferences.get("space_before_pt")
        if space_before is not None:
            fmt.space_before = Pt(space_before)

        space_after = style_preferences.get("space_after_pt")
        if space_after is not None:
            fmt.space_after = Pt(space_after)

        line_spacing = style_preferences.get("line_spacing_pt")
        if line_spacing is not None:
            fmt.line_spacing = Pt(line_spacing)

        indent_left = style_preferences.get("indent_left_inches")
        if indent_left is not None:
            fmt.left_indent = Inches(indent_left)

        logger.debug("Applied paragraph properties to output")
    except Exception as e:
        logger.warning(f"Failed to apply paragraph properties: {e}")


# ======================================================================================
# PHASE 7: QUALITY VALIDATION & LEARNING
# ======================================================================================

def _score_output_quality(
    blueprint: Dict,
    rendered_path: str,
    original_donor_path: str
) -> Dict[str, Any]:
    """
    Score the quality of the rendered output.
    
    METRICS:
        - Layout fidelity (margins, columns preserved correctly)
        - Font fidelity (same font family/size)
        - Branding fidelity (logo, header, footer)
        - Content completeness (all questions present)
        - Answer space appropriateness
    
    RETURNS:
        {
            "overall_score": float (0-100),
            "layout_score": float,
            "font_score": float,
            "branding_score": float,
            "completeness_score": float,
            "warnings": List[str],
            "suggestions": List[str]
        }
    """
    if not DOCX_AVAILABLE:
        logger.warning("Cannot score output quality: python-docx not available")
        return {
            "overall_score": 0,
            "layout_score": 0,
            "font_score": 0,
            "branding_score": 0,
            "completeness_score": 0,
            "warnings": ["python-docx not available for quality scoring"],
            "suggestions": ["Install python-docx for quality validation"]
        }
    
    try:
        from docx import Document
        
        scores = {}
        warnings = []
        suggestions = []
        
        # 1. Content completeness score
        question_blocks = blueprint.get("question_blocks", [])
        expected_questions = len(question_blocks)
        
        completeness_score = 100
        if expected_questions == 0:
            completeness_score = 0
            warnings.append("No questions found in blueprint")
            suggestions.append("Add questions to your draft")
        else:
            # Try to count questions in rendered document
            try:
                rendered_doc = Document(rendered_path)
                rendered_questions = 0
                for para in rendered_doc.paragraphs:
                    if re.match(r"^(Question|Q\.?)\s+\d+", para.text, re.I):
                        rendered_questions += 1
                
                if rendered_questions < expected_questions:
                    completeness_score = int((rendered_questions / expected_questions) * 100)
                    warnings.append(f"Only {rendered_questions}/{expected_questions} questions found in output")
                    suggestions.append("Check question numbering in your draft")
            except Exception as e:
                logger.warning(f"Could not verify content completeness: {e}")
        
        scores["completeness_score"] = completeness_score

        # 1.5 Subject leakage detection (prevents donor subject contamination)
        donor_subject_terms = [
            "mathematics", "maths", "math", "algebra", "calculus", "geometry",
            "biology", "chemistry", "physics", "science",
            "history", "geography", "economics", "business", "accounting"
        ]
        try:
            from docx import Document
            rendered_doc = Document(rendered_path)
            full_text = " ".join([p.text for p in rendered_doc.paragraphs]).lower()
            leakage_terms = [term for term in donor_subject_terms if term in full_text]
            if leakage_terms:
                warnings.append(f"Potential donor subject leakage detected: {', '.join(leakage_terms[:5])}")
                suggestions.append("Review output for unintended donor subject content (e.g., math terms in biology worksheet).")
                # Penalize quality score for leakage
                completeness_score = max(0, completeness_score - 20)
        except Exception as e:
            logger.warning(f"Subject leakage detection failed: {e}")
        
        # 2. Layout score (margins, columns)
        layout_score = 70  # Default moderate score
        layout_warnings = []
        
        page_setup = blueprint.get("page_setup", {})
        columns_allowed = blueprint.get("appropriateness_decisions", {}).get("columns", {}).get("apply", False)
        
        if columns_allowed:
            layout_score = 90
            suggestions.append("Columns applied based on question structure analysis")
        else:
            layout_score = 85
            suggestions.append("Single column layout used (recommended for this document type)")
        
        if page_setup.get("orientation") == "landscape":
            suggestions.append("Landscape orientation used - ensure printing setup matches")
        
        scores["layout_score"] = layout_score
        
        # 3. Font score (if donor had font info)
        font_score = 75  # Default
        style_prefs = blueprint.get("style_preferences", {})
        donor_font = style_prefs.get("font_family", "")
        
        if donor_font and donor_font != "Unknown":
            font_score = 85
            suggestions.append(f"Using donor font: {donor_font}")
        else:
            warnings.append("No donor font information available - using defaults")
            suggestions.append("Upload a DOCX donor file for better font preservation")
        
        scores["font_score"] = font_score
        
        # 4. Branding score (header, footer)
        branding_score = 70
        header_block = blueprint.get("header_block", {})
        institution_name = header_block.get("institution_name", "")
        
        if institution_name:
            branding_score = 90
            suggestions.append(f"Institution branding preserved: {institution_name}")
        else:
            warnings.append("No institution branding detected in donor")
            suggestions.append("Ensure donor document contains school name in header/footer")
        
        if header_block.get("show_header", False):
            branding_score = min(branding_score + 5, 100)
        
        scores["branding_score"] = branding_score
        
        # 5. Answer space appropriateness score
        answer_score = 80
        question_blocks = blueprint.get("question_blocks", [])
        
        answer_styles_used = set()
        for q in question_blocks:
            style = q.get("answer_style", "short_response")
            answer_styles_used.add(style)
        
        if "paragraph_response" in answer_styles_used:
            answer_score = 85
            suggestions.append("Paragraph answer spaces included for explanation questions")
        
        if "show_working" in answer_styles_used:
            answer_score = 85
            suggestions.append("Show-working spaces included for calculation questions")
        
        if "label_diagram" in answer_styles_used:
            answer_score = 90
            suggestions.append("Diagram labeling spaces included")
        
        scores["answer_score"] = answer_score
        
        # Calculate overall score (weighted average)
        weights = {
            "completeness_score": 0.30,
            "layout_score": 0.25,
            "font_score": 0.15,
            "branding_score": 0.20,
            "answer_score": 0.10
        }
        
        overall_score = 0
        for metric, weight in weights.items():
            overall_score += scores.get(metric, 0) * weight
        
        overall_score = int(overall_score)
        
        # Final suggestions based on overall score
        if overall_score < 60:
            suggestions.insert(0, "Consider re-analyzing with a different donor document")
        elif overall_score < 80:
            suggestions.insert(0, "Good quality - for best results, ensure donor has clear branding")
        else:
            suggestions.insert(0, "Excellent quality - document is print-ready")
        
        logger.info(f"Quality score: {overall_score}/100 ({len(warnings)} warnings)")
        
        return {
            "overall_score": overall_score,
            "layout_score": scores.get("layout_score", 0),
            "font_score": scores.get("font_score", 0),
            "branding_score": scores.get("branding_score", 0),
            "completeness_score": scores.get("completeness_score", 0),
            "answer_score": scores.get("answer_score", 0),
            "warnings": warnings[:10],
            "suggestions": suggestions[:10]
        }
        
    except Exception as e:
        logger.error(f"Quality scoring failed: {e}")
        return {
            "overall_score": 0,
            "layout_score": 0,
            "font_score": 0,
            "branding_score": 0,
            "completeness_score": 0,
            "answer_score": 0,
            "warnings": [f"Quality scoring failed: {str(e)}"],
            "suggestions": ["Run the template application again"]
        }


def _record_teacher_correction(
    template_id: str,
    decision_name: str,
    engine_decision: bool,
    teacher_decision: bool,
    draft_content: str,
    document_type: Optional[str] = None,
    subject: Optional[str] = None,
) -> None:
    """
    Record teacher overrides for future learning.
    
    PURPOSE:
        - Improve confidence scoring over time
        - Build institution-specific rules
        - Detect edge cases automatically
    """
    try:
        # Create correction record
        correction = {
            "template_id": template_id,
            "decision_name": decision_name,
            "engine_decision": engine_decision,
            "teacher_decision": teacher_decision,
            "draft_preview": draft_content[:500] if draft_content else "",
            "document_type": document_type or "",
            "subject": subject or "",
            "created_at": datetime.now().isoformat(),
            "applied_count": 1
        }
        
        # Try to store in database (if table exists)
        try:
            result = supabase.table("template_corrections").insert(correction).execute()
            logger.info(f"Recorded teacher correction for {decision_name}: engine={engine_decision} → teacher={teacher_decision}")
        except Exception as db_error:
            logger.warning(f"Could not save correction to DB: {db_error}")
            # Fallback: save to local file
            corrections_file = os.path.join(tempfile.gettempdir(), "template_corrections.jsonl")
            with open(corrections_file, "a") as f:
                f.write(json.dumps(correction) + "\n")
            logger.info(f"Correction saved to {corrections_file}")
        
        # Update aggregated learning rules
        _update_learning_rules_from_correction(correction)
        
    except Exception as e:
        logger.error(f"Failed to record teacher correction: {e}")


def _update_learning_rules_from_correction(correction: Dict) -> None:
    """Update aggregated learning rules incrementally."""
    doc_type = correction.get("document_type", "")
    feature = correction.get("decision_name", "")
    teacher_decision = correction.get("teacher_decision")
    if not doc_type or not feature or teacher_decision is None:
        return
    rule_key = f"{doc_type}|{feature}"
    try:
        # Try to fetch existing rule
        result = supabase.table("learning_rules").select("*").eq("rule_key", rule_key).execute()
        rows = result.data or []
        if rows:
            rule = rows[0]
            total = rule.get("total_decisions", 0) + 1
            override_count = rule.get("override_count", 0) + (1 if teacher_decision else 0)
            suggested_apply = (override_count / total) > 0.8  # adopt if override rate > 80%
            supabase.table("learning_rules").update({
                "total_decisions": total,
                "override_count": override_count,
                "suggested_apply": suggested_apply,
                "updated_at": datetime.now().isoformat()
            }).eq("rule_key", rule_key).execute()
        else:
            supabase.table("learning_rules").insert({
                "rule_key": rule_key,
                "document_type": doc_type,
                "feature_name": feature,
                "suggested_apply": teacher_decision,
                "total_decisions": 1,
                "override_count": 1 if teacher_decision else 0,
            }).execute()
    except Exception as e:
        logger.debug(f"Could not update learning rules: {e}")


def _get_learned_rule(document_type: DocumentType, feature_name: str) -> Optional[Dict]:
    """Return a learned rule override if it exists and is strong enough."""
    doc_type_str = document_type.value if hasattr(document_type, "value") else str(document_type)
    rule_key = f"{doc_type_str}|{feature_name}"
    try:
        result = supabase.table("learning_rules").select("*").eq("rule_key", rule_key).execute()
        rows = result.data or []
        if rows:
            rule = rows[0]
            total = rule.get("total_decisions", 0)
            if total >= 5 and rule.get("override_count", 0) / total > 0.8:
                return {
                    "apply": rule.get("suggested_apply", False),
                    "confidence": 0.9,
                    "reason": f"Learned from {total} previous teacher decisions: {'apply' if rule['suggested_apply'] else 'suppress'} {feature_name} for {doc_type_str}."
                }
    except Exception:
        pass
    return None


def _apply_learned_rules(template_id: str, decisions: Dict) -> Dict:
    """
    Apply previously learned teacher corrections to current decisions.
    
    PRIORITY:
        - Teacher corrections from same template > default rules
        - Corrections from similar document types > generic rules
    """
    if not decisions:
        return decisions
    
    try:
        # Try to load corrections from database
        corrections = []
        
        try:
            result = supabase.table("template_corrections").select("*").eq("template_id", template_id).execute()
            rows = getattr(result, "data", None) or []
            corrections = rows
        except Exception:
            # Fallback: load from local file
            corrections_file = os.path.join(tempfile.gettempdir(), "template_corrections.jsonl")
            if os.path.exists(corrections_file):
                with open(corrections_file, "r") as f:
                    for line in f:
                        try:
                            corr = json.loads(line.strip())
                            if corr.get("template_id") == template_id:
                                corrections.append(corr)
                        except:
                            pass
        
        if not corrections:
            return decisions
        
        # Apply corrections to decisions
        modified_decisions = decisions.copy()
        
        for corr in corrections:
            decision_name = corr.get("decision_name")
            teacher_decision = corr.get("teacher_decision", False)
            
            if decision_name in modified_decisions:
                old_decision = modified_decisions[decision_name]
                
                # Override with teacher's decision
                modified_decisions[decision_name] = {
                    "apply": teacher_decision,
                    "confidence": 0.95,  # High confidence for learned rule
                    "reason": f"Learned from teacher correction: {old_decision.get('reason', 'unknown')}",
                    "source": "learned",
                    "requires_confirmation": False
                }
                
                logger.info(f"Applied learned rule: {decision_name} → {teacher_decision}")
        
        return modified_decisions
        
    except Exception as e:
        logger.warning(f"Failed to apply learned rules: {e}")
        return decisions


# Optional: Helper to initialize corrections table in Supabase
def _init_corrections_table() -> bool:
    """
    Initialize the template_corrections table in Supabase if it doesn't exist.
    Returns True if successful or table exists.
    """
    try:
        # This is a placeholder - actual table creation would need to be done
        # via Supabase dashboard or migration script
        logger.info("Corrections table should be created in Supabase with: "
                    "CREATE TABLE template_corrections (id UUID PRIMARY KEY DEFAULT gen_random_uuid(), "
                    "template_id TEXT, decision_name TEXT, engine_decision BOOLEAN, "
                    "teacher_decision BOOLEAN, draft_preview TEXT, created_at TIMESTAMP, applied_count INT)")
        return True
    except Exception as e:
        logger.error(f"Failed to initialize corrections table: {e}")
        return False


# Optional: Helper to get correction statistics
def _get_correction_stats(template_id: str) -> Dict[str, Any]:
    """
    Get statistics about teacher corrections for a template.
    
    RETURNS:
        {
            "total_corrections": int,
            "most_overridden": List[str],
            "agreement_rate": float  # How often teacher agreed with engine
        }
    """
    try:
        corrections = []
        
        try:
            result = supabase.table("template_corrections").select("*").eq("template_id", template_id).execute()
            corrections = getattr(result, "data", None) or []
        except Exception:
            corrections_file = os.path.join(tempfile.gettempdir(), "template_corrections.jsonl")
            if os.path.exists(corrections_file):
                with open(corrections_file, "r") as f:
                    for line in f:
                        try:
                            corr = json.loads(line.strip())
                            if corr.get("template_id") == template_id:
                                corrections.append(corr)
                        except:
                            pass
        
        if not corrections:
            return {
                "total_corrections": 0,
                "most_overridden": [],
                "agreement_rate": 1.0
            }
        
        # Count overrides by decision
        override_counts = {}
        agreements = 0
        
        for corr in corrections:
            decision = corr.get("decision_name", "unknown")
            engine = corr.get("engine_decision", False)
            teacher = corr.get("teacher_decision", False)
            
            override_counts[decision] = override_counts.get(decision, 0) + 1
            
            if engine == teacher:
                agreements += 1
        
        # Sort by frequency
        most_overridden = sorted(override_counts.items(), key=lambda x: x[1], reverse=True)
        most_overridden = [f"{d} ({c}x)" for d, c in most_overridden[:5]]
        
        agreement_rate = agreements / len(corrections) if corrections else 1.0
        
        return {
            "total_corrections": len(corrections),
            "most_overridden": most_overridden,
            "agreement_rate": agreement_rate
        }
        
    except Exception as e:
        logger.error(f"Failed to get correction stats: {e}")
        return {
            "total_corrections": 0,
            "most_overridden": [],
            "agreement_rate": 1.0
        }


# ======================================================================================
# PUBLIC APP ENTRY POINTS (STABLE INTERFACE - DO NOT CHANGE SIGNATURES)
# ======================================================================================

def action_analyze_template_upload(upload_obj, model_name: str = "gpt-5"):
    """
    Analyze uploaded donor file.
    
    Returns:
        status_message: str
        clean_preview: str
        pretty_profile_json: str
        packed_bundle: str
        button_update: gr.update
    """
    file_path, filename = _resolve_upload_path(upload_obj)
    
    if not file_path or not filename:
        return (
            "⚠️ No template file was provided.",
            "",
            "{}",
            "",
            gr.update(interactive=False),
        )
    
    ext = _safe_ext(filename)
    if ext not in SUPPORTED_TEMPLATE_EXTS:
        return (
            f"⚠️ Unsupported template type: {ext or 'unknown'}",
            "",
            "{}",
            "",
            gr.update(interactive=False),
        )
    
    # Extract using appropriate analyzer
    if ext == ".docx":
        if DOCX_AVAILABLE:
            raw_profile = _extract_donor_with_styles(file_path)
        else:
            raw_profile = _extract_donor_fallback(file_path)
    elif ext == ".pdf":
        raw_profile = analyze_pdf_template(file_path)
    elif ext in {".ppt", ".pptx"}:
        raw_profile = analyze_pptx_template(file_path)
    else:
        raw_profile = {}

    # --------------------------------------
    # TEMP DEBUG: expose raw donor identity signals
    # so we can inspect the real live upload path
    # inside the JSON preview.
    # --------------------------------------
    try:
        raw_profile = raw_profile or {}
        if not isinstance(raw_profile, dict):
            raw_profile = {}

        raw_profile["_debug_identity_from_raw"] = _extract_institution_identity(raw_profile)
        raw_profile["_debug_title_block_signals"] = raw_profile.get("title_block_signals", {})
        raw_profile["_debug_header_footer"] = raw_profile.get("header_footer", {})
        raw_profile["_debug_extracted_images"] = raw_profile.get("extracted_images", [])
        raw_profile["_debug_suppressed_images"] = raw_profile.get("suppressed_images", [])
        raw_profile["_debug_paragraphs_head"] = (raw_profile.get("paragraphs", []) or [])[:25]
    except Exception as e:
        raw_profile["_debug_identity_error"] = f"{type(e).__name__}: {e}"
    
    # Normalize the profile
    normalized_profile = _normalize_template_profile(raw_profile)
    
    # Build human-readable preview
    clean_preview = _build_clean_template_md_from_profile(normalized_profile)
    
    # Pretty JSON for display
    safe_normalized_profile = _make_profile_json_safe(normalized_profile)
    pretty_profile_json = json.dumps(safe_normalized_profile, ensure_ascii=False, indent=2)
    
    # Pack bundle for storage
    packed_bundle = pack_template_bundle(clean_preview, normalized_profile)

    # --------------------------------------
    # TEMP DEBUG: SAVE FIRST BRANDING IMAGE
    # --------------------------------------
    try:
        raw_core = (normalized_profile.get("normalized_profile", {}) or {}).get("raw", {}) or {}
        imgs = raw_core.get("extracted_images", []) or []
        debug_saved = False

        for img in imgs:
            if img.get("image_role") == "branding" and img.get("binary_data"):
                blob = img.get("binary_data")

                if isinstance(blob, str):
                    try:
                        import base64
                        raw = base64.b64decode(blob)
                    except Exception:
                        raw = blob.encode("utf-8")
                else:
                    raw = bytes(blob)

                import os
                safe_name = os.path.splitext(os.path.basename(filename))[0]
                out_path = f"/tmp/debug_branding_{safe_name}.bin"

                with open(out_path, "wb") as f:
                    f.write(raw)

                clean_preview += f"\n\n[DEBUG] Branding image saved to: {out_path}\n"
                debug_saved = True
                break

        if not debug_saved:
            clean_preview += "\n\n[DEBUG] No branding image found.\n"

    except Exception as e:
        clean_preview += f"\n\n[DEBUG] Branding image save failed: {type(e).__name__}: {e}\n"

    except Exception as e:
        clean_preview += f"\n\n[DEBUG] Branding image save failed: {type(e).__name__}: {e}\n"
    
    source_type = normalized_profile.get("source_type", "unknown").upper()
    status = (
        f"✅ Donor analyzed as institution style profile: {filename} ({source_type}). "
        f"Layout features classified for intelligent application."
    )
    
    return (
        status,
        clean_preview,
        pretty_profile_json,
        packed_bundle,
        gr.update(interactive=True),
    )


def save_template_record(
    sess,
    template_name: str,
    template_description: str,
    category: str,
    template_bundle: str,
    source_file=None,
    share_template: bool = False,
):
    """
    Save template to database.
    
    Returns status message string.
    """
    try:
        if not sess:
            return "⚠️ You must be logged in to save a template."

        access_token, refresh_token, user_id, err = _require_session(sess)
        if err:
            return f"⚠️ {err}"

        supabase.postgrest.auth(access_token)
        
        title = (template_name or "").strip()
        description = (template_description or "").strip()
        category = (category or "").strip() or "Custom"
        bundle = (template_bundle or "").strip()
        
        if not title:
            return "⚠️ Please enter a template name."
        
        if not bundle:
            return "⚠️ No analyzed template bundle is available to save."
        
        # Validate and repack the bundle
        clean_md, profile = unpack_template_bundle(bundle)
        
        # If unpack returned empty profile, try to treat bundle as direct profile
        if not profile or not isinstance(profile, dict):
            try:
                direct = json.loads(bundle)
                if isinstance(direct, dict):
                    profile = direct
                    clean_md = _build_clean_template_md_from_profile(profile)
            except:
                pass
        
        # Final validation - create minimal profile if needed
        if not profile or not isinstance(profile, dict):
            profile = {
                "source_type": "manual",
                "source_filename": title,
                "template_title": title,
                "style_preferences": {},
                "layout_features": [
                    {
                        "feature_name": "margins",
                        "category": FeatureCategory.ALWAYS_INHERIT,
                        "extracted_value": {"fallback": True}
                    }
                ],
                "page_setup": {"columns": {"count": 1}},
                "header_footer": {},
                "text_stats": {},
                "style_notes": ["Manually created from preview"]
            }
        
        # Ensure minimum required fields
        if "source_type" not in profile:
            profile["source_type"] = "unknown"
        if "template_title" not in profile:
            profile["template_title"] = title
        if "layout_features" not in profile:
            profile["layout_features"] = []
        if "page_setup" not in profile:
            profile["page_setup"] = {"columns": {"count": 1}}
        if "style_preferences" not in profile:
            profile["style_preferences"] = {}
        if "header_footer" not in profile:
            profile["header_footer"] = {}
        
        # Create proper preview if clean_md is empty
        if not clean_md:
            clean_md = _build_clean_template_md_from_profile(profile)
        
        # Create proper bundle
        proper_bundle = pack_template_bundle(clean_md, profile)
        
        # Verify the bundle is valid
        test_clean, test_profile = unpack_template_bundle(proper_bundle)
        if not test_profile:
            return "⚠️ Failed to create valid template bundle. Please re-analyze."
        
        payload = {
            "user_id": user_id,
            "name": title,
            "description": description,
            "category": category,
            "template_md": proper_bundle,
            "is_public": bool(share_template),
        }
        
        res = supabase.table("templates").insert(payload).execute()
        rows = getattr(res, "data", None) or []
        
        if rows:
            source_type = profile.get("source_type", "unknown")
            suffix = f" ({source_type} donor)" if source_type != "unknown" else ""
            return f"✅ Institution style profile saved: **{title}**{suffix}"
        
        return f"⚠️ Template save may have failed for **{title}**."
    
    except Exception as e:
        logger.error(f"Save template failed: {e}")
        return f"⚠️ Saving template failed: {type(e).__name__}: {e}"


def load_template_bundle_from_db(template_id: str, session_state=None) -> Tuple[str, Dict[str, Any], str]:
    """
    Load template from database.
    
    Returns:
        preview_text: str
        profile_dict: Dict
        status_message: str
    """
    try:
        if not (template_id or "").strip():
            return "", {}, "⚠️ No template selected."
        
        query = supabase.table("templates").select(
            "id,name,template_md,template_ppt,description,category"
        ).eq("id", template_id)

        if session_state is not None:
            access_token, refresh_token, user_id, err = _require_session(session_state)
            if err:
                return "", {}, f"⚠️ {err}"
            query = query.eq("user_id", user_id)

        res = query.limit(1).execute()
        rows = getattr(res, "data", None) or []
        
        if not rows:
            return "", {}, "⚠️ Template not found."
        
        row = rows[0]
        title = (row.get("name") or "Template").strip()
        raw_template_md = row.get("template_md") or ""
        
        # Unpack using the robust unpacker
        clean_md, profile = unpack_template_bundle(raw_template_md)
        
        # If no profile found, try emergency extraction
        if not profile or not isinstance(profile, dict) or len(profile.keys()) == 0:
            # Try to find any JSON in the raw content
            try:
                json_pattern = r'\{[^{}]*"source_type"[^{}]*\}'
                json_match = re.search(json_pattern, raw_template_md, re.DOTALL)
                if json_match:
                    profile = json.loads(json_match.group(1))
                    logger.debug("Emergency extraction found profile")
            except:
                pass
        
        # If still no profile, create minimal one from template name
        if not profile or not isinstance(profile, dict) or len(profile.keys()) == 0:
            profile = {
                "source_type": "legacy",
                "source_filename": title,
                "template_title": title,
                "style_preferences": {
                    "font_family": "Not extracted (legacy template)",
                    "font_size_pt": "Unknown"
                },
                "layout_features": [
                    {
                        "feature_name": "margins",
                        "category": FeatureCategory.ALWAYS_INHERIT,
                        "extracted_value": {"fallback": True}
                    }
                ],
                "page_setup": {
                    "left_margin_inches": 1.0,
                    "right_margin_inches": 1.0,
                    "top_margin_inches": 1.0,
                    "bottom_margin_inches": 1.0,
                    "columns": {"count": 1}
                },
                "header_footer": {},
                "text_stats": {},
                "style_notes": ["Loaded from legacy template format - re-analyze donor for full features"]
            }
            logger.debug("Created minimal fallback profile")
        
        # Normalize the profile
        normalized_profile = _normalize_template_profile(profile)
        
        # Build preview text
        preview_text = _build_clean_template_md_from_profile(normalized_profile)
        
        # Determine status message
        is_legacy = profile.get("source_type") == "legacy" or "fallback" in str(profile.get("style_notes", []))
        if is_legacy:
            status = f"⚠️ Loaded legacy template: **{title}** (re-analyze donor for full style extraction)"
        else:
            status = f"✅ Loaded style profile: **{title}**"
        
        return preview_text, normalized_profile, status
    
    except Exception as e:
        logger.error(f"Load template failed: {e}")
        return "", {}, f"⚠️ Loading template failed: {type(e).__name__}: {e}"


def apply_template_to_draft(
    current_draft_md: str,
    template_bundle: str,
    model_name: str = "gpt-5",
    year_level: str = "",
    subject: str = "",
):
    """
    Apply template to draft.
    
    INTELLIGENT PIPELINE:
        1. Unpack template bundle
        2. Normalize profile
        3. Analyze draft (document type, structure, hints)
        4. Detect education level
        5. Apply appropriateness rules (with confidence)
        6. Apply learned rules from past corrections
        7. Build blueprint
        8. Render to DOCX
        9. Score output quality
        10. Return output path or confirmation prompt if uncertain
    
    Returns either:
        - DOCX file path (if rendering successful)
        - Confirmation prompt string (if uncertain)
        - Markdown string (if fallback)
        - Error message string
    """
    try:
        draft_md = (current_draft_md or "").strip()
        raw_bundle = (template_bundle or "").strip()
        
        if not draft_md:
            return "⚠️ No current draft content was provided."
        
        if not raw_bundle:
            return "⚠️ No template bundle was provided."
        
        # ============================================================
        # Step 1: Unpack template bundle
        # ============================================================
        clean_md, profile = unpack_template_bundle(raw_bundle)
        
        if not profile or not isinstance(profile, dict):
            # Try to parse as direct JSON
            try:
                direct = json.loads(raw_bundle)
                if isinstance(direct, dict):
                    profile = direct
            except:
                pass
        
        if not profile:
            return "⚠️ Could not load template profile. Please re-analyze the donor document."
        
        # ============================================================
        # Step 2: Normalize profile
        # ============================================================
        normalized_profile = _normalize_template_profile(profile)
        
        # ============================================================
        # Step 3: Analyze draft content
        # ============================================================
        draft_model = analyze_draft_content(
            draft_md,
            year_level=year_level or "",
            subject=subject or "",
        )
        
        # ============================================================
        # Step 4: Get education level and document type
        # ============================================================
        education_level = draft_model.get("education_level", EducationLevel.UNKNOWN)
        document_type = draft_model.get("document_type", DocumentType.CUSTOM)
        
        # ============================================================
        # Step 5: Apply appropriateness rules
        # ============================================================
        layout_features = normalized_profile.get("layout_features", [])
        questions = draft_model.get("question_map", {}).get("questions", [])
        user_hints = draft_model.get("user_layout_hints", {})
        
        appropriateness_result = _apply_appropriateness_rules(
        layout_features=layout_features,
        draft_questions=questions,
        education_level=education_level,
        document_type=document_type,
        user_hints=user_hints,
        layout_needs=draft_model.get("layout_needs", {})
    )
        
        # ============================================================
        # Step 6: Check if teacher confirmation is needed
        # ============================================================
        uncertain_decisions = {
            name: decision for name, decision in appropriateness_result.decisions.items()
            if decision.requires_teacher_confirmation
        }
        
        # If there are uncertain decisions and no user hints override them,
        # return a confirmation prompt instead of rendering
        if uncertain_decisions and not user_hints:
            confirmation_prompt = _build_teacher_confirmation_prompt(uncertain_decisions)
            return confirmation_prompt
        
        # ============================================================
        # Step 7: Build document blueprint
        # ============================================================
        blueprint = build_document_blueprint(
            template_profile=normalized_profile,
            draft_model=draft_model,
            year_level=year_level or "",
            subject=subject or "",
            appropriateness_result=appropriateness_result
        )
        
        # Add appropriateness decisions to blueprint
        blueprint["appropriateness_decisions"] = {
            name: {
                "apply": decision.apply,
                "confidence": decision.confidence,
                "reason": decision.reason,
                "source": decision.source
            }
            for name, decision in appropriateness_result.decisions.items()
        }
        
        # ============================================================
        # Step 8: Render to DOCX
        # ============================================================
        output_path = render_blueprint_to_docx(blueprint, normalized_profile)
        
        if not output_path or not os.path.exists(output_path):
            # Fallback to markdown
            markdown_output = render_blueprint_to_markdown(blueprint)
            return markdown_output
        
        # ============================================================
        # Step 9: Score output quality (passive, doesn't block)
        # ============================================================
        try:
            quality_score = _score_output_quality(blueprint, output_path, "")
            logger.info(f"Quality score: {quality_score.get('overall_score', 0)}/100")
            if quality_score.get('overall_score', 100) < 60:
                logger.warning(f"Low quality score warnings: {quality_score.get('warnings', [])}")
        except Exception as qe:
            logger.warning(f"Quality scoring failed: {qe}")
        
        # ============================================================
        # Step 10: Return output path
        # ============================================================
        return output_path
    
    except NotImplementedError as e:
        return f"⚠️ Feature not yet implemented: {e}"
    
    except Exception as e:
        logger.error(f"Template application failed: {e}")
        return f"⚠️ Template application failed: {type(e).__name__}: {e}"


# ======================================================================================
# TEACHER CONFIRMATION HANDLER (for uncertain decisions)
# ======================================================================================

def handle_teacher_confirmation(
    confirmation_response: str,
    uncertain_decisions: Dict[str, Decision],
    draft_md: str,
    template_bundle: str
) -> str:
    """
    Handle teacher's response to uncertainty confirmation prompt.
    
    This is called after the engine returns a confirmation prompt.
    Teacher's choice (A, B, or C) is used to override decisions.
    
    Expected response formats:
        - "A" or "a" → Accept engine's recommendation
        - "B" or "b" → Override - Apply the feature
        - "C" or "c" → Override - Do not apply the feature
        - "D" or "d" → Keep donor's original layout
    
    Returns:
        - DOCX file path after applying teacher's choice
        - Error message if something fails
    """
    try:
        if not confirmation_response:
            return "⚠️ No confirmation response provided."
        
        if not uncertain_decisions:
            return "⚠️ No uncertain decisions to resolve."
        
        if not draft_md:
            return "⚠️ No draft content provided."
        
        if not template_bundle:
            return "⚠️ No template bundle provided."
        
        # ============================================================
        # Step 1: Parse teacher's choice
        # ============================================================
        response = (confirmation_response or "").strip().upper()

        # Support multi-feature override format:
        # Example: "columns=B, tables=C"
        feature_overrides = {}

        if "=" in response:
            parts = [p.strip() for p in response.split(",")]
            for part in parts:
                if "=" in part:
                    name, choice = part.split("=")
                    feature_overrides[name.strip()] = choice.strip().upper()
        
        # Map response to decision overrides
        # Default: accept engine's recommendation
        apply_override = None  # None means use engine's recommendation
        
        if response in ["B", "OVERRIDE_APPLY", "APPLY"]:
            apply_override = True
        elif response in ["C", "OVERRIDE_REJECT", "REJECT"]:
            apply_override = False
        elif response in ["D", "KEEP_DONOR", "DONOR"]:
            apply_override = "donor"  # Special: keep donor's original layout
        
        # A or anything else means accept engine's recommendation (no override)
        
        # ============================================================
        # Step 2: Unpack template bundle
        # ============================================================
        clean_md, profile = unpack_template_bundle(template_bundle)
        
        if not profile or not isinstance(profile, dict):
            try:
                direct = json.loads(template_bundle)
                if isinstance(direct, dict):
                    profile = direct
            except:
                pass
        
        if not profile:
            return "⚠️ Could not load template profile. Please re-analyze the donor document."
        
        # ============================================================
        # Step 3: Normalize profile
        # ============================================================
        normalized_profile = _normalize_template_profile(profile)
        
        # ============================================================
        # Step 4: Analyze draft content
        # ============================================================
        draft_model = analyze_draft_content(draft_md)
        
        # ============================================================
        # Step 5: Apply overrides to uncertain decisions
        # ============================================================
        # Create user hints from teacher's response
        user_hints = {}
        
        for decision_name, decision in uncertain_decisions.items():

            # Check if feature-specific override exists
            feature_choice = feature_overrides.get(decision_name)

            if feature_choice:
                if feature_choice in ["B"]:
                    user_hints[decision_name] = True
                elif feature_choice in ["C"]:
                    user_hints[decision_name] = False
                elif feature_choice in ["D"]:
                    donor_features = normalized_profile.get("layout_features", [])
                    donor_value = any(f.get("feature_name") == decision_name for f in donor_features)
                    user_hints[decision_name] = donor_value
                else:
                    user_hints[decision_name] = decision.apply

                continue

            # Fallback to global override
            if apply_override is not None:
                if apply_override == "donor":
                    donor_features = normalized_profile.get("layout_features", [])
                    donor_value = any(f.get("feature_name") == decision_name for f in donor_features)
                    user_hints[decision_name] = donor_value
                else:
                    user_hints[decision_name] = apply_override
            else:
                user_hints[decision_name] = decision.apply
        
        # Add user hints to draft model
        draft_model["user_layout_hints"] = user_hints
        
        # ============================================================
        # Step 6: Re-apply appropriateness rules with user hints
        # ============================================================
        layout_features = normalized_profile.get("layout_features", [])
        questions = draft_model.get("question_map", {}).get("questions", [])
        education_level = draft_model.get("education_level", EducationLevel.UNKNOWN)
        document_type = draft_model.get("document_type", DocumentType.CUSTOM)
        
        appropriateness_result = _apply_appropriateness_rules(
        layout_features=layout_features,
        draft_questions=questions,
        education_level=education_level,
        document_type=document_type,
        user_hints=user_hints,
        layout_needs=draft_model.get("layout_needs", {})
    )
        
        # ============================================================
        # Step 7: Record the correction for future learning
        # ============================================================
        template_id = "temp_" + str(uuid.uuid4())[:8]
        
        # Extract document type and subject for learning
        doc_type_obj = draft_model.get("document_type")
        doc_type_str = doc_type_obj.value if hasattr(doc_type_obj, "value") else str(doc_type_obj)
        subject_str = draft_model.get("subject", "")

        for decision_name, decision in uncertain_decisions.items():
            engine_decision = decision.apply
            teacher_decision = user_hints.get(decision_name, engine_decision)

            _record_teacher_correction(
                template_id=template_id,
                decision_name=decision_name,
                engine_decision=engine_decision,
                teacher_decision=teacher_decision,
                draft_content=draft_md[:500],
                document_type=doc_type_str,
                subject=subject_str
            )
        
        # ============================================================
        # Step 8: Build blueprint with confirmed decisions
        # ============================================================
        blueprint = build_document_blueprint(
            template_profile=normalized_profile,
            draft_model=draft_model,
            year_level=draft_model.get("year_level", ""),
            subject=draft_model.get("subject", ""),
            appropriateness_result=appropriateness_result
        )
        
        # Add appropriateness decisions to blueprint
        blueprint["appropriateness_decisions"] = {
            name: {
                "apply": decision.apply,
                "confidence": decision.confidence,
                "reason": decision.reason,
                "source": decision.source
            }
            for name, decision in appropriateness_result.decisions.items()
        }
        
        # ============================================================
        # Step 9: Render to DOCX
        # ============================================================
        output_path = render_blueprint_to_docx(blueprint, normalized_profile)
        
        if not output_path or not os.path.exists(output_path):
            # Fallback to markdown
            markdown_output = render_blueprint_to_markdown(blueprint)
            return markdown_output
        
        # ============================================================
        # Step 10: Return output path
        # ============================================================
        logger.info(f"Teacher confirmation applied. Overrides: {user_hints}")
        return output_path
    
    except Exception as e:
        logger.error(f"Teacher confirmation handling failed: {e}")
        return f"⚠️ Failed to apply teacher confirmation: {type(e).__name__}: {e}"


# Optional: Helper function to format confirmation choices for UI
def _format_confirmation_choices(uncertain_decisions: Dict[str, Decision]) -> Dict[str, str]:
    """
    Format the available choices for the teacher confirmation UI.
    
    Returns:
        {
            "A": "Accept engine's recommendation",
            "B": "Apply this feature",
            "C": "Do not apply this feature",
            "D": "Keep donor's original layout"
        }
    """
    return {
        "A": "Accept engine's recommendation",
        "B": "Apply this feature",
        "C": "Do not apply this feature",
        "D": "Keep donor's original layout"
    }


# Optional: Helper to parse teacher response from various input formats
def _parse_teacher_response(response: str) -> Tuple[str, Dict[str, bool]]:
    """
    Parse teacher response from text, button value, or radio selection.
    
    Returns:
        (action, overrides) where action is one of: "accept", "apply", "reject", "donor"
    """
    response = (response or "").strip().upper()
    
    if response in ["A", "ACCEPT", "ACCEPT_ENGINE"]:
        return ("accept", {})
    elif response in ["B", "APPLY", "OVERRIDE_APPLY"]:
        return ("apply", {})
    elif response in ["C", "REJECT", "OVERRIDE_REJECT"]:
        return ("reject", {})
    elif response in ["D", "DONOR", "KEEP_DONOR"]:
        return ("donor", {})
    else:
        return ("accept", {})


# ======================================================================================
# LEGACY COMPATIBILITY / DEBUG ONLY
# ======================================================================================
# These functions are DEPRECATED. They exist only for backward compatibility.
# DO NOT USE in new code. They will raise NotImplementedError.

def extract_docx_question_map(docx_path: str) -> Dict[str, Any]:
    """DEPRECATED. Use _extract_donor_with_styles + draft analysis instead."""
    raise NotImplementedError("DEPRECATED: Use intelligent extraction pipeline.")

def align_template_and_draft_questions(
    template_question_map: Dict[str, Any],
    draft_question_map: Dict[str, Any],
) -> Dict[str, Any]:
    """DEPRECATED. Use appropriateness rules engine instead."""
    raise NotImplementedError("DEPRECATED: Use _apply_appropriateness_rules.")

def _build_docx_alignment_debug_payload(
    template_question_map: Dict[str, Any],
    draft_question_map: Dict[str, Any],
    alignment_result: Dict[str, Any],
) -> Dict[str, Any]:
    """DEPRECATED. Debug function no longer needed."""
    raise NotImplementedError("DEPRECATED: Use blueprint debugging instead.")

def _inject_matched_questions_into_docx(
    source_docx_path: str,
    alignment_result: Dict[str, Any],
    draft_question_map: Dict[str, Any],
) -> str:
    """DEPRECATED. Use render_blueprint_to_docx instead."""
    raise NotImplementedError("DEPRECATED: Use intelligent rendering pipeline.")

# ... other legacy functions ...


# ======================================================================================
# SELF-TEST (only runs when file is executed directly, not when imported)
# ======================================================================================

if __name__ == "__main__":
    # Test basic helpers
    print("Running template_engine.py self-tests...")
    
    # Test _resolve_upload_path
    assert _resolve_upload_path(None) == ("", "")
    assert _resolve_upload_path({"/tmp/file.docx"}) == ("/tmp/file.docx", "file.docx")
    assert _resolve_upload_path("/tmp/file.docx") == ("/tmp/file.docx", "file.docx")
    
    # Test _safe_ext
    assert _safe_ext("document.docx") == ".docx"
    assert _safe_ext("file.PDF") == ".pdf"
    assert _safe_ext("noextension") == ""
    
    # Test _norm_spaces
    assert _norm_spaces("Hello\r\nWorld") == "Hello\nWorld"
    assert _norm_spaces("Too   many   spaces") == "Too many spaces"
    assert _norm_spaces("Line1\n\n\n\nLine2") == "Line1\n\nLine2"
    
    # Test _short
    assert _short("Short text", 100) == "Short text"
    truncated = _short("This is a very long string that needs truncation", 20)
    assert truncated.startswith("This is a very long")
    assert truncated.endswith("...")
    
    # Test _try_int
    assert _try_int("123") == 123
    assert _try_int(3.7) == 3
    assert _try_int("abc", 0) == 0
    
    # Test _try_float
    assert _try_float("123.45") == 123.45
    assert _try_float(5) == 5.0
    assert _try_float("abc", 0.0) == 0.0
    
    # Test _dominant
    from collections import Counter
    c = Counter(['a', 'a', 'b', 'b', 'b', 'c'])
    result = _dominant(c, 2)
    assert len(result) == 2
    assert result[0]["value"] == "b"
    assert result[0]["count"] == 3
    assert result[1]["value"] == "a"
    assert result[1]["count"] == 2
    
    print("✅ All basic helper tests passed!")




def diagnose_first_page_layout_contract(donor_docx_path: str, output_docx_path: Optional[str] = None) -> Dict[str, Any]:
    """
    Lightweight Colab/manual diagnostic for BS5 first-page layout contract support.

    Loads a donor DOCX, prints its first-page contract, renders a tiny draft-like
    blueprint through the normal renderer, saves the output DOCX, and returns
    simple pass/fail checks. No external services are required.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for this diagnostic")

    raw_profile = _extract_donor_with_styles(donor_docx_path)
    profile = _normalize_template_profile(raw_profile)
    contract = profile.get("first_page_layout_contract", {}) or {}

    print("## First Page Layout Contract")
    print(json.dumps(contract, indent=2, default=str))

    blueprint = {
        "title": "First Page Contract Diagnostic",
        "page_setup": profile.get("page_setup", {}),
        "appropriateness_decisions": {},
        "header_block": {"show_header": True, "header_lines": profile.get("header_footer", {}).get("header_texts", [])},
        "question_blocks": [
            {"q_no": "1", "marks_label": "(2 marks)", "question_lines": ["This is a diagnostic question."], "marks_value": 2}
        ],
        "visual_slots": [],
        "page_breaks": [],
    }

    rendered_path = render_blueprint_to_docx(blueprint, profile)
    if output_docx_path:
        import shutil
        shutil.copyfile(rendered_path, output_docx_path)
        rendered_path = output_docx_path

    rendered_doc = Document(rendered_path)
    rendered_section = rendered_doc.sections[0]
    donor_diff = bool(contract.get("different_first_page_header_footer", False))
    out_diff = bool(getattr(rendered_section, "different_first_page_header_footer", False))
    out_first_header = _container_text(rendered_section.first_page_header)
    out_default_header = _container_text(rendered_section.header)

    checks = {
        "output_path": rendered_path,
        "same_different_first_page_header_footer": out_diff == donor_diff,
        "first_page_header_not_duplicated_into_default_unless_donor_did": (
            not contract.get("first_page_header_text")
            or contract.get("first_page_header_text") == contract.get("default_header_text")
            or contract.get("first_page_header_text") not in out_default_header
        ),
        "default_header_remains_present_when_donor_had_one": (
            not contract.get("default_header_text") or bool(out_default_header)
        ),
        "question_heading_style_signature_detected": bool(contract.get("dominant_question_heading_style_signature")),
        "output_first_page_header_text": out_first_header,
        "output_default_header_text": out_default_header,
    }
    print("## Diagnostic Checks")
    print(json.dumps(checks, indent=2, default=str))
    return {"contract": contract, "checks": checks}
