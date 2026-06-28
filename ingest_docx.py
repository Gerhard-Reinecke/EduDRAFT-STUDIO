# ======================================================================================
# ingest_docx.py
# ======================================================================================
# Module: DOCX Ingestion & Draft Generation Engine
#
# System: EduDraft Studio (Marike App)
# Version: 4.4.10
#
# --------------------------------------------------------------------------------------
# PURPOSE
# --------------------------------------------------------------------------------------
# This module handles the ingestion, parsing, and transformation of teacher-uploaded
# DOCX files into structured, editable educational drafts.
#
# It extracts textual content (paragraphs and tables), reconstructs classroom-ready
# materials using LLM processing, and integrates seamlessly into the Drafts system
# with automatic versioning and preview generation.
#
# --------------------------------------------------------------------------------------
# CORE RESPONSIBILITIES
# --------------------------------------------------------------------------------------
# 1. DOCX Text Extraction
#    - Reads .docx files using python-docx
#    - Extracts paragraphs and table content
#    - Normalises whitespace and structure
#    - Applies character limits to protect token usage
#
# 2. Document Analysis
#    - Validates uploaded DOCX files
#    - Provides user feedback (file size, extracted content)
#    - Enables/disables generation workflow
#
# 3. Draft Reconstruction (LLM Integration)
#    - Converts extracted DOCX text into structured Markdown drafts
#    - Preserves assessment structure (questions, marks, headings)
#    - Supports memo/answer key generation
#
# 4. Draft Creation & Persistence
#    - Creates new Draft records in Supabase
#    - Saves Version 1 automatically
#    - Stores document content and metadata
#
# 5. Preview Generation
#    - Generates HTML preview using MathJax rendering
#    - Supports immediate visual validation of output
#
# --------------------------------------------------------------------------------------
# DEPENDENCIES
# --------------------------------------------------------------------------------------
# - python-docx → DOCX parsing
# - llm.py → Prompt construction and LLM interaction
# - exports.py → HTML preview rendering
# - auth.py → Session validation
# - rate_limit.py → Usage control
# - config.py → Supabase client
# - Gradio → UI interaction layer
#
# --------------------------------------------------------------------------------------
# DESIGN PRINCIPLES
# --------------------------------------------------------------------------------------
# - Teacher-first workflow (minimal friction, maximum clarity)
# - Safe token usage (controlled extraction limits)
# - Deterministic draft creation (Version 1 auto-save)
# - Clean separation between ingestion, generation, and persistence
#
# --------------------------------------------------------------------------------------
# NOTES
# --------------------------------------------------------------------------------------
# - Only .docx files are supported
# - Extraction is best-effort and may not perfectly preserve formatting
# - Large documents are truncated to protect LLM performance
# - This module mirrors PDF ingestion behavior for consistency
#
# ======================================================================================

import os
import re
import uuid
import tempfile
from datetime import datetime, timezone

import gradio as gr
from docx import Document  # python-docx

from llm import build_user_request, call_llm, split_sections, combine_doc_and_memo
from exports import build_mathjax_html
from auth import _require_session
from rate_limit import check_rate_limit
from config import supabase


# =============================
# DOCX INGESTION (Teacher-first)
# =============================
def extract_text_from_docx(docx_path: str, max_chars: int = 120_000) -> str:
    """
    Extracts text from a DOCX (paragraphs + tables). Best-effort.
    Truncates to max_chars to protect token usage.
    """
    if not docx_path or not os.path.exists(docx_path):
        return ""

    doc = Document(docx_path)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    cell_text = re.sub(r"\s+", " ", cell_text)
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    text = "\n\n".join(parts).strip()
    text = re.sub(r"\n{3,}", "\n\n", text)

    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[TRUNCATED]"
    return text


def action_analyze_docx(docx_file):
    """
    Called when a DOCX is uploaded.
    Returns:
      1) markdown status string
      2) gr.update(interactive=True/False) for the generate button
    """
    try:
        if not docx_file:
            return "No DOCX uploaded yet.", gr.update(interactive=False)

        if isinstance(docx_file, dict):
            docx_path = docx_file.get("name") or docx_file.get("path") or ""
            docx_name = os.path.basename(docx_path) if docx_path else "uploaded.docx"
        else:
            docx_path = str(docx_file)
            docx_name = os.path.basename(docx_path)

        if not docx_path or not os.path.exists(docx_path):
            return "⚠️ DOCX uploaded, but the file path was not found in the container.", gr.update(interactive=False)

        size_kb = os.path.getsize(docx_path) / 1024.0
        txt = extract_text_from_docx(docx_path)
        n = len(txt or "")

        if n == 0:
            msg = (
                f"📄 **DOCX detected**\n\n"
                f"- File: `{docx_name}`\n"
                f"- Size: **{size_kb:.1f} KB**\n"
                f"- Extracted chars: **0** ❌\n\n"
                f"This DOCX seems empty or text could not be read."
            )
            return msg, gr.update(interactive=False)

        msg = (
            f"📄 **DOCX detected**\n\n"
            f"- File: `{docx_name}`\n"
            f"- Size: **{size_kb:.1f} KB**\n"
            f"- Extracted chars: **{n}** ✅"
        )
        return msg, gr.update(interactive=True)

    except Exception as e:
        return f"⚠️ DOCX analysis failed: {type(e).__name__}: {e}", gr.update(interactive=False)


def docx_to_draft_prompt(docx_text: str, output_type: str, include_memo: bool, docx_mode: str = "Convert to editable draft") -> str:
    memo_line = "Yes" if include_memo else "No"
    mode = (docx_mode or "Convert to editable draft").strip()

    if mode == "Generate similar new paper":
        return f"""
You are given text extracted from an existing teacher DOCX exam/test/worksheet.

The teacher does NOT want a copy of the same paper.

The teacher wants a NEW paper that closely follows the uploaded document's:
- assessment type
- section structure
- question numbering
- mark allocation pattern
- difficulty level
- topic spread
- wording style
- student-facing instructions
- answer-space expectations

But the new paper must use DIFFERENT content, values, names, scenarios, examples, and answers.

Important limitations:
- The extracted DOCX text may be missing diagrams, graphs, tables, formula layout, headers, footers, or visual details.
- If a visual is clearly required, insert a structured placeholder like:
  [[VISUAL id="v1" kind="diagram" where="Q1" prompt="Describe exactly what diagram/graph/table is needed" notes="black-and-white, exam-ready" data="MISSING"]]
- Do not invent decorative images.
- Only create visuals that are educationally required by the question.

Your job:
- Create a clean classroom-ready Markdown draft with LaTeX where useful.
- Preserve the model paper's assessment structure, but not its exact question content.
- Keep marks realistic and aligned to the model.
- Include clear answer spaces where needed.
- If include memo is Yes, include an answer key/memo after the student paper.

Target output type: {output_type}
Include answer key/memo: {memo_line}

DOCX_TEXT_START
<<<
{docx_text}
>>>
DOCX_TEXT_END

Now produce the required output following the OUTPUT RULE exactly.
""".strip()

    return f"""
You are given the text extracted from a teacher's DOCX. The teacher wants an EDITABLE draft.
Your job:
- Reconstruct the document as a clean classroom-ready editable draft in Markdown with LaTeX.
- Preserve question numbering, marks, and headings as best as possible.
- If the DOCX is an exam/test/worksheet, keep the assessment structure.
- If some items are unclear, make minimal reasonable assumptions and list 1–3 short assumptions at the end.
Target output type: {output_type}
Include answer key/memo: {memo_line}
DOCX_TEXT_START
<<<
{docx_text}
>>>
DOCX_TEXT_END
Now produce the required output following the OUTPUT RULE exactly.
""".strip()


def action_generate_from_docx(
    supabase_session,
    docx_file,
    docx_draft_title: str,
    docx_mode: str,
    education_level,
    country, state_province,
    uni_country, university_name, faculty, module_code,
    year_level, course, typed_subject, course_stream, output_type,
    include_memo, model_name
):
    """
    Upload DOCX -> generate NEW editable draft + auto-save v1.
    Mirrors PDF behavior.
    """
    try:
        access_token, refresh_token, user_id, err = _require_session(supabase_session)
        if err:
            return "", "", "", None, f"❌ {err}", "", "DOCX upload", "", "", 0

        can_proceed, limit_msg = check_rate_limit(supabase_session, action="generate")
        if not can_proceed:
            return "", "", "", None, limit_msg, "", "DOCX upload", "", "", 0

        if not docx_file:
            return "", "", "", None, f"❌ Please upload a DOCX first. {limit_msg}", "", "DOCX upload", "", "", 0

        if isinstance(docx_file, dict):
            docx_path = docx_file.get("name") or docx_file.get("path") or ""
            docx_name = os.path.basename(docx_path) if docx_path else "uploaded.docx"
        else:
            docx_path = str(docx_file)
            docx_name = os.path.basename(docx_path)

        if not docx_path or not os.path.exists(docx_path):
            return "", "", "", None, f"❌ Uploaded DOCX path not found. {limit_msg}", "", "DOCX upload", "", "", 0

        if os.path.splitext(docx_name)[1].lower() != ".docx":
            return "", "", "", None, f"❌ That file is not a DOCX. {limit_msg}", "", "DOCX upload", "", "", 0

        docx_text = extract_text_from_docx(docx_path)
        if not (docx_text or "").strip():
            return "", "", "", None, (
                "❌ I couldn't extract any readable text from this DOCX.\n\n"
                f"File: {docx_name}\n"
                f"Extracted chars: {len(docx_text or '')}\n\n"
                f"{limit_msg}"
            ), "", "DOCX upload", "", "", 0

        # build effective course/subject
        if education_level == "University / Tertiary":
            eff_country = (uni_country or "").strip() or "Not specified"
            eff_state = ""
        else:
            eff_country = (country or "").strip() or "Not specified"
            eff_state = (state_province or "").strip()

        if education_level == "School (Primary / Secondary)":
            chosen = (course or "").strip()
        
            if chosen == "Other (type it)":
                ts = (typed_subject or "").strip()
                if not ts:
                    return "", "", "", None, f"❌ Subject is required. Please type your subject. {limit_msg}", "", "DOCX upload", "", "", 0
                effective_course = ts
            else:
                effective_course = chosen or "Not specified"
        else:
            effective_course = (module_code or "").strip() or "Not specified"
        
        curriculum_stream = (course_stream or "").strip()

        docx_mode = (docx_mode or "Convert to editable draft").strip()

        if docx_mode == "Generate similar new paper":
            instruction_text = (
                f"Generate a new similar exam/test paper using the uploaded DOCX as the model structure: {docx_name}"
            ).strip()
        else:
            instruction_text = (
                f"Generate an editable draft from the uploaded DOCX: {docx_name}"
            ).strip()

        teacher_context = build_user_request(
            instruction_text=instruction_text,
            education_level=education_level,
            country=eff_country,
            state_province=eff_state,
            year_level=year_level,
            course=effective_course,
            output_type=output_type,
            include_memo=include_memo,
            university_name=university_name,
            faculty=faculty,
            module_code=module_code,
            curriculum_stream=curriculum_stream,
        )

        prompt = teacher_context + "\n\n" + docx_to_draft_prompt(docx_text, output_type, include_memo, docx_mode)

        llm_text = call_llm(prompt, model_name, edit_mode=False)
        doc_md, ppt_outline, answer_key = split_sections(llm_text)
        combined_md = combine_doc_and_memo(doc_md, answer_key, include_memo)

        tmpdir = tempfile.mkdtemp()
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        html_path = os.path.join(tmpdir, f"Preview_{stamp}.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(build_mathjax_html(combined_md))

        now = datetime.now(timezone.utc).isoformat()
        title = (docx_draft_title or "").strip()
        if not title:
            title = os.path.splitext(docx_name)[0].strip() or "DOCX Draft"

        draft_id = str(uuid.uuid4())

        draft_row = {
            "id": draft_id,
            "user_id": user_id,
            "title": title,
            "subject": (effective_course or "").strip() or None,
            "curriculum_stream": (curriculum_stream or "").strip() or None,   # ✅ ADD
            "education_level": (education_level or "").strip() or None,
            "country": (eff_country or "").strip() or None,
            "state_province": (eff_state or "").strip() or None,
            "year_level": (year_level or "").strip() or None,
            "course": (effective_course or "").strip() or None,
            "created_at": now,
            "updated_at": now,
        }
        supabase.table("drafts").insert(draft_row).execute()

        version_row = {
            "draft_id": draft_id,
            "version": 1,
            "doc_md": combined_md or "",
            "ppt_outline": ppt_outline or "",
            "created_at": now,
        }
        supabase.table("draft_versions").insert(version_row).execute()

        status = (
            "✅ DOCX converted into a NEW editable draft and auto-saved as Version 1.\n"
            f"Draft name: {title}\n"
            f"{limit_msg}\n\n"
            "You can now edit it and Save NEW versions anytime."
        )

        return (
            f"DOCX used: {docx_name}",
            combined_md,
            ppt_outline,
            html_path,
            status,
            instruction_text,
            "DOCX upload",
            f"DOCX used: {docx_name}",
            draft_id,
            1
        )

    except Exception as e:
        from llm import safe_err
        return "", "", "", None, safe_err("DOCX → Draft failed.", e), "", "DOCX upload", "", "", 0
